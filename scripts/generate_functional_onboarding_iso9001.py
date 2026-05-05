from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DOCX = ROOT / "plantilla_iso_9001_proceso.docx"
LOGO_PATH = ROOT / "logo_simplia.png"
DOCX_OUT = ROOT / "Onboarding_Funcional_Simplia_Chatbot_ISO9001.docx"

DOC_CODE = "PRO-ONB-FUNC-001"
DOC_VERSION = "1.0"
ISSUE_DATE = "04/05/2026"
NEXT_REVIEW = "04/11/2026"
DOC_STATUS = "Borrador / Vigente al aprobarse"

BLANK = "______________________________"
LONG_BLANK = "____________________________________________________________"

COLORS = {
    "navy": "0f2344",
    "blue": "274690",
    "soft_blue": "eaf1ff",
    "light": "f8fafc",
    "green": "0a9b6f",
    "yellow": "fef3c7",
    "orange": "f59e0b",
    "red": "ef4444",
    "slate": "64748b",
    "line": "d9e2ef",
    "white": "ffffff",
}

GUIDE_LINKS = {
    "cuentas": "https://docs.google.com/document/d/13LigK2TaIFLSGggJfEXY2PmioYfw8LCE/edit?usp=sharing&ouid=102450533900336054092&rtpof=true&sd=true",
    "plantillas": "https://drive.google.com/file/d/1yTxQ0lrj5Yp1vBQU7mQzv1HnURZZmz_R/view?usp=sharing",
    "ubicaciones": "https://drive.google.com/file/d/1Mh-6TUZY9uJL00DKoSiaiwFTX4u5z_0Q/view?usp=sharing",
    "flujo": "https://drive.google.com/file/d/1IE1pB5VgpOoc4UoCzKh8PvRVtP4YWVsb/view?usp=sharing",
    "levantamiento": "https://docs.google.com/document/d/14X8QjAzFkEX0bCDdB2frfdYNWIKJlOlm/edit?usp=sharing&ouid=102450533900336054092&rtpof=true&sd=true",
}

URGENCY_FILL = {
    "Alta": "fee2e2",
    "Media": "fef3c7",
    "Baja": "dcfce7",
}


ONBOARDING_SECTIONS = [
    {
        "title": "1. WhatsApp Business y perfil del canal",
        "urgency": "Media",
        "guide": None,
        "need": "Número nuevo o apto, SIM activa, capacidad de recibir SMS, logo, nombre público, descripción, horarios, sitio web y política de privacidad.",
        "how": "Completar los datos del canal y pegar links de capturas, logo o documentos. No avanzar si el número no puede verificarse.",
        "deliverable": "Link a carpeta con evidencias de WhatsApp, logo y datos del perfil.",
    },
    {
        "title": "2. Cuentas y accesos",
        "urgency": "Alta",
        "guide": GUIDE_LINKS["cuentas"],
        "need": "Correos, usuarios, permisos o accesos necesarios para conectar canales, agenda, CRM, Chatwoot, n8n y documentos.",
        "how": "Registrar cuenta, encargado y método seguro de entrega. No escribir contraseñas en este Word; usar gestor, invitación o link protegido.",
        "deliverable": "Link compartido con accesos o confirmación de invitaciones realizadas.",
    },
    {
        "title": "3. Información de la empresa, servicios y sedes",
        "urgency": "Alta",
        "guide": GUIDE_LINKS["ubicaciones"],
        "need": "Servicios, costos, beneficios, restricciones, horarios, sedes, direcciones, links de Maps y condiciones comerciales.",
        "how": "Pegar el link del documento oficial que el bot usará para responder. La información debe estar aprobada por la empresa.",
        "deliverable": "Documento o carpeta con información institucional, servicios, precios y ubicaciones.",
    },
    {
        "title": "4. Plantillas de conversación",
        "urgency": "Alta",
        "guide": GUIDE_LINKS["plantillas"],
        "need": "Mensajes exactos para bienvenida, información, interés, solicitud de datos, confirmación de cita, cierre y derivación humana.",
        "how": "Escribir los textos aprobados o pegar un link donde estén. Indicar tono, emojis permitidos y cuándo se usa cada plantilla.",
        "deliverable": "Documento compartido con plantillas aprobadas.",
    },
    {
        "title": "5. Flujo conversacional de inicio a fin",
        "urgency": "Alta",
        "guide": GUIDE_LINKS["flujo"],
        "need": "Cómo atiende hoy la empresa: desde el primer mensaje hasta información, dudas, agenda, venta o seguimiento humano.",
        "how": "Pegar el flujo actual, ejemplos de conversaciones o notas de cómo debe actuar el bot en cada caso.",
        "deliverable": "Link al documento de flujo o carpeta de ejemplos.",
    },
    {
        "title": "6. Etiquetas y categorización de leads",
        "urgency": "Alta",
        "guide": None,
        "need": "Confirmar qué representa cada etiqueta y cuáles las pone el bot o una persona.",
        "how": "Revisar la tabla de etiquetas de este documento y completar observaciones si la empresa usa reglas especiales.",
        "deliverable": "Aprobación de etiquetas y reglas de uso.",
    },
    {
        "title": "7. Datos para agendar",
        "urgency": "Alta",
        "guide": None,
        "need": "Datos que el bot debe guardar para crear cita: nombre, fecha, hora, agencia, celular, correo, ciudad, edad u otros.",
        "how": "Marcar qué datos son obligatorios, cuáles opcionales y en qué momento se piden.",
        "deliverable": "Lista aprobada de datos obligatorios y opcionales.",
    },
    {
        "title": "8. Score y vocabulario del negocio",
        "urgency": "Media",
        "guide": None,
        "need": "Palabras y frases que indican interés, compra, agenda, asesor, rechazo, fuera de negocio o riesgo.",
        "how": "Completar vocabulario propio de la empresa para que el score y la intención funcionen correctamente.",
        "deliverable": "Lista de palabras/frases y ejemplos reales.",
    },
    {
        "title": "9. Ejemplos reales para entrenamiento",
        "urgency": "Media",
        "guide": None,
        "need": "Capturas o exportaciones de conversaciones reales con clientes: buenos leads, dudas, ventas, rechazos y casos fuera de negocio.",
        "how": "Subir 80 a 100 ejemplos si es posible. Si no se tiene esa cantidad, dejar los mejores ejemplos disponibles.",
        "deliverable": "Carpeta compartida con capturas o exportaciones.",
    },
    {
        "title": "10. Preguntas frecuentes",
        "urgency": "Media",
        "guide": None,
        "need": "Preguntas comunes y respuestas aprobadas sobre servicios, costos, sedes, horarios, pagos, beneficios y restricciones.",
        "how": "Completar la tabla FAQ o pegar link del documento ya aprobado.",
        "deliverable": "Documento de preguntas frecuentes.",
    },
    {
        "title": "11. Campañas de retoma 30 min y 3 h",
        "urgency": "Alta",
        "guide": GUIDE_LINKS["levantamiento"],
        "need": "Script exacto de retoma a los 30 minutos y script exacto de retoma a las 3 horas.",
        "how": "Definir texto, tono, emojis, medios permitidos y qué pasa si el lead responde o no responde.",
        "deliverable": "Scripts aprobados para campañas automáticas.",
    },
]


def rgb(hex_color: str) -> RGBColor:
    value = hex_color.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def clear_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    for style_name, size, color in [
        ("Heading 1", 15, COLORS["blue"]),
        ("Heading 2", 12.5, COLORS["navy"]),
        ("Heading 3", 10.5, COLORS["green"]),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"{DOC_CODE} | v{DOC_VERSION} | {DOC_STATUS}")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["slate"])

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Simplia Chatbot - Onboarding Funcional ISO 9001")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["slate"])


def set_cell_shading(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def set_cell_width(cell, width: float) -> None:
    cell.width = Inches(width)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color: str = COLORS["line"]) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement("w:" + edge)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    table._tbl.tblPr.append(borders)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), COLORS["blue"])
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)

    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def clear_cell(cell) -> None:
    cell.text = ""
    for paragraph in cell.paragraphs:
        paragraph.text = ""


def write_cell(cell, value, *, bold: bool = False, color: str | None = None, size: float = 8.0) -> None:
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if isinstance(value, dict):
        add_hyperlink(paragraph, value["label"], value["url"])
    else:
        for idx, part in enumerate(str(value).split("\n")):
            if idx:
                paragraph.add_run().add_break()
            run = paragraph.add_run(part)
            run.bold = bold
            run.font.name = "Arial"
            run.font.size = Pt(size)
            if color:
                run.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def link(label: str, url: str | None):
    return {"label": label, "url": url} if url else "No aplica"


def add_table(
    document: Document,
    headers: list[str],
    rows: list[list[object]],
    widths: list[float] | None = None,
    *,
    font_size: float = 7.8,
    urgency_col: int | None = None,
) -> None:
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])

    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        write_cell(cell, header, bold=True, color=COLORS["white"], size=8.2)
        set_cell_shading(cell, COLORS["blue"])
        if widths and col < len(widths):
            set_cell_width(cell, widths[col])

    for row_idx, row in enumerate(rows, start=1):
        fill = COLORS["light"] if row_idx % 2 == 0 else COLORS["white"]
        for col, value in enumerate(row):
            cell = table.cell(row_idx, col)
            write_cell(cell, value, size=font_size)
            if urgency_col is not None and col == urgency_col:
                set_cell_shading(cell, URGENCY_FILL.get(str(value), fill))
            else:
                set_cell_shading(cell, fill)
            if widths and col < len(widths):
                set_cell_width(cell, widths[col])
    document.add_paragraph()


def add_paragraph(document: Document, text: str, *, bold: bool = False, color: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9.2)
    if color:
        run.font.color.rgb = rgb(color)


def add_note_box(document: Document, title: str, text: str, fill: str = COLORS["soft_blue"]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(8.8)
    run.font.color.rgb = rgb(COLORS["navy"])
    paragraph.add_run().add_break()
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(8.4)
    run.font.color.rgb = rgb(COLORS["navy"])
    document.add_paragraph()


def add_cover(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        paragraph.add_run().add_picture(str(LOGO_PATH), width=Inches(1.45))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ONBOARDING FUNCIONAL - SIMPLIA CHATBOT")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(19)
    run.font.color.rgb = rgb(COLORS["blue"])

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Documento para llenar la información necesaria para que el agente funcione correctamente de inicio a fin")
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.color.rgb = rgb(COLORS["slate"])

    add_table(
        document,
        ["Campo", "Valor"],
        [
            ["Código", DOC_CODE],
            ["Versión", DOC_VERSION],
            ["Fecha de emisión", ISSUE_DATE],
            ["Próxima revisión", NEXT_REVIEW],
            ["Estado", DOC_STATUS],
            ["Empresa cliente", BLANK],
            ["Encargado principal de la empresa", BLANK],
            ["Correo / teléfono", BLANK],
        ],
        widths=[2.1, 5.2],
    )

    add_note_box(
        document,
        "Uso del documento",
        "Este documento se completa por secciones. En cada apartado la empresa debe dejar el link compartido con la información, el encargado, la fecha compromiso y cualquier observación. No está pensado para explicar detalles técnicos; sirve para recoger lo que el bot necesita para operar bien.",
    )


def add_inputs_outputs(document: Document) -> None:
    document.add_heading("1. Entradas y salidas del onboarding", level=1)
    add_table(
        document,
        ["Qué entrega la empresa", "Para qué sirve en el bot", "Resultado esperado"],
        [
            ["Datos de WhatsApp, perfil y canal", "Conectar y presentar correctamente el agente.", "Canal listo para pruebas."],
            ["Servicios, sedes, costos, beneficios y restricciones", "Responder preguntas reales sin improvisar.", "Base de respuestas aprobada."],
            ["Plantillas y ejemplos de conversación", "Mantener el tono y guiar cada escenario.", "Flujo conversacional usable."],
            ["Datos de agenda", "Pedir la información correcta antes de crear cita.", "Citas completas y trazables."],
            ["Vocabulario y score", "Detectar intención y calidad del lead.", "Leads clasificados en Frío, Tibio o Caliente."],
            ["Scripts de retoma", "Recuperar conversaciones que se quedan sin respuesta.", "Campañas de 30 min y 3 h listas."],
        ],
        widths=[2.35, 2.65, 2.2],
        font_size=7.6,
    )


def add_section_summary(document: Document) -> None:
    document.add_heading("2. Secciones que debe completar la empresa", level=1)
    rows = []
    for item in ONBOARDING_SECTIONS:
        rows.append(
            [
                item["title"],
                item["need"],
                link("Documento guía", item["guide"]),
                "Pegar link aquí: " + BLANK,
                "Encargado: " + BLANK,
                item["urgency"],
                "Fecha: dd/mm/aaaa",
            ]
        )
    add_table(
        document,
        ["Sección", "Qué se necesita", "Guía", "Link de información", "Encargado empresa", "Urgencia", "Fecha"],
        rows,
        widths=[1.35, 2.35, 0.9, 1.35, 1.1, 0.75, 0.75],
        font_size=6.8,
        urgency_col=5,
    )


def add_detailed_sections(document: Document) -> None:
    document.add_heading("3. Levantamiento por sección", level=1)
    for item in ONBOARDING_SECTIONS:
        document.add_heading(item["title"], level=2)
        add_table(
            document,
            ["Campo", "Completar"],
            [
                ["Qué se necesita", item["need"]],
                ["Cómo llenarlo", item["how"]],
                ["Documento guía", link("Abrir guía", item["guide"])],
                ["Link compartido de la empresa", LONG_BLANK],
                ["Encargado de la empresa", BLANK],
                ["Fecha compromiso", "dd/mm/aaaa"],
                ["Urgencia", item["urgency"]],
                ["Estado", "Pendiente / En revisión / Completo / No aplica"],
                ["Entregable esperado", item["deliverable"]],
                ["Observaciones", LONG_BLANK],
            ],
            widths=[1.9, 5.4],
            font_size=7.5,
            urgency_col=None,
        )


def add_labels(document: Document) -> None:
    document.add_heading("4. Etiquetas y categorización aprobada", level=1)
    add_note_box(
        document,
        "Regla simple",
        "Estas son las etiquetas operativas del proyecto. Algunas las puede asignar el bot y otras se asignan manualmente por una persona.",
    )
    add_table(
        document,
        ["Etiqueta", "Quién la asigna", "Qué significa", "Qué debe pasar"],
        [
            ["bienvenida", "Bot", "Primer contacto o lead sin intención clara todavía.", "El bot saluda, pregunta qué necesita y espera respuesta."],
            ["solicita_informacion", "Bot", "El lead pide información, costos, servicios, beneficios, sedes u horarios.", "El bot responde con información aprobada e invita a avanzar."],
            ["interesado", "Bot", "El lead muestra intención de avanzar o agendar.", "El bot pide datos necesarios y valida disponibilidad."],
            ["desinteresado", "Bot", "El lead rechaza, insulta, está fuera del negocio o manda contenido irrelevante.", "El bot cierra de forma educada y no insiste."],
            ["cita_agendada", "Bot", "Cita creada automáticamente por el bot.", "Se guardan fecha, hora, agencia y datos de contacto."],
            ["seguimiento_humano", "Bot", "El bot ya no debe seguir respondiendo o el caso requiere persona.", "Se marca para que un humano tome el caso."],
            ["venta_exitosa", "Manual", "La venta u operación se concretó.", "La persona debe llenar `monto_operacion` y `fecha_monto_operacion`."],
            ["cita_agendado_humano", "Manual", "Una persona agenda o modifica una cita manualmente.", "La persona debe llenar `responsable` con su nombre."],
        ],
        widths=[1.35, 0.9, 2.65, 2.35],
        font_size=7.2,
    )


def add_attributes(document: Document) -> None:
    document.add_heading("5. Datos que se guardan del lead", level=1)
    add_table(
        document,
        ["Atributo", "Tipo", "Quién lo llena", "Cuándo se usa"],
        [
            ["nombre_completo", "string", "Bot", "Cuando el lead entrega su nombre o se necesita para agenda."],
            ["fecha_visita", "string", "Bot", "Cuando se agenda una cita."],
            ["hora_visita", "string", "Bot", "Cuando se agenda una cita."],
            ["agencia", "string", "Bot", "Cuando el lead elige sede/agencia."],
            ["celular", "string", "Bot", "Desde WhatsApp o cuando el lead lo escribe."],
            ["correo", "string", "Bot", "Cuando el lead lo entrega o el flujo lo solicita."],
            ["campana", "string", "Bot", "Para identificar campaña o fuente cuando aplique."],
            ["ciudad", "string", "Bot", "Cuando el lead la entrega o se necesita para segmentar."],
            ["edad", "string", "Bot", "Cuando el negocio la necesita y el lead la entrega."],
            ["canal", "string", "Bot", "Canal por donde entró el lead."],
            ["score_interes", "number", "Bot", "Puntaje acumulado de calidad e intención."],
            ["agente", "checkbox", "Manual", "Se marca cuando el bot ya no debe responder a ese lead."],
            ["monto_operacion", "string", "Manual", "Se llena cuando se marca `venta_exitosa`."],
            ["fecha_monto_operacion", "Date", "Manual", "Se llena cuando se marca `venta_exitosa`."],
            ["responsable", "string", "Manual", "Se llena cuando se marca `cita_agendado_humano`."],
        ],
        widths=[1.7, 0.9, 0.95, 3.7],
        font_size=7.2,
    )


def add_booking_and_scripts(document: Document) -> None:
    document.add_heading("6. Datos para agenda y scripts", level=1)
    add_table(
        document,
        ["Dato", "¿Obligatorio?", "Cuándo se pide", "Texto o regla aprobada"],
        [
            ["nombre_completo", "Sí / No", "Inicio / Al agendar", LONG_BLANK],
            ["celular", "Sí / No", "Inicio / Al agendar", LONG_BLANK],
            ["correo", "Sí / No", "Inicio / Al agendar", LONG_BLANK],
            ["agencia", "Sí / No", "Al agendar", LONG_BLANK],
            ["fecha_visita", "Sí / No", "Al agendar", LONG_BLANK],
            ["hora_visita", "Sí / No", "Al agendar", LONG_BLANK],
            ["ciudad", "Sí / No", "Inicio / Al agendar", LONG_BLANK],
            ["edad", "Sí / No", "Inicio / Al agendar", LONG_BLANK],
        ],
        widths=[1.35, 1.0, 1.5, 3.4],
        font_size=7.3,
    )

    add_table(
        document,
        ["Script", "Cuándo se envía", "Texto exacto aprobado", "Si responde", "Si no responde"],
        [
            ["Retoma 30 min", "Lead en `bienvenida`, `solicita_informacion` o `interesado` sin responder por 30 minutos.", LONG_BLANK, "Continuar flujo según intención.", "Esperar retoma de 3 horas."],
            ["Retoma 3 h", "Lead sigue sin responder luego de 3 horas.", LONG_BLANK, "Continuar flujo según intención.", "Pasar a `seguimiento_humano`."],
        ],
        widths=[1.1, 2.0, 2.25, 1.0, 1.15],
        font_size=7.0,
    )


def add_score_and_vocab(document: Document) -> None:
    document.add_heading("7. Score y vocabulario", level=1)
    add_table(
        document,
        ["Nivel", "Rango", "Interpretación"],
        [
            ["FRÍO", "Menor a 45 o sin puntaje", "Señal inicial, mensaje genérico o lead todavía sin intención clara."],
            ["TIBIO", "45 a 69", "Tiene fit o señales concretas: precio, disponibilidad, servicio específico o continuidad."],
            ["CALIENTE", "70 o más", "Alta intención: quiere comprar, reservar, agendar, pagar o hablar con asesor."],
        ],
        widths=[1.1, 1.7, 4.4],
        font_size=7.5,
    )
    add_table(
        document,
        ["Vocabulario", "Palabras/frases del negocio", "Ejemplos reales"],
        [
            ["Servicios / productos", LONG_BLANK, LONG_BLANK],
            ["Precios / costos", LONG_BLANK, LONG_BLANK],
            ["Agenda / cita", LONG_BLANK, LONG_BLANK],
            ["Asesor humano", LONG_BLANK, LONG_BLANK],
            ["Rechazo / desinterés", LONG_BLANK, LONG_BLANK],
            ["Fuera del negocio", LONG_BLANK, LONG_BLANK],
            ["Riesgo / fraude / insultos", LONG_BLANK, LONG_BLANK],
        ],
        widths=[1.6, 2.85, 2.85],
        font_size=7.1,
    )


def add_checklist(document: Document) -> None:
    document.add_heading("8. Checklist final antes de construir", level=1)
    items = [
        "WhatsApp y perfil del canal completos.",
        "Cuentas y accesos compartidos por método seguro.",
        "Servicios, sedes, costos y restricciones aprobados.",
        "Plantillas principales aprobadas.",
        "Etiquetas y responsabilidades entendidas.",
        "Datos de agenda definidos como obligatorios u opcionales.",
        "Score con 3 niveles y vocabulario del negocio completo.",
        "Ejemplos reales y FAQs compartidos.",
        "Scripts de 30 min y 3 h aprobados.",
        "Cliente aprueba que la información está lista para implementación.",
    ]
    add_table(
        document,
        ["#", "Verificación", "Resultado", "Observaciones"],
        [[str(i), item, "OK / No OK / N.A.", BLANK] for i, item in enumerate(items, start=1)],
        widths=[0.35, 4.6, 1.1, 1.3],
        font_size=7.2,
    )


def add_control(document: Document) -> None:
    document.add_heading("9. Control de cambios y aprobación", level=1)
    add_table(
        document,
        ["Versión", "Fecha", "Cambio", "Responsable"],
        [
            ["1.0", ISSUE_DATE, "Onboarding funcional reorganizado por secciones rellenables para implementación del agente.", "Simplia"],
            ["1.x", "dd/mm/aaaa", "Resumen del cambio.", BLANK],
        ],
        widths=[0.8, 1.1, 4.5, 0.9],
        font_size=7.5,
    )
    add_table(
        document,
        ["Elaborado por", "Revisado por", "Aprobado por"],
        [
            ["Simplia / Responsable onboarding", "Encargado de la empresa", "Responsable autorizado"],
            ["Firma / fecha", "Firma / fecha", "Firma / fecha"],
        ],
        widths=[2.4, 2.4, 2.4],
        font_size=7.8,
    )


def build_document() -> None:
    document = Document(TEMPLATE_DOCX)
    clear_body(document)
    configure_document(document)
    add_cover(document)
    add_inputs_outputs(document)
    add_section_summary(document)
    add_detailed_sections(document)
    add_labels(document)
    add_attributes(document)
    add_booking_and_scripts(document)
    add_score_and_vocab(document)
    add_checklist(document)
    add_control(document)
    document.save(DOCX_OUT)


if __name__ == "__main__":
    build_document()
    print(f"Documento generado: {DOCX_OUT}")
