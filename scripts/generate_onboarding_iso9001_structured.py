from __future__ import annotations

import os
import re
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "ONBOARDING TECNICO -SIMPLIA CHATBOT .docx"
TEMPLATE_DOCX = ROOT / "plantilla_iso_9001_proceso.docx"
LOGO_PATH = ROOT / "logo_simplia.png"
ASSET_DIR = ROOT / "Onboarding_Tecnico_Simplia_Chatbot_assets"
DOCX_OUT = ROOT / "Onboarding_Tecnico_Simplia_Chatbot_ISO9001.docx"
PDF_OUT = ROOT / "Onboarding_Tecnico_Simplia_Chatbot_ISO9001.pdf"

DOC_CODE = "PRO-ONB-CHATBOT-001"
DOC_VERSION = "1.0"
DOC_STATUS = "Vigente/Borrador"
ISSUE_DATE = "04/05/2026"
NEXT_REVIEW = "04/11/2026"

COLORS = {
    "blue": "274690",
    "navy": "0f2344",
    "slate": "64748b",
    "line": "d9e2ef",
    "light": "f8fafc",
    "soft_blue": "eaf1ff",
    "green": "0a9b6f",
    "orange": "f59e0b",
    "red": "ef4444",
}


@dataclass
class Block:
    kind: str
    text: str = ""
    images: list[int] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)


@dataclass
class OnboardingSection:
    title: str
    blocks: list[Block] = field(default_factory=list)


def rgb(hex_color: str) -> tuple[int, int, int]:
    value = hex_color.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def clean_text(text: str) -> str:
    replacements = {
        "cuneta": "cuenta",
        "antiguedad": "antigüedad",
        "contrasena": "contraseña",
        "dirigmos": "dirigimos",
        "aplasatr": "aplastar",
        "coloar": "colocar",
        "fcebook": "Facebook",
        "Facebok": "Facebook",
        "facebookk": "Facebook",
        "Ahroa": "Ahora",
        "tegamos": "tengamos",
        "configuarcion": "configuración",
        "configuracion": "configuración",
        "Configuracion": "Configuración",
        "CONFIGURACION": "CONFIGURACIÓN",
        "verificacion": "verificación",
        "VERIFICACION": "VERIFICACIÓN",
        "soluciionarse": "solucionarse",
        "tavez": "través",
        "Incio": "Inicio",
        "inicar": "iniciar",
        "javascrpt": "JavaScript",
        "COnfiguarcion": "configuración",
        "Messanger": "Messenger",
        "adress": "address",
        "OpenAi": "OpenAI",
    }
    value = text.replace("\u00a0", " ").strip()
    for wrong, right in replacements.items():
        value = value.replace(wrong, right)
    value = re.sub(r"[ \t]+", " ", value)
    value = value.replace(" :", ":")
    return value


def image_font(size: int, bold: bool = False):
    for path in [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]:
        if os.path.exists(path):
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def extract_blocks_and_images(source: Document) -> tuple[list[Block], dict[int, Path]]:
    ASSET_DIR.mkdir(exist_ok=True)

    blocks: list[Block] = []
    images: dict[int, Path] = {}
    img_idx = 0
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    for child in source.element.body.iterchildren():
        if child.tag == qn("w:p"):
            text = clean_text("".join(node.text or "" for node in child.iter(qn("w:t"))))
            block_images = []
            for blip in child.xpath(".//a:blip"):
                img_idx += 1
                rid = blip.get(qn("r:embed"))
                rel = source.part.rels.get(rid)
                if not rel:
                    continue
                content_type = rel.target_part.content_type
                ext = ".jpg" if "jpeg" in content_type else ".png"
                path = ASSET_DIR / f"seccion_evidencia_{img_idx:03d}{ext}"
                path.write_bytes(rel.target_part.blob)
                images[img_idx] = path
                block_images.append(img_idx)
            if text or block_images:
                blocks.append(Block(kind="p", text=text, images=block_images))
        elif child.tag == qn("w:tbl"):
            rows = []
            for tr in child.findall(".//w:tr", ns):
                cells = []
                for tc in tr.findall("./w:tc", ns):
                    cell_text = clean_text(" / ".join(t.text or "" for t in tc.findall(".//w:t", ns)))
                    cells.append(cell_text)
                if cells:
                    rows.append(cells)
            if rows:
                blocks.append(Block(kind="table", rows=rows))
    return blocks, images


SECTION_STARTS = [
    "1. Creación del Portafolio en Facebook Business Manager",
    "2. Creación de la App en Meta for Developers",
    "3. Creación y Conexión de Railway",
    "4. Creación de la Instancia de n8n",
    "5. Creación de la Instancia de Chatwoot",
    "5.5. Cambio de configuración en railway",
    "6. Configuración del Inbox en Chatwoot",
    "Configuración para messenger/Instagram",
    "FACEBOOK VERIFICACIÓN NEGOCIO",
    "CONFIGURACIÓN INSTAGRAM",
    "7. Integración con OpenAI",
    "8. Reinicio de n8n o Chatwoot",
    "9. Configuración n8n nodo drive",
    "10. configuración n8n nodo gmail",
    "11. Configuración plantillas de respuesta",
    "12.- Control remoto de pc",
    "Agregar Métodos de pago",
    "CREDENCIALES Cuentas",
]


def normalize_title(text: str) -> str:
    title = clean_text(text)
    title = title.replace(" WHATSAPP", " WhatsApp")
    title = title.replace("CHATWOOT", "Chatwoot")
    title = title.replace("OpenAi", "OpenAI")
    title = title.replace("FACEBOOK VERIFICACIÓN NEGOCIO", "Facebook - verificación de negocio")
    title = title.replace("CONFIGURACIÓN INSTAGRAM", "Configuración Instagram")
    title = title.replace("messenger/Instagram", "Messenger / Instagram")
    return title


def is_section_start(text: str, current_title: str | None = None) -> bool:
    if not text:
        return False
    normalized = clean_text(text)
    if current_title and current_title.startswith("Agregar Métodos de pago") and not normalized.startswith("CREDENCIALES Cuentas"):
        return False
    normalized_lower = normalized.lower()
    return any(normalized_lower.startswith(prefix.lower()) for prefix in SECTION_STARTS)


def group_sections(blocks: list[Block]) -> list[OnboardingSection]:
    sections: list[OnboardingSection] = []
    current: OnboardingSection | None = None
    for block in blocks:
        if block.kind == "p" and block.text in {"ONBOARDING", "Documento de requisitos"}:
            continue
        if block.kind == "p" and is_section_start(block.text, current.title if current else None):
            current = OnboardingSection(title=normalize_title(block.text))
            sections.append(current)
            continue
        if current is None:
            continue
        current.blocks.append(block)
    return sections


def section_fields(section: OnboardingSection) -> list[list[str]]:
    rows = [["Campo a completar", "Valor", "Observaciones"]]
    seen = set()
    for block in section.blocks:
        if block.kind != "p":
            continue
        text = block.text
        if "____________________" in text:
            parts = text.split("____________________")[:-1]
            for part in parts:
                label = part.strip(" :/-")
                if ":" in label:
                    label = label.split(":")[-1].strip()
                label = label[-80:].strip(" :/-")
                if not label:
                    label = "Dato requerido"
                if label not in seen:
                    seen.add(label)
                    rows.append([label, "____________________", ""])
    if len(rows) == 1:
        rows.append(["Observaciones del cliente", "____________________", ""])
        rows.append(["Responsable de validación", "____________________", ""])
    return rows


def make_process_map() -> Path:
    path = ASSET_DIR / "mapa_proceso_onboarding.png"
    width, height = 1800, 520
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = image_font(34, True)
    box_font = image_font(23, True)
    small_font = image_font(17)
    draw.text((60, 38), "Mapa del proceso de onboarding técnico Simplia Chatbot", font=title_font, fill=rgb(COLORS["navy"]))
    steps = [
        ("1", "Levantamiento", "Empresa,\ncanales y accesos"),
        ("2", "Cuentas base", "Meta, Railway,\nn8n y Chatwoot"),
        ("3", "Canales", "WhatsApp,\nMessenger e Instagram"),
        ("4", "IA y flujos", "OpenAI,\nDrive y Gmail"),
        ("5", "Pruebas", "Mensajes,\nwebhooks y evidencias"),
        ("6", "Cierre", "Credenciales,\nchecklist y aprobación"),
    ]
    x, y = 70, 165
    box_w, box_h, gap = 245, 180, 42
    for idx, (num, name, desc) in enumerate(steps):
        bx = x + idx * (box_w + gap)
        draw.rounded_rectangle((bx, y, bx + box_w, y + box_h), radius=22, fill=rgb(COLORS["soft_blue"]), outline=rgb(COLORS["blue"]), width=3)
        draw.ellipse((bx + 18, y + 18, bx + 68, y + 68), fill=rgb(COLORS["blue"]))
        draw.text((bx + 43, y + 43), num, font=image_font(22, True), fill="white", anchor="mm")
        draw.text((bx + 82, y + 26), name, font=box_font, fill=rgb(COLORS["navy"]))
        draw.multiline_text((bx + 26, y + 88), desc, font=small_font, fill=rgb(COLORS["slate"]), spacing=6)
        if idx < len(steps) - 1:
            ax = bx + box_w + 7
            ay = y + box_h // 2
            draw.line((ax, ay, ax + gap - 14, ay), fill=rgb(COLORS["blue"]), width=4)
            draw.polygon([(ax + gap - 14, ay - 10), (ax + gap + 4, ay), (ax + gap - 14, ay + 10)], fill=rgb(COLORS["blue"]))
    draw.text((70, 420), "Salida esperada: chatbot conectado, canales probados, evidencias guardadas y datos de operación completos.", font=small_font, fill=rgb(COLORS["navy"]))
    image.save(path)
    return path


def set_cell_shading(cell, fill: str):
    props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def set_table_borders(table, color="d9e2ef"):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement("w:" + edge)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    table._tbl.tblPr.append(borders)


def set_cell_text(cell, text, bold=False, color="0f2344"):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(8.7)
    run.font.color.rgb = RGBColor(*rgb(color))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_doc_table(doc: Document, rows, widths=None, header=True):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            is_header = header and i == 0
            cell = table.cell(i, j)
            set_cell_text(cell, value, bold=is_header, color=("ffffff" if is_header else COLORS["navy"]))
            set_cell_shading(cell, COLORS["blue"] if is_header else ("f8fafc" if i % 2 == 0 else "ffffff"))
            if widths and j < len(widths):
                cell.width = Inches(widths[j])
    return table


def add_heading(doc: Document, text: str, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(*rgb(COLORS["blue"] if level == 1 else COLORS["navy"]))
    return paragraph


def add_para(doc: Document, text: str, style=None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.1)
    run.font.color.rgb = RGBColor(*rgb(COLORS["navy"]))
    return paragraph


def add_bullets(doc: Document, items):
    for item in items:
        add_para(doc, item, style="List Bullet")


def add_callout(doc: Document, title: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, "b8c7e8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS["soft_blue"])
    p = cell.paragraphs[0]
    run = p.add_run(title + ": ")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(*rgb(COLORS["blue"]))
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(*rgb(COLORS["navy"]))
    doc.add_paragraph()


def add_image_docx(doc: Document, path: Path):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(6.55))


PROCEDURE_SUMMARY = [
    ["1", "Crear portafolio en Facebook Business Manager", "Cliente / Técnico Simplia", "Portafolio creado y cuenta apta", "Captura / guía"],
    ["2", "Crear app en Meta y registrar número nuevo", "Técnico Simplia", "App creada y número registrado", "Número y token"],
    ["3", "Crear Railway, n8n y Chatwoot", "Técnico Simplia", "Instancias operativas", "URLs y credenciales"],
    ["4", "Conectar WhatsApp, Messenger e Instagram", "Técnico Simplia", "Canales conectados y probados", "Capturas de conexión"],
    ["5", "Configurar OpenAI, Drive y Gmail cuando aplique", "Desarrollador / Técnico", "Flujos listos para operar", "Tokens/placeholders y evidencia"],
    ["6", "Completar credenciales, checklist y aprobación", "Responsable onboarding", "Documento completo", "Registro controlado"],
]


def build_docx(sections: list[OnboardingSection], images: dict[int, Path], map_path: Path):
    doc = Document(TEMPLATE_DOCX)
    # Keep the plantilla as source of style, then clear body content to populate the same structure cleanly.
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    section = doc.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.1)
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"{DOC_CODE} | v{DOC_VERSION} | {DOC_STATUS}")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(*rgb(COLORS["slate"]))
    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Simplia Chatbot - Onboarding Técnico ISO 9001")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(*rgb(COLORS["slate"]))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.55))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ONBOARDING TÉCNICO - SIMPLIA CHATBOT")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(*rgb(COLORS["blue"]))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Documento de requisitos y proceso alineado al formato ISO 9001")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(*rgb(COLORS["slate"]))
    doc.add_paragraph()

    add_doc_table(
        doc,
        [
            ["Código del documento", DOC_CODE, "Versión", DOC_VERSION],
            ["Nombre del proceso", "Onboarding técnico Simplia Chatbot", "Fecha de emisión", "04/05/2026"],
            ["Dueño del proceso", "Responsable de onboarding / Simplia", "Próxima revisión", "04/11/2026"],
            ["Aprobado por", "Gerencia / Responsable del cliente", "Estado", DOC_STATUS],
        ],
        widths=[1.45, 2.0, 1.35, 1.55],
    )
    add_callout(doc, "Cómo usar este documento", "Completar los campos en blanco durante el levantamiento. Las capturas se mantienen junto al paso correspondiente para guiar la ejecución y conservar evidencia.")

    add_heading(doc, "1. Información general del proceso", 1)
    add_doc_table(
        doc,
        [
            ["Campo", "Valor / espacio para completar"],
            ["Objetivo del proceso", "Recopilar información, accesos y requisitos de la empresa para configurar el chatbot Simplia."],
            ["Alcance", "Desde la validación de cuentas base hasta la conexión de canales, pruebas y cierre del onboarding."],
            ["Cliente interno o externo", "Empresa cliente / equipo técnico Simplia"],
            ["Frecuencia", "Por cada nuevo cliente, nuevo canal o nueva instancia técnica."],
            ["Empresa / cliente", "____________________"],
            ["Responsable del cliente", "____________________"],
            ["Correo de contacto", "____________________"],
            ["Teléfono de contacto", "____________________"],
            ["Canales solicitados", "WhatsApp: ___  Messenger/Facebook: ___  Instagram: ___  Otro: ____________________"],
            ["Sistemas o herramientas usadas", "Meta Business, Meta Developers, WhatsApp Business, Railway, n8n, Chatwoot, OpenAI, Google Drive, Gmail, AnyDesk y CRM si aplica."],
        ],
        widths=[2.0, 4.8],
    )

    add_heading(doc, "2. Contexto del proceso", 1)
    add_bullets(
        doc,
        [
            "Necesidad que resuelve: evita configuraciones incompletas, pérdida de accesos y falta de evidencia durante la puesta en marcha del chatbot.",
            "Procesos relacionados: levantamiento comercial, configuración técnica, pruebas de canales, despliegue y soporte inicial.",
            "Requisitos aplicables: accesos autorizados, método de pago activo cuando aplique, datos completos del negocio y aprobación del cliente.",
        ],
    )

    add_heading(doc, "3. Entradas y salidas del proceso", 1)
    add_doc_table(
        doc,
        [
            ["Proveedor / origen", "Entrada", "Salida", "Cliente / destino"],
            ["Cliente / empresa", "Datos de empresa, canales, correos, números y accesos", "Requisitos completos para configurar", "Equipo técnico Simplia"],
            ["Meta / Facebook", "Business Manager, app, número, tokens, webhooks y método de pago", "WhatsApp, Messenger e Instagram conectados", "Chatwoot / Cliente"],
            ["Railway", "Proyecto, servicios y credenciales", "Chatwoot y n8n disponibles", "Equipo técnico Simplia"],
            ["OpenAI / Google", "API Key, Drive, Gmail OAuth y permisos", "IA y automatizaciones listas cuando aplique", "n8n / Chatbot"],
        ],
        widths=[1.35, 2.05, 2.0, 1.35],
    )

    add_heading(doc, "4. Mapa del proceso", 1)
    add_image_docx(doc, map_path)

    add_heading(doc, "5. Procedimiento operativo estándar (SOP)", 1)
    add_para(doc, "La tabla resume el flujo completo. Debajo se mantiene el contenido original reorganizado por apartado, con espacios para completar y evidencias visuales junto a cada paso.")
    add_doc_table(
        doc,
        [["Paso", "Actividad", "Responsable", "Criterio de aceptación", "Registro / evidencia"]] + PROCEDURE_SUMMARY,
        widths=[0.45, 2.55, 1.25, 1.55, 1.35],
    )

    for idx, section in enumerate(sections, start=1):
        add_heading(doc, f"5.{idx}. {section.title}", 2)
        add_doc_table(doc, section_fields(section), widths=[2.5, 2.2, 2.1])
        add_para(doc, "Contenido operativo y evidencias:")
        for block in section.blocks:
            if block.kind == "p" and block.text:
                add_para(doc, block.text)
            elif block.kind == "table" and block.rows:
                width_count = len(block.rows[0])
                widths = [6.6 / width_count] * width_count
                add_doc_table(doc, block.rows, widths=widths, header=False)
            for image_id in block.images:
                if image_id in images:
                    caption = doc.add_paragraph()
                    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = caption.add_run(f"Evidencia visual {image_id:03d} - {section.title}")
                    run.bold = True
                    run.font.name = "Arial"
                    run.font.size = Pt(8.5)
                    run.font.color.rgb = RGBColor(*rgb(COLORS["blue"]))
                    add_image_docx(doc, images[image_id])

    add_heading(doc, "6. Roles y responsabilidades (RACI)", 1)
    add_doc_table(
        doc,
        [
            ["Actividad clave", "R", "A", "C", "I", "Observaciones"],
            ["Levantamiento de información", "Técnico Simplia", "Responsable onboarding", "Cliente", "Desarrollador", "Validar datos antes de configurar."],
            ["Creación de cuentas y accesos", "Cliente / Técnico", "Cliente", "Simplia", "Gerencia", "No registrar secretos reales en copias no autorizadas."],
            ["Configuración Meta/Canales", "Técnico Simplia", "Responsable técnico", "Cliente", "Operación", "Requiere método de pago activo si aplica."],
            ["Configuración IA y automatizaciones", "Desarrollador", "Responsable técnico", "Cliente", "Operación", "Aplicar límites de crédito."],
            ["Pruebas y cierre", "Técnico Simplia", "Responsable onboarding", "Cliente", "Gerencia", "Cerrar con checklist y evidencias."],
        ],
        widths=[1.6, 0.9, 1.0, 0.9, 0.8, 1.6],
    )

    add_heading(doc, "7. Riesgos, desvíos y controles", 1)
    add_doc_table(
        doc,
        [
            ["Riesgo o falla posible", "Causa probable", "Control preventivo", "Control de verificación", "Acción ante desvío"],
            ["Cuenta Facebook no apta", "Cuenta nueva o sin antigüedad suficiente", "Validar cuenta antes de iniciar", "Revisión Business Manager", "Usar cuenta apta o esperar aprobación"],
            ["WhatsApp no funciona", "Método de pago no configurado", "Agregar método antes de pruebas", "Mensaje entrante en Chatwoot", "Configurar facturación y repetir prueba"],
            ["Webhook no verifica", "URL o token incorrecto", "Copiar datos desde Chatwoot", "Verificar callback y suscripciones", "Regenerar token y guardar"],
            ["Tokens caducan", "App Google queda en testing", "Publicar app cuando aplique", "Revisar OAuth", "Pasar a producción y reconectar"],
            ["IA deja de responder", "Créditos OpenAI agotados", "Configurar límites", "Revisar consumo", "Recargar o ajustar presupuesto"],
        ],
        widths=[1.3, 1.25, 1.35, 1.4, 1.3],
    )

    add_heading(doc, "8. Registros e información documentada", 1)
    add_doc_table(
        doc,
        [
            ["Registro / evidencia", "Ubicación", "Responsable", "Retención", "Control de acceso / edición"],
            ["Documento de onboarding", "Carpeta del cliente", "Responsable onboarding", "Proyecto + 1 año", "Acceso autorizado"],
            ["Credenciales y accesos", "Gestor o copia controlada", "Cliente / Técnico Simplia", "Según política cliente", "Restringido"],
            ["Capturas de configuración", "Sección correspondiente del documento", "Técnico Simplia", "Proyecto + 1 año", "Lectura controlada"],
            ["Checklist de pruebas", "Sección 10", "Responsable onboarding", "Proyecto + 1 año", "Cliente y Simplia"],
        ],
        widths=[1.45, 1.55, 1.35, 1.25, 1.45],
    )

    add_heading(doc, "9. Indicadores del proceso", 1)
    add_doc_table(
        doc,
        [
            ["KPI", "Fórmula", "Frecuencia", "Fuente", "Dueño", "Meta / umbral"],
            ["Completitud de información", "Campos completados / campos requeridos", "Por onboarding", "Documento", "Responsable onboarding", "100% antes de configurar"],
            ["Canales conectados", "Canales probados / canales solicitados", "Por onboarding", "Chatwoot / Meta", "Técnico Simplia", "100% antes de entrega"],
            ["Pruebas exitosas", "Pruebas OK / pruebas ejecutadas", "Por onboarding", "Checklist", "Técnico Simplia", "100% OK o N.A. justificado"],
        ],
        widths=[1.2, 1.6, 1.0, 1.0, 1.1, 1.2],
    )

    add_heading(doc, "10. Checklist de ejecución", 1)
    add_doc_table(
        doc,
        [
            ["#", "Verificación", "Resultado", "Observaciones"],
            ["1", "Portafolio Facebook Business Manager creado y validado", "OK / No OK / N.A.", ""],
            ["2", "App Meta creada y número WhatsApp nuevo registrado", "OK / No OK / N.A.", ""],
            ["3", "Railway, n8n y Chatwoot creados con accesos documentados", "OK / No OK / N.A.", ""],
            ["4", "Inbox WhatsApp conectado y probado", "OK / No OK / N.A.", ""],
            ["5", "Messenger/Facebook conectado si fue solicitado", "OK / No OK / N.A.", ""],
            ["6", "Instagram conectado si fue solicitado", "OK / No OK / N.A.", ""],
            ["7", "OpenAI configurado con límites de crédito", "OK / No OK / N.A.", ""],
            ["8", "Drive/Gmail OAuth configurado si aplica", "OK / No OK / N.A.", ""],
            ["9", "Plantillas Chatwoot configuradas si aplica", "OK / No OK / N.A.", ""],
            ["10", "Credenciales, evidencias y aprobación completas", "OK / No OK / N.A.", ""],
        ],
        widths=[0.45, 4.0, 1.35, 1.2],
    )

    add_heading(doc, "11. Control de cambios", 1)
    add_doc_table(
        doc,
        [["Versión", "Fecha", "Cambio realizado", "Responsable"], ["1.0", "04/05/2026", "Transformación del onboarding técnico al formato ISO 9001, conservando contenido e imágenes por sección.", "Simplia"]],
        widths=[0.9, 1.2, 3.8, 1.3],
    )

    add_heading(doc, "12. Aprobación final", 1)
    add_doc_table(doc, [["Elaborado por", "Revisado por", "Aprobado por"], ["Simplia / Responsable onboarding", "", ""], ["Firma / fecha", "Firma / fecha", "Firma / fecha"]], widths=[2.2, 2.2, 2.2])

    doc.core_properties.title = "Onboarding Técnico Simplia Chatbot ISO 9001"
    doc.core_properties.subject = "Documento de requisitos y proceso de onboarding técnico"
    doc.core_properties.author = "Simplia"
    doc.core_properties.keywords = "onboarding, ISO 9001, Simplia Chatbot, Chatwoot, n8n, Meta, OpenAI"
    doc.save(DOCX_OUT)


def pdf_escape(text) -> str:
    return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_pdf(sections: list[OnboardingSection], images: dict[int, Path], map_path: Path):
    try:
        pdfmetrics.registerFont(TTFont("Arial", "C:/Windows/Fonts/arial.ttf"))
        pdfmetrics.registerFont(TTFont("Arial-Bold", "C:/Windows/Fonts/arialbd.ttf"))
        base_font = "Arial"
        bold_font = "Arial-Bold"
    except Exception:
        base_font = "Helvetica"
        bold_font = "Helvetica-Bold"

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleX", parent=styles["Title"], fontName=bold_font, fontSize=22, leading=27, alignment=TA_CENTER, textColor=colors.HexColor("#274690")))
    styles.add(ParagraphStyle(name="SubX", parent=styles["Normal"], fontName=base_font, fontSize=10, leading=13, alignment=TA_CENTER, textColor=colors.HexColor("#64748b")))
    styles.add(ParagraphStyle(name="H1X", parent=styles["Heading1"], fontName=bold_font, fontSize=15, leading=18, textColor=colors.HexColor("#274690"), spaceBefore=12, spaceAfter=8))
    styles.add(ParagraphStyle(name="H2X", parent=styles["Heading2"], fontName=bold_font, fontSize=12, leading=15, textColor=colors.HexColor("#0f2344"), spaceBefore=8, spaceAfter=5))
    styles.add(ParagraphStyle(name="BodyX", parent=styles["BodyText"], fontName=base_font, fontSize=8.6, leading=11.5, textColor=colors.HexColor("#0f2344"), spaceAfter=5))
    styles.add(ParagraphStyle(name="CellX", parent=styles["BodyText"], fontName=base_font, fontSize=7.1, leading=8.8, textColor=colors.HexColor("#0f2344")))
    styles.add(ParagraphStyle(name="HeadCellX", parent=styles["BodyText"], fontName=bold_font, fontSize=7.1, leading=8.8, textColor=colors.white))
    styles.add(ParagraphStyle(name="CaptionX", parent=styles["BodyText"], fontName=bold_font, fontSize=8, leading=10, alignment=TA_CENTER, textColor=colors.HexColor("#274690")))

    def p(text, style="BodyX"):
        return Paragraph(pdf_escape(text), styles[style])

    def table(rows, widths=None):
        data = [[p(cell, "HeadCellX" if i == 0 else "CellX") for cell in row] for i, row in enumerate(rows)]
        item = Table(data, colWidths=[w * inch for w in widths] if widths else None, repeatRows=1, hAlign="LEFT")
        item.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#274690")),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#d9e2ef")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]
            )
        )
        return item

    def scaled_image(path: Path, max_w=6.45, max_h=8.1):
        with Image.open(path) as img:
            w, h = img.size
        ratio = min((max_w * inch) / w, (max_h * inch) / h, 1.0)
        return RLImage(str(path), width=w * ratio, height=h * ratio, hAlign="CENTER")

    def header(canvas, doc_obj):
        canvas.saveState()
        canvas.setFont(base_font, 7)
        canvas.setFillColor(colors.HexColor("#64748b"))
        canvas.drawRightString(letter[0] - 0.5 * inch, letter[1] - 0.35 * inch, f"{DOC_CODE} | v{DOC_VERSION} | {DOC_STATUS}")
        canvas.drawCentredString(letter[0] / 2, 0.33 * inch, f"Simplia Chatbot - Onboarding Técnico ISO 9001 | Página {doc_obj.page}")
        canvas.restoreState()

    story = []
    if LOGO_PATH.exists():
        story.append(RLImage(str(LOGO_PATH), width=1.45 * inch, height=0.62 * inch, kind="proportional", hAlign="CENTER"))
        story.append(Spacer(1, 0.12 * inch))
    story.append(Paragraph("ONBOARDING TÉCNICO - SIMPLIA CHATBOT", styles["TitleX"]))
    story.append(Paragraph("Documento de requisitos y proceso alineado al formato ISO 9001", styles["SubX"]))
    story.append(Spacer(1, 0.1 * inch))
    story.append(table([["Código del documento", DOC_CODE, "Versión", DOC_VERSION], ["Nombre del proceso", "Onboarding técnico Simplia Chatbot", "Fecha de emisión", "04/05/2026"], ["Dueño del proceso", "Responsable de onboarding / Simplia", "Próxima revisión", "04/11/2026"], ["Aprobado por", "Gerencia / Responsable del cliente", "Estado", DOC_STATUS]], widths=[1.45, 2.0, 1.35, 1.55]))
    story.append(PageBreak())

    story.append(Paragraph("1. Información general del proceso", styles["H1X"]))
    story.append(table([["Campo", "Valor / espacio para completar"], ["Objetivo del proceso", "Recopilar información, accesos y requisitos de la empresa para configurar el chatbot Simplia."], ["Empresa / cliente", "____________________"], ["Responsable del cliente", "____________________"], ["Correo de contacto", "____________________"], ["Canales solicitados", "WhatsApp: ___  Messenger/Facebook: ___  Instagram: ___  Otro: ____________________"]], widths=[1.9, 4.5]))
    story.append(Paragraph("2. Contexto del proceso", styles["H1X"]))
    story.append(p("El onboarding técnico evita configuraciones incompletas, pérdida de accesos y falta de evidencia durante la puesta en marcha del chatbot."))
    story.append(Paragraph("3. Entradas y salidas del proceso", styles["H1X"]))
    story.append(table([["Proveedor / origen", "Entrada", "Salida", "Cliente / destino"], ["Cliente / empresa", "Datos de empresa, canales, correos, números y accesos", "Requisitos completos para configurar", "Equipo técnico Simplia"], ["Meta / Facebook", "Business Manager, app, número, tokens, webhooks y método de pago", "Canales conectados", "Chatwoot / Cliente"], ["Railway", "Proyecto, servicios y credenciales", "Chatwoot y n8n disponibles", "Equipo técnico Simplia"]], widths=[1.35, 2.05, 2.0, 1.35]))
    story.append(Paragraph("4. Mapa del proceso", styles["H1X"]))
    story.append(scaled_image(map_path, max_w=6.6, max_h=2.0))
    story.append(Paragraph("5. Procedimiento operativo estándar (SOP)", styles["H1X"]))
    story.append(table([["Paso", "Actividad", "Responsable", "Criterio de aceptación", "Registro / evidencia"]] + PROCEDURE_SUMMARY, widths=[0.35, 2.3, 1.1, 1.45, 1.15]))

    for idx, section in enumerate(sections, start=1):
        story.append(PageBreak())
        story.append(Paragraph(f"5.{idx}. {section.title}", styles["H2X"]))
        story.append(table(section_fields(section), widths=[2.25, 2.05, 1.95]))
        story.append(p("Contenido operativo y evidencias:"))
        for block in section.blocks:
            if block.kind == "p" and block.text:
                story.append(p(block.text))
            elif block.kind == "table" and block.rows:
                cols = max(len(block.rows[0]), 1)
                story.append(table(block.rows, widths=[6.2 / cols] * cols))
            for image_id in block.images:
                path = images.get(image_id)
                if path:
                    story.append(Paragraph(f"Evidencia visual {image_id:03d} - {section.title}", styles["CaptionX"]))
                    story.append(scaled_image(path))
                    story.append(Spacer(1, 0.08 * inch))

    for title, rows, widths in [
        ("6. Roles y responsabilidades (RACI)", [["Actividad clave", "R", "A", "C", "I", "Observaciones"], ["Levantamiento de información", "Técnico Simplia", "Responsable onboarding", "Cliente", "Desarrollador", "Validar datos antes de configurar."], ["Configuración Meta/Canales", "Técnico Simplia", "Responsable técnico", "Cliente", "Operación", "Requiere método de pago activo si aplica."], ["Pruebas y cierre", "Técnico Simplia", "Responsable onboarding", "Cliente", "Gerencia", "Cerrar con checklist y evidencias."]], [1.6, 0.9, 1.0, 0.9, 0.8, 1.6]),
        ("7. Riesgos, desvíos y controles", [["Riesgo o falla posible", "Causa probable", "Control preventivo", "Control de verificación", "Acción ante desvío"], ["Cuenta Facebook no apta", "Cuenta nueva o sin antigüedad suficiente", "Validar cuenta antes de iniciar", "Revisión Business Manager", "Usar cuenta apta o esperar aprobación"], ["WhatsApp no funciona", "Método de pago no configurado", "Agregar método antes de pruebas", "Mensaje entrante en Chatwoot", "Configurar facturación y repetir prueba"], ["Webhook no verifica", "URL o token incorrecto", "Copiar datos desde Chatwoot", "Verificar callback", "Regenerar token y guardar"]], [1.3, 1.25, 1.35, 1.4, 1.3]),
        ("8. Registros e información documentada", [["Registro / evidencia", "Ubicación", "Responsable", "Retención", "Control de acceso / edición"], ["Documento de onboarding", "Carpeta del cliente", "Responsable onboarding", "Proyecto + 1 año", "Acceso autorizado"], ["Credenciales y accesos", "Gestor o copia controlada", "Cliente / Técnico Simplia", "Según política cliente", "Restringido"], ["Capturas de configuración", "Sección correspondiente", "Técnico Simplia", "Proyecto + 1 año", "Lectura controlada"]], [1.45, 1.55, 1.35, 1.25, 1.45]),
        ("9. Indicadores del proceso", [["KPI", "Fórmula", "Frecuencia", "Fuente", "Dueño", "Meta / umbral"], ["Completitud de información", "Campos completados / campos requeridos", "Por onboarding", "Documento", "Responsable onboarding", "100% antes de configurar"], ["Canales conectados", "Canales probados / canales solicitados", "Por onboarding", "Chatwoot / Meta", "Técnico Simplia", "100% antes de entrega"]], [1.2, 1.6, 1.0, 1.0, 1.1, 1.2]),
        ("10. Checklist de ejecución", [["#", "Verificación", "Resultado", "Observaciones"], ["1", "Portafolio Facebook Business Manager creado y validado", "OK / No OK / N.A.", ""], ["2", "App Meta creada y número WhatsApp nuevo registrado", "OK / No OK / N.A.", ""], ["3", "Railway, n8n y Chatwoot creados con accesos documentados", "OK / No OK / N.A.", ""], ["4", "Canales conectados y probados", "OK / No OK / N.A.", ""], ["5", "Credenciales, evidencias y aprobación completas", "OK / No OK / N.A.", ""]], [0.45, 4.0, 1.35, 1.2]),
    ]:
        story.append(PageBreak())
        story.append(Paragraph(title, styles["H1X"]))
        story.append(table(rows, widths=widths))

    story.append(Paragraph("11. Control de cambios", styles["H1X"]))
    story.append(table([["Versión", "Fecha", "Cambio realizado", "Responsable"], ["1.0", "04/05/2026", "Transformación del onboarding técnico al formato ISO 9001, conservando contenido e imágenes por sección.", "Simplia"]], widths=[0.9, 1.2, 3.8, 1.3]))
    story.append(Paragraph("12. Aprobación final", styles["H1X"]))
    story.append(table([["Elaborado por", "Revisado por", "Aprobado por"], ["Simplia / Responsable onboarding", "", ""], ["Firma / fecha", "Firma / fecha", "Firma / fecha"]], widths=[2.2, 2.2, 2.2]))

    pdf = SimpleDocTemplate(str(PDF_OUT), pagesize=letter, rightMargin=0.5 * inch, leftMargin=0.5 * inch, topMargin=0.52 * inch, bottomMargin=0.52 * inch, title="Onboarding Técnico Simplia Chatbot ISO 9001", author="Simplia")
    pdf.build(story, onFirstPage=header, onLaterPages=header)


def main():
    source = Document(SOURCE_DOCX)
    blocks, images = extract_blocks_and_images(source)
    sections = group_sections(blocks)
    map_path = make_process_map()
    build_docx(sections, images, map_path)
    build_pdf(sections, images, map_path)
    print(f"DOCX={DOCX_OUT}")
    print(f"PDF={PDF_OUT}")
    print(f"ASSETS={ASSET_DIR}")
    print(f"SECTIONS={len(sections)}")
    print(f"IMAGES={len(images)}")


if __name__ == "__main__":
    main()
