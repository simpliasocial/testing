from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
LOGO_PATH = ROOT / "logo_simplia.png"
DOCX_PATH = ROOT / "Sistema_Flujo_Conversacional_Simplia_Chatbot_ISO10013.docx"

DOC_CODE = "MAN-FLUJO-BOT-001"
DOC_VERSION = "1.1"
DOC_DATE = "06/05/2026"
NEXT_REVIEW = "06/11/2026"
DOC_STATUS = "Borrador / Vigente al aprobarse"

COLORS = {
    "navy": "0f2344",
    "blue": "274690",
    "soft_blue": "eaf1ff",
    "light": "f8fafc",
    "green": "0a9b6f",
    "slate": "64748b",
    "line": "d9e2ef",
    "white": "ffffff",
}


@dataclass
class Event:
    kind: str
    payload: tuple


class Manual:
    def __init__(self) -> None:
        self.events: list[Event] = []

    def h(self, level: int, text: str) -> None:
        self.events.append(Event("h", (level, text)))

    def p(self, text: str) -> None:
        self.events.append(Event("p", (text,)))

    def bullets(self, items: Iterable[str]) -> None:
        self.events.append(Event("bullets", (list(items),)))

    def table(self, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
        self.events.append(Event("table", (headers, rows, widths)))

    def note(self, title: str, text: str) -> None:
        self.events.append(Event("note", (title, text)))


def rgb(hex_color: str) -> RGBColor:
    value = hex_color.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_cell_shading(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def set_table_borders(table, color: str = COLORS["line"]) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement("w:" + edge)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    table._tbl.tblPr.append(borders)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def write_cell(cell, text: str, *, bold: bool = False, color: str | None = None, size: float = 8.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for index, part in enumerate(str(text).split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def style_docx_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    set_table_borders(table)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if row_index == 0:
                set_cell_shading(cell, COLORS["blue"])
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = rgb(COLORS["white"])
            else:
                set_cell_shading(cell, COLORS["light"] if row_index % 2 == 0 else COLORS["white"])
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(7.8 if row_index else 8.2)


def build_manual() -> Manual:
    manual = Manual()

    manual.h(1, "1. Control documental")
    manual.table(
        ["Campo", "Valor"],
        [
            ["Codigo", DOC_CODE],
            ["Nombre", "Sistema de Flujo Conversacional Simplia Chatbot"],
            ["Version", DOC_VERSION],
            ["Fecha", DOC_DATE],
            ["Estado", DOC_STATUS],
            ["Proxima revision", NEXT_REVIEW],
            ["Tipo de documento", "Manual ISO 10013 de informacion documentada"],
        ],
        [2.0, 5.1],
    )

    manual.h(1, "2. Que hace el agente")
    manual.p(
        "El agente atiende leads de inicio a fin. Recibe mensajes, identifica la intencion, responde con informacion aprobada, categoriza el lead, guarda datos, calcula score, guia al usuario hacia una cita, envia follow-ups y deriva a una persona cuando corresponde."
    )
    manual.note(
        "Idea central",
        "El bot atiende lo repetible y ordena la conversacion. La persona interviene cuando hay una excepcion, una cita gestionada manualmente, una venta registrada o un caso que requiere seguimiento_humano.",
    )

    manual.h(1, "3. Arquitectura general del flujo")
    manual.table(
        ["Capa", "Para que sirve", "Resultado"],
        [
            ["Canal de entrada", "WhatsApp, Instagram, Facebook/Messenger, TikTok u otro canal conectado.", "Mensaje del cliente recibido."],
            ["Motor conversacional", "Lee el mensaje y decide si corresponde bienvenida, informacion, agenda, rechazo o humano.", "Intencion identificada."],
            ["Plantillas", "Usa mensajes aprobados para bienvenida, datos de cita, confirmacion, derivacion humana y no aplica.", "Respuesta consistente."],
            ["Etiquetas", "Ordenan el estado operativo del lead.", "Lead categorizado."],
            ["contact_attributes", "Guardan datos del lead, agenda, score y gestion manual.", "Informacion disponible para cita y seguimiento."],
            ["Score", "Mide calidad e intencion del lead con FRIO, TIBIO y CALIENTE.", "Prioridad comercial entendible."],
            ["Agenda", "Recoge datos, valida fecha/hora/agencia y confirma cita si aplica.", "Cita agendada o pendiente."],
            ["Follow-up", "Retoma conversaciones sin respuesta a los 20 minutos y a las 3 horas.", "Lead recuperado o enviado a humano."],
            ["Gestion humana", "Una persona toma casos manuales, ventas o excepciones.", "Seguimiento_humano, cita manual o venta registrada."],
        ],
        [1.45, 4.0, 1.9],
    )

    manual.h(1, "4. Etiquetas/categorias del sistema")
    manual.p("Las etiquetas indican en que parte del flujo esta el lead y que debe pasar despues.")
    manual.table(
        ["Etiqueta", "Quien la asigna", "Para que sirve", "Siguiente accion"],
        [
            ["bienvenida", "Bot", "Primer contacto o lead sin intencion clara todavia.", "Saludar, entender necesidad y esperar respuesta."],
            ["solicita_informacion", "Bot", "El lead pide informacion, costos, servicios, beneficios, sedes u horarios.", "Responder con informacion aprobada e invitar al siguiente paso."],
            ["interesado", "Bot", "El lead muestra intencion de avanzar, comprar o agendar.", "Pedir datos necesarios y orientar a cita."],
            ["desinteresado", "Bot", "El lead rechaza, esta fuera del negocio o envia contenido irrelevante.", "Cerrar de forma educada y evitar insistencia."],
            ["cita_agendada", "Bot", "Cita creada automaticamente por el bot.", "Confirmar cita y guardar datos de agenda."],
            ["seguimiento_humano", "Bot", "El bot deja el caso para una persona.", "Una persona revisa y decide si contacta, agenda o cierra."],
            ["venta_exitosa", "Manual", "Venta u operacion concretada.", "La persona registra monto_operacion y fecha_monto_operacion."],
            ["cita_agendado_humano", "Manual", "Cita creada o modificada por una persona.", "La persona registra responsable y datos de cita."],
        ],
        [1.35, 1.0, 2.65, 2.3],
    )

    manual.h(1, "5. Datos/contact attributes que guarda el bot")
    manual.p(
        "Los contact_attributes permiten que el agente recuerde datos del lead y que el equipo vea la informacion necesaria para agenda, seguimiento y venta."
    )
    manual.table(
        ["Atributo", "Tipo", "Quien lo llena", "Uso dentro del flujo"],
        [
            ["nombre_completo", "string", "Bot", "Identificar al lead y crear cita."],
            ["fecha_visita", "string", "Bot", "Guardar fecha solicitada o confirmada para la cita."],
            ["hora_visita", "string", "Bot", "Guardar hora solicitada o confirmada para la cita."],
            ["agencia", "string", "Bot", "Guardar sede/agencia elegida."],
            ["celular", "string", "Bot", "Contactar al lead; si entra por WhatsApp puede venir del canal."],
            ["correo", "string", "Bot", "Dato de contacto o registro si el negocio lo pide."],
            ["campana", "string", "Bot", "Identificar fuente, anuncio o campana cuando aplique."],
            ["ciudad", "string", "Bot", "Segmentacion, disponibilidad o validacion territorial."],
            ["edad", "string", "Bot", "Dato del lead si el negocio lo necesita."],
            ["canal", "string", "Bot", "Canal por donde entro el lead."],
            ["score_interes", "number", "Bot", "Puntaje acumulado de calidad e intencion."],
            ["agente", "checkbox", "Manual", "Se marca cuando el bot ya no debe responder."],
            ["monto_operacion", "string", "Manual", "Se llena cuando se marca venta_exitosa."],
            ["fecha_monto_operacion", "Date", "Manual", "Se llena cuando se marca venta_exitosa."],
            ["responsable", "string", "Manual", "Se llena cuando una persona agenda o modifica una cita."],
        ],
        [1.65, 0.9, 1.1, 3.55],
    )

    manual.h(1, "6. Opciones de recoleccion de datos")
    manual.table(
        ["Modo", "Como funciona", "Datos sugeridos", "Regla operativa"],
        [
            [
                "Filtro inicial obligatorio",
                "El bot pide datos desde la primera interaccion antes de continuar con otras intenciones.",
                "Obligatorios sugeridos: nombre_completo y celular.\nOpcionales sugeridos: correo, ciudad y edad.",
                "Si falta un dato obligatorio, el bot insiste de forma educada hasta completar la informacion minima.",
            ],
            [
                "Sin filtro inicial",
                "El bot saluda, responde preguntas y solo pide datos completos cuando el usuario quiere agendar.",
                "Para agendar: fecha_visita, hora_visita, agencia y los datos extra que defina la empresa.",
                "Si faltan datos para cita, el bot los pide uno por uno antes de confirmar.",
            ],
        ],
        [1.5, 2.4, 2.1, 1.75],
    )
    manual.note(
        "Decision del cliente",
        "La empresa define si desea filtro obligatorio al inicio o si prefiere que todos los datos se recojan al momento de agendar.",
    )

    manual.h(1, "7. Flujo conversacional de inicio a fin")
    manual.table(
        ["Paso", "Que pasa", "Que hace el bot", "Resultado esperado"],
        [
            ["1", "Entra un mensaje del cliente.", "Identifica canal, contacto y mensaje valido del usuario.", "Conversacion lista para analizar."],
            ["2", "El usuario saluda o pregunta algo.", "Responde bienvenida o informacion directa segun intencion.", "Lead en bienvenida o solicita_informacion."],
            ["3", "El bot detecta intencion.", "Clasifica si pide informacion, muestra interes, quiere cita, rechaza o pide humano.", "Etiqueta asignada."],
            ["4", "El usuario pide informacion.", "Responde servicios, costos, sedes, requisitos, horarios o beneficios con informacion aprobada.", "Lead informado e invitado a avanzar."],
            ["5", "El usuario muestra interes.", "Marca interesado y orienta a agendar o resolver una duda concreta.", "Lead listo para siguiente paso."],
            ["6", "El usuario quiere agendar.", "Solicita o valida fecha_visita, hora_visita, agencia y datos definidos por la empresa.", "Datos de cita completos."],
            ["7", "Hay disponibilidad y datos completos.", "Confirma cita usando la plantilla aprobada.", "cita_agendada."],
            ["8", "El usuario no responde.", "Envia follow-up a los 20 minutos y luego a las 3 horas.", "Lead recuperado o pendiente de humano."],
            ["9", "No responde despues de la segunda retoma.", "Pasa el caso a seguimiento_humano.", "Una persona revisa el lead."],
            ["10", "Una persona agenda o registra venta.", "El bot no lo hace automaticamente; queda como gestion manual.", "cita_agendado_humano o venta_exitosa."],
        ],
        [0.45, 2.2, 2.9, 1.9],
    )

    manual.h(1, "8. Agendamiento de cita")
    manual.bullets(
        [
            "El bot agenda solo si tiene los datos necesarios definidos por la empresa.",
            "Los datos base para cita son fecha_visita, hora_visita y agencia.",
            "La empresa puede pedir datos adicionales como nombre_completo, celular, correo, ciudad, edad u otro dato propio del negocio.",
            "Si el lead entra por WhatsApp, el celular puede venir del canal; si entra por otro canal, el bot puede pedirlo.",
            "La cita automatica queda como cita_agendada.",
            "Si una persona agenda o modifica la cita, se usa cita_agendado_humano y debe llenarse responsable.",
            "La confirmacion final usa la plantilla aprobada con fecha, hora, agencia y datos del cliente.",
        ]
    )

    manual.h(1, "9. Score de calidad con 3 estados")
    manual.p(
        "El score mide solo mensajes del cliente. No se calcula con respuestas de IA, bot o asesor. Es acumulativo, reversible y sin limite tecnico; sube o baja segun senales del mensaje."
    )
    manual.table(
        ["Estado", "Rango", "Interpretacion", "Accion sugerida"],
        [
            ["FRIO", "Menor a 45 o sin puntaje", "Lead con senal inicial, baja o todavia sin intencion comercial clara.", "Responder informacion aprobada e invitar a avanzar."],
            ["TIBIO", "45 a 69", "Lead con senales accionables: precio, disponibilidad, servicio especifico, requisitos o continuidad.", "Priorizar respuesta, resolver duda e invitar a cita."],
            ["CALIENTE", "70 o mas", "Lead con intencion fuerte: quiere comprar, reservar, agendar, pagar o hablar con asesor.", "Solicitar datos de agenda, confirmar cita o derivar si corresponde."],
        ],
        [1.0, 1.35, 3.1, 2.0],
    )
    manual.table(
        ["Que observa el score", "Ejemplos"],
        [
            ["Interes inicial", "info, precio, vi el anuncio, me interesa."],
            ["Intencion comercial", "quiero agendar, quiero comprar, cuanto cuesta, llamenme."],
            ["Fit con el negocio", "Pregunta por servicios, sedes, horarios o disponibilidad real."],
            ["Contactabilidad", "Deja telefono, correo, nombre o pide contacto."],
            ["Engagement", "Sigue la conversacion, responde preguntas o confirma interes."],
            ["Senales negativas", "No me interesa, insultos, fuera de negocio, fraude o spam."],
        ],
        [2.2, 5.0],
    )

    manual.h(1, "10. Vocabulario del negocio")
    manual.p(
        "El bot necesita vocabulario real de cada empresa para entender mejor la intencion. Esto incluye nombres de servicios, frases comerciales, formas de pedir precio, terminos de agenda, dudas frecuentes y senales de rechazo."
    )
    manual.table(
        ["Tipo de vocabulario", "Ejemplos que debe entregar la empresa"],
        [
            ["Servicios / productos", "Nombres de servicios, tratamientos, planes, productos o paquetes."],
            ["Costos / promociones", "precio, valor, cuanto cuesta, promocion, descuento, formas de pago."],
            ["Agenda", "quiero agendar, reservar, cita, cuando puedo ir, disponibilidad."],
            ["Asesor humano", "asesor, persona, llamenme, quiero hablar con alguien."],
            ["Rechazo", "no me interesa, despues veo, no deseo informacion."],
            ["Fuera de negocio", "busco trabajo, tareas, memes, consultas ajenas al servicio."],
            ["Riesgo", "fraude, amenaza, insulto, phishing o conducta critica."],
        ],
        [2.1, 5.1],
    )

    manual.h(1, "11. Follow-ups de 20 minutos y 3 horas")
    manual.table(
        ["Momento", "A quien aplica", "Que hace el bot", "Si responde", "Si no responde"],
        [
            ["20 minutos", "bienvenida, solicita_informacion o interesado sin respuesta.", "Envia primer script de retoma aprobado.", "Continua el flujo segun intencion.", "Espera hasta la retoma de 3 horas."],
            ["3 horas", "Lead que sigue sin responder despues de la primera retoma.", "Envia segundo script de retoma aprobado.", "Continua el flujo segun intencion.", "Pasa a seguimiento_humano."],
        ],
        [1.0, 2.0, 1.75, 1.35, 1.35],
    )
    manual.note(
        "Regla de retoma",
        "Despues de la retoma de 3 horas, si el lead no responde, el caso queda en seguimiento_humano para revision de una persona.",
    )

    manual.h(1, "12. Cuando pasa a seguimiento humano")
    manual.bullets(
        [
            "El usuario pide hablar con una persona o asesor.",
            "El bot no puede resolver la solicitud con informacion aprobada.",
            "El caso necesita decision comercial o validacion manual.",
            "El lead no responde despues de las retomas de 20 minutos y 3 horas.",
            "La empresa marca agente para que el bot no siga respondiendo.",
            "Hay rechazo fuerte, riesgo, insultos, fraude o contenido fuera del negocio que requiere revision.",
        ]
    )

    manual.h(1, "13. Que hace el bot vs que hace una persona")
    manual.table(
        ["Actividad", "Bot", "Persona"],
        [
            ["Responder bienvenida", "Si, con plantilla aprobada.", "Solo si el bot no entiende o se requiere trato manual."],
            ["Responder preguntas frecuentes", "Si, con informacion aprobada.", "Actualiza la informacion cuando cambie."],
            ["Clasificar lead", "Si, con etiquetas automaticas.", "Puede corregir o complementar manualmente."],
            ["Calcular score_interes", "Si, solo con mensajes del cliente.", "Lo usa para priorizar revision."],
            ["Agendar cita automatica", "Si, si tiene datos y disponibilidad.", "Interviene si hay excepcion o agenda manual."],
            ["Registrar cita manual", "No.", "Si, con cita_agendado_humano y responsable."],
            ["Registrar venta exitosa", "No.", "Si, con venta_exitosa, monto_operacion y fecha_monto_operacion."],
            ["Tomar seguimiento_humano", "Marca el caso cuando corresponde.", "Revisa, responde, agenda o cierra el lead."],
        ],
        [2.1, 2.55, 2.55],
    )

    manual.h(1, "14. Checklist de entendimiento")
    manual.table(
        ["#", "Verificacion", "Resultado"],
        [
            ["1", "Se entiende que hace el agente de inicio a fin.", "OK / No OK"],
            ["2", "Se entienden las etiquetas/categorias y para que sirven.", "OK / No OK"],
            ["3", "Se entienden los contact_attributes y datos de agenda.", "OK / No OK"],
            ["4", "Se entiende si se usara filtro inicial o recoleccion al agendar.", "OK / No OK"],
            ["5", "Se entiende que el score tiene solo FRIO, TIBIO y CALIENTE.", "OK / No OK"],
            ["6", "Se entienden los follow-ups de 20 minutos y 3 horas.", "OK / No OK"],
            ["7", "Se entiende cuando pasa a seguimiento_humano.", "OK / No OK"],
            ["8", "Se entiende que venta_exitosa y cita_agendado_humano son manuales.", "OK / No OK"],
        ],
        [0.35, 5.4, 1.4],
    )

    return manual


def add_docx_header(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"{DOC_CODE} | v{DOC_VERSION} | {DOC_STATUS}")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["slate"])

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Simplia Chatbot - Sistema de Flujo Conversacional ISO 10013")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["slate"])


def add_docx_title(document: Document) -> None:
    if LOGO_PATH.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(LOGO_PATH), width=Inches(1.45))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Sistema de Flujo Conversacional Simplia Chatbot")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.color.rgb = rgb(COLORS["blue"])

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Manual para entender como funciona el agente de inicio a fin")
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.color.rgb = rgb(COLORS["slate"])


def configure_styles(document: Document) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    for style_name, size, color in [
        ("Heading 1", 15, COLORS["blue"]),
        ("Heading 2", 12.5, COLORS["navy"]),
        ("Heading 3", 10.5, COLORS["green"]),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def render_docx(manual: Manual) -> None:
    document = Document()
    configure_styles(document)
    add_docx_header(document)
    add_docx_title(document)

    for event in manual.events:
        if event.kind == "h":
            level, text = event.payload
            document.add_heading(text, level=level)
        elif event.kind == "p":
            (text,) = event.payload
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(9.4)
        elif event.kind == "bullets":
            (items,) = event.payload
            for item in items:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.space_after = Pt(2)
                run = paragraph.add_run(item)
                run.font.name = "Arial"
                run.font.size = Pt(9.1)
        elif event.kind == "note":
            title, text = event.payload
            table = document.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            set_table_borders(table)
            cell = table.cell(0, 0)
            set_cell_shading(cell, COLORS["soft_blue"])
            cell.text = ""
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(title)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(8.8)
            paragraph.add_run().add_break()
            run = paragraph.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(8.4)
            document.add_paragraph()
        elif event.kind == "table":
            headers, rows, widths = event.payload
            table = document.add_table(rows=1, cols=len(headers))
            for idx, header in enumerate(headers):
                write_cell(table.cell(0, idx), header, bold=True, color=COLORS["white"], size=8.2)
            set_repeat_table_header(table.rows[0])
            for row in rows:
                cells = table.add_row().cells
                for idx, value in enumerate(row):
                    write_cell(cells[idx], value, size=7.8)
            style_docx_table(table)
            if widths:
                for row in table.rows:
                    for idx, width in enumerate(widths):
                        if idx < len(row.cells):
                            row.cells[idx].width = Inches(width)
            document.add_paragraph()

    document.core_properties.title = "Sistema de Flujo Conversacional Simplia Chatbot ISO 10013"
    document.core_properties.subject = "Manual de funcionamiento del agente conversacional de inicio a fin"
    document.core_properties.author = "Simplia"
    document.core_properties.keywords = "Simplia, chatbot, flujo conversacional, ISO 10013, etiquetas, score, agenda"
    try:
        document.save(DOCX_PATH)
    except PermissionError as exc:
        raise SystemExit(f"No se pudo guardar {DOCX_PATH.name}. Cierre el archivo en Word y vuelva a ejecutar este script.") from exc


def main() -> None:
    manual = build_manual()
    render_docx(manual)
    print(f"DOCX={DOCX_PATH}")


if __name__ == "__main__":
    main()
