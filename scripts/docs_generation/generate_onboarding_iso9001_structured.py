from __future__ import annotations

import textwrap
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DOCX = ROOT / "plantilla_iso_9001_proceso.docx"
LOGO_PATH = ROOT / "logo_simplia.png"
ASSET_DIR = ROOT / "Onboarding_Tecnico_Simplia_Chatbot_assets"
DOCX_OUT = ROOT / "Onboarding_Tecnico_Simplia_Chatbot_ISO9001.docx"
PDF_OUT = ROOT / "Onboarding_Tecnico_Simplia_Chatbot_ISO9001.pdf"

DOC_CODE = "PRO-ONB-CHATBOT-001"
DOC_VERSION = "1.1"
DOC_STATUS = "Vigente/Borrador"
ISSUE_DATE = "05/05/2026"
NEXT_REVIEW = "05/11/2026"

COLORS = {
    "blue": "274690",
    "navy": "0f2344",
    "slate": "64748b",
    "line": "d9e2ef",
    "light": "f8fafc",
    "soft_blue": "eaf1ff",
    "soft_green": "eaf8f1",
    "soft_orange": "fff7e6",
    "green": "0a9b6f",
    "orange": "b45309",
    "red": "dc2626",
}

OFFICIAL_SOURCES = [
    ["Chatwoot Cloud - precios", "https://www.chatwoot.com/pricing/", "Plan Startups: USD 19 por agente/mes, facturado anual. Confirmar precio vigente antes de contratar."],
    ["Chatwoot WhatsApp Cloud", "https://www.chatwoot.com/docs/product/channels/whatsapp/whatsapp-cloud", "Chatwoot Cloud permite conectar WhatsApp Cloud por Embedded Signup o configuración manual."],
    ["Chatwoot Facebook Messenger", "https://www.chatwoot.com/hc/user-guide/articles/1677778588-how-to-setup-a-facebook-channel", "Para Cloud se crea inbox desde Settings > Inboxes > Add Inbox > Messenger y se autoriza con Facebook."],
    ["Chatwoot Instagram", "https://www.chatwoot.com/hc/user-guide/articles/1744361165-how-to-setup-an-instagram-channel-via-instagram-login", "Para Cloud se recomienda Instagram Business Login desde Settings > Inboxes > Add Inbox."],
    ["Chatwoot TikTok", "https://developers.chatwoot.com/self-hosted/configuration/features/integrations/tiktok", "La guía oficial disponible contiene pasos self-hosted; en Cloud se debe validar disponibilidad o habilitación con Chatwoot."],
    ["Railway pricing", "https://railway.com/pricing", "Railway cobra por uso. Hobby tiene mínimo de uso bajo; Pro parte de USD 20 de uso incluido."],
]

COMMON_GUIDE_LINKS = {
    "meta_business": [
        "Meta Business Suite: https://business.facebook.com/",
        "Configuración del negocio: https://business.facebook.com/settings/",
    ],
    "meta_developers": [
        "Meta for Developers: https://developers.facebook.com/apps/",
        "WhatsApp Cloud API: https://developers.facebook.com/docs/whatsapp/cloud-api/get-started",
    ],
    "railway": [
        "Railway pricing: https://railway.com/pricing",
        "Railway n8n guide: https://docs.railway.com/guides/n8n",
    ],
    "chatwoot": [
        "Chatwoot: https://www.chatwoot.com/",
        "Chatwoot pricing: https://www.chatwoot.com/pricing/",
    ],
    "whatsapp": [
        "Chatwoot WhatsApp Cloud: https://www.chatwoot.com/docs/product/channels/whatsapp/whatsapp-cloud",
        "WhatsApp Cloud API: https://developers.facebook.com/docs/whatsapp/cloud-api/get-started",
    ],
    "facebook": [
        "Chatwoot Facebook channel: https://www.chatwoot.com/hc/user-guide/articles/1677778588-how-to-setup-a-facebook-channel",
        "Messenger Platform: https://developers.facebook.com/docs/messenger-platform/getting-started",
    ],
    "instagram": [
        "Chatwoot Instagram Business Login: https://www.chatwoot.com/hc/user-guide/articles/1744361165-how-to-setup-an-instagram-channel-via-instagram-login",
        "Instagram Platform: https://developers.facebook.com/docs/instagram-platform",
    ],
    "tiktok": [
        "Chatwoot TikTok integration: https://developers.chatwoot.com/self-hosted/configuration/features/integrations/tiktok",
        "TikTok for Developers: https://developers.tiktok.com/",
    ],
    "openai": [
        "OpenAI Platform: https://platform.openai.com/",
        "API keys: https://platform.openai.com/api-keys",
    ],
    "google": [
        "n8n Google credentials: https://docs.n8n.io/integrations/builtin/credentials/google/",
        "n8n Google Service Account: https://docs.n8n.io/integrations/builtin/credentials/google/service-account/",
    ],
    "anydesk": [
        "AnyDesk: https://anydesk.com/",
    ],
}

STEPS = [
    {
        "title": "Creación del Portafolio en Facebook Business Manager",
        "objective": "Dejar listo el portafolio empresarial desde donde se administrarán WhatsApp, Facebook, Instagram, pagos y permisos.",
        "actions": [
            "Ingresar con la cuenta Meta autorizada por la empresa.",
            "Crear o seleccionar el portafolio de negocio correcto.",
            "Validar nombre legal, país, sitio web o perfil público y datos de contacto.",
            "Agregar al desarrollador o responsable Simplia con permisos suficientes para configurar activos.",
        ],
        "save": ["Nombre del portafolio", "ID del negocio", "Correo administrador", "Captura del portafolio", "Estado de verificación"],
        "evidence": "Captura de Business Manager con el portafolio creado o seleccionado.",
        "guide_links": COMMON_GUIDE_LINKS["meta_business"],
    },
    {
        "title": "Creación de App en Meta for Developers y registro de número nuevo para WhatsApp",
        "objective": "Crear la app de Meta y asociar un chip/número nuevo que nunca haya sido usado con WhatsApp.",
        "actions": [
            "Crear app tipo negocio en Meta for Developers.",
            "Agregar producto WhatsApp y aceptar condiciones.",
            "Registrar número nuevo con código de verificación por SMS o llamada.",
            "Guardar Phone Number ID, WhatsApp Business Account ID y App ID si se usa configuración manual.",
        ],
        "save": ["App ID", "Business Account ID", "Phone Number ID", "Número WhatsApp", "Token/API key si aplica", "Capturas de verificación"],
        "evidence": "Captura de la app Meta y del número registrado.",
        "guide_links": COMMON_GUIDE_LINKS["meta_developers"],
    },
    {
        "title": "Creación y conexión de Railway",
        "objective": "Crear el espacio de hosting que se usará solo para n8n. Chatwoot ya no se instala en Railway.",
        "actions": [
            "Crear o ingresar a la cuenta Railway del cliente o de la operación definida.",
            "Registrar método de pago si el proyecto requiere despliegue estable.",
            "Crear proyecto para n8n y guardar la URL del proyecto.",
            "Usar Railway para n8n. Plan referencial: Hobby para operación básica, estimando hasta USD 20/mes según consumo real; confirmar precio vigente antes de desplegar.",
            "Aclarar que Railway cobra por uso, por lo que el valor final depende de recursos, logs, almacenamiento y tráfico.",
        ],
        "save": ["Correo Railway", "Proyecto Railway", "Plan", "Método de pago configurado", "URL del proyecto", "Responsable de facturación"],
        "evidence": "Captura de proyecto Railway y facturación activa si aplica.",
        "guide_links": COMMON_GUIDE_LINKS["railway"],
    },
    {
        "title": "Creación de instancia n8n en Railway",
        "objective": "Desplegar n8n como motor de automatizaciones del agente conversacional.",
        "actions": [
            "Crear servicio n8n dentro del proyecto Railway.",
            "Configurar dominio o URL pública del servicio.",
            "Guardar usuario administrador y método seguro de contraseña.",
            "Validar que n8n abra correctamente y pueda ejecutar workflows.",
        ],
        "save": ["URL n8n", "Correo admin", "Método seguro de contraseña", "Variables configuradas", "Estado de despliegue"],
        "evidence": "Captura de n8n abierto y servicio activo en Railway.",
        "guide_links": COMMON_GUIDE_LINKS["railway"],
    },
    {
        "title": "Cambio de configuración en Railway por cada nodo creado",
        "objective": "Actualizar variables y configuración técnica cada vez que se cree un nodo o conexión que lo requiera.",
        "actions": [
            "Revisar variables de entorno requeridas por cada workflow.",
            "Guardar secretos solo en Railway/n8n o gestor seguro, no en texto plano del documento.",
            "Redeploy o restart del servicio cuando el cambio lo requiera.",
            "Probar que el workflow afectado sigue funcionando.",
        ],
        "save": ["Variable o nodo creado", "Servicio afectado", "Fecha de cambio", "Responsable", "Prueba realizada"],
        "evidence": "Captura de variables sin exponer secretos y prueba OK.",
        "guide_links": COMMON_GUIDE_LINKS["railway"],
    },
    {
        "title": "Creación de cuenta Chatwoot Cloud y contratación del plan",
        "objective": "Crear la cuenta Cloud donde se administrarán las conversaciones del chatbot.",
        "actions": [
            "Ingresar a https://www.chatwoot.com con correo empresarial.",
            "Crear workspace de la empresa.",
            "Contratar plan Startups USD 19 por agente/mes, confirmando precio vigente en la página oficial.",
            "Guardar responsable, correo y método seguro para credenciales.",
        ],
        "save": ["Correo empresarial", "Workspace Chatwoot", "Plan contratado", "Cantidad de agentes", "Método de pago", "Responsable"],
        "evidence": "Captura del workspace Chatwoot y plan activo.",
        "guide_links": COMMON_GUIDE_LINKS["chatwoot"],
    },
    {
        "title": "Configuración del Inbox WhatsApp en Chatwoot con Meta",
        "objective": "Conectar WhatsApp Cloud en Chatwoot para recibir y responder conversaciones.",
        "actions": [
            "En Chatwoot ir a Settings > Inboxes > Add Inbox > WhatsApp.",
            "Usar WhatsApp Embedded Signup cuando sea posible para que Meta configure número, webhook y credenciales.",
            "Si Embedded Signup no aplica, usar configuración manual con los datos de Meta.",
            "Agregar agentes al inbox y enviar una prueba entrante/saliente.",
        ],
        "save": ["Inbox WhatsApp", "Número conectado", "WABA", "Agentes asignados", "Prueba entrante", "Prueba saliente"],
        "evidence": "Captura del inbox WhatsApp conectado y conversación de prueba.",
        "guide_links": COMMON_GUIDE_LINKS["whatsapp"],
    },
    {
        "title": "Configuración de método de pago para WhatsApp",
        "objective": "Evitar que WhatsApp falle por falta de facturación en Meta.",
        "actions": [
            "Ir a WhatsApp > Configuración > Centro de facturación > Métodos de pago > Agregar.",
            "Como alternativa, agregar el método cuando Meta lo solicite en la prueba de API.",
            "Confirmar con la empresa qué tarjeta o método se usará para cargos de Meta.",
            "No iniciar pruebas finales de WhatsApp sin confirmar este punto.",
        ],
        "save": ["Método agregado", "Responsable de pago", "Fecha", "Captura de facturación", "Observaciones"],
        "evidence": "Captura de centro de facturación o confirmación de método agregado.",
        "critical": True,
        "guide_links": [
            "Centro de facturación Meta: https://business.facebook.com/billing_hub/",
            "WhatsApp Manager: https://business.facebook.com/wa/manage/",
        ],
    },
    {
        "title": "Configuración Messenger",
        "objective": "Conectar Facebook Messenger en Chatwoot Cloud cuando la empresa lo solicite.",
        "actions": [
            "En Chatwoot ir a Settings > Inboxes > Add Inbox > Messenger.",
            "Autorizar con la cuenta Facebook administradora de la página.",
            "Seleccionar la página correcta y habilitar permisos solicitados.",
            "Agregar agentes y probar mensaje desde la página.",
        ],
        "save": ["Página Facebook", "Inbox Messenger", "Cuenta autorizadora", "Agentes", "Prueba de mensaje"],
        "evidence": "Captura de Messenger conectado en Chatwoot.",
        "guide_links": COMMON_GUIDE_LINKS["facebook"],
    },
    {
        "title": "Configuración Instagram",
        "objective": "Conectar Instagram Business en Chatwoot Cloud.",
        "actions": [
            "Validar que la cuenta sea Instagram Business.",
            "Usar Instagram Business Login desde Chatwoot cuando esté disponible.",
            "Autorizar mensajes y permisos solicitados por Meta.",
            "Agregar agentes y probar DM entrante.",
        ],
        "save": ["Usuario Instagram", "Inbox Instagram", "Cuenta autorizadora", "Permisos aceptados", "Prueba DM"],
        "evidence": "Captura de Instagram conectado en Chatwoot.",
        "guide_links": COMMON_GUIDE_LINKS["instagram"],
    },
    {
        "title": "Configuración TikTok",
        "objective": "Validar y conectar TikTok Business Messaging si el plan/cuenta Cloud lo permite.",
        "actions": [
            "Crear cuenta TikTok Developer si aplica.",
            "Registrar app y solicitar acceso a Business Messaging API.",
            "Validar con Chatwoot Cloud si el canal TikTok está disponible o requiere habilitación.",
            "Si se habilita, conectar cuenta TikTok Business y probar mensaje.",
        ],
        "save": ["Cuenta TikTok Developer", "App ID", "App Secret en gestor seguro", "Estado de aprobación", "Estado en Chatwoot"],
        "evidence": "Captura de disponibilidad o respuesta de Chatwoot/estado de integración.",
        "note": "La documentación oficial disponible incluye pasos de super admin/self-hosted; en Cloud se debe validar disponibilidad antes de prometer la conexión.",
        "guide_links": COMMON_GUIDE_LINKS["tiktok"],
    },
    {
        "title": "Integración con OpenAI / ChatGPT",
        "objective": "Conectar el modelo de IA usado por el agente conversacional.",
        "actions": [
            "Crear o usar cuenta OpenAI autorizada por la empresa.",
            "Agregar método de pago y límites de uso.",
            "Crear API key y guardarla solo en n8n/Railway o gestor seguro.",
            "Ejecutar prueba de respuesta desde el workflow.",
        ],
        "save": ["Cuenta OpenAI", "Método de pago", "Límite mensual", "API key guardada en seguro", "Prueba OK"],
        "evidence": "Captura de configuración sin exponer API key.",
        "guide_links": COMMON_GUIDE_LINKS["openai"],
    },
    {
        "title": "Reinicio o validación final de n8n y canales Chatwoot",
        "objective": "Confirmar que todo responde después de cambios de configuración.",
        "actions": [
            "Reiniciar o redeploy n8n si se cambiaron variables.",
            "Validar que Chatwoot Cloud reciba mensajes en cada inbox configurado.",
            "Enviar prueba de extremo a extremo: cliente > Chatwoot > n8n > respuesta.",
            "Registrar errores pendientes antes del cierre.",
        ],
        "save": ["Workflow probado", "Canal probado", "Resultado", "Fecha", "Responsable"],
        "evidence": "Capturas de pruebas por canal.",
        "guide_links": [
            "Railway n8n guide: https://docs.railway.com/guides/n8n",
            "Chatwoot inboxes: https://www.chatwoot.com/docs/product/channels/",
        ],
    },
    {
        "title": "Configuración n8n nodo Drive - Service Account",
        "objective": "Configurar Google Drive solo si el flujo requiere leer o guardar documentos.",
        "actions": [
            "Preguntar al desarrollador si este nodo aplica al proyecto.",
            "Crear service account y compartir carpetas necesarias.",
            "Guardar JSON/credencial en n8n de forma segura.",
            "Probar lectura o escritura en Drive.",
        ],
        "save": ["Aplica: Sí / No", "Cuenta de servicio", "Carpeta Drive", "Permisos", "Prueba OK"],
        "evidence": "Captura de credencial creada sin exponer secretos.",
        "guide_links": COMMON_GUIDE_LINKS["google"],
    },
    {
        "title": "Configuración n8n nodo Gmail - OAuth",
        "objective": "Configurar Gmail solo si el flujo requiere enviar o leer correos.",
        "actions": [
            "Preguntar al desarrollador si este nodo aplica al proyecto.",
            "Crear OAuth consent y credenciales si corresponde.",
            "Conectar cuenta Gmail desde n8n.",
            "Probar envío o lectura según el flujo.",
        ],
        "save": ["Aplica: Sí / No", "Correo Gmail", "OAuth client", "Permisos", "Prueba OK"],
        "evidence": "Captura de conexión Gmail en n8n.",
        "guide_links": COMMON_GUIDE_LINKS["google"],
    },
    {
        "title": "Configuración plantillas de respuesta Chatwoot",
        "objective": "Crear plantillas rápidas si la empresa necesita respuestas manuales estandarizadas.",
        "actions": [
            "Recolectar textos aprobados por la empresa.",
            "Crear plantillas en Chatwoot.",
            "Validar nombre, texto, tono y uso interno.",
            "Probar uso desde una conversación.",
        ],
        "save": ["Plantilla", "Escenario", "Texto aprobado", "Responsable de aprobación", "Estado"],
        "evidence": "Captura de plantillas creadas.",
        "guide_links": [
            "Chatwoot macros: https://www.chatwoot.com/docs/product/features/macros/",
            "Chatwoot canned responses: https://www.chatwoot.com/docs/product/features/canned-responses/",
        ],
    },
    {
        "title": "Control remoto de PC - AnyDesk",
        "objective": "Permitir acompañamiento remoto durante creación de cuentas, permisos y pagos.",
        "actions": [
            "Solicitar AnyDesk cuando el cliente deba operar desde su sesión.",
            "Confirmar autorización antes de tomar control.",
            "No guardar contraseñas visibles ni datos bancarios en el documento.",
            "Registrar solo evidencia operativa sin exponer información sensible.",
        ],
        "save": ["ID AnyDesk", "Responsable conectado", "Fecha", "Acción realizada", "Observaciones"],
        "evidence": "Registro de sesión o nota de acompañamiento.",
        "guide_links": COMMON_GUIDE_LINKS["anydesk"],
    },
    {
        "title": "Cierre, evidencias y credenciales",
        "objective": "Dejar documentado qué quedó creado, dónde están los accesos y qué pruebas fueron aprobadas.",
        "actions": [
            "Revisar que cada paso tenga responsable, fecha, estado y evidencia.",
            "Completar la tabla de credenciales sin escribir contraseñas ni tokens completos.",
            "Confirmar que el acceso al Dashboard queda registrado con URL, usuario, responsable y estado.",
            "Cerrar con aprobación del cliente o responsable interno.",
        ],
        "save": ["Checklist final", "Tabla de credenciales", "Links de evidencias", "Aprobación", "Pendientes si existen"],
        "evidence": "Documento completo con aprobación final.",
        "guide_links": [
            "Carpeta de evidencias del cliente: ____________________",
            "Dashboard Simplia: ____________________",
        ],
    },
]

CREDENTIALS = [
    ["Chatwoot Cloud", "Correo empresarial", "Plan / workspace", "Método seguro de contraseña", "Responsable", "Estado"],
    ["Railway", "Correo", "Proyecto n8n", "Método seguro de contraseña", "Responsable", "Estado"],
    ["n8n", "URL", "Usuario admin", "Método seguro de contraseña", "Responsable", "Estado"],
    ["Meta Business", "Correo admin", "Business ID / Portafolio", "Permisos", "Responsable", "Estado"],
    ["Meta Developers", "App ID", "WABA / Phone Number ID", "Tokens en gestor seguro", "Responsable", "Estado"],
    ["OpenAI", "Correo", "Proyecto / API", "API key en gestor seguro", "Responsable", "Estado"],
    ["Google Drive / Gmail", "Correo", "Service Account / OAuth", "Permisos", "Responsable", "Estado"],
    ["AnyDesk", "ID", "Equipo", "Autorización", "Responsable", "Estado"],
    ["Dashboard", "URL", "Usuario", "Método seguro de contraseña", "Responsable", "Estado"],
]


def rgb(hex_color: str) -> tuple[int, int, int]:
    value = hex_color.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def image_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def make_process_map() -> Path:
    ASSET_DIR.mkdir(exist_ok=True)
    path = ASSET_DIR / "mapa_proceso_onboarding_tecnico_cloud.png"
    width, height = 1900, 560
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 42), "Mapa del onboarding técnico Simplia Chatbot", font=image_font(34, True), fill=rgb(COLORS["navy"]))
    steps = [
        ("1", "Meta", "Portafolio,\napp y número"),
        ("2", "Railway + n8n", "Hosting solo\npara automatizar"),
        ("3", "Chatwoot Cloud", "Cuenta, plan\ne inboxes"),
        ("4", "Canales", "WhatsApp,\nMessenger, IG, TikTok"),
        ("5", "IA y datos", "OpenAI,\nDrive, Gmail"),
        ("6", "Pruebas", "Mensajes,\nevidencias y cierre"),
    ]
    x, y = 65, 170
    box_w, box_h, gap = 255, 182, 42
    for idx, (num, name, desc) in enumerate(steps):
        bx = x + idx * (box_w + gap)
        draw.rounded_rectangle((bx, y, bx + box_w, y + box_h), radius=22, fill=rgb(COLORS["soft_blue"]), outline=rgb(COLORS["blue"]), width=3)
        draw.ellipse((bx + 18, y + 18, bx + 68, y + 68), fill=rgb(COLORS["blue"]))
        draw.text((bx + 43, y + 43), num, font=image_font(22, True), fill="white", anchor="mm")
        draw.text((bx + 82, y + 26), name, font=image_font(23, True), fill=rgb(COLORS["navy"]))
        draw.multiline_text((bx + 26, y + 88), desc, font=image_font(18), fill=rgb(COLORS["slate"]), spacing=6)
        if idx < len(steps) - 1:
            ax = bx + box_w + 7
            ay = y + box_h // 2
            draw.line((ax, ay, ax + gap - 14, ay), fill=rgb(COLORS["blue"]), width=4)
            draw.polygon([(ax + gap - 14, ay - 10), (ax + gap + 4, ay), (ax + gap - 14, ay + 10)], fill=rgb(COLORS["blue"]))
    draw.text((70, 440), "Salida esperada: cuentas creadas, pagos listos, canales probados, n8n funcionando y credenciales documentadas de forma segura.", font=image_font(18), fill=rgb(COLORS["navy"]))
    image.save(path)
    return path


def wrapped_lines(text: str, width: int) -> list[str]:
    lines: list[str] = []
    for raw_line in str(text).splitlines():
        if not raw_line.strip():
            lines.append("")
            continue
        lines.extend(textwrap.wrap(raw_line, width=width, break_long_words=False, replace_whitespace=False))
    return lines


def make_step_guide_image(index: int, step: dict) -> Path:
    ASSET_DIR.mkdir(exist_ok=True)
    ascii_title = unicodedata.normalize("NFKD", step["title"]).encode("ascii", "ignore").decode("ascii")
    safe_name = "".join(ch.lower() if ch.isalnum() else "_" for ch in ascii_title)[:54].strip("_")
    path = ASSET_DIR / f"guia_paso_{index:02d}_{safe_name}.png"
    width, height = 1500, 760
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = image_font(34, True)
    subtitle_font = image_font(21, True)
    body_font = image_font(19)
    small_font = image_font(17)

    draw.rounded_rectangle((32, 32, width - 32, height - 32), radius=30, fill=rgb(COLORS["light"]), outline=rgb(COLORS["line"]), width=3)
    draw.rounded_rectangle((58, 58, width - 58, 150), radius=20, fill=rgb(COLORS["blue"]))
    draw.text((88, 82), f"Referencia visual - Paso {index}", font=subtitle_font, fill="white")
    draw.text((88, 112), step["title"], font=title_font, fill="white")

    left_x, right_x = 86, 790
    y = 188
    draw.text((left_x, y), "Dónde entrar", font=subtitle_font, fill=rgb(COLORS["navy"]))
    y += 35
    for link in step.get("guide_links", [])[:3]:
        for line in wrapped_lines(f"- {link}", 62):
            draw.text((left_x, y), line, font=small_font, fill=rgb(COLORS["slate"]))
            y += 24
    if not step.get("guide_links"):
        draw.text((left_x, y), "- Link guía pendiente de completar", font=small_font, fill=rgb(COLORS["slate"]))
        y += 24

    y += 20
    draw.text((left_x, y), "Qué revisar en pantalla", font=subtitle_font, fill=rgb(COLORS["navy"]))
    y += 35
    for item in step["actions"][:4]:
        for line in wrapped_lines(f"- {item}", 64):
            draw.text((left_x, y), line, font=small_font, fill=rgb(COLORS["navy"]))
            y += 24

    draw.rounded_rectangle((right_x, 188, width - 86, 618), radius=24, fill="white", outline=rgb(COLORS["line"]), width=2)
    draw.text((right_x + 34, 220), "Datos que deben quedar guardados", font=subtitle_font, fill=rgb(COLORS["navy"]))
    yy = 265
    for item in step["save"][:7]:
        draw.rounded_rectangle((right_x + 34, yy, width - 126, yy + 38), radius=15, fill=rgb(COLORS["soft_blue"]), outline=rgb(COLORS["line"]))
        draw.text((right_x + 52, yy + 9), item, font=body_font, fill=rgb(COLORS["navy"]))
        yy += 52

    footer_fill = COLORS["soft_orange"] if step.get("critical") else COLORS["soft_green"]
    footer_outline = "facc15" if step.get("critical") else "86efac"
    draw.rounded_rectangle((86, 650, width - 86, 710), radius=18, fill=rgb(footer_fill), outline=rgb(footer_outline), width=2)
    footer = step.get("note") or "Pegar la evidencia real del cliente en la fila 'Evidencia del cliente'. No escribir contraseñas ni tokens completos."
    for i, line in enumerate(wrapped_lines(footer, 118)[:2]):
        draw.text((112, 668 + i * 23), line, font=small_font, fill=rgb(COLORS["navy"]))

    image.save(path)
    return path


def set_cell_shading(cell, fill: str):
    props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def set_table_borders(table, color="d9e2ef"):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement("w:" + edge)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    table._tbl.tblPr.append(borders)


def set_cell_text(cell, text, bold=False, color="0f2344", size=8.7):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(*rgb(color))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_doc_table(doc: Document, rows, widths=None, header=True, first_col=False):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            is_header = header and i == 0
            cell = table.cell(i, j)
            set_cell_text(cell, value, bold=is_header or (first_col and j == 0), color=("ffffff" if is_header else COLORS["navy"]))
            fill = COLORS["blue"] if is_header else ("f8fafc" if i % 2 == 0 else "ffffff")
            set_cell_shading(cell, fill)
            if widths and j < len(widths):
                cell.width = Inches(widths[j])
    return table


def add_heading(doc: Document, text: str, level=1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(*rgb(COLORS["blue"] if level == 1 else COLORS["navy"]))
    return paragraph


def add_para(doc: Document, text: str, style=None):
    paragraph = doc.add_paragraph(style=style)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.1)
    run.font.color.rgb = RGBColor(*rgb(COLORS["navy"]))
    return paragraph


def add_bullets(doc: Document, items):
    for item in items:
        add_para(doc, item, style="List Bullet")


def add_callout(doc: Document, title: str, text: str, tone="blue"):
    fill = COLORS["soft_blue"]
    border = "b8c7e8"
    title_color = COLORS["blue"]
    if tone == "warning":
        fill = COLORS["soft_orange"]
        border = "facc15"
        title_color = COLORS["orange"]
    elif tone == "ok":
        fill = COLORS["soft_green"]
        border = "86efac"
        title_color = COLORS["green"]
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, border)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    run = p.add_run(title + ": ")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(*rgb(title_color))
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(*rgb(COLORS["navy"]))
    doc.add_paragraph()


def add_image_docx(doc: Document, path: Path, width=6.55):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))


def step_table(step):
    rows = [
        ["Documento / link guía", "\n".join(f"- {item}" for item in step.get("guide_links", ["____________________"]))],
        ["Foto / captura guía", "Referencia visual incluida debajo de este paso."],
        ["Qué se debe hacer", "\n".join(f"- {item}" for item in step["actions"])],
        ["Qué se debe guardar", "\n".join(f"- {item}" for item in step["save"])],
        ["Evidencia esperada", step["evidence"]],
        ["Responsable", "____________________"],
        ["Fecha de ejecución", "____ / ____ / ______"],
        ["Estado", "Pendiente / En proceso / Completo / No aplica"],
        ["Evidencia del cliente", "____________________"],
        ["Observaciones", "____________________"],
    ]
    if step.get("note"):
        rows.insert(3, ["Nota importante", step["note"]])
    return rows


def add_step_docx(doc: Document, index: int, step):
    add_heading(doc, f"5.{index}. {step['title']}", 2)
    add_callout(doc, "Objetivo", step["objective"], "warning" if step.get("critical") else "blue")
    add_doc_table(doc, step_table(step), widths=[1.75, 5.0], first_col=True)
    guide_image = step.get("guide_image")
    if guide_image:
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(f"Referencia visual - {step['title']}")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(8.8)
        run.font.color.rgb = RGBColor(*rgb(COLORS["blue"]))
        add_image_docx(doc, Path(guide_image), width=6.2)


def build_docx(map_path: Path):
    doc = Document(TEMPLATE_DOCX)
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)

    section = doc.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.1)
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"{DOC_CODE} | v{DOC_VERSION} | {DOC_STATUS}")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(*rgb(COLORS["slate"]))

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Simplia Chatbot - Onboarding Técnico ISO 9001")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(*rgb(COLORS["slate"]))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.55))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ONBOARDING TÉCNICO - SIMPLIA CHATBOT")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(*rgb(COLORS["blue"]))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Guía paso a paso para crear cuentas, permisos, pagos, canales e integraciones")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(*rgb(COLORS["slate"]))
    doc.add_paragraph()

    add_doc_table(
        doc,
        [
            ["Código", DOC_CODE, "Versión", DOC_VERSION],
            ["Nombre", "Onboarding técnico Simplia Chatbot", "Fecha de emisión", ISSUE_DATE],
            ["Uso", "Guía operativa de configuración", "Próxima revisión", NEXT_REVIEW],
            ["Estado", DOC_STATUS, "Formato", "ISO 9001"],
        ],
        widths=[1.1, 2.35, 1.35, 1.8],
    )
    add_callout(
        doc,
        "Objetivo del documento",
        "Guiar al desarrollador y al encargado de la empresa para crear cuentas, activar pagos, otorgar permisos, conectar canales y dejar el agente conversacional funcionando. No reemplaza el gestor seguro de contraseñas: aquí solo se registra dónde quedó guardado cada acceso.",
    )

    add_heading(doc, "1. Información general", 1)
    add_doc_table(
        doc,
        [
            ["Campo", "Completar durante la reunión"],
            ["Empresa / cliente", "____________________"],
            ["Responsable de cuentas y permisos", "____________________"],
            ["Responsable de pagos", "____________________"],
            ["Desarrollador / técnico Simplia", "____________________"],
            ["Correo empresarial para Chatwoot", "____________________"],
            ["Canales a configurar", "WhatsApp: ___  Messenger: ___  Instagram: ___  TikTok: ___"],
            ["Carpeta de evidencias / links", "____________________"],
        ],
        widths=[2.2, 4.5],
    )

    add_heading(doc, "2. Alcance práctico", 1)
    add_bullets(
        doc,
        [
            "Railway se usa solo para n8n y sus automatizaciones.",
            "Chatwoot se configura en Cloud, con plan Startups o el plan vigente aprobado por la empresa.",
            "Meta Business y Meta Developers se usan para WhatsApp, Facebook/Messenger e Instagram.",
            "TikTok se valida antes de comprometer activación porque la documentación oficial disponible tiene pasos self-hosted.",
            "OpenAI, Google Drive y Gmail se configuran únicamente si el flujo del agente lo requiere.",
        ],
    )

    add_heading(doc, "3. Fuentes oficiales consultadas", 1)
    add_doc_table(doc, [["Fuente", "Link", "Uso en este onboarding"]] + OFFICIAL_SOURCES, widths=[1.65, 2.35, 2.75])

    add_heading(doc, "4. Mapa del proceso", 1)
    add_image_docx(doc, map_path)

    add_heading(doc, "5. Procedimiento paso a paso", 1)
    add_para(doc, "Cada paso incluye link guía y referencia visual. Usa el link guía para ejecutar el paso y pega en evidencia del cliente el link o captura de lo que quedó creado. Los secretos reales deben quedar en un gestor seguro o dentro de la herramienta correspondiente.")
    for idx, step in enumerate(STEPS, start=1):
        add_step_docx(doc, idx, step)

    add_heading(doc, "6. Validación final de funcionamiento", 1)
    add_doc_table(
        doc,
        [
            ["Prueba", "Qué debe comprobarse", "Resultado", "Evidencia"],
            ["WhatsApp", "Mensaje entrante y respuesta desde Chatwoot/n8n", "OK / No OK / N.A.", "____________________"],
            ["Messenger", "Mensaje entrante y asignación de agente", "OK / No OK / N.A.", "____________________"],
            ["Instagram", "DM entrante y respuesta", "OK / No OK / N.A.", "____________________"],
            ["TikTok", "Disponibilidad validada o canal probado", "OK / No OK / N.A.", "____________________"],
            ["n8n", "Workflow activo, credenciales conectadas y ejecución sin error", "OK / No OK / N.A.", "____________________"],
            ["OpenAI", "Respuesta generada con límites de uso definidos", "OK / No OK / N.A.", "____________________"],
            ["Dashboard", "Acceso operativo y datos mínimos visibles", "OK / No OK / N.A.", "____________________"],
        ],
        widths=[1.1, 3.4, 1.25, 1.5],
    )

    add_heading(doc, "7. Checklist de cierre", 1)
    add_doc_table(
        doc,
        [
            ["#", "Validación", "Resultado", "Observaciones"],
            ["1", "Portafolio Meta creado o seleccionado", "OK / No OK / N.A.", ""],
            ["2", "App Meta y número WhatsApp nuevo registrados", "OK / No OK / N.A.", ""],
            ["3", "Método de pago Meta/WhatsApp configurado", "OK / No OK / N.A.", ""],
            ["4", "Railway creado para n8n, sin Chatwoot self-hosted", "OK / No OK / N.A.", ""],
            ["5", "n8n desplegado y probado", "OK / No OK / N.A.", ""],
            ["6", "Chatwoot Cloud creado con plan aprobado", "OK / No OK / N.A.", ""],
            ["7", "Canales solicitados conectados o marcados como no aplica", "OK / No OK / N.A.", ""],
            ["8", "OpenAI / Drive / Gmail configurados si aplica", "OK / No OK / N.A.", ""],
            ["9", "Credenciales registradas sin secretos en texto plano", "OK / No OK / N.A.", ""],
            ["10", "Evidencias y aprobación final completas", "OK / No OK / N.A.", ""],
        ],
        widths=[0.45, 4.05, 1.35, 1.15],
    )

    add_heading(doc, "8. Credenciales y cuentas", 1)
    add_callout(doc, "Importante", "No escribir contraseñas, API keys ni tokens completos en este documento. Usar gestor seguro, n8n, Railway, Chatwoot o el método aprobado por la empresa.", "warning")
    add_doc_table(doc, [["Cuenta / sistema", "Dato principal", "Dato secundario", "Dónde quedó el acceso", "Responsable", "Estado"]] + CREDENTIALS, widths=[1.25, 1.15, 1.35, 1.55, 1.05, 0.95])

    add_heading(doc, "9. Control de cambios", 1)
    add_doc_table(
        doc,
        [
            ["Versión", "Fecha", "Cambio realizado", "Responsable"],
            ["1.1", ISSUE_DATE, "Actualización a Chatwoot Cloud, Railway solo para n8n, pagos WhatsApp, canales Cloud y cierre práctico de credenciales.", "Simplia"],
        ],
        widths=[0.9, 1.2, 3.8, 1.3],
    )

    add_heading(doc, "10. Aprobación final", 1)
    add_doc_table(
        doc,
        [
            ["Elaborado por", "Revisado por", "Aprobado por"],
            ["Simplia / Desarrollador", "", ""],
            ["Firma / fecha", "Firma / fecha", "Firma / fecha"],
        ],
        widths=[2.2, 2.2, 2.2],
    )

    doc.core_properties.title = "Onboarding Técnico Simplia Chatbot ISO 9001"
    doc.core_properties.subject = "Guía de creación de cuentas, permisos, pagos, canales e integraciones"
    doc.core_properties.author = "Simplia"
    doc.core_properties.keywords = "onboarding, ISO 9001, Simplia Chatbot, Chatwoot Cloud, n8n, Railway, Meta, OpenAI"
    doc.save(DOCX_OUT)


def pdf_escape(text) -> str:
    return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_pdf(map_path: Path):
    try:
        pdfmetrics.registerFont(TTFont("Arial", "C:/Windows/Fonts/arial.ttf"))
        pdfmetrics.registerFont(TTFont("Arial-Bold", "C:/Windows/Fonts/arialbd.ttf"))
        base_font = "Arial"
        bold_font = "Arial-Bold"
    except Exception:
        base_font = "Helvetica"
        bold_font = "Helvetica-Bold"

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleX", parent=styles["Title"], fontName=bold_font, fontSize=21, leading=25, alignment=TA_CENTER, textColor=colors.HexColor("#274690")))
    styles.add(ParagraphStyle(name="SubX", parent=styles["Normal"], fontName=base_font, fontSize=9.5, leading=12, alignment=TA_CENTER, textColor=colors.HexColor("#64748b")))
    styles.add(ParagraphStyle(name="H1X", parent=styles["Heading1"], fontName=bold_font, fontSize=14, leading=17, textColor=colors.HexColor("#274690"), spaceBefore=12, spaceAfter=7))
    styles.add(ParagraphStyle(name="H2X", parent=styles["Heading2"], fontName=bold_font, fontSize=11.5, leading=14, textColor=colors.HexColor("#0f2344"), spaceBefore=8, spaceAfter=5))
    styles.add(ParagraphStyle(name="BodyX", parent=styles["BodyText"], fontName=base_font, fontSize=8.4, leading=11, textColor=colors.HexColor("#0f2344"), spaceAfter=4))
    styles.add(ParagraphStyle(name="CellX", parent=styles["BodyText"], fontName=base_font, fontSize=6.9, leading=8.5, textColor=colors.HexColor("#0f2344")))
    styles.add(ParagraphStyle(name="HeadCellX", parent=styles["BodyText"], fontName=bold_font, fontSize=6.9, leading=8.5, textColor=colors.white))

    def p(text, style="BodyX"):
        return Paragraph(pdf_escape(text), styles[style])

    def table(rows, widths=None):
        data = [[p(cell, "HeadCellX" if i == 0 else "CellX") for cell in row] for i, row in enumerate(rows)]
        item = Table(data, colWidths=[w * inch for w in widths] if widths else None, repeatRows=1, hAlign="LEFT")
        item.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#274690")),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#d9e2ef")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]
            )
        )
        return item

    def scaled_image(path: Path, max_w=6.45, max_h=2.1):
        with Image.open(path) as img:
            w, h = img.size
        ratio = min((max_w * inch) / w, (max_h * inch) / h, 1.0)
        return RLImage(str(path), width=w * ratio, height=h * ratio, hAlign="CENTER")

    def header(canvas, doc_obj):
        canvas.saveState()
        canvas.setFont(base_font, 7)
        canvas.setFillColor(colors.HexColor("#64748b"))
        canvas.drawRightString(letter[0] - 0.5 * inch, letter[1] - 0.35 * inch, f"{DOC_CODE} | v{DOC_VERSION} | {DOC_STATUS}")
        canvas.drawCentredString(letter[0] / 2, 0.33 * inch, f"Simplia Chatbot - Onboarding Técnico ISO 9001 | Página {doc_obj.page}")
        canvas.restoreState()

    story = []
    if LOGO_PATH.exists():
        story.append(RLImage(str(LOGO_PATH), width=1.45 * inch, height=0.62 * inch, kind="proportional", hAlign="CENTER"))
        story.append(Spacer(1, 0.12 * inch))
    story.append(Paragraph("ONBOARDING TÉCNICO - SIMPLIA CHATBOT", styles["TitleX"]))
    story.append(Paragraph("Guía paso a paso para crear cuentas, permisos, pagos, canales e integraciones", styles["SubX"]))
    story.append(Spacer(1, 0.1 * inch))
    story.append(table([["Código", DOC_CODE, "Versión", DOC_VERSION], ["Nombre", "Onboarding técnico Simplia Chatbot", "Fecha de emisión", ISSUE_DATE], ["Uso", "Guía operativa de configuración", "Próxima revisión", NEXT_REVIEW], ["Estado", DOC_STATUS, "Formato", "ISO 9001"]], widths=[1.1, 2.35, 1.35, 1.8]))
    story.append(PageBreak())

    story.append(Paragraph("1. Información general", styles["H1X"]))
    story.append(table([["Campo", "Completar durante la reunión"], ["Empresa / cliente", "____________________"], ["Responsable de cuentas y permisos", "____________________"], ["Responsable de pagos", "____________________"], ["Desarrollador / técnico Simplia", "____________________"], ["Correo empresarial para Chatwoot", "____________________"], ["Canales a configurar", "WhatsApp: ___  Messenger: ___  Instagram: ___  TikTok: ___"]], widths=[2.2, 4.4]))
    story.append(Paragraph("2. Alcance práctico", styles["H1X"]))
    story.append(p("Railway se usa solo para n8n. Chatwoot se configura en Cloud. Meta se usa para WhatsApp, Messenger e Instagram. TikTok se valida antes de comprometer activación porque la documentación oficial disponible contiene pasos self-hosted."))
    story.append(Paragraph("3. Fuentes oficiales consultadas", styles["H1X"]))
    story.append(table([["Fuente", "Link", "Uso"]] + OFFICIAL_SOURCES, widths=[1.45, 2.45, 2.4]))
    story.append(Paragraph("4. Mapa del proceso", styles["H1X"]))
    story.append(scaled_image(map_path))
    story.append(Paragraph("5. Procedimiento paso a paso", styles["H1X"]))
    story.append(p("Cada paso incluye link guía y referencia visual. Usa el link guía para ejecutar el paso y pega en evidencia del cliente el link o captura de lo que quedó creado."))

    for idx, step in enumerate(STEPS, start=1):
        story.append(PageBreak() if idx in {1, 7, 13} else Spacer(1, 0.05 * inch))
        story.append(Paragraph(f"5.{idx}. {step['title']}", styles["H2X"]))
        story.append(p(f"Objetivo: {step['objective']}"))
        rows = [
            ["Documento / link guía", "\n".join(f"- {item}" for item in step.get("guide_links", ["____________________"]))],
            ["Foto / captura guía", "Referencia visual incluida debajo de este paso."],
            ["Qué se debe hacer", "\n".join(f"- {item}" for item in step["actions"])],
            ["Qué se debe guardar", "\n".join(f"- {item}" for item in step["save"])],
            ["Evidencia esperada", step["evidence"]],
            ["Evidencia del cliente", "____________________"],
            ["Responsable", "____________________"],
            ["Fecha", "____ / ____ / ______"],
            ["Estado", "Pendiente / En proceso / Completo / No aplica"],
            ["Observaciones", "____________________"],
        ]
        if step.get("note"):
            rows.insert(3, ["Nota", step["note"]])
        story.append(table(rows, widths=[1.55, 4.95]))
        guide_image = step.get("guide_image")
        if guide_image:
            story.append(p(f"Referencia visual - {step['title']}"))
            story.append(scaled_image(Path(guide_image), max_w=6.35, max_h=3.25))

    story.append(PageBreak())
    story.append(Paragraph("6. Validación final de funcionamiento", styles["H1X"]))
    story.append(table([["Prueba", "Qué debe comprobarse", "Resultado", "Evidencia"], ["WhatsApp", "Mensaje entrante y respuesta desde Chatwoot/n8n", "OK / No OK / N.A.", "____________________"], ["Messenger", "Mensaje entrante y asignación de agente", "OK / No OK / N.A.", "____________________"], ["Instagram", "DM entrante y respuesta", "OK / No OK / N.A.", "____________________"], ["TikTok", "Disponibilidad validada o canal probado", "OK / No OK / N.A.", "____________________"], ["n8n", "Workflow activo y ejecución sin error", "OK / No OK / N.A.", "____________________"], ["Dashboard", "Acceso operativo y datos mínimos visibles", "OK / No OK / N.A.", "____________________"]], widths=[1.1, 3.4, 1.25, 1.5]))
    story.append(Paragraph("7. Checklist de cierre", styles["H1X"]))
    story.append(table([["#", "Validación", "Resultado", "Observaciones"], ["1", "Portafolio Meta creado o seleccionado", "OK / No OK / N.A.", ""], ["2", "Método de pago Meta/WhatsApp configurado", "OK / No OK / N.A.", ""], ["3", "Railway creado para n8n, sin Chatwoot self-hosted", "OK / No OK / N.A.", ""], ["4", "Chatwoot Cloud creado con plan aprobado", "OK / No OK / N.A.", ""], ["5", "Credenciales registradas sin secretos en texto plano", "OK / No OK / N.A.", ""]], widths=[0.45, 4.05, 1.35, 1.15]))
    story.append(Paragraph("8. Credenciales y cuentas", styles["H1X"]))
    story.append(table([["Cuenta / sistema", "Dato principal", "Dato secundario", "Dónde quedó el acceso", "Responsable", "Estado"]] + CREDENTIALS, widths=[1.15, 1.05, 1.25, 1.45, 0.95, 0.8]))
    story.append(Paragraph("9. Control de cambios", styles["H1X"]))
    story.append(table([["Versión", "Fecha", "Cambio realizado", "Responsable"], ["1.1", ISSUE_DATE, "Actualización a Chatwoot Cloud, Railway solo para n8n, pagos WhatsApp, canales Cloud y cierre práctico de credenciales.", "Simplia"]], widths=[0.9, 1.2, 3.8, 1.3]))
    story.append(Paragraph("10. Aprobación final", styles["H1X"]))
    story.append(table([["Elaborado por", "Revisado por", "Aprobado por"], ["Simplia / Desarrollador", "", ""], ["Firma / fecha", "Firma / fecha", "Firma / fecha"]], widths=[2.2, 2.2, 2.2]))

    pdf = SimpleDocTemplate(str(PDF_OUT), pagesize=letter, rightMargin=0.5 * inch, leftMargin=0.5 * inch, topMargin=0.52 * inch, bottomMargin=0.52 * inch, title="Onboarding Técnico Simplia Chatbot ISO 9001", author="Simplia")
    pdf.build(story, onFirstPage=header, onLaterPages=header)


def main():
    map_path = make_process_map()
    for idx, step in enumerate(STEPS, start=1):
        step["guide_image"] = str(make_step_guide_image(idx, step))
    build_docx(map_path)
    build_pdf(map_path)
    print(f"DOCX={DOCX_OUT}")
    print(f"PDF={PDF_OUT}")
    print(f"ASSETS={ASSET_DIR}")
    print(f"STEPS={len(STEPS)}")


if __name__ == "__main__":
    main()
