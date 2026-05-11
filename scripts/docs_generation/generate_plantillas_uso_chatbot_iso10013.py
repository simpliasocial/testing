from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "Plantillas uso chatbot.docx"
PROMPT_TXT = ROOT / "promt_chatbot.txt"
LOGO_PATH = ROOT / "logo_simplia.png"
DOCX_OUT = ROOT / "Plantillas_Uso_Chatbot_ISO10013.docx"

DOC_CODE = "MAN-PLANTILLAS-BOT-001"
DOC_VERSION = "1.0"
DOC_DATE = "07/05/2026"
NEXT_REVIEW = "07/11/2026"
DOC_STATUS = "Borrador / Vigente al aprobarse"

HEADER_RE = re.compile(r"^(PLANTILLA_[A-ZÁÉÍÓÚÑ0-9_]+)\b")

COLORS = {
    "navy": "0f2344",
    "blue": "274690",
    "soft_blue": "eaf1ff",
    "light": "f8fafc",
    "green": "0a9b6f",
    "slate": "64748b",
    "line": "d9e2ef",
    "white": "ffffff",
    "code": "f3f6fb",
    "warning": "fff7ed",
}


@dataclass
class TemplateBlock:
    original_name: str
    name: str
    body: str
    scenario: str
    suggested_label: str


SCENARIOS: dict[str, str] = {
    "PLANTILLA_NO_APLICA": "Se usa cuando la consulta no corresponde a los servicios del negocio o esta fuera del alcance del agente.",
    "PLANTILLA_VALOR_IMPLANTE": "Se usa cuando el usuario pregunta por el valor de un implante y aun necesita valoracion para recibir un plan correcto.",
    "PLANTILLA_DESCONFIANZA": "Se usa cuando el usuario expresa desconfianza, miedo o duda sobre la valoracion, el tratamiento o el proceso.",
    "PLANTILLA_NO_INTERES": "Se usa cuando el usuario indica que no desea continuar, no tiene interes o prefiere dejar la conversacion para otro momento.",
    "PLANTILLA_VALOR_SERVICIO": "Se usa cuando el usuario pregunta por costos o valor de un servicio sin tener aun una valoracion clinica.",
    "PLANTILLA_INTENCION_SERVICIO": "Se usa cuando el usuario manifiesta interes general por un servicio y se necesita identificar que tratamiento o necesidad tiene.",
    "PLANTILLA_CLIENTE_INTERESADO": "Se usa cuando el usuario ya muestra interes claro y el bot puede guiarlo hacia el agendamiento.",
    "PLANTILLA_INTERES_CON_CIUDAD": "Se usa cuando el usuario esta interesado y ya se conoce su ciudad, pero todavia debe elegir agencia o sede.",
    "PLANTILLA_INTERES_CON_AGENCIA": "Se usa cuando el usuario esta interesado y ya se conoce la agencia, por lo que se puede avanzar a datos para agendar.",
    "PLANTILLA_BIENVENIDA": "Se usa en la primera respuesta o inicio de la conversacion para saludar, presentar la atencion y abrir el flujo.",
    "PLANTILLA_CIUDADES": "Se usa cuando el usuario quiere agendar o conocer sedes, pero aun no se sabe la ciudad donde desea atenderse.",
    "PLANTILLA_SECTOR_MATCH_CIUDAD_DEDUCIDA": "Se usa cuando el bot puede deducir la ciudad o sector del usuario y ofrece sedes relacionadas.",
    "PLANTILLA_UBICACION_DESCONOCIDA": "Se usa cuando el usuario pregunta por ubicacion, pero no entrega ciudad, sector o referencia suficiente.",
    "PLANTILLA_CIUDADAGENCIA_NO_REGISTRADA": "Se usa cuando la ciudad o agencia mencionada no existe dentro de las sedes registradas.",
    "PLANTILLA_CIUDAD_SIN_AGENCIA": "Se usa cuando el usuario indica una ciudad donde la empresa no tiene agencia disponible.",
    "PLANTILLA_CIUDAD_SIN_AGENCIA_INSISTE": "Se usa cuando el usuario insiste en una ciudad sin agencia y se debe cerrar o redirigir sin inventar sedes.",
    "PLANTILLA_UBICACIONES_SIN_CIUDAD": "Se usa cuando el usuario pide ubicaciones en general y el bot debe pedir ciudad antes de mostrar sedes.",
    "PLANTILLA_UBICACIONES_CON_CIUDAD": "Se usa cuando el usuario pide ubicaciones y ya se conoce la ciudad para listar sedes aplicables.",
    "PLANTILLA_REFERENCIA_MATCH": "Se usa cuando el usuario entrega una referencia geografica que coincide con una sede o zona registrada.",
    "PLANTILLA_DATOS_CITA": "Se usa cuando el usuario ya eligio agencia o sede y el bot debe solicitar los datos necesarios para agendar.",
    "PLANTILLA_FECHA_PASADA": "Se usa cuando el usuario intenta agendar en una fecha anterior a la fecha actual.",
    "PLANTILLA_NO_HOY_EMPRESA": "Se usa cuando el usuario pide cita para hoy y la regla del negocio no permite agendar para el mismo dia.",
    "PLANTILLA_DATOS_INCOMPLETOS": "Se usa cuando faltan datos obligatorios para agendar y el bot debe insistir hasta completar informacion.",
    "PLANTILLA_HORA_NO_EN_PUNTO": "Se usa cuando el usuario entrega una hora no valida para agenda, por ejemplo minutos no permitidos.",
    "PLANTILLA_HORA_FUERA_DE_ATENCION": "Se usa cuando la hora solicitada esta fuera del horario de atencion de la agencia.",
    "PLANTILLA_HORA_OCUPADA": "Se usa cuando la hora solicitada no tiene cupo disponible y se deben ofrecer alternativas.",
    "PLANTILLA_VISITA_SIN_CITA": "Se usa cuando el usuario pregunta si puede asistir sin cita previa o quiere llegar directamente.",
    "PLANTILLA_CONFIRMACION": "Se usa unicamente cuando todos los datos y validaciones de agendamiento ya estan completos.",
    "PLANTILLA_INFO_INSTITUCIONAL": "Se usa cuando el usuario pregunta por la empresa, trayectoria, confianza o informacion institucional.",
    "PLANTILLA_INFO_SERVICIOS_OFRECEN": "Se usa cuando el usuario pregunta que servicios ofrece la empresa.",
    "PLANTILLA_VALOR_CONSULTA": "Se usa cuando el usuario pregunta por el costo de la consulta o valoracion.",
    "PLANTILLA_INFO_CUOTA_CREDITO": "Se usa cuando el usuario pregunta por cuotas, financiamiento o credito.",
    "PLANTILLA_PARQUEADEROS": "Se usa cuando el usuario pregunta si existe parqueadero o facilidad para llegar a la sede.",
    "PLANTILLA_EN_CAMINO": "Se usa cuando el usuario avisa que ya esta en camino a la agencia o cita.",
    "PLANTILLA_CLIENTE_TIENE_DUDAS": "Se usa cuando el usuario indica dudas generales y necesita que el bot lo acompane antes de avanzar.",
    "PLANTILLA_EXPERIENCIA_DOCTOR": "Se usa cuando el usuario pregunta por la experiencia del doctor o del equipo profesional.",
    "PLANTILLA_CLIENTE_ASESOR_CONSULTA": "Se usa cuando el usuario solicita hablar con asesor o requiere ayuda humana para resolver una consulta.",
    "PLANTILLA_AGRADECIMIENTO_UTIL": "Se usa cuando el usuario agradece y la conversacion fue util o quedo bien orientada.",
    "PLANTILLA_AGRADECIMIENTO_LIMITADO": "Se usa cuando el usuario agradece, pero aun no hay una accion concreta o no se debe avanzar mas.",
    "PLANTILLA_TECNOLOGIA_EMPRESA": "Se usa cuando el usuario pregunta por tecnologia, equipos, tomografia o herramientas usadas por la empresa.",
    "PLANTILLA_LENGUAJE_OFENSIVO_CIERRE": "Se usa cuando el usuario utiliza lenguaje ofensivo y se debe cerrar con respeto.",
    "PLANTILLA_NO_AYUDA": "Se usa cuando el usuario rechaza ayuda o indica que no necesita nada mas.",
    "PLANTILLA_EQUIVOCACION_CHAT": "Se usa cuando el usuario escribio por error o confundio el chat con otro destino.",
    "PLANTILLA_GARANTIA_IMPLANTE": "Se usa cuando el usuario pregunta por garantia de implantes o respaldo del tratamiento.",
    "PLANTILLA_CONFIABILIDAD_DOCTOR": "Se usa cuando el usuario pregunta por seguridad, confianza o calidad profesional del doctor.",
    "PLANTILLA_METODO_DE_PAGO": "Se usa cuando el usuario pregunta por formas de pago aceptadas.",
    "PLANTILLA_RADIOGARFIAS": "Se usa cuando el usuario pregunta por radiografias o examenes relacionados.",
    "PLANTILLA_EVITAR_ALIMENTOS": "Se usa cuando el usuario pregunta si debe evitar alimentos antes o despues de un procedimiento.",
    "PLANTILLA_TOMOGRAFIAS": "Se usa cuando el usuario pregunta por tomografias o examenes necesarios para valoracion.",
    "PLANTILLA_VALOR_PRECIO_IMPLANTE": "Se usa cuando el usuario pregunta directamente por precio de implante y se debe evitar inventar valores.",
    "PLANTILLA_MATERIAL_IMPLANTE": "Se usa cuando el usuario pregunta por el material del implante o caracteristicas del tratamiento.",
    "PLANTILLA_MARCA_IMPLANTE": "Se usa cuando el usuario pregunta por la marca del implante.",
    "PLANTILLA_ENDODONCIA": "Se usa cuando el usuario pregunta por endodoncia u otro servicio relacionado.",
    "PLANTILLA_SIN_CITAS": "Se usa cuando el usuario consulta sus citas, pero no existen citas registradas para mostrar.",
    "PLANTILLA_LISTADO_CITAS_PARA_ELIMINAR": "Se usa cuando el usuario desea cancelar, eliminar o mover una cita y el bot debe listar opciones.",
    "PLANTILLA_NUMERO_OBLIGATORIO": "Se usa cuando el canal no entrega telefono y el bot necesita pedir numero obligatorio para gestionar la cita.",
    "PLANTILLA_MIS_CITAS": "Se usa cuando el usuario pide revisar sus citas registradas.",
}

SUGGESTED_LABELS: dict[str, str] = {
    "PLANTILLA_BIENVENIDA": "bienvenida",
    "PLANTILLA_CLIENTE_INTERESADO": "interesado",
    "PLANTILLA_INTERES_CON_CIUDAD": "interesado",
    "PLANTILLA_INTERES_CON_AGENCIA": "interesado",
    "PLANTILLA_DATOS_CITA": "interesado",
    "PLANTILLA_CONFIRMACION": "cita_agendada",
    "PLANTILLA_EN_CAMINO": "cita_agendada",
    "PLANTILLA_VISITA_SIN_CITA": "interesado",
    "PLANTILLA_NO_INTERES": "desinteresado",
    "PLANTILLA_NO_APLICA": "desinteresado",
    "PLANTILLA_LENGUAJE_OFENSIVO_CIERRE": "desinteresado",
    "PLANTILLA_NO_AYUDA": "desinteresado",
    "PLANTILLA_EQUIVOCACION_CHAT": "desinteresado",
    "PLANTILLA_DESCONFIANZA": "tiene_dudas",
    "PLANTILLA_CLIENTE_TIENE_DUDAS": "tiene_dudas",
    "PLANTILLA_CLIENTE_ASESOR_CONSULTA": "tiene_dudas",
    "PLANTILLA_AGRADECIMIENTO_UTIL": "tiene_dudas",
    "PLANTILLA_AGRADECIMIENTO_LIMITADO": "tiene_dudas",
    "PLANTILLA_DATOS_INCOMPLETOS": "interesado",
    "PLANTILLA_FECHA_PASADA": "interesado",
    "PLANTILLA_NO_HOY_EMPRESA": "interesado",
    "PLANTILLA_HORA_NO_EN_PUNTO": "interesado",
    "PLANTILLA_HORA_FUERA_DE_ATENCION": "interesado",
    "PLANTILLA_HORA_OCUPADA": "interesado",
    "PLANTILLA_SIN_CITAS": "cita_agendada",
    "PLANTILLA_LISTADO_CITAS_PARA_ELIMINAR": "cita_agendada",
    "PLANTILLA_NUMERO_OBLIGATORIO": "interesado",
    "PLANTILLA_MIS_CITAS": "cita_agendada",
}


def rgb(hex_color: str) -> RGBColor:
    value = hex_color.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def normalize_brand(text: str) -> str:
    return re.sub(r"osnex", "EMPRESA", text, flags=re.IGNORECASE)


def clean_lines(lines: list[str]) -> list[str]:
    cleaned = [normalize_brand(line.rstrip()) for line in lines]
    while cleaned and not cleaned[0].strip():
        cleaned.pop(0)
    while cleaned and not cleaned[-1].strip():
        cleaned.pop()
    return cleaned


def infer_scenario(name: str) -> str:
    readable = name.replace("PLANTILLA_", "").replace("_", " ").lower()
    return f"Se usa cuando el flujo detecta la situacion relacionada con {readable}."


def infer_label(name: str) -> str:
    if name in SUGGESTED_LABELS:
        return SUGGESTED_LABELS[name]
    if any(term in name for term in ["VALOR", "INFO", "UBICACION", "CIUDAD", "SERVICIO", "PARQUEADEROS", "DOCTOR", "GARANTIA", "PAGO", "RADIO", "TOMOGRAF", "MATERIAL", "MARCA", "ENDODONCIA", "TECNOLOGIA"]):
        return "solicita_informacion"
    if any(term in name for term in ["DATOS", "HORA", "FECHA", "CITA", "REFERENCIA"]):
        return "interesado"
    return "por_definir"


def extract_templates() -> list[TemplateBlock]:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(f"No existe el documento fuente: {SOURCE_DOCX}")
    if not PROMPT_TXT.exists():
        raise FileNotFoundError(f"No existe el prompt de referencia: {PROMPT_TXT}")

    # The prompt is intentionally read so this generator depends on the same
    # reference material used to define scenarios, without copying internal
    # prompt instructions into the client-facing Word.
    _prompt_reference = PROMPT_TXT.read_text(encoding="utf-8", errors="ignore")

    source = Document(SOURCE_DOCX)
    templates: list[TemplateBlock] = []
    current_name: str | None = None
    current_lines: list[str] = []

    def close_current() -> None:
        nonlocal current_name, current_lines
        if not current_name:
            return
        normalized_name = normalize_brand(current_name)
        body = "\n".join(clean_lines(current_lines))
        scenario = SCENARIOS.get(normalized_name, infer_scenario(normalized_name))
        suggested_label = infer_label(normalized_name)
        templates.append(
            TemplateBlock(
                original_name=current_name,
                name=normalized_name,
                body=body,
                scenario=scenario,
                suggested_label=suggested_label,
            )
        )
        current_name = None
        current_lines = []

    for paragraph in source.paragraphs:
        raw = paragraph.text
        stripped = raw.strip()
        match = HEADER_RE.match(stripped)
        if match:
            close_current()
            current_name = match.group(1)
            remainder = stripped[len(match.group(1)) :].lstrip()
            current_lines = [remainder] if remainder else []
            continue
        if current_name is not None:
            current_lines.append(raw)

    close_current()
    return templates


def set_cell_shading(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def set_table_borders(table, color: str = COLORS["line"]) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement("w:" + edge)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    table._tbl.tblPr.append(borders)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def write_cell(cell, text: str, *, bold: bool = False, color: str | None = None, size: float = 8.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for index, part in enumerate(str(text).split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    set_table_borders(table)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if row_index == 0:
                set_cell_shading(cell, COLORS["blue"])
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = rgb(COLORS["white"])
            else:
                set_cell_shading(cell, COLORS["light"] if row_index % 2 == 0 else COLORS["white"])
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(7.8 if row_index else 8.2)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        write_cell(table.cell(0, idx), header, bold=True, color=COLORS["white"], size=8.2)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row[: len(headers)]):
            write_cell(cells[idx], value, size=7.8)
    style_table(table)
    document.add_paragraph()


def add_note(document: Document, title: str, text: str) -> None:
    table = document.add_table(rows=1, cols=1)
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS["soft_blue"])
    write_cell(cell, f"{title}: {text}", bold=False, size=8.4)
    document.add_paragraph()


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    for style_name, size, color in [
        ("Heading 1", 15, COLORS["blue"]),
        ("Heading 2", 12.5, COLORS["navy"]),
        ("Heading 3", 10.5, COLORS["green"]),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"{DOC_CODE} | v{DOC_VERSION} | {DOC_STATUS}")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["slate"])

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Simplia Chatbot - Plantillas de Uso ISO 10013")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["slate"])


def add_title(document: Document) -> None:
    if LOGO_PATH.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.2))

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SIMPLIA")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = rgb(COLORS["blue"])

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Plantillas de Uso del Chatbot")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(21)
    run.font.color.rgb = rgb(COLORS["navy"])

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Manual ISO 10013 para revision operativa de scripts conversacionales")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = rgb(COLORS["slate"])

    document.add_paragraph()
    add_table(
        document,
        ["Campo", "Valor"],
        [
            ["Codigo", DOC_CODE],
            ["Version", DOC_VERSION],
            ["Fecha", DOC_DATE],
            ["Estado", DOC_STATUS],
            ["Documento fuente", SOURCE_DOCX.name],
            ["Referencia de intencion", PROMPT_TXT.name],
        ],
    )


def add_template_section(document: Document, template: TemplateBlock, index: int) -> None:
    heading = document.add_heading(f"{index}. {template.name}", level=2)
    heading.paragraph_format.space_before = Pt(8)
    heading.paragraph_format.space_after = Pt(3)

    add_table(
        document,
        ["Elemento", "Detalle"],
        [
            ["Situacion en que se usa", template.scenario],
            ["Etiqueta sugerida por validar", template.suggested_label],
            ["Etiqueta final", "____________________________________________"],
            ["Color de etiqueta", "____________________________________________"],
        ],
    )

    table = document.add_table(rows=1, cols=1)
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS["code"])
    write_cell(cell, f"Texto actual de la plantilla\n\n{template.body}", size=8.4)
    document.add_paragraph()

    add_table(
        document,
        ["Revision del cliente", "Respuesta"],
        [
            ["Aprobado", "[ ] Si    [ ] No"],
            ["Solicita cambio", "[ ] Si    [ ] No"],
            ["Nuevo texto sugerido", "\n\n\n\n\n"],
            ["Observaciones", "\n\n\n"],
        ],
    )


def build_document(templates: list[TemplateBlock]) -> Document:
    document = Document()
    configure_document(document)
    add_title(document)

    document.add_heading("1. Objetivo y alcance", level=1)
    document.add_paragraph(
        "Este documento permite revisar todas las plantillas que puede usar el agente conversacional. "
        "Cada bloque explica cuando se usa la plantilla, muestra el texto actual y deja espacios para validar etiqueta, color y cambios."
    )
    document.add_paragraph(
        "El documento esta preparado para que el cliente pueda leer, aprobar o editar directamente los mensajes antes de dejarlos como version final del chatbot."
    )

    document.add_heading("2. Instrucciones de revision", level=1)
    add_note(
        document,
        "Uso",
        "Revise cada plantilla en orden. Si el texto esta correcto, marque Aprobado. Si desea cambio, marque Solicita cambio y escriba el texto final en el campo Nuevo texto sugerido.",
    )
    add_table(
        document,
        ["Campo", "Como llenarlo"],
        [
            ["Etiqueta final", "Indicar la categoria que debe activarse para esa situacion."],
            ["Color de etiqueta", "Indicar el color que se usara para identificar esa categoria en el pipeline."],
            ["Nuevo texto sugerido", "Escribir la version final si el mensaje actual debe cambiar."],
            ["Observaciones", "Anotar condiciones, excepciones o comentarios del equipo."],
        ],
    )

    document.add_heading("3. Etiquetas base sugeridas", level=1)
    add_table(
        document,
        ["Etiqueta", "Uso general"],
        [
            ["bienvenida", "Inicio de conversacion o primer saludo."],
            ["solicita_informacion", "Consultas sobre servicios, precios, ubicaciones, requisitos o detalles del negocio."],
            ["interesado", "Usuario con intencion de avanzar, elegir sede o agendar."],
            ["desinteresado", "Usuario fuera de alcance, sin interes o cierre de conversacion."],
            ["cita_agendada", "Confirmacion, consulta o gestion de citas."],
            ["tiene_dudas", "Usuario con objeciones, dudas o solicitud de apoyo antes de decidir."],
        ],
    )

    document.add_section(WD_SECTION_START.NEW_PAGE)
    document.add_heading("4. Plantillas para revision", level=1)
    document.add_paragraph(f"Total de plantillas incluidas desde el documento fuente: {len(templates)}.")

    for index, template in enumerate(templates, start=1):
        add_template_section(document, template, index)

    document.add_heading("5. Cierre de validacion", level=1)
    add_table(
        document,
        ["Campo", "Respuesta"],
        [
            ["Responsable de revision", "____________________________________________"],
            ["Fecha de revision", "____ / ____ / ______"],
            ["Version aprobada para chatbot", "____________________________________________"],
            ["Comentarios finales", "\n\n\n\n"],
        ],
    )

    return document


def main() -> None:
    templates = extract_templates()
    document = build_document(templates)
    document.save(DOCX_OUT)
    print(f"Generated {DOCX_OUT.name} with {len(templates)} templates")


if __name__ == "__main__":
    main()
