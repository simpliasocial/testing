from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DOCX = ROOT / "plantilla_iso_9001_proceso.docx"
LOGO_PATH = ROOT / "logo_simplia.png"
SOURCE_DOCX = ROOT / "documentos.docx"

ONBOARDING_DOCX = ROOT / "Onboarding_Funcional_Simplia_Chatbot_ISO9001.docx"
GUIDE_PLANTILLAS = ROOT / "Guia_Simplia_Plantillas_Conversacion.docx"
GUIDE_DATOS = ROOT / "Guia_Simplia_Datos_A_Recoger.docx"
GUIDE_CATEGORIAS = ROOT / "Guia_Simplia_Categorizacion_Leads.docx"
GUIDE_FOLLOWUP = ROOT / "Guia_Simplia_Followup_Plantillas.docx"

DOC_CODE = "PRO-ONB-FUNC-001"
DOC_VERSION = "1.1"
ISSUE_DATE = "06/05/2026"
NEXT_REVIEW = "06/11/2026"
DOC_STATUS = "Borrador / Vigente al aprobarse"

BLANK = "______________________________"
LONG_BLANK = "____________________________________________________________"

COLORS = {
    "navy": "0f2344",
    "blue": "274690",
    "soft_blue": "eaf1ff",
    "light": "f8fafc",
    "green": "0a9b6f",
    "yellow": "fef3c7",
    "line": "d9e2ef",
    "white": "ffffff",
    "slate": "64748b",
}


def rgb(hex_color: str) -> RGBColor:
    value = hex_color.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def clear_body(document: Document) -> None:
    body = document._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    for style_name, size, color in [
        ("Heading 1", 15, COLORS["blue"]),
        ("Heading 2", 12.5, COLORS["navy"]),
        ("Heading 3", 10.5, COLORS["green"]),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def configure_iso_header_footer(document: Document) -> None:
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"{DOC_CODE} | v{DOC_VERSION} | {DOC_STATUS}")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["slate"])

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Simplia Chatbot - Onboarding Funcional ISO 9001")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["slate"])


def set_cell_shading(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def set_table_borders(table, color: str = COLORS["line"]) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement("w:" + edge)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    table._tbl.tblPr.append(borders)


def clear_cell(cell) -> None:
    cell.text = ""
    for paragraph in cell.paragraphs:
        paragraph.text = ""


def write_cell(cell, value: str, *, bold: bool = False, color: str | None = None, size: float = 8.0) -> None:
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    for index, part in enumerate(str(value).split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(document: Document, headers: list[str], rows: list[list[str]], *, font_size: float = 8.0) -> None:
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    set_table_borders(table)

    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        write_cell(cell, header, bold=True, color=COLORS["white"], size=8.2)
        set_cell_shading(cell, COLORS["blue"])

    for row_index, row in enumerate(rows, start=1):
        fill = COLORS["light"] if row_index % 2 == 0 else COLORS["white"]
        for col, value in enumerate(row):
            cell = table.cell(row_index, col)
            write_cell(cell, value, size=font_size)
            set_cell_shading(cell, fill)
    document.add_paragraph()


def add_paragraph(document: Document, text: str, *, bold: bool = False, color: str | None = None) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    if color:
        run.font.color.rgb = rgb(color)


def add_bullet(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9.2)


def add_numbered_line(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.15)
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9.2)


def add_template_block(document: Document, title: str, text: str) -> None:
    document.add_heading(title, level=2)
    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS["light"])
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    for index, part in enumerate(text.strip().split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        run.font.name = "Arial"
        run.font.size = Pt(9.0)
    document.add_paragraph()


def add_simple_cover(document: Document, title: str, subtitle: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        paragraph.add_run().add_picture(str(LOGO_PATH), width=Inches(1.25))

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(18)
    run.font.color.rgb = rgb(COLORS["blue"])

    subtitle_p = document.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle_p.add_run(subtitle)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = rgb(COLORS["slate"])
    document.add_paragraph()


def build_guide_plantillas() -> None:
    document = Document()
    configure_document(document)
    add_simple_cover(
        document,
        "Simplia - Plantillas de Conversacion",
        "Mensajes base para revisar y aprobar antes de configurar el chatbot.",
    )
    add_paragraph(
        document,
        "Usar EMPRESA para el nombre del negocio y reemplazar los campos entre corchetes con datos reales del lead o de la cita.",
    )

    add_template_block(
        document,
        "PLANTILLA_BIENVENIDA",
        """
Hola! Bienvenido a EMPRESA.
Somos [descripcion breve del negocio].
Podemos ayudarle con informacion, dudas y agendamiento de cita.
Le gustaria agendar una cita?
""",
    )
    add_template_block(
        document,
        "PLANTILLA_DERIVACION_HUMANA",
        """
Gracias, [nombre]. Para ayudarle mejor, voy a derivar su caso con un asesor de EMPRESA.
En breve una persona del equipo revisara su solicitud y se comunicara con usted por este mismo medio.
""",
    )
    add_template_block(
        document,
        "PLANTILLA_DATOS_CITA",
        """
EMPRESA, encantados de atenderle [nombre].
Para agendar su cita y brindarle una atencion segura y personalizada, necesitamos estos datos:
- Nombre completo
- Fecha de visita
- Hora de visita
- Agencia
- Celular
- Correo
- Ciudad
- Edad

Estamos atentos a sus datos para agendar su cita lo antes posible.

Al registrar sus datos, usted acepta que sean tratados conforme a la Ley Organica de Proteccion de Datos Personales de Ecuador. Puede solicitar acceso, correccion o eliminacion de sus datos en cualquier momento.
""",
    )
    add_template_block(
        document,
        "PLANTILLA_CONFIRMACION_CITA",
        """
Hemos recibido todos sus datos, muchas gracias.
Su cita ha sido agendada con exito, [nombre].
Confirmamos su cita en EMPRESA el [fecha] a las [hora] en la sede [agencia].
Numero celular: [celular]

Le sugerimos llegar unos minutos antes. Si no puede asistir, le agradecemos notificarnos con anticipacion para liberar el cupo.
Le esperamos!
""",
    )
    add_template_block(
        document,
        "PLANTILLA_NO_APLICA",
        """
Hola [nombre], gracias por su mensaje.
Parece que su consulta no corresponde a los servicios de EMPRESA.
De igual manera, quedamos atentos si mas adelante necesita informacion sobre nuestros servicios.
Hasta pronto.
""",
    )

    safe_save(document, GUIDE_PLANTILLAS)


def build_guide_datos() -> None:
    document = Document()
    configure_document(document)
    add_simple_cover(
        document,
        "Simplia - Datos a Recoger",
        "Guia simple para definir que datos pide el bot y en que momento los solicita.",
    )

    document.add_heading("Datos base para cita", level=1)
    for item in [
        "nombre_completo",
        "fecha_visita",
        "hora_visita",
        "agencia",
        "celular",
        "correo",
        "ciudad",
        "edad",
        "otro: ______________________________",
    ]:
        add_bullet(document, item)

    document.add_heading("Modo 1: filtro obligatorio en primera interaccion", level=1)
    add_paragraph(
        document,
        "El bot pide datos desde el inicio antes de continuar con otras intenciones. Si falta un dato obligatorio, el bot insiste de forma educada hasta completar la informacion minima.",
    )
    add_paragraph(document, "Datos obligatorios sugeridos:", bold=True)
    add_bullet(document, "nombre_completo")
    add_bullet(document, "celular")
    add_paragraph(document, "Datos opcionales sugeridos:", bold=True)
    add_bullet(document, "correo")
    add_bullet(document, "ciudad")
    add_bullet(document, "edad")
    add_template_block(
        document,
        "PLANTILLA_BIENVENIDA_CON_FILTRO",
        """
Hola! Gracias por escribirnos a EMPRESA.
Para una atencion personalizada, por favor indiquenos:
- Nombre completo (obligatorio)
- Numero de telefono (obligatorio)
- Correo electronico (opcional)
- Ciudad (opcional)
- Edad (opcional)

Cuando tengamos los datos obligatorios, continuamos con su solicitud.
""",
    )

    document.add_heading("Modo 2: sin filtro inicial", level=1)
    add_paragraph(
        document,
        "El bot saluda, responde informacion y solo recoge datos completos cuando el usuario quiere agendar una cita.",
    )
    add_paragraph(document, "Datos minimos al agendar:", bold=True)
    add_bullet(document, "fecha_visita")
    add_bullet(document, "hora_visita")
    add_bullet(document, "agencia")
    add_paragraph(document, "Datos extra al agendar si no se tienen todavia:", bold=True)
    add_bullet(document, "nombre_completo")
    add_bullet(document, "celular")
    add_bullet(document, "correo")
    add_bullet(document, "ciudad")
    add_bullet(document, "edad")
    add_bullet(document, "otro: ______________________________")

    document.add_heading("Decision que debe confirmar la empresa", level=1)
    add_bullet(document, "Usar filtro obligatorio al inicio: Si / No")
    add_bullet(document, "Datos obligatorios al inicio: ______________________________")
    add_bullet(document, "Datos opcionales al inicio: ______________________________")
    add_bullet(document, "Datos obligatorios para agendar cita: ______________________________")

    safe_save(document, GUIDE_DATOS)


def build_guide_categorias() -> None:
    document = Document()
    configure_document(document)
    add_simple_cover(
        document,
        "Simplia - Categorizacion de Leads",
        "Etiquetas base para ordenar leads y saber si las asigna el bot o una persona.",
    )
    add_table(
        document,
        ["Etiqueta", "Quien la asigna", "Que significa"],
        [
            ["bienvenida", "Bot", "Primer contacto o lead sin intencion clara todavia."],
            ["solicita_informacion", "Bot", "El lead pide informacion, costos, servicios, beneficios, sedes u horarios."],
            ["interesado", "Bot", "El lead muestra intencion de avanzar, comprar o agendar."],
            ["desinteresado", "Bot", "El lead rechaza, esta fuera del negocio o envia contenido irrelevante."],
            ["cita_agendada", "Bot", "Cita creada automaticamente por el bot."],
            ["seguimiento_humano", "Bot", "Caso que debe tomar una persona o que ya no debe responder el bot."],
            ["venta_exitosa", "Manual", "La venta u operacion se concreto."],
            ["cita_agendado_humano", "Manual", "Una persona agenda o modifica una cita manualmente."],
        ],
        font_size=7.8,
    )
    document.add_heading("Confirmacion de la empresa", level=1)
    add_bullet(document, "Etiquetas aprobadas: Si / No")
    add_bullet(document, "Etiquetas que desea quitar: ______________________________")
    add_bullet(document, "Etiquetas que desea agregar: ______________________________")
    add_bullet(document, "Observaciones: ______________________________")

    safe_save(document, GUIDE_CATEGORIAS)


def build_guide_followup() -> None:
    document = Document()
    configure_document(document)
    add_simple_cover(
        document,
        "Simplia - Followup Plantillas",
        "Scripts de retoma para leads que dejan de responder.",
    )
    add_paragraph(
        document,
        "Estos mensajes se usan cuando el lead queda sin respuesta. Si responde, el bot continua el flujo segun la intencion. Si no responde despues del segundo intento, pasa a seguimiento_humano.",
    )
    add_template_block(
        document,
        "FOLLOWUP_20_MIN",
        """
Hola [nombre], seguimos atentos para ayudarle en EMPRESA.
Si desea agendar su cita, puede indicarnos fecha, hora y agencia de preferencia.
Tambien podemos resolver cualquier duda antes de avanzar.
""",
    )
    add_template_block(
        document,
        "FOLLOWUP_3_HORAS",
        """
Hola [nombre], le escribimos nuevamente de EMPRESA.
Si aun desea informacion o quiere agendar una cita, respondanos por este medio y con gusto le ayudamos.
Si no recibimos respuesta, dejaremos su caso para seguimiento humano.
""",
    )
    document.add_heading("Regla operativa", level=1)
    add_bullet(document, "Enviar retoma a los 20 minutos si el lead no responde.")
    add_bullet(document, "Enviar retoma a las 3 horas si sigue sin responder.")
    add_bullet(document, "Si no responde despues de la retoma de 3 horas, cambiar estado a seguimiento_humano.")

    safe_save(document, GUIDE_FOLLOWUP)


def add_onboarding_cover(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        paragraph.add_run().add_picture(str(LOGO_PATH), width=Inches(1.45))

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ONBOARDING FUNCIONAL - SIMPLIA CHATBOT")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(19)
    run.font.color.rgb = rgb(COLORS["blue"])

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Documento para levantar la informacion funcional necesaria para que el agente atienda, clasifique y agende correctamente"
    )
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.color.rgb = rgb(COLORS["slate"])

    add_table(
        document,
        ["Campo", "Valor"],
        [
            ["Codigo", DOC_CODE],
            ["Version", DOC_VERSION],
            ["Fecha de emision", ISSUE_DATE],
            ["Proxima revision", NEXT_REVIEW],
            ["Estado", DOC_STATUS],
            ["Empresa cliente", BLANK],
            ["Encargado principal de la empresa", BLANK],
            ["Correo / telefono", BLANK],
        ],
        font_size=7.8,
    )

    add_paragraph(
        document,
        "Uso del documento",
        bold=True,
        color=COLORS["navy"],
    )
    add_paragraph(
        document,
        "Completar cada bloque con links, responsables y observaciones. Este documento recoge solo informacion funcional para configurar la atencion, clasificacion, agenda y seguimiento del bot.",
    )


def add_requirement_block(
    document: Document,
    number: int,
    title: str,
    description: str,
    bullets: list[str],
    guide: Path | None = None,
) -> None:
    document.add_heading(f"{number}. {title}", level=1)
    add_paragraph(document, description)
    if guide:
        add_bullet(document, f"Documento guia Simplia: {guide.name}")
    for item in bullets:
        add_bullet(document, item)
    add_bullet(document, f"Link compartido por la empresa: {LONG_BLANK}")
    add_bullet(document, f"Encargado de la empresa: {BLANK}")
    add_bullet(document, "Fecha compromiso: dd/mm/aaaa")
    add_bullet(document, "Estado: Pendiente / En revision / Completo / No aplica")


def build_onboarding() -> None:
    document = Document(TEMPLATE_DOCX)
    clear_body(document)
    configure_document(document)
    configure_iso_header_footer(document)
    add_onboarding_cover(document)

    add_requirement_block(
        document,
        1,
        "Informacion sedes/agencias",
        "Brindar informacion de agencias con sus horarios de atencion de lunes a domingo, direcciones y links de Maps por cada agencia.",
        [
            "Incluir nombre de cada agencia o sede.",
            "Incluir horarios por dia o aclarar si todos los dias manejan el mismo horario.",
            "Incluir direccion exacta y link de Google Maps por agencia.",
        ],
    )

    add_requirement_block(
        document,
        2,
        "Plantillas de conversacion",
        "Mensajes exactos que el bot debe usar para bienvenida, derivacion humana, agendar cita, confirmacion de cita y no aplica.",
        [
            "Aprobar tono, uso de emojis si aplica y campos variables.",
            "Usar placeholders claros como EMPRESA, [nombre], [fecha], [hora], [agencia] y [celular].",
        ],
        GUIDE_PLANTILLAS,
    )

    add_requirement_block(
        document,
        3,
        "Flujo conversacional de inicio a fin",
        "Explicar como atiende hoy la empresa: desde el primer mensaje hasta informacion, dudas, agenda, venta o seguimiento humano.",
        [
            "Compartir ejemplos con fotos de conversaciones reales de inicio a fin.",
            "Incluir casos donde el cliente pide informacion, agenda, no responde, rechaza o requiere asesor humano.",
            "El objetivo es entender como se viene manejando el proceso completo.",
        ],
    )

    add_requirement_block(
        document,
        4,
        "Datos para agendar",
        "Definir que datos debe guardar el bot para crear una cita y en que momento debe pedirlos.",
        [
            "Datos base: nombre_completo, fecha_visita, hora_visita, agencia, celular, correo, ciudad y edad.",
            "Confirmar si existe filtro obligatorio desde la primera interaccion o si solo se usa plantilla de bienvenida.",
            "Si hay filtro obligatorio, definir datos obligatorios y opcionales. Si falta un dato obligatorio, el bot insiste hasta completar la informacion minima.",
            "Si no hay filtro inicial, el bot recoge los datos cuando el usuario quiere agendar la cita.",
            "Para agendar, confirmar si ademas de fecha_visita, hora_visita y agencia se requiere otro dato extra.",
        ],
        GUIDE_DATOS,
    )

    add_requirement_block(
        document,
        5,
        "Etiquetas y categorizacion de leads",
        "Confirmar que representa cada etiqueta y cuales las asigna el bot o una persona.",
        [
            "Revisar si las etiquetas base se aprueban tal como estan.",
            "Indicar si desea agregar, quitar o cambiar alguna etiqueta.",
            "La guia explica para que sirve cada categorizacion y cuando se usa.",
        ],
        GUIDE_CATEGORIAS,
    )

    add_requirement_block(
        document,
        6,
        "Score y vocabulario del negocio",
        "Confirmar palabras y frases que indican interes, compra, agenda, asesor, rechazo, fuera de negocio o riesgo.",
        [
            "Referencia interna: Score Simplia leads.pdf.",
            "Agregar vocabulario propio del negocio para mejorar la clasificacion.",
            "Incluir ejemplos reales de frases que usa el cliente.",
        ],
    )

    add_requirement_block(
        document,
        7,
        "Ejemplos reales para entrenamiento",
        "Compartir capturas de conversaciones reales con clientes para entrenar y validar el comportamiento del bot.",
        [
            "Objetivo sugerido: 100 fotos reales.",
            "Incluir conversaciones completas de inicio a fin.",
            "Cubrir buenos leads, dudas frecuentes, agenda, rechazo, fuera de negocio y seguimiento humano.",
        ],
    )

    add_requirement_block(
        document,
        8,
        "Preguntas frecuentes",
        "Documentar FAQ del negocio con preguntas y respuestas aprobadas.",
        [
            "Incluir servicios que ofrece la empresa.",
            "Incluir costos, requisitos necesarios, horarios, sedes, beneficios y restricciones.",
            "Las respuestas deben estar aprobadas por la empresa antes de configurar el bot.",
        ],
    )

    add_requirement_block(
        document,
        9,
        "Campanas de retoma de follow up a los 20 minutos y a las 3 horas",
        "Definir el script exacto de retoma a los 20 minutos y el script exacto de retoma a las 3 horas.",
        [
            "Si el lead responde, el bot continua el flujo segun la intencion.",
            "Si no responde despues del follow up de 3 horas, el lead pasa a seguimiento_humano.",
        ],
        GUIDE_FOLLOWUP,
    )

    document.add_heading("Checklist final antes de construir", level=1)
    for item in [
        "Informacion de sedes/agencias completa.",
        "Plantillas de conversacion aprobadas.",
        "Flujo conversacional de inicio a fin compartido.",
        "Datos de agenda definidos como obligatorios u opcionales.",
        "Filtro inicial definido: si aplica o no aplica.",
        "Etiquetas y responsables aprobados.",
        "Score y vocabulario del negocio confirmados.",
        "Ejemplos reales y FAQs compartidos.",
        "Follow up de 20 minutos y 3 horas aprobado.",
        "Cliente aprueba que la informacion esta lista para implementacion.",
    ]:
        add_bullet(document, f"OK / No OK / N.A. - {item}")

    document.add_heading("Control de cambios y aprobacion", level=1)
    add_table(
        document,
        ["Version", "Fecha", "Cambio", "Responsable"],
        [
            ["1.1", ISSUE_DATE, "Se eliminan requisitos de WhatsApp/cuentas y se ordena el levantamiento funcional en bloques simples.", "Simplia"],
            ["1.x", "dd/mm/aaaa", "Resumen del cambio.", BLANK],
        ],
        font_size=7.6,
    )
    add_table(
        document,
        ["Elaborado por", "Revisado por", "Aprobado por"],
        [
            ["Simplia / Responsable onboarding", "Encargado de la empresa", "Responsable autorizado"],
            ["Firma / fecha", "Firma / fecha", "Firma / fecha"],
        ],
        font_size=7.8,
    )

    safe_save(document, ONBOARDING_DOCX)


def safe_save(document: Document, path: Path) -> None:
    try:
        document.save(path)
    except PermissionError as exc:
        raise SystemExit(
            f"No se pudo guardar {path.name}. Cierre el archivo en Word y vuelva a ejecutar este script."
        ) from exc


def verify_source_exists() -> None:
    if not SOURCE_DOCX.exists():
        raise SystemExit(f"No se encontro el documento fuente: {SOURCE_DOCX}")


def build_all() -> None:
    verify_source_exists()
    build_guide_plantillas()
    build_guide_datos()
    build_guide_categorias()
    build_guide_followup()
    build_onboarding()


if __name__ == "__main__":
    build_all()
    for path in [GUIDE_PLANTILLAS, GUIDE_DATOS, GUIDE_CATEGORIAS, GUIDE_FOLLOWUP, ONBOARDING_DOCX]:
        print(f"Documento generado: {path}")
