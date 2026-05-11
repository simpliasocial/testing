from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
PDF_OUT = ROOT / "Score Simplia leads.pdf"
DOCX_OUT = ROOT / "Score Simplia leads.docx"
XLSX_OUT = ROOT / "Score Simplia leads.xlsx"

BLUE = colors.HexColor("#274690")
NAVY = colors.HexColor("#0f2344")
GREEN = colors.HexColor("#0a9b6f")
LIGHT = colors.HexColor("#f8fafc")
SOFT_BLUE = colors.HexColor("#eaf1ff")
LINE = colors.HexColor("#d9e2ef")
YELLOW = colors.HexColor("#fef3c7")
RED = colors.HexColor("#fee2e2")


STATES = [
    [
        "FRIO",
        "Menor a 45 o sin puntaje",
        "Lead con señal inicial, baja o todavia sin intencion comercial clara. Puede incluir saludo, informacion general, rechazo leve o poca data.",
        "Responder con informacion aprobada, invitar a avanzar y mantener automatizacion.",
    ],
    [
        "TIBIO",
        "45 a 69",
        "Lead con señales accionables: pregunta por precio, disponibilidad, servicio especifico, requisitos o continuidad.",
        "Priorizar respuesta, resolver duda concreta e invitar a cita, datos o siguiente paso.",
    ],
    [
        "CALIENTE",
        "70 o mas",
        "Lead con intencion fuerte: quiere comprar, contratar, reservar, agendar, pagar o hablar con asesor.",
        "Prioridad alta. Solicitar datos de agenda, confirmar cita o derivar si el flujo lo requiere.",
    ],
]

POSITIVE_MATRIX = [
    ["Interes inicial", "Pregunta por un producto o servicio.", "+20"],
    ["Interes inicial", "Escribe desde campaña, anuncio o publicacion especifica.", "+18"],
    ["Interes inicial", 'Escribe "info", "precio", "disponible", "me interesa" o similar.', "+15"],
    ["Interes inicial", "Solo saluda sin contexto comercial.", "+5"],
    ["Interes inicial", "Mensaje irrelevante o sin relacion comercial clara.", "0"],
    ["Intencion comercial", "Quiere comprar, reservar, aplicar, contratar o iniciar el proceso.", "+35"],
    ["Intencion comercial", "Pide agendar una cita, reunion, valoracion, llamada o demo.", "+35"],
    ["Intencion comercial", "Pide hablar con asesor, ejecutivo, vendedor o persona del equipo.", "+30"],
    ["Intencion comercial", "Pregunta por precio, costo, valor, cotizacion, promocion o formas de pago.", "+25"],
    ["Intencion comercial", "Pregunta por disponibilidad, stock, cupos, horarios o agenda disponible.", "+22"],
    ["Intencion comercial", "Pide mas informacion sobre un producto o servicio especifico.", "+18"],
    ["Intencion comercial", "Pregunta general sobre el negocio, producto o servicio.", "+12"],
    ["Intencion comercial", 'Solo dice "info" sin mas contexto.', "+8"],
    ["Fit con producto o servicio", "Pregunta por un producto o servicio que la empresa si ofrece.", "+20"],
    ["Fit con producto o servicio", "Pregunta por una categoria relevante del negocio.", "+15"],
    ["Fit con producto o servicio", "Su necesidad parece alineada, aunque no mencione producto especifico.", "+10"],
    ["Fit con producto o servicio", "Pregunta por algo relacionado, pero no prioritario para el negocio.", "+5"],
    ["Fit con producto o servicio", "Pregunta por algo que la empresa no ofrece.", "0"],
    ["Fit con producto o servicio", "Pregunta totalmente fuera del negocio.", "-10"],
    ["Contactabilidad", "Deja telefono, WhatsApp, correo o dato concreto de contacto.", "+15"],
    ["Contactabilidad", "Responde por el mismo canal social o canal de entrada.", "+10"],
    ["Contactabilidad", "Deja nombre o identificador claro.", "+5"],
    ["Contactabilidad", "No deja datos, pero sigue conversando por el canal.", "+5"],
    ["Contactabilidad", "No deja datos y no hay otra señal adicional.", "0"],
    ["Engagement conversacional", "Responde despues de la intervencion del agente IA.", "+5"],
    ["Engagement conversacional", "Hace mas de una pregunta en la conversacion.", "+3"],
    ["Engagement conversacional", "Acepta recibir mas informacion, propuesta, cotizacion o seguimiento.", "+2"],
    ["Engagement conversacional", "Confirma que entendio, agradece o cierra positivo.", "+1"],
    ["Engagement conversacional", "No hay nueva interaccion del cliente.", "0"],
]

NEGATIVE_SIGNALS = [
    ["Rechazo temporal o bajo interes", '"por ahora no", "despues veo", "no estoy seguro"', "-10 a -15", "Baja prioridad, pero el lead puede recuperarse si vuelve a preguntar."],
    ["Rechazo claro", '"no me interesa", "ya no quiero", "no deseo informacion"', "-20", "Reduce prioridad porque el cliente cierra la conversacion."],
    ["Fuera del negocio", "Pregunta por temas no relacionados.", "-10", "No elimina el lead, pero baja la calidad de esa interaccion."],
    ["Agresividad o insulto", "Burla, insulto o trato ofensivo.", "-25", "Debe bajar fuerte por calidad de conversacion."],
    ["Riesgo", "Amenaza, fraude, acoso, phishing o conducta critica.", "-50", "Puede dejar el score muy bajo o negativo."],
    ["Spam o autopromocion", "Intenta vender algo al negocio sin relacion con el embudo.", "-20", "No es lead comercial valido en ese momento."],
    ["Links", "Un enlace sin contexto comercial.", "0", "El link es neutral; se puntua solo el texto del cliente."],
]

RULES = [
    ["Evaluar solo mensaje del cliente", "No sumar ni restar por respuestas generadas por la IA."],
    ["Calculo por interaccion", "score_final = score_actual + puntos_positivos - puntos_negativos."],
    ["Score acumulativo y reversible", "Puede subir o bajar en cada mensaje segun la nueva señal."],
    ["Sin limite tecnico", "No forzar minimo 0 ni maximo 100; la clasificacion visual usa rangos."],
    ["Evitar doble conteo excesivo", "Dentro de una misma dimension usar la señal de mayor peso; entre dimensiones si se puede sumar."],
    ["Seguimiento humano separado", "El seguimiento humano se decide por etiquetas operativas, no por score."],
]

SCENARIOS = [
    ['"Precio?"', "Interes +15, precio +25, fit +15, contactabilidad por canal +10.", "+65", "TIBIO"],
    ['"Hola, tienen disponible este producto?"', "Interes +20, disponibilidad +22, fit +20, contactabilidad +10, engagement +3.", "+75", "CALIENTE"],
    ['"Info"', "Interes +15, intencion +8, fit +10, contactabilidad +10.", "+43", "FRIO"],
    ['"Me interesa, como puedo agendar?"', "Interes +20, agenda +35, fit +20, contactabilidad +10, engagement +5.", "+90", "CALIENTE"],
    ['"No me interesa por ahora"', "Rechazo temporal.", "-10 a -15", "FRIO o baja el estado actual"],
    ['"Eso es una estafa, no molesten"', "Agresividad o rechazo claro segun contexto.", "-20 a -25", "FRIO o score muy bajo"],
    ['"Hola, quiero informacion" + link', "Se puntua por quiero informacion. El link vale 0.", "Segun contexto", "FRIO/TIBIO"],
    ['"Quiero comprar, cuanto cuesta y tienen para hoy?"', "Compra, precio, disponibilidad, fit y contactabilidad.", "+90 o mas", "CALIENTE"],
    ['"Ustedes venden algo fuera del negocio?"', "Fit 0 o fuera del negocio -10.", "0 a -10", "FRIO"],
]

N8N_CHANGES = [
    ["Campo", "Mantener el mismo campo numerico actual de score."],
    ["Fuente", "Tomar como base el mensaje entrante del cliente."],
    ["Formula", "score_final = score_actual + delta_de_la_interaccion."],
    ["Delta", "Puede ser positivo, negativo o cero."],
    ["Actualizacion", "Guardar siempre score_final, aunque sea menor que el anterior."],
    ["Vocabulario", "Usar vocabulario global y permitir terminos propios por negocio."],
]


DOCX_COLORS = {
    "blue": "274690",
    "navy": "0f2344",
    "green": "0a9b6f",
    "light": "f8fafc",
    "line": "d9e2ef",
    "white": "ffffff",
    "slate": "64748b",
}


def pdf_styles():
    base = getSampleStyleSheet()
    base.add(
        ParagraphStyle(
            name="TitleCenter",
            parent=base["Title"],
            alignment=TA_CENTER,
            textColor=BLUE,
            fontName="Helvetica-Bold",
            fontSize=19,
            leading=23,
            spaceAfter=8,
        )
    )
    base.add(
        ParagraphStyle(
            name="SubTitle",
            parent=base["Normal"],
            alignment=TA_CENTER,
            textColor=colors.HexColor("#64748b"),
            fontSize=10,
            leading=13,
            spaceAfter=16,
        )
    )
    base.add(
        ParagraphStyle(
            name="Section",
            parent=base["Heading1"],
            textColor=BLUE,
            fontName="Helvetica-Bold",
            fontSize=13,
            leading=16,
            spaceBefore=8,
            spaceAfter=6,
        )
    )
    base.add(
        ParagraphStyle(
            name="Small",
            parent=base["Normal"],
            fontSize=7.3,
            leading=9,
        )
    )
    base["Normal"].fontName = "Helvetica"
    base["Normal"].fontSize = 9
    base["Normal"].leading = 12
    return base


def p(text: str, style) -> Paragraph:
    safe = str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
    return Paragraph(safe, style)


def make_table(headers: list[str], rows: list[list[str]], styles, widths: list[float] | None = None, font_size: int = 7):
    data = [[p(header, styles["Small"]) for header in headers]]
    data.extend([[p(value, styles["Small"]) for value in row] for row in rows])
    table = Table(data, colWidths=widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), BLUE),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), font_size),
                ("GRID", (0, 0), (-1, -1), 0.35, LINE),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("BACKGROUND", (0, 1), (-1, -1), colors.white),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT]),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
            ]
        )
    )
    return table


def footer(canvas, doc):
    canvas.saveState()
    canvas.setFont("Helvetica", 7)
    canvas.setFillColor(colors.HexColor("#64748b"))
    canvas.drawRightString(doc.pagesize[0] - 1.2 * cm, 0.7 * cm, f"Simplia Leads - Score 3 estados | Pagina {doc.page}")
    canvas.restoreState()


def build_pdf() -> None:
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(PDF_OUT),
        pagesize=landscape(A4),
        rightMargin=1.1 * cm,
        leftMargin=1.1 * cm,
        topMargin=1.0 * cm,
        bottomMargin=1.0 * cm,
        title="Score de Calidad - Simplia Leads",
    )
    story = [
        p("Score de Calidad - Simplia Leads", styles["TitleCenter"]),
        p("Version actualizada con 3 estados visuales: FRIO, TIBIO y CALIENTE.", styles["SubTitle"]),
        p("1. Objetivo", styles["Section"]),
        p(
            "El score mide la calidad, prioridad e interes del lead a lo largo de la conversacion. "
            "No decide por si solo el seguimiento humano: el seguimiento se maneja por etiquetas operativas definidas con cada negocio.",
            styles["Normal"],
        ),
        Spacer(1, 6),
        make_table(["Regla", "Aplicacion"], RULES, styles, widths=[6.0 * cm, 18.0 * cm]),
        p("2. Estados visuales del lead", styles["Section"]),
        make_table(
            ["Estado", "Rango", "Interpretacion", "Accion sugerida"],
            STATES,
            styles,
            widths=[3.0 * cm, 4.0 * cm, 11.0 * cm, 8.0 * cm],
        ),
        p("3. Matriz de puntuacion positiva", styles["Section"]),
        make_table(
            ["Dimension", "Senal detectada en el mensaje del cliente", "Puntaje"],
            POSITIVE_MATRIX,
            styles,
            widths=[5.0 * cm, 17.5 * cm, 3.0 * cm],
        ),
        PageBreak(),
        p("4. Senales negativas", styles["Section"]),
        p(
            "Las senales negativas no descalifican de forma definitiva. Solo restan score; si luego el cliente vuelve a mostrar interes real, puede recuperarse.",
            styles["Normal"],
        ),
        make_table(
            ["Senal", "Ejemplo", "Puntaje", "Como interpretarlo"],
            NEGATIVE_SIGNALS,
            styles,
            widths=[4.3 * cm, 7.0 * cm, 3.0 * cm, 11.0 * cm],
        ),
        p("5. Escenarios concretos", styles["Section"]),
        make_table(
            ["Mensaje del cliente", "Calculo aproximado", "Delta", "Estado esperado"],
            SCENARIOS,
            styles,
            widths=[6.0 * cm, 13.0 * cm, 3.0 * cm, 4.0 * cm],
        ),
        p("6. Cambios minimos en n8n", styles["Section"]),
        make_table(["Punto", "Cambio"], N8N_CHANGES, styles, widths=[5.0 * cm, 19.0 * cm]),
        p("7. Conclusion recomendada", styles["Section"]),
        p(
            "La clasificacion queda simplificada en tres estados: FRIO, TIBIO y CALIENTE. "
            "El score sigue siendo acumulativo, reversible y sin limite tecnico, pero la lectura comercial se hace con estos tres rangos. "
            "Esto evita mezclar calidad del lead con seguimiento humano y deja una logica clara para ventas y automatizacion.",
            styles["Normal"],
        ),
    ]
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def docx_rgb(hex_color: str) -> RGBColor:
    value = hex_color.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def configure_docx(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.2)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    for style_name, size, color in [
        ("Heading 1", 14, DOCX_COLORS["blue"]),
        ("Heading 2", 11.5, DOCX_COLORS["navy"]),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = docx_rgb(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    footer_paragraph = section.footer.paragraphs[0]
    footer_paragraph.text = ""
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_paragraph.add_run("Simplia Leads - Score 3 estados")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = docx_rgb(DOCX_COLORS["slate"])


def set_docx_cell_shading(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def set_docx_table_borders(table, color: str = DOCX_COLORS["line"]) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement("w:" + edge)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    table._tbl.tblPr.append(borders)


def clear_docx_cell(cell) -> None:
    cell.text = ""
    for paragraph in cell.paragraphs:
        paragraph.text = ""


def write_docx_cell(cell, value: str, *, bold: bool = False, color: str | None = None, size: float = 8.0) -> None:
    clear_docx_cell(cell)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for index, part in enumerate(str(value).split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = docx_rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_docx_paragraph(document: Document, text: str, *, bold: bool = False) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(9.2)


def add_docx_table(document: Document, headers: list[str], rows: list[list[str]], *, font_size: float = 7.5) -> None:
    table = document.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    set_docx_table_borders(table)

    for col, header in enumerate(headers):
        cell = table.cell(0, col)
        write_docx_cell(cell, header, bold=True, color=DOCX_COLORS["white"], size=8.0)
        set_docx_cell_shading(cell, DOCX_COLORS["blue"])

    for row_index, row in enumerate(rows, start=1):
        fill = DOCX_COLORS["light"] if row_index % 2 == 0 else DOCX_COLORS["white"]
        for col, value in enumerate(row):
            cell = table.cell(row_index, col)
            write_docx_cell(cell, value, size=font_size)
            set_docx_cell_shading(cell, fill)
    document.add_paragraph()


def build_docx() -> None:
    document = Document()
    configure_docx(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Score de Calidad - Simplia Leads")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(19)
    run.font.color.rgb = docx_rgb(DOCX_COLORS["blue"])

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Version actualizada con 3 estados visuales: FRIO, TIBIO y CALIENTE.")
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.color.rgb = docx_rgb(DOCX_COLORS["slate"])

    document.add_heading("1. Objetivo", level=1)
    add_docx_paragraph(
        document,
        "El score mide la calidad, prioridad e interes del lead a lo largo de la conversacion. "
        "No decide por si solo el seguimiento humano: el seguimiento se maneja por etiquetas operativas definidas con cada negocio.",
    )
    add_docx_table(document, ["Regla", "Aplicacion"], RULES, font_size=7.6)

    document.add_heading("2. Estados visuales del lead", level=1)
    add_docx_table(document, ["Estado", "Rango", "Interpretacion", "Accion sugerida"], STATES, font_size=7.3)

    document.add_heading("3. Matriz de puntuacion positiva", level=1)
    add_docx_table(document, ["Dimension", "Senal detectada en el mensaje del cliente", "Puntaje"], POSITIVE_MATRIX, font_size=7.1)

    document.add_heading("4. Senales negativas", level=1)
    add_docx_paragraph(
        document,
        "Las senales negativas no descalifican de forma definitiva. Solo restan score; si luego el cliente vuelve a mostrar interes real, puede recuperarse.",
    )
    add_docx_table(document, ["Senal", "Ejemplo", "Puntaje", "Como interpretarlo"], NEGATIVE_SIGNALS, font_size=7.1)

    document.add_heading("5. Escenarios concretos", level=1)
    add_docx_table(document, ["Mensaje del cliente", "Calculo aproximado", "Delta", "Estado esperado"], SCENARIOS, font_size=7.1)

    document.add_heading("6. Cambios minimos en n8n", level=1)
    add_docx_table(document, ["Punto", "Cambio"], N8N_CHANGES, font_size=7.5)

    document.add_heading("7. Conclusion recomendada", level=1)
    add_docx_paragraph(
        document,
        "La clasificacion queda simplificada en tres estados: FRIO, TIBIO y CALIENTE. "
        "El score sigue siendo acumulativo, reversible y sin limite tecnico, pero la lectura comercial se hace con estos tres rangos. "
        "Esto evita mezclar calidad del lead con seguimiento humano y deja una logica clara para ventas y automatizacion.",
    )

    try:
        document.save(DOCX_OUT)
    except PermissionError as exc:
        raise SystemExit(f"No se pudo guardar {DOCX_OUT.name}. Cierre el archivo en Word y vuelva a ejecutar este script.") from exc


def style_sheet(ws, widths: list[int] | None = None) -> None:
    header_fill = PatternFill("solid", fgColor="274690")
    header_font = Font(color="FFFFFF", bold=True)
    thin = Side(style="thin", color="D9E2EF")
    border = Border(left=thin, right=thin, top=thin, bottom=thin)

    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            cell.border = border
            if cell.row == 1:
                cell.fill = header_fill
                cell.font = header_font

    ws.freeze_panes = "A2"
    if widths:
        for index, width in enumerate(widths, start=1):
            ws.column_dimensions[get_column_letter(index)].width = width
    else:
        for column in range(1, ws.max_column + 1):
            ws.column_dimensions[get_column_letter(column)].width = 28


def add_sheet(wb: Workbook, title: str, headers: list[str], rows: list[list[str]], widths: list[int] | None = None):
    ws = wb.create_sheet(title)
    ws.append(headers)
    for row in rows:
        ws.append(row)
    style_sheet(ws, widths)
    return ws


def build_xlsx() -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "Resumen"
    rows = [
        ["Documento", "Score de Calidad - Simplia Leads"],
        ["Version", "Actualizada a 3 estados"],
        ["Estados", "FRIO, TIBIO, CALIENTE"],
        ["Formula", "score_final = score_actual + puntos_positivos - puntos_negativos"],
        ["Seguimiento humano", "Se maneja por etiquetas operativas, no por score"],
        ["Nota", "El score no tiene limite tecnico; los rangos solo sirven para interpretacion visual"],
    ]
    ws.append(["Campo", "Detalle"])
    for row in rows:
        ws.append(row)
    style_sheet(ws, [26, 90])

    add_sheet(wb, "Estados_3", ["Estado", "Rango", "Interpretacion", "Accion sugerida"], STATES, [18, 24, 70, 65])
    add_sheet(wb, "Matriz_Positiva", ["Dimension", "Senal detectada", "Puntaje"], POSITIVE_MATRIX, [32, 90, 14])
    add_sheet(wb, "Senales_Negativas", ["Senal", "Ejemplo", "Puntaje", "Como interpretarlo"], NEGATIVE_SIGNALS, [32, 55, 14, 70])
    add_sheet(wb, "Reglas", ["Regla", "Aplicacion"], RULES, [36, 90])
    add_sheet(wb, "Escenarios", ["Mensaje del cliente", "Calculo aproximado", "Delta", "Estado esperado"], SCENARIOS, [44, 85, 18, 24])
    add_sheet(wb, "n8n", ["Punto", "Cambio"], N8N_CHANGES, [28, 90])
    wb.save(XLSX_OUT)


def build_all() -> None:
    build_pdf()
    build_docx()
    build_xlsx()


if __name__ == "__main__":
    build_all()
    print(f"PDF generado: {PDF_OUT}")
    print(f"Word generado: {DOCX_OUT}")
    print(f"Excel generado: {XLSX_OUT}")
