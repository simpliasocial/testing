from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
DOCS_DIR = ROOT / "docs"

TODAY = date(2026, 5, 14)
TIMEZONE = "America/Guayaquil"

DECISION_MD = DOCS_DIR / "Decision_Estrategica_Arquitectura_SaaS_SimpliaLeads.md"
DECISION_DOCX = DOCS_DIR / "Decision_Estrategica_Arquitectura_SaaS_SimpliaLeads.docx"
IMPLEMENTATION_MD = DOCS_DIR / "Plan_Implementacion_SaaS_Multiempresa_SimpliaLeads.md"
IMPLEMENTATION_DOCX = DOCS_DIR / "Plan_Implementacion_SaaS_Multiempresa_SimpliaLeads.docx"

COLORS = {
    "navy": "0F2344",
    "blue": "274690",
    "soft_blue": "EAF1FF",
    "green": "0A9B6F",
    "mint": "DFF7EC",
    "orange": "F59E0B",
    "red": "B91C1C",
    "soft_red": "FEE2E2",
    "slate": "475569",
    "line": "D9E2EF",
    "light": "F8FAFC",
    "white": "FFFFFF",
}

SOURCES = [
    [
        "Supabase Pricing",
        "https://supabase.com/pricing",
        "Pro desde USD 25/mes, 8 GB de disco incluidos por proyecto, excedente de disco desde USD 0.125/GB-mes, egress y storage con overage, daily backups en Pro.",
    ],
    [
        "Supabase Row Level Security",
        "https://supabase.com/docs/guides/database/postgres/row-level-security",
        "RLS debe estar habilitado en tablas de schemas expuestos y se usa para reglas granulares de autorizacion.",
    ],
    [
        "Supabase Hardening Data API",
        "https://supabase.com/docs/guides/database/hardening-data-api",
        "El schema public suele estar expuesto por la Data API; las tablas accesibles por API deben tener RLS y se recomiendan schemas privados para datos internos.",
    ],
    [
        "Supabase Database Size",
        "https://supabase.com/docs/guides/platform/database-size",
        "El uso de disco incluye datos, indices, WAL y archivos internos; en Pro el disco puede autoescalar y hay riesgo de modo read-only si se supera la capacidad.",
    ],
    [
        "Supabase Backups",
        "https://supabase.com/docs/guides/platform/backups",
        "Daily backups y PITR dependen del plan y de la retencion; PITR se cotiza adicionalmente.",
    ],
    [
        "Vercel Pricing",
        "https://vercel.com/pricing",
        "Pro desde USD 20/mes, incluye credito de uso y cobra consumo adicional por compute, edge requests, transferencia y otros recursos.",
    ],
    [
        "AWS SaaS Lens - Silo, Pool and Bridge",
        "https://docs.aws.amazon.com/wellarchitected/latest/saas-lens/silo-pool-and-bridge-models.html",
        "Define los modelos SaaS silo, pool y bridge para aislamiento y eficiencia multi-tenant.",
    ],
    [
        "AWS RDS for PostgreSQL Pricing",
        "https://aws.amazon.com/rds/postgresql/pricing/",
        "RDS cobra por instancia, almacenamiento, backups, I/O y transferencia segun region y configuracion.",
    ],
    [
        "Azure SaaS Tenancy Patterns",
        "https://learn.microsoft.com/en-us/azure/azure-sql/database/saas-tenancy-app-design-patterns",
        "Describe patrones de tenancy para apps SaaS, incluyendo base compartida, base por tenant y modelos hibridos.",
    ],
    [
        "Azure App Service Costs",
        "https://learn.microsoft.com/en-us/azure/app-service/overview-manage-costs",
        "App Service se cobra por plan, tier, instancias y recursos relacionados.",
    ],
    [
        "Google Cloud SQL Pricing",
        "https://cloud.google.com/sql/pricing",
        "Cloud SQL se cobra por vCPU, memoria, almacenamiento, backups y red segun region y configuracion.",
    ],
    [
        "Neon Pricing",
        "https://neon.com/pricing",
        "Neon cobra por compute-hour, storage, time-travel/restore y transferencia segun plan.",
    ],
    [
        "Railway Pricing",
        "https://docs.railway.com/reference/pricing/plans",
        "Railway cobra plan base y uso de RAM, CPU, egress y volumen; Pro desde USD 20/mes mas uso.",
    ],
    [
        "Render Pricing",
        "https://render.com/pricing/",
        "Render cobra servicios y bases de datos por plan/recurso, util como opcion PaaS pero menos integrada que Supabase para Auth/RLS/Edge.",
    ],
    [
        "n8n Pricing",
        "https://n8n.io/pricing/",
        "Referencia oficial para planes cloud/self-hosted; el documento usa el supuesto comercial indicado: USD 20/mes por cliente.",
    ],
    [
        "Chatwoot Pricing",
        "https://www.chatwoot.com/pricing/",
        "Startups desde USD 19/agente/mes facturado anualmente; el documento usa USD 19/mes por cliente como supuesto operativo.",
    ],
    [
        "OpenAI API Pricing",
        "https://openai.com/api/pricing/",
        "La facturacion API depende de modelo y tokens; el documento usa el supuesto interno indicado: USD 70/mes por cliente.",
    ],
]


@dataclass(frozen=True)
class TableBlock:
    title: str
    headers: list[str]
    rows: list[list[str]]
    note: str | None = None


@dataclass(frozen=True)
class Section:
    title: str
    blocks: list[tuple[str, object]]


@dataclass(frozen=True)
class DocSpec:
    title: str
    subtitle: str
    code: str
    version: str
    status: str
    owner: str
    audience: str
    sections: list[Section]


def rgb(hex_color: str) -> RGBColor:
    value = hex_color.replace("#", "")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def money(value: float) -> str:
    return f"USD {value:,.0f}".replace(",", ",")


def monthly(value: float) -> str:
    return f"USD {value:,.2f}/mes".replace(",", ",")


def annual(value: float) -> str:
    return f"USD {value:,.2f}/ano".replace(",", ",")


def active_clients(year: int, retention: float) -> float:
    return sum(26 * (retention**age) for age in range(year))


def md_escape(value: object) -> str:
    return str(value).replace("|", "\\|").replace("\n", "<br>")


def md_table(headers: Sequence[str], rows: Sequence[Sequence[str]]) -> str:
    header_line = "| " + " | ".join(md_escape(h) for h in headers) + " |"
    separator = "| " + " | ".join("---" for _ in headers) + " |"
    body = ["| " + " | ".join(md_escape(cell) for cell in row) + " |" for row in rows]
    return "\n".join([header_line, separator, *body])


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str | None = None, size: float = 8.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_docx_table(table) -> None:
    table.style = "Table Grid"
    table.autofit = True
    repeat_table_header(table.rows[0])
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if row_index == 0:
                set_cell_shading(cell, COLORS["blue"])
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(8.5)
                        run.bold = True
                        run.font.color.rgb = rgb(COLORS["white"])
            else:
                set_cell_shading(cell, COLORS["white"] if row_index % 2 else COLORS["light"])
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(8)


def configure_doc(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    for name in ["Title", "Heading 1", "Heading 2", "Heading 3"]:
        styles[name].font.name = "Arial"
        styles[name].font.color.rgb = rgb(COLORS["navy"])
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 3"].font.size = Pt(11)


def add_header_footer(document: Document, spec: DocSpec) -> None:
    section = document.sections[0]
    header = section.header.paragraphs[0]
    header.text = f"{spec.code} | {spec.version} | SimpliaLeads"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = rgb(COLORS["slate"])
    footer = section.footer.paragraphs[0]
    footer.text = f"Documento generado el {TODAY.isoformat()} ({TIMEZONE})"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = rgb(COLORS["slate"])


def add_cover(document: Document, spec: DocSpec) -> None:
    for _ in range(3):
        document.add_paragraph()
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(spec.title)
    title_run.bold = True
    title_run.font.name = "Arial"
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = rgb(COLORS["navy"])

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(spec.subtitle)
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = rgb(COLORS["slate"])

    document.add_paragraph()
    control = document.add_table(rows=1, cols=2)
    control.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rows = [
        ["Codigo", spec.code],
        ["Version", spec.version],
        ["Estado", spec.status],
        ["Fecha", TODAY.isoformat()],
        ["Zona horaria", TIMEZONE],
        ["Responsable", spec.owner],
        ["Audiencia", spec.audience],
    ]
    for label, value in rows:
        row = control.add_row().cells
        set_cell_text(row[0], label, bold=True, color=COLORS["navy"], size=9)
        set_cell_text(row[1], value, size=9)
    control._tbl.remove(control.rows[0]._tr)
    style_docx_table(control)

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Documento de trabajo para decision ejecutiva y tecnica. Los costos son estimaciones "
        "basadas en precios oficiales consultados el dia de generacion y deben validarse con "
        "calculadoras de cada proveedor antes de contratar."
    )
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(8.5)
    run.font.color.rgb = rgb(COLORS["slate"])
    document.add_page_break()


def add_callout(document: Document, title: str, body: str, fill: str = "EAF1FF") -> None:
    table = document.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = rgb(COLORS["navy"])
    p.add_run("\n")
    body_run = p.add_run(body)
    body_run.font.name = "Arial"
    body_run.font.size = Pt(9)
    body_run.font.color.rgb = rgb(COLORS["navy"])


def add_blocks(document: Document, blocks: Iterable[tuple[str, object]]) -> None:
    for block_type, payload in blocks:
        if block_type == "p":
            paragraph = document.add_paragraph(str(payload))
            paragraph.paragraph_format.space_after = Pt(5)
        elif block_type == "bullets":
            for item in payload:  # type: ignore[union-attr]
                paragraph = document.add_paragraph(str(item), style="List Bullet")
                paragraph.paragraph_format.space_after = Pt(2)
        elif block_type == "numbered":
            for item in payload:  # type: ignore[union-attr]
                paragraph = document.add_paragraph(str(item), style="List Number")
                paragraph.paragraph_format.space_after = Pt(2)
        elif block_type == "table":
            table_block: TableBlock = payload  # type: ignore[assignment]
            document.add_heading(table_block.title, level=3)
            table = document.add_table(rows=1, cols=len(table_block.headers))
            for idx, header in enumerate(table_block.headers):
                set_cell_text(table.rows[0].cells[idx], header, bold=True, color=COLORS["white"], size=8.5)
            for row_data in table_block.rows:
                cells = table.add_row().cells
                for idx, value in enumerate(row_data):
                    set_cell_text(cells[idx], value, size=8)
            style_docx_table(table)
            if table_block.note:
                note = document.add_paragraph(table_block.note)
                note.runs[0].italic = True
                note.runs[0].font.size = Pt(8)
                note.runs[0].font.color.rgb = rgb(COLORS["slate"])
        elif block_type == "code":
            title, code = payload  # type: ignore[misc]
            document.add_heading(str(title), level=3)
            code_table = document.add_table(rows=1, cols=1)
            code_cell = code_table.rows[0].cells[0]
            set_cell_shading(code_cell, COLORS["light"])
            paragraph = code_cell.paragraphs[0]
            run = paragraph.add_run(str(code))
            run.font.name = "Consolas"
            run.font.size = Pt(8)
        elif block_type == "callout":
            title, body, fill = payload  # type: ignore[misc]
            add_callout(document, str(title), str(body), str(fill))
        elif block_type == "page_break":
            document.add_page_break()


def write_docx(spec: DocSpec, output: Path) -> None:
    document = Document()
    configure_doc(document)
    add_header_footer(document, spec)
    add_cover(document, spec)
    document.add_heading("Control Documental", level=1)
    table = document.add_table(rows=1, cols=4)
    headers = ["Version", "Fecha", "Cambio", "Responsable"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=COLORS["white"])
    row = table.add_row().cells
    for idx, value in enumerate([spec.version, TODAY.isoformat(), "Documento base generado para revision interna.", spec.owner]):
        set_cell_text(row[idx], value)
    style_docx_table(table)

    document.add_heading("Indice Ejecutivo", level=1)
    for idx, section in enumerate(spec.sections, start=1):
        document.add_paragraph(f"{idx}. {section.title}")

    for section in spec.sections:
        document.add_heading(section.title, level=1)
        add_blocks(document, section.blocks)

    output.parent.mkdir(parents=True, exist_ok=True)
    document.save(output)


def block_to_markdown(block: tuple[str, object]) -> str:
    block_type, payload = block
    if block_type == "p":
        return str(payload)
    if block_type == "bullets":
        return "\n".join(f"- {item}" for item in payload)  # type: ignore[union-attr]
    if block_type == "numbered":
        return "\n".join(f"{idx}. {item}" for idx, item in enumerate(payload, start=1))  # type: ignore[arg-type]
    if block_type == "table":
        table_block: TableBlock = payload  # type: ignore[assignment]
        text = f"### {table_block.title}\n\n{md_table(table_block.headers, table_block.rows)}"
        if table_block.note:
            text += f"\n\n> {table_block.note}"
        return text
    if block_type == "callout":
        title, body, _fill = payload  # type: ignore[misc]
        return f"> **{title}**\n>\n> {body}"
    if block_type == "code":
        title, code = payload  # type: ignore[misc]
        return f"### {title}\n\n```sql\n{code}\n```"
    if block_type == "page_break":
        return "\n---\n"
    return ""


def write_markdown(spec: DocSpec, output: Path) -> None:
    lines = [
        f"# {spec.title}",
        "",
        spec.subtitle,
        "",
        md_table(
            ["Campo", "Valor"],
            [
                ["Codigo", spec.code],
                ["Version", spec.version],
                ["Estado", spec.status],
                ["Fecha", TODAY.isoformat()],
                ["Zona horaria", TIMEZONE],
                ["Responsable", spec.owner],
                ["Audiencia", spec.audience],
            ],
        ),
        "",
        "## Control Documental",
        "",
        md_table(
            ["Version", "Fecha", "Cambio", "Responsable"],
            [[spec.version, TODAY.isoformat(), "Documento base generado para revision interna.", spec.owner]],
        ),
        "",
    ]
    for section in spec.sections:
        lines.append(f"## {section.title}")
        lines.append("")
        for block in section.blocks:
            rendered = block_to_markdown(block)
            if rendered:
                lines.append(rendered)
                lines.append("")
    output.parent.mkdir(parents=True, exist_ok=True)
    output.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def build_projection_tables() -> tuple[TableBlock, TableBlock, TableBlock]:
    retention_rows: list[list[str]] = []
    for year in [1, 3, 5]:
        before_churn = 26 * year
        row = [f"Ano {year}", str(before_churn)]
        for retention in [0.90, 0.95, 0.98]:
            active = active_clients(year, retention)
            row.append(f"{active:.1f}")
        retention_rows.append(row)
    retention_table = TableBlock(
        "Proyeccion de clientes activos con retencion anual",
        ["Periodo", "Acumulados antes de churn", "Activos con 90%", "Activos con 95%", "Activos con 98%"],
        retention_rows,
        "Formula: 26 clientes nuevos por ano y cohortes retenidas anualmente. No sustituye el forecast comercial real.",
    )

    external_rows: list[list[str]] = []
    for year in [1, 3, 5]:
        clients = 26 * year
        monthly_cost = clients * 109
        external_rows.append(
            [
                f"Ano {year}",
                str(clients),
                monthly(monthly_cost),
                annual(monthly_cost * 12),
                "Pagado por el cliente: Chatwoot 19 + n8n 20 + OpenAI 70.",
            ]
        )
    external_table = TableBlock(
        "Costo externo operativo pagado por clientes antes de churn",
        ["Periodo", "Clientes acumulados", "Costo mensual total", "Costo anual total", "Notas"],
        external_rows,
    )

    internal_rows = [
        [
            "Ano 1",
            "26",
            "55 a 350/mes",
            "Supabase Pro pequeno/medio, Vercel Pro 1-3 seats, backups exportados simples, monitoreo basico.",
        ],
        [
            "Ano 3",
            "70 a 76 activos aprox.",
            "120 a 770/mes",
            "Sube almacenamiento, indices, egress, logs, compute y soporte operativo. Aun conviene pooled si hay RLS fuerte.",
        ],
        [
            "Ano 5",
            "106 a 125 activos aprox.",
            "240 a 1,530/mes",
            "Puede requerir compute mayor, PITR, observabilidad dedicada, particionado y politicas de archivo.",
        ],
    ]
    internal_table = TableBlock(
        "Costo interno Simplia estimado para modelo pooled",
        ["Periodo", "Base de clientes", "Rango interno estimado", "Que incluye"],
        internal_rows,
        "Rangos orientativos para decision. La cifra real depende de filas por cliente, archivos, egress, dashboards, retencion de logs y numero de usuarios internos en Vercel/Supabase.",
    )
    return retention_table, external_table, internal_table


def option_detail(
    name: str,
    how: str,
    github: str,
    vercel: str,
    database: str,
    cw_n8n: str,
    benefits: str,
    limits: str,
    risks: str,
    cost: str,
    applies: str,
    not_applies: str,
) -> TableBlock:
    return TableBlock(
        name,
        ["Criterio", "Detalle"],
        [
            ["Como funcionaria", how],
            ["GitHub", github],
            ["Vercel", vercel],
            ["Supabase/base de datos", database],
            ["cw y tablas n8n", cw_n8n],
            ["Beneficios", benefits],
            ["Limitantes", limits],
            ["Riesgos", risks],
            ["Costos aproximados", cost],
            ["Cuando si aplica", applies],
            ["Cuando no aplica", not_applies],
        ],
    )


def build_decision_doc() -> DocSpec:
    retention_table, external_table, internal_table = build_projection_tables()

    summary_matrix = TableBlock(
        "Matriz ejecutiva de opciones",
        ["Opcion", "Modelo", "Ventaja principal", "Riesgo principal", "Decision recomendada"],
        [
            [
                "Repo + Supabase + Vercel por cliente",
                "Silo completo manual",
                "Aislamiento alto y facil de entender al inicio.",
                "Mantenimiento N veces, versiones divergentes, migraciones repetidas, errores por prompts repetidos.",
                "No usar como estandar. Solo posible para enterprise con precio alto y automatizacion fuerte.",
            ],
            [
                "Branch por cliente",
                "Fork logico dentro del mismo repo",
                "Permite cambios especificos rapidamente.",
                "Deuda tecnica explosiva, merges permanentes, bugs distintos por cliente.",
                "No usar. Branches solo temporales feature/* o hotfix/*.",
            ],
            [
                "Un repo + un Vercel + Supabase compartido con company_id",
                "Pool SaaS",
                "Un producto, una version, costos compartidos, onboarding rapido.",
                "Exige RLS, tenant isolation, indices y disciplina de configuracion.",
                "Recomendado como modelo base para SimpliaLeads.",
            ],
            [
                "Schema por cliente",
                "Pool parcial por proyecto",
                "Aisla nombres y tablas por cliente.",
                "Migrations por schema, duplicacion de tablas, consultas/reportes complejos.",
                "No como base. Usar solo para modulos muy excepcionales.",
            ],
            [
                "Tablas por cliente",
                "Fragmentacion fisica por tabla",
                "Separacion visible.",
                "Antipattern para SaaS: DDL infinito, codigo condicional, indices duplicados.",
                "Descartar.",
            ],
            [
                "Supabase/proyecto por cliente",
                "Silo de base/Backend gestionado",
                "Aislamiento de datos fuerte y costo trasladable a enterprise.",
                "Migraciones, backups, secrets y jobs multiplicados.",
                "Modelo bridge/enterprise futuro, no modelo base.",
            ],
            [
                "Backend unico + DB por cliente",
                "Bridge como los diagramas enviados",
                "Mismo codigo con aislamiento de base.",
                "Routing por tenant, pool de conexiones, migraciones N DBs.",
                "Valido para enterprise si hay equipo DevOps y precio adicional.",
            ],
            [
                "AWS/Azure/GCP nativo",
                "Cloud enterprise",
                "Control, compliance, redes privadas, escalabilidad avanzada.",
                "Mayor costo y complejidad operativa que Supabase/Vercel.",
                "Evaluar cuando haya requisitos enterprise, no para acelerar 26 clientes/ano.",
            ],
            [
                "Neon",
                "Postgres serverless",
                "Branching y compute elastico de base de datos.",
                "No reemplaza por si solo Auth, Storage, Edge Functions y RLS operacional de Supabase.",
                "Alternativa DB futura si Supabase limita costos/performance.",
            ],
            [
                "Railway/Render",
                "PaaS generalista",
                "Rapido para servicios, workers o self-host de n8n.",
                "Menos nativo para Auth/RLS/tenant isolation de producto SaaS.",
                "Complemento posible, no recomendacion principal para core de datos.",
            ],
        ],
    )

    strategic_options = [
        option_detail(
            "Opcion 1 - Repo + Supabase + Vercel por cliente",
            "Por cada cliente se clona el repositorio, se crea un proyecto Supabase separado y se conecta un proyecto Vercel a su main.",
            "Habria N repos o N clones. Cada feature debe replicarse manualmente en cada cliente.",
            "Habria N proyectos/deployments. Cada entorno tiene variables propias y cada release se hace N veces.",
            "Cada cliente tiene su proyecto Supabase. Aislamiento alto, pero sin economia de escala operativa.",
            "Cada cliente mantiene su propio schema cw y sus tablas public/n8n. Las migraciones de cw, automatizaciones y correcciones se repiten por cliente.",
            "Aislamiento fuerte, bajo riesgo de fuga entre clientes y facil de razonar en los primeros 1-2 clientes.",
            "No escala con 26 clientes/ano; el esfuerzo crece linealmente con clientes y cada feature se vuelve una operacion manual.",
            "Versiones divergentes, errores humanos, prompts diferentes, deuda tecnica por cliente, costos y secretos dispersos.",
            "Supabase Pro desde 25/mes por proyecto si cada cliente requiere Pro, Vercel Pro/uso por entorno, mas tiempo humano. Con 26 clientes podria ser 650/mes solo Supabase base, sin contar compute/egress ni Vercel.",
            "Cliente enterprise que paga aislamiento dedicado, contrato especial, SLA propio y presupuesto de mantenimiento.",
            "Producto SaaS estandar con 26 clientes nuevos por ano.",
        ),
        option_detail(
            "Opcion 2 - Branch por cliente",
            "Un repo con una branch permanente por empresa, por ejemplo client/acme, client/empresa-b.",
            "GitHub parece ordenado al inicio, pero cada cambio global debe mergearse a N branches y resolver conflictos.",
            "Vercel podria apuntar cada proyecto a una branch distinta, generando deployments por cliente.",
            "Podria usarse una base por cliente o una base compartida, pero el problema de versiones sigue vivo.",
            "Las tablas cw y n8n no quedan normalizadas por producto; cada branch puede esperar columnas o schemas distintos.",
            "Permite personalizaciones rapidas y visibilidad por cliente.",
            "Convierte el producto en N productos. La branch deja de ser flujo de desarrollo y se vuelve contrato tecnico.",
            "Merges imposibles, regresiones distintas, QA N veces, imposibilidad de garantizar que todos tengan la misma seguridad.",
            "El costo cloud puede no subir mucho al inicio, pero el costo humano explota. Es el peor costo oculto.",
            "Pruebas temporales, hotfixes o migraciones controladas con fecha de cierre.",
            "Clientes permanentes, features globales o producto SaaS.",
        ),
        option_detail(
            "Opcion 3 - Un repo + un Vercel + Supabase compartido con company_id",
            "Una sola aplicacion SaaS. Cada fila de negocio tiene company_id. El login resuelve membresia y RLS filtra los datos.",
            "Un repo. main es produccion, staging/previews para QA, feature/* para desarrollo. No hay ramas por cliente.",
            "Un Vercel global, por ejemplo dashboard.simplia.com, dashboard.simplia.com/acme o acme.dashboard.simplia.com.",
            "Un proyecto Supabase inicial con schemas app, cw y automation. RLS por company_id y membresia activa.",
            "cw mantiene datos de Chatwoot/dashboard con company_id. n8n pasa a automation con company_id, source y raw_payload. public queda sin datos de negocio.",
            "Un producto, releases globales, onboarding rapido, costos compartidos, control de seguridad centralizado, configuracion por empresa.",
            "Requiere diseno serio: RLS, indices compuestos, auditoria, pruebas de aislamiento, limites de consumo y observabilidad por tenant.",
            "Fuga de datos si RLS falla, noisy neighbor si un cliente consume demasiado, tablas muy grandes si no se particiona/archiva.",
            "Supabase Pro desde 25/mes mas compute/disco/egress; Vercel Pro desde 20/mes mas uso. Rango interno estimado: 55-350/mes ano 1, 240-1,530/mes ano 5 antes de optimizaciones enterprise.",
            "Modelo base recomendado para SimpliaLeads y para 26 clientes nuevos por ano.",
            "Clientes con requisitos contractuales de base dedicada, region dedicada o compliance fuerte.",
        ),
        option_detail(
            "Opcion 4 - Schema por cliente dentro del mismo Supabase",
            "Se crea un schema por cliente: cw_acme, automation_acme, app_acme o similares.",
            "Un repo, pero el codigo debe construir queries dinamicas por schema o usar clientes con search_path variable.",
            "Puede seguir siendo un Vercel global.",
            "Un proyecto Supabase, muchos schemas duplicados. La base comparte compute y disco.",
            "Cada cliente tiene copia de las tablas cw y n8n. Las migraciones deben aplicarse a todos los schemas.",
            "Mas aislamiento nominal que company_id y facilita borrar/exportar todo un cliente.",
            "No elimina el problema de migraciones N veces. Complica tipos, queries, reportes globales y Edge Functions.",
            "Schemas olvidados, migraciones parciales, errores de search_path, permisos inconsistentes.",
            "Menor costo cloud que proyecto por cliente, pero alto costo de mantenimiento. A 130 clientes seria una base con cientos o miles de objetos duplicados.",
            "Casos donde un subconjunto de tablas realmente necesita aislamiento fisico dentro de la misma instancia.",
            "Modelo base de SaaS con features globales y crecimiento anual constante.",
        ),
        option_detail(
            "Opcion 5 - Tablas por cliente",
            "Se crean tablas como conversations_acme, conversations_empresa_b o n8n_chat_histories_acme.",
            "El repo debe generar nombres de tabla dinamicos o duplicar repositorios de consultas.",
            "Vercel no resuelve el problema; solo despliega el codigo.",
            "Un proyecto con muchas tablas duplicadas por cliente.",
            "cw y n8n se fragmentan por tabla. Reportes multiempresa y migraciones se vuelven muy fragiles.",
            "Separacion visual rapida para operadores tecnicos.",
            "Es un antipattern: DDL por cliente, indices duplicados, tipos imposibles, codigo complejo.",
            "Perdida de consistencia, operaciones manuales, errores de tabla equivocada, costos por indices duplicados.",
            "Cloud bajo al inicio, costo humano alto y creciente. No recomendable.",
            "Practicamente nunca para SaaS. Solo podria servir para dumps temporales o staging de importacion.",
            "Producto multiempresa real.",
        ),
        option_detail(
            "Opcion 6 - Proyecto Supabase por cliente",
            "Mismo producto, pero cada cliente tiene su proyecto Supabase y secrets propios.",
            "Un repo puede desplegar el mismo codigo, pero se necesita orquestacion de migraciones a N proyectos.",
            "Puede haber un Vercel global con backend que enrute a proyectos, o un Vercel por cliente.",
            "Aislamiento fuerte por proyecto, pero N instancias, N backups, N secretos, N cron jobs.",
            "Cada proyecto tendria cw y automation propios. Las tablas public de n8n ya no deberian existir sin RLS.",
            "Buen aislamiento, buena historia comercial para clientes enterprise, facil exportar/borrar un tenant.",
            "Multiplica operacion, monitoreo, migraciones, costos base y soporte.",
            "Proyectos desactualizados, secretos expuestos, cron jobs fallidos, falta de visibilidad central.",
            "Supabase Pro desde 25/mes por cliente si cada uno requiere Pro. Con 26 clientes: desde 650/mes; con 130: desde 3,250/mes antes de uso.",
            "Clientes grandes que pagan aislamiento dedicado o requisitos regulatorios.",
            "Clientes estandar con margen bajo/medio.",
        ),
        option_detail(
            "Opcion 7 - Backend unico + base de datos por cliente",
            "El front y backend son unicos, pero el backend selecciona la base del cliente despues del login.",
            "Un repo y una version del backend. Migrations deben correr contra cada DB.",
            "Un Vercel global para front/backend serverless o un backend separado en cloud. Las credenciales de DB nunca van al frontend.",
            "DB dedicada por cliente en Supabase, Neon, RDS, Cloud SQL o Azure SQL.",
            "cw y automation viven completos por DB. Jobs recorren clientes y conectan a cada DB.",
            "Combina version unica con aislamiento de datos alto. Es el modelo bridge de los diagramas enviados.",
            "Mas complejo que pooled: routing, pool de conexiones, migraciones N DBs, health checks por tenant.",
            "Un cliente roto puede requerir rollback propio; alto riesgo operacional si no hay plataforma de provisioning.",
            "Mas caro que pooled. Puede arrancar en cientos/mes y crecer a miles/mes con 130 clientes si cada DB usa recursos dedicados.",
            "Plan enterprise, clientes con alto volumen o necesidad contractual de aislamiento.",
            "Primer modelo base sin equipo DevOps dedicado.",
        ),
        option_detail(
            "Opcion 8 - AWS nativo",
            "Front en CloudFront/Amplify o Vercel, backend en Lambda/ECS/App Runner, DB en RDS/Aurora, storage en S3, colas en SQS.",
            "Un repo, CI/CD mas robusto, infraestructura como codigo recomendada.",
            "Vercel puede mantenerse para frontend o migrarse a AWS.",
            "RDS/Aurora Postgres permite pool, silo o bridge. Auth/RLS/app layer debe implementarse con mas piezas.",
            "cw y automation se modelan igual, pero jobs y workers migran a Lambda/EventBridge/SQS.",
            "Enterprise grade, redes privadas, compliance, backups, observabilidad y aislamiento avanzados.",
            "Mayor complejidad, mas servicios, mas FinOps, mas DevOps.",
            "Sobrecosto por arquitectura sobredimensionada y mayor tiempo de implementacion.",
            "RDS cobra instancia, storage, backups, I/O y red. Un entorno pequeno real puede empezar aprox. 80-300/mes y subir segun HA/volumen.",
            "Cuando se necesite compliance, VPC, SLA, clientes enterprise o volumen que justifique equipo cloud.",
            "MVP SaaS que necesita lanzar rapido y centralizar producto.",
        ),
        option_detail(
            "Opcion 9 - Azure",
            "App Service/Container Apps/Functions con Azure SQL o PostgreSQL Flexible Server, Key Vault, Monitor.",
            "Un repo con pipelines. Bueno si la empresa ya opera Microsoft.",
            "Vercel puede continuar o migrarse a Static Web Apps/App Service.",
            "Azure documenta patrones SaaS de base compartida, base por tenant e hibridos.",
            "cw y automation quedan igual a nivel modelo; jobs pueden pasar a Functions/Logic Apps.",
            "Muy fuerte para empresas Microsoft, identity corporativa, governance y compliance.",
            "Curva de aprendizaje, costos dispersos y mas complejidad que Supabase/Vercel.",
            "Sobredimensionar App Service/DB; costos dificiles si no se presupuestan planes e instancias.",
            "App Service cobra por plan/tier/instancias; DB por compute/storage/backup. Usar calculadora por region.",
            "Clientes corporativos Microsoft, SSO/Azure AD, compliance.",
            "Primer lanzamiento si el equipo aun esta consolidando producto.",
        ),
        option_detail(
            "Opcion 10 - GCP",
            "Cloud Run/Functions, Cloud SQL Postgres, Cloud Scheduler, Secret Manager, Cloud Storage, Monitoring.",
            "Un repo con Cloud Build/GitHub Actions.",
            "Vercel puede mantenerse o migrarse a Cloud Run/Hosting.",
            "Cloud SQL puede usar pooled, silo o bridge. Auth/RLS se implementa con Postgres/app.",
            "cw y automation migran a Postgres; jobs a Scheduler/Run/Functions.",
            "Buen serverless, escalado flexible y pricing por uso para compute.",
            "Cloud SQL puede ser costoso 24/7; requiere mas arquitectura que Supabase.",
            "Egress, conexiones y cold starts mal disenados; sobrecostos por HA.",
            "Cloud SQL cobra vCPU, memoria, storage y red. Referencia oficial: vCPU y memoria por hora segun region.",
            "Cuando se necesite Cloud Run, BigQuery/analytics o equipo con GCP.",
            "Cuando Supabase ya cubre Auth, Postgres, Edge y velocidad de producto.",
        ),
        option_detail(
            "Opcion 11 - Neon",
            "Postgres serverless con branching, compute autoscaling y storage separado.",
            "Un repo. Neon seria DB; se necesita backend/auth aparte o conservar Supabase para auth.",
            "Vercel encaja bien con Neon, pero las credenciales deben estar server-side.",
            "DB en Neon. No incluye el mismo paquete integrado de Supabase Auth/Storage/Edge/Data API.",
            "cw y automation pueden vivir como schemas en Neon con company_id; jobs requieren backend propio.",
            "Excelente para Postgres, branching, ambientes preview y costos elastico segun compute.",
            "Hay que construir o integrar Auth, storage, functions y RLS operacional.",
            "Complejidad si se migra antes de que Supabase sea una limitante real.",
            "Launch usage-based con compute-hour y storage; storage indicado oficialmente a USD 0.35/GB-mes.",
            "Futuro si el core necesita Postgres mas flexible y el equipo acepta separar backend.",
            "Si se quiere rapidez integrada con Supabase Auth/Edge/RLS ahora.",
        ),
        option_detail(
            "Opcion 12 - Railway/Render",
            "PaaS para desplegar servicios, workers, APIs o Postgres gestionado de forma sencilla.",
            "Un repo y despliegues simples.",
            "Puede reemplazar o complementar Vercel, pero Vercel sigue mejor para frontend moderno.",
            "Postgres gestionado disponible, aunque sin el mismo stack Auth/RLS/Edge de Supabase.",
            "cw y automation podrian migrarse, pero se pierde integracion Supabase.",
            "Rapido, simple y util para servicios auxiliares o self-host de n8n.",
            "Menos alineado con seguridad multiempresa basada en Supabase Auth + RLS.",
            "Costos por RAM/CPU/egress y menor governance enterprise.",
            "Railway Pro desde 20/mes mas RAM/CPU/egress/volumen. Render depende de plan por servicio/base.",
            "Servicios auxiliares, prototipos, workers, n8n self-host si se decide salir de n8n Cloud.",
            "Core de datos multiempresa con aislamiento fuerte.",
        ),
    ]

    storage_table = TableBlock(
        "Modelo de almacenamiento a 5 anos",
        ["Escenario", "Supuesto", "Ano 1", "Ano 5", "Lectura"],
        [
            [
                "Conservador",
                "0.5 GB por cliente por ano entre datos, indices y overhead.",
                "13 GB aprox.",
                "195 GB aprox.",
                "Supabase pooled sigue siendo barato; el foco es RLS e indices.",
            ],
            [
                "Medio",
                "2 GB por cliente por ano.",
                "52 GB aprox.",
                "780 GB aprox.",
                "Se requiere gobernar indices, historicos, particionado y backups.",
            ],
            [
                "Alto",
                "5 GB por cliente por ano.",
                "130 GB aprox.",
                "1.95 TB aprox.",
                "Necesario archivar historicos, revisar PITR, particionar y posiblemente separar analytics.",
            ],
        ],
        "Ano 5 suma cinco cohortes con 1 a 5 anos de datos. La cifra real debe medirse con pg_total_relation_size por tabla.",
    )

    cost_formula_table = TableBlock(
        "Formulas de costo para validar con el equipo",
        ["Componente", "Formula orientativa", "Decision que habilita"],
        [
            [
                "Costo externo cliente",
                "clientes_activos x USD 109/mes. Incluye Chatwoot 19 + n8n 20 + OpenAI 70 segun supuesto interno.",
                "Separar claramente lo que paga el cliente de lo que asume Simplia.",
            ],
            [
                "Supabase pooled",
                "plan_base + compute + max(disco_provisionado - 8 GB, 0) x 0.125 + egress_extra + storage_extra + PITR/log drains si aplica.",
                "Proyectar margen interno y saber cuando subir compute o vender enterprise.",
            ],
            [
                "Supabase por cliente",
                "N clientes x plan/proyecto + compute/disco/egress de cada uno + tiempo humano de migraciones N veces.",
                "Mostrar por que no debe ser el modelo base.",
            ],
            [
                "Vercel global",
                "seats Pro x USD 20/mes + uso excedente de bandwidth/edge/compute/builds.",
                "Mantener una sola app global y controlar uso.",
            ],
            [
                "AWS/Azure/GCP",
                "compute 24/7 o serverless + DB gestionada + storage + backups + egress + observabilidad + soporte.",
                "Comparar cuando aparezcan requisitos enterprise o compliance.",
            ],
            [
                "Bridge/enterprise",
                "costo_base_pooled + entorno dedicado del cliente + margen de soporte + SLA + backups/restores propios.",
                "Definir precio minimo para aceptar aislamiento dedicado.",
            ],
        ],
    )

    risk_table = TableBlock(
        "Riesgos principales y mitigacion",
        ["Riesgo", "Impacto", "Mitigacion requerida"],
        [
            ["Fuga de datos entre empresas", "Critico", "RLS por company_id, pruebas automaticas de aislamiento, policies revisadas y service_role solo en Edge Functions."],
            ["Personalizaciones convertidas en forks", "Alto", "Contrato: features globales; cambios por configuracion, feature flags o modulo pago."],
            ["Noisy neighbor", "Medio/Alto", "Usage events por company_id, limites por tenant, indices, rate limits y plan enterprise dedicado si supera umbrales."],
            ["Crecimiento de storage 5 anos", "Medio/Alto", "Particionado por fecha/company_id, retencion de logs, archivo frio y monitoreo de disk growth."],
            ["Dependencia n8n", "Medio", "Webhooks idempotentes, colas, error logs por company_id y plan de migrar automatizaciones criticas a codigo."],
            ["Secrets en frontend", "Critico", "Eliminar tokens VITE_* de clientes; las integraciones viven en app.company_integrations y se usan server-side."],
        ],
    )

    enterprise_gate_table = TableBlock(
        "Gates para pasar de pooled a bridge/enterprise",
        ["Condicion", "Respuesta recomendada", "Por que"],
        [
            ["Cliente exige base dedicada por contrato", "Ofrecer plan enterprise con Supabase/DB dedicada.", "Es requisito comercial/legal, no una preferencia tecnica."],
            ["Cliente supera umbral de volumen o afecta performance", "Moverlo a bridge o aislar recursos de lectura/analytics.", "Evita noisy neighbor sin cambiar el producto para todos."],
            ["Requiere region, backup, RPO/RTO o compliance distinto", "Entorno dedicado con SLA y costo separado.", "El pooled no debe cargar requisitos especiales de un solo cliente."],
            ["Pide custom code no reusable", "Convertir a modulo pago o enterprise dedicado; no branch permanente.", "Protege el producto global y evita forks encubiertos."],
            ["Necesita integraciones muy sensibles", "Secrets y workers dedicados, pero manteniendo repo unico.", "Aisla riesgo sin duplicar la logica del dashboard."],
        ],
    )

    decision_sections = [
        Section(
            "Resumen ejecutivo y recomendacion",
            [
                (
                    "callout",
                    (
                        "Decision recomendada",
                        "Construir SimpliaLeads como SaaS pooled: un repositorio, una aplicacion Vercel global, un proyecto Supabase compartido inicialmente, tablas compartidas con company_id, RLS estricto y configuracion por empresa. Mantener un modelo bridge/enterprise como excepcion pagada para clientes que requieran aislamiento dedicado.",
                        COLORS["mint"],
                    ),
                ),
                (
                    "p",
                    "El problema actual no es solamente de infraestructura. Es un problema de producto: si cada cliente tiene su propio repo, Supabase y Vercel, cada feature se transforma en una tarea repetida por cliente. Con 26 clientes nuevos por ano, ese modelo deja de ser mantenible incluso si al inicio funciona con dos clientes.",
                ),
                (
                    "p",
                    "La decision recomendada separa codigo, datos y configuracion. El codigo se despliega una vez para todos. Los datos se separan por company_id y RLS. Las diferencias entre empresas se expresan como configuracion, prompts, plantillas, integraciones, labels, feature flags y permisos, no como branches ni forks permanentes.",
                ),
                (
                    "bullets",
                    [
                        "No usar branch por cliente como estrategia permanente.",
                        "No crear Vercel por cliente como estandar.",
                        "No crear tablas por cliente.",
                        "No dejar datos de negocio en public sin RLS.",
                        "Si un cliente necesita aislamiento fuerte, venderlo como plan enterprise/bridge con precio y mantenimiento adicional.",
                    ],
                ),
            ],
        ),
        Section(
            "Diagnostico del problema actual",
            [
                (
                    "p",
                    "El flujo actual parte de un dashboard en GitHub conectado a Vercel desde main. Para cada negocio se replica el repositorio y se crea un proyecto Supabase con sus propias tablas. En el cliente actual, el schema cw concentra los backups y datos derivados de Chatwoot, mientras public contiene tablas operativas de automatizaciones n8n y configuraciones puntuales.",
                ),
                (
                    "p",
                    "Ese modelo parece seguro porque cada cliente esta separado, pero no escala como producto. Cada cambio funcional, ajuste visual, correccion de seguridad, nueva tabla, funcion Edge o cron debe repetirse en todos los clientes. Si el cambio se genero por iteraciones con prompts, la probabilidad de inconsistencias sube con cada cliente.",
                ),
                (
                    "bullets",
                    [
                        "A 2 clientes todavia se puede replicar manualmente; a 5 ya es una operacion fragil; a 26 nuevos por ano se vuelve inviable.",
                        "La calidad del producto se rompe porque cada cliente puede quedar en una version distinta.",
                        "El soporte se vuelve mas caro porque no existe una sola fuente de verdad.",
                        "La seguridad se vuelve dificil de auditar porque policies, tablas y secrets viven en lugares repetidos.",
                        "La venta de cambios personalizados sin arquitectura de configuracion crea forks encubiertos.",
                    ],
                ),
            ],
        ),
        Section(
            "Proyeccion comercial y costos externos",
            [
                ("p", "El supuesto corregido es 26 clientes nuevos por ano. El objetivo real no es solo vender 26 al ano, sino retenerlos, porque la data se conserva durante 5 anos por cliente. Por eso, el almacenamiento y los historicos crecen por cohortes, aunque la venta anual se mantenga constante."),
                ("table", retention_table),
                ("table", external_table),
                (
                    "p",
                    "El costo externo indicado por negocio es USD 109/mes por cliente: Chatwoot USD 19, n8n USD 20 y OpenAI USD 70. Ese costo deberia ser asumido o trasladado al cliente como parte de su operacion. No debe confundirse con el costo interno de Simplia por operar el SaaS.",
                ),
            ],
        ),
        Section(
            "Costos internos Simplia",
            [
                (
                    "p",
                    "Simplia asumiria principalmente Supabase, Vercel, almacenamiento, backups, logs, monitoreo, Edge Functions, mantenimiento del producto y tiempo de soporte. En un modelo pooled estos costos se comparten entre clientes; en un modelo silo se multiplican por cliente.",
                ),
                ("table", internal_table),
                ("table", storage_table),
                ("table", cost_formula_table),
                (
                    "p",
                    "Supabase Pro parte desde USD 25/mes segun pricing oficial e incluye 8 GB de disco por proyecto. El disco adicional se cotiza oficialmente desde USD 0.125/GB-mes para gp3. Vercel Pro parte desde USD 20/mes mas uso adicional. En ambos casos, el costo real depende del consumo: egress, compute, storage, logs, usuarios, Edge Functions y retencion.",
                ),
                (
                    "callout",
                    (
                        "Como leer estos rangos",
                        "Los rangos no son una promesa de factura. Sirven para comparar modelos. La factura real debe medirse con datos: filas por tabla, mensajes por dia, tamano de raw_payload, cantidad de reportes, archivos almacenados, ejecuciones Edge y egress.",
                        COLORS["soft_blue"],
                    ),
                ),
            ],
        ),
        Section(
            "Comparativa completa de alternativas",
            [
                ("table", summary_matrix),
                (
                    "p",
                    "La comparativa se basa en tres dimensiones: velocidad de producto, aislamiento de datos y costo operativo. Para SimpliaLeads, el cuello de botella actual es la repeticion manual y la divergencia de versiones. Por eso, la solucion base debe maximizar una sola version del producto y mover la personalizacion a configuracion.",
                ),
                *[("table", option) for option in strategic_options],
            ],
        ),
        Section(
            "Politica de producto y personalizaciones",
            [
                (
                    "p",
                    "La politica comercial debe decir explicitamente que SimpliaLeads es un producto SaaS con features globales. El cliente no compra un repositorio, una branch ni un Vercel propio; compra acceso al producto. Las diferencias permitidas deben vivir en configuracion por empresa.",
                ),
                (
                    "bullets",
                    [
                        "Permitido como configuracion: prompts, plantillas, contexto empresarial, labels, columnas visibles, dashboards visibles, campos personalizados, webhooks, URLs de integracion, horarios, reportes y feature flags.",
                        "Permitido como modulo: funcionalidad reutilizable que puede activarse para uno o varios clientes sin forkear el codigo.",
                        "No permitido como estandar: cambiar codigo solo para un cliente mediante branch permanente.",
                        "Excepcion enterprise: entorno dedicado, base dedicada o deployment aislado con precio, SLA, soporte y mantenimiento adicional.",
                    ],
                ),
                (
                    "p",
                    "Si un cliente pide cambiar el nombre de una columna visible, no se debe alterar la columna fisica de base de datos. Se resuelve con traduccion/label de UI en app.dashboard_settings o app.company_settings. El estandar tecnico se mantiene en ingles y snake_case; la presentacion al usuario puede estar en espanol o en el lenguaje comercial de cada cliente.",
                ),
            ],
        ),
        Section(
            "Recomendacion final",
            [
                (
                    "callout",
                    (
                        "Arquitectura objetivo",
                        "Un repo GitHub, un deployment Vercel global, un proyecto Supabase compartido inicialmente, schemas app/cw/automation, company_id en toda tabla de negocio, RLS por membresia activa, roles por empresa y configuracion por tenant.",
                        COLORS["mint"],
                    ),
                ),
                (
                    "p",
                    "Esta opcion se adapta mejor porque equilibra velocidad, costo y escalabilidad. Permite vender 26 clientes nuevos por ano sin que cada cliente duplique el ciclo de desarrollo. Tambien preserva una ruta enterprise: si un cliente paga aislamiento adicional, se puede mover a un proyecto/base dedicado manteniendo el mismo codigo.",
                ),
                (
                    "bullets",
                    [
                        "Primero: reestructurar el dashboard para multiempresa.",
                        "Segundo: mover tablas operativas fuera de public y agregar company_id.",
                        "Tercero: activar RLS real y pruebas automaticas de aislamiento.",
                        "Cuarto: transformar prompts, plantillas e integraciones en configuracion por empresa.",
                        "Quinto: definir contrato comercial de features globales y enterprise dedicado como excepcion.",
                    ],
                ),
                ("table", enterprise_gate_table),
            ],
        ),
        Section("Riesgos, limites y mitigaciones", [("table", risk_table)]),
        Section(
            "Fuentes oficiales consultadas",
            [
                (
                    "table",
                    TableBlock(
                        "Fuentes",
                        ["Fuente", "URL", "Uso en el documento"],
                        SOURCES,
                    ),
                )
            ],
        ),
    ]

    return DocSpec(
        title="Decision Estrategica de Arquitectura SaaS Multiempresa",
        subtitle="SimpliaLeads - evaluacion de alternativas, costos, riesgos y recomendacion para escalar a 26 clientes nuevos por ano",
        code="SIMPLIA-SAA-DEC-001",
        version="1.0",
        status="Borrador formal para decision interna",
        owner="Simplia - Producto y Tecnologia",
        audience="Direccion, Gerencia, Producto, Tecnologia y Operaciones",
        sections=decision_sections,
    )


def build_implementation_doc() -> DocSpec:
    current_state_table = TableBlock(
        "Estado actual observado en repo y Supabase",
        ["Area", "Estado actual", "Implicacion para SaaS"],
        [
            ["GitHub/Vercel", "Un dashboard conectado desde main hacia Vercel.", "Debe pasar a ser un producto global, no un deployment por cliente."],
            ["Schema cw", "Concentra datos core del dashboard y backups desde Chatwoot.", "Debe mantenerse como dominio comercial, pero con company_id en tablas de negocio."],
            ["public", "Contiene n8n_followups_sent, N8N_error_Logs, n8n_chat_histories, citas_agendadas, scheduled_appointments, dashboard_tag_settings y users.", "No debe alojar datos de negocio multiempresa sin RLS ni separacion por company_id."],
            ["dashboard_tag_settings", "Existe fallback desde codigo y Edge Functions a public.dashboard_tag_settings.", "Debe centralizarse en app.dashboard_settings o cw.dashboard_tag_settings con company_id; recomendado app.dashboard_settings."],
            ["Roles", "Ya aparecen platform_admin, company_admin y operator en codigo.", "Falta unirlos a empresa/membresia real."],
            ["Auth", "auth.users identifica usuarios, y user_profiles se consulta para rol.", "Debe migrarse a app.user_profiles + app.company_members para multiempresa."],
            ["Jobs", "chatwoot-sync diario 12:01 a. m. Ecuador y send-scheduled-reports cada 5 minutos.", "Ambos deben iterar o filtrar por company_id y registrar runs por empresa."],
            ["Secrets", "El frontend referencia VITE_CHATWOOT_ACCOUNT_ID y VITE_CHATWOOT_API_TOKEN.", "Los tokens de clientes no pueden vivir en VITE_*; deben usarse server-side."],
        ],
    )

    github_table = TableBlock(
        "Modelo GitHub recomendado",
        ["Rama/flujo", "Uso", "Regla"],
        [
            ["main", "Produccion", "Siempre deployable. No se hacen cambios directos sin PR/revision."],
            ["staging", "QA integrado", "Opcional si el equipo necesita ambiente estable previo a main."],
            ["feature/*", "Desarrollo de funcionalidades", "Temporal. Se elimina despues de merge."],
            ["hotfix/*", "Correcciones urgentes", "Temporal y con merge rapido a main/staging."],
            ["client/*", "No recomendado", "No crear branches permanentes por cliente."],
        ],
    )

    vercel_table = TableBlock(
        "Modelo Vercel recomendado",
        ["Decision", "Como se maneja", "Motivo"],
        [
            ["Un dashboard global", "dashboard.simplia.com o dashboard.simplia.com/<slug> o <slug>.dashboard.simplia.com.", "Un solo producto, una sola version y menos operacion."],
            ["Identificacion de empresa", "Por sesion y membresia; opcionalmente slug/subdominio para experiencia.", "La URL no sustituye la seguridad. RLS filtra por company_id."],
            ["Variables de entorno", "Solo variables globales no secretas en VITE_*. Tokens por cliente guardados server-side.", "Todo VITE_* se expone al navegador."],
            ["Previews", "Cada PR puede tener preview Vercel conectado a Supabase staging.", "QA sin afectar clientes."],
            ["Vercel dedicado", "Solo enterprise con entorno aislado y precio adicional.", "Evita duplicar deployments para clientes estandar."],
        ],
    )

    schema_table = TableBlock(
        "Schemas Supabase objetivo",
        ["Schema", "Responsabilidad", "Ejemplos", "Exposicion API"],
        [
            ["app", "Core SaaS: empresas, usuarios, roles, configuracion, prompts, integraciones, auditoria.", "companies, company_members, company_settings, dashboard_settings, prompt_templates.", "Preferir no exponer completo; usar RPC/API controlada o RLS estricta."],
            ["cw", "Datos derivados de Chatwoot y dashboard comercial.", "contacts_current, conversations_current, messages, inboxes, reports, sync_runs.", "Solo lo necesario para el dashboard, con RLS por company_id."],
            ["automation", "n8n, webhooks, citas, colas, logs operativos, raw payloads.", "chat_histories, followups_sent, scheduled_appointments, error_logs.", "Privado por defecto; escritura mediante Edge Functions/webhooks server-side."],
            ["public", "Idealmente sin datos de negocio.", "Extensiones o views publicas estrictamente necesarias.", "Minimizar superficie; cualquier tabla expuesta requiere RLS."],
        ],
    )

    new_tables = TableBlock(
        "Nuevas tablas SaaS requeridas",
        ["Tabla", "Proposito", "Campos minimos recomendados"],
        [
            ["app.companies", "Tenant/empresa cliente.", "id, slug, legal_name, display_name, status, plan, created_at, updated_at."],
            ["app.company_members", "Membresia usuario-empresa y rol.", "company_id, user_id, role, status, invited_by, created_at."],
            ["app.user_profiles", "Perfil de usuario complementario a auth.users.", "user_id, full_name, default_company_id, locale, created_at."],
            ["app.company_settings", "Configuracion general por empresa.", "company_id, key, value_json, updated_by, updated_at."],
            ["app.company_integrations", "Credenciales y endpoints por empresa.", "company_id, provider, encrypted_secret_ref, settings_json, status."],
            ["app.company_feature_flags", "Activacion de modulos/features por empresa.", "company_id, flag_key, enabled, config_json."],
            ["app.dashboard_settings", "Settings visuales/comerciales del dashboard.", "company_id, profile_key, settings_json, updated_by."],
            ["app.prompt_templates", "Prompts por empresa o globales versionados.", "company_id nullable, template_key, version, content, status."],
            ["app.message_templates", "Plantillas de mensajes por empresa.", "company_id, channel, template_key, content, variables_json."],
            ["app.company_audit_logs", "Auditoria funcional y administrativa.", "company_id, actor_user_id, action, entity, metadata, created_at."],
            ["app.usage_events", "Medicion de consumo por empresa.", "company_id, event_type, quantity, metadata, created_at."],
        ],
    )

    relocation_table = TableBlock(
        "Reubicacion de tablas public actuales",
        ["Actual", "Destino recomendado", "Accion tecnica", "Por que"],
        [
            ["public.n8n_chat_histories", "automation.chat_histories", "Crear tabla con company_id, source, raw_payload, session_id protegido, created_at.", "Es historial operativo de automatizacion, no dato publico."],
            ["public.n8n_followups_sent", "automation.followups_sent", "Migrar filas y backfill company_id desde contexto de flujo.", "Evita mezclar followups entre empresas."],
            ["public.N8N_error_Logs", "automation.error_logs", "Normalizar nombre a snake_case y registrar company_id cuando exista.", "Logs por empresa permiten soporte y auditoria."],
            ["public.citas_agendadas", "automation.scheduled_appointments", "Unificar con scheduled_appointments o mapear a entidad de citas.", "Evita doble fuente de verdad en citas."],
            ["public.scheduled_appointments", "automation.scheduled_appointments", "Definir columnas canonicas y migrar.", "Debe ser operacional y multiempresa."],
            ["public.dashboard_tag_settings", "app.dashboard_settings", "Eliminar fallback inseguro/duplicado y guardar por company_id/profile_key.", "Es configuracion del producto por empresa."],
            ["public.users", "app.user_profiles o eliminar", "Si duplica auth.users, migrar solo perfil y rol a company_members.", "La identidad real debe ser auth.users."],
        ],
    )

    automation_variability_table = TableBlock(
        "Regla para tablas operativas variables por cliente",
        ["Caso", "Solucion recomendada", "Evitar"],
        [
            [
                "Workflow reutilizable para varios clientes",
                "Crear tabla canonica en automation con company_id, workflow_key, status, timestamps y payload tipado.",
                "Crear una tabla por cliente o por nombre de empresa.",
            ],
            [
                "Workflow unico de un cliente, pero con datos que el dashboard no consulta",
                "Guardar en automation.workflow_events o automation.raw_events con company_id, workflow_key, event_type, raw_payload jsonb.",
                "Ensuciar public con tablas sueltas o duplicar schemas.",
            ],
            [
                "Workflow unico que luego sera feature del producto",
                "Modelarlo desde el inicio como tabla canonica y feature flag por empresa.",
                "Codificarlo como excepcion privada que no pueda versionarse.",
            ],
            [
                "Citas, followups, chat histories o errores",
                "Usar tablas comunes: scheduled_appointments, followups_sent, chat_histories, error_logs.",
                "Mantener nombres en espanol, mayusculas o duplicados public/cw.",
            ],
            [
                "Datos crudos para auditoria o debug",
                "Guardar payload en jsonb con retencion definida y posibilidad de purga/archivo.",
                "Guardar raw_payload infinito sin politica de retencion.",
            ],
            [
                "Necesidad de analitica historica",
                "Promover campos consultados a columnas indexadas y dejar el resto en jsonb.",
                "Consultar masivamente jsonb sin indices durante 5 anos de datos.",
            ],
        ],
        "Esta regla responde al caso donde un cliente tenga 3 tablas n8n y otro 8: se modela por tipo de dato/workflow, no por empresa.",
    )

    cw_company_table = TableBlock(
        "cw multiempresa",
        ["Grupo de tablas", "Cambio requerido", "Indices recomendados"],
        [
            ["Conversaciones", "Agregar company_id a conversations_current, conversation_attribute_history, conversation_label_history, conversation_label_events.", "(company_id, updated_at), (company_id, conversation_id), (company_id, business_stage)."],
            ["Contactos", "Agregar company_id a contacts_current, contact_inboxes, contact_attribute_history.", "(company_id, contact_id), (company_id, phone_number), (company_id, updated_at)."],
            ["Mensajes", "Agregar company_id a messages.", "(company_id, conversation_id, created_at), (company_id, created_at)."],
            ["Inboxes/equipos", "Agregar company_id a inboxes, teams y configuracion relacionada.", "(company_id, inbox_id), (company_id, name)."],
            ["Reportes", "Agregar company_id a automated_reports, automated_report_runs, reporting_events, daily_metrics.", "(company_id, scheduled_at), (company_id, run_at), (company_id, metric_date)."],
            ["Sync/import", "Agregar company_id a sync_runs, raw_ingest, import_batches, import_batch_errors.", "(company_id, started_at), (company_id, batch_id)."],
            ["Auditoria", "Agregar company_id a commercial_audit_events y eventos derivados.", "(company_id, created_at), (company_id, entity_type, entity_id)."],
        ],
    )

    rls_table = TableBlock(
        "Reglas RLS base",
        ["Caso", "Regla", "Notas"],
        [
            ["platform_admin", "Puede operar todas las empresas.", "Debe identificarse por membership interna o claim controlado, no por email hardcodeado."],
            ["company_admin", "Solo ve/administra su company_id.", "Puede gestionar settings, usuarios del cliente, reportes e integraciones permitidas."],
            ["operator", "Solo ve modulos autorizados de su company_id.", "No accede a configuracion global ni datos de otras empresas."],
            ["service_role", "Puede escribir procesos server-side.", "Solo en Edge Functions/jobs; nunca en frontend."],
            ["Webhooks n8n", "Si no resuelven company_id, se rechaza la escritura.", "Resolver por secret, integration_id, account_id, inbox_id o webhook path firmado."],
        ],
    )

    onboarding_table = TableBlock(
        "Flujo de entrega a cliente",
        ["Paso", "Responsable", "Resultado"],
        [
            ["Crear empresa", "platform_admin", "Registro en app.companies con slug, plan, estado activo y parametros base."],
            ["Configurar integraciones", "platform_admin", "Chatwoot/n8n/OpenAI/webhooks guardados en app.company_integrations con secrets server-side."],
            ["Configurar dashboard", "platform_admin/company_admin", "Settings, prompts, plantillas, labels, campos visibles y reportes por empresa."],
            ["Invitar usuarios", "platform_admin/company_admin", "auth.users + app.user_profiles + app.company_members."],
            ["Login", "Usuario cliente", "El sistema resuelve empresas disponibles y rol."],
            ["Seleccion de empresa", "Sistema/usuario", "active_company_id se guarda en sesion o se deriva de slug/subdominio."],
            ["Acceso a datos", "Supabase RLS", "Toda consulta queda filtrada por company_id y rol."],
        ],
    )

    jobs_table = TableBlock(
        "Jobs y backups multiempresa",
        ["Job", "Estado actual", "Objetivo multiempresa"],
        [
            ["chatwoot-sync diario 12:01 a. m. Ecuador", "Cron a las 05:01 UTC llama chatwoot-sync.", "Recorrer app.companies activas con integracion Chatwoot; cada run escribe company_id en cw.sync_runs y tablas destino."],
            ["send-scheduled-reports cada 5 minutos", "Cron llama Edge Function que consulta reportes.", "Buscar reportes por company_id, respetar timezone/estado, crear automated_report_runs con company_id."],
            ["n8n webhooks", "Insertan en tablas public operativas.", "Insertar en automation.* resolviendo empresa por webhook/integration secret; rechazar sin company_id."],
            ["Backups de datos", "Supabase daily backups segun plan y backups logicos si se configuran.", "Plan de 5 anos: daily backup operativo, exports periodicos por tenant, prueba de restore y archivo historico."],
            ["Errores", "Logs dispersos o N8N_error_Logs.", "automation.error_logs con company_id, provider, severity, raw_payload y correlation_id."],
        ],
    )

    migration_table = TableBlock(
        "Plan de migracion por fases",
        ["Fase", "Objetivo", "Tareas", "Criterio de salida"],
        [
            ["0. Medicion", "Conocer volumen real.", "Medir filas y tamanos por tabla; listar RLS/policies; inventariar webhooks y secrets.", "Inventario aprobado y backup previo."],
            ["1. Base SaaS", "Crear app.companies, company_members, user_profiles y settings.", "Crear tenant inicial para cliente actual, roles y seeds.", "Usuarios actuales pueden loguear con empresa inicial."],
            ["2. company_id", "Preparar datos existentes.", "Agregar company_id nullable, backfill, indices compuestos, constraints.", "Todas las tablas de negocio tienen company_id poblado."],
            ["3. Reubicar public", "Mover operativas a app/automation.", "Crear nuevas tablas, copiar datos, views temporales si se necesita compatibilidad.", "No hay escrituras nuevas a public para datos de negocio."],
            ["4. Actualizar codigo", "Frontend y Edge Functions multiempresa.", "Resolver active_company_id, filtrar queries, quitar fallback a public.dashboard_tag_settings.", "Dashboard funciona para empresa inicial."],
            ["5. RLS real", "Aislamiento por tenant.", "Activar policies restrictivas, quitar using(true), pruebas A/B.", "Empresa A no puede leer/escribir B."],
            ["6. Jobs multiempresa", "Backups, sync y reportes por empresa.", "Iterar companies activas, registrar runs por company_id, logs por empresa.", "Jobs pasan pruebas con dos empresas simuladas."],
            ["7. Hardening", "Seguridad y operacion.", "Mover secrets, monitoreo, alertas, limites de uso, restore drill.", "Checklist de salida a produccion aprobado."],
        ],
    )

    testing_table = TableBlock(
        "Pruebas obligatorias",
        ["Prueba", "Resultado esperado"],
        [
            ["Empresa A consulta conversaciones", "Solo recibe company_id A."],
            ["Empresa A intenta consultar B por id conocido", "RLS devuelve 0 filas o error controlado."],
            ["company_admin de A accede a settings de B", "Acceso denegado."],
            ["operator accede a configuracion global", "Acceso denegado."],
            ["Webhook n8n sin empresa", "HTTP 400/401; no inserta datos."],
            ["Webhook n8n con secret de A", "Inserta solo company_id A."],
            ["Sync nocturno", "Crea sync_run por empresa e inserta company_id en destino."],
            ["Reporte programado", "Genera runs por empresa sin mezclar settings/prompts."],
            ["dashboard_settings por empresa", "Cambiar labels de A no afecta B."],
            ["Consulta historica 5 anos", "Usa indices/particiones y no escanea toda la base."],
        ],
    )

    sql_example = """-- Ejemplo conceptual, no ejecutar sin adaptarlo a migraciones reales.
create schema if not exists app;
create schema if not exists automation;

create table if not exists app.companies (
  id uuid primary key default gen_random_uuid(),
  slug text not null unique,
  display_name text not null,
  status text not null default 'active',
  created_at timestamptz not null default now()
);

create table if not exists app.company_members (
  company_id uuid not null references app.companies(id),
  user_id uuid not null references auth.users(id),
  role text not null check (role in ('platform_admin', 'company_admin', 'operator')),
  status text not null default 'active',
  created_at timestamptz not null default now(),
  primary key (company_id, user_id)
);

alter table app.company_members enable row level security;
-- Policies finales deben revisarse con SECURITY DEFINER helpers auditados."""

    implementation_sections = [
        Section(
            "Resumen tecnico",
            [
                (
                    "callout",
                    (
                        "Objetivo de implementacion",
                        "Convertir el dashboard actual de un cliente en una plataforma SaaS multiempresa: misma app, mismo codigo, datos aislados por company_id y RLS, configuracion por empresa, jobs multiempresa y sin secretos de cliente en frontend.",
                        COLORS["mint"],
                    ),
                ),
                (
                    "p",
                    "Este blueprint no ejecuta la migracion. Define como deberia hacerse para que el siguiente documento o sprint pueda convertirlo en tickets, migraciones SQL, cambios de frontend, Edge Functions y pruebas de seguridad.",
                ),
            ],
        ),
        Section("Estado actual real", [("table", current_state_table)]),
        Section(
            "Arquitectura objetivo",
            [
                (
                    "p",
                    "La aplicacion debe operar como SaaS pooled. El navegador nunca decide por si solo que datos puede ver. El usuario inicia sesion, el sistema determina sus empresas en app.company_members, selecciona active_company_id y cada consulta queda protegida por RLS y/o funciones server-side.",
                ),
                (
                    "bullets",
                    [
                        "Una sola version de codigo para todos los clientes.",
                        "Una sola app Vercel global para clientes estandar.",
                        "Un Supabase compartido inicial con schemas app, cw y automation.",
                        "company_id obligatorio en toda tabla que contenga datos de cliente.",
                        "Configuracion por empresa para prompts, plantillas, labels, campos visibles, reportes y webhooks.",
                        "Modelo bridge/dedicado solo para enterprise.",
                    ],
                ),
            ],
        ),
        Section("GitHub", [("table", github_table)]),
        Section("Vercel", [("table", vercel_table)]),
        Section(
            "Supabase y schemas",
            [
                ("table", schema_table),
                (
                    "p",
                    "La razon de mover datos fuera de public es de seguridad y claridad. En Supabase, public suele estar expuesto por la Data API. En multiempresa, cualquier tabla accesible sin RLS correcta es un riesgo de fuga. Los schemas app, cw y automation permiten separar producto, dominio Chatwoot y automatizaciones.",
                ),
                ("table", new_tables),
                ("table", relocation_table),
                ("table", cw_company_table),
            ],
        ),
        Section(
            "Reglas de datos operativos n8n",
            [
                (
                    "p",
                    "Las tablas operativas no tienen que vivir en cw si no son parte del dashboard comercial ni backups de Chatwoot. Deben vivir en automation porque pertenecen a flujos, webhooks, errores, colas e historiales tecnicos. Eso permite que cw se mantenga como modelo de negocio y analytics del dashboard.",
                ),
                (
                    "bullets",
                    [
                        "Toda tabla automation con datos de cliente debe tener company_id.",
                        "Toda tabla de log/historial debe tener created_at, source, raw_payload y correlation_id cuando aplique.",
                        "Los webhooks de n8n deben resolver company_id antes de insertar, usando integration_id, token firmado, path con slug validado server-side o mapping de Chatwoot account/inbox.",
                        "Si no se puede resolver empresa, la solicitud se rechaza y se registra error sin mezclar datos.",
                        "Los nombres deben normalizarse a ingles y snake_case: error_logs, chat_histories, followups_sent, scheduled_appointments.",
                    ],
                ),
                ("table", automation_variability_table),
            ],
        ),
        Section(
            "Roles, login y membresia",
            [
                ("table", rls_table),
                (
                    "p",
                    "auth.users debe ser la identidad. app.user_profiles guarda perfil complementario. app.company_members es la tabla que responde que empresas puede ver un usuario y con que rol. Un usuario puede pertenecer a una o varias empresas; por eso el rol no debe ser un unico campo global sin company_id.",
                ),
                (
                    "bullets",
                    [
                        "platform_admin: usuario interno Simplia con acceso multiempresa y operaciones de soporte.",
                        "company_admin: gerencia del cliente; administra solo su empresa.",
                        "operator: usuario operativo; ve modulos permitidos y datos de su empresa.",
                        "La UI puede ocultar modulos, pero la seguridad real vive en RLS/Edge Functions.",
                    ],
                ),
                ("table", onboarding_table),
            ],
        ),
        Section(
            "Configuracion por empresa",
            [
                (
                    "p",
                    "La personalizacion permitida debe convertirse en datos, no en codigo por cliente. Eso permite que una empresa tenga prompts, plantillas, labels, reportes o webhooks distintos sin romper el producto global.",
                ),
                (
                    "bullets",
                    [
                        "Prompts por empresa: app.prompt_templates con version y estado.",
                        "Plantillas: app.message_templates por canal y variables.",
                        "Labels y nombres visibles: app.dashboard_settings o app.company_settings.",
                        "Campos visibles y orden de columnas: settings_json por perfil/dashboard.",
                        "Webhooks y URLs de integracion: app.company_integrations, secrets server-side.",
                        "Feature flags: app.company_feature_flags para activar modulos sin branch.",
                        "Contexto empresarial para IA: company_settings/prompt_templates versionados.",
                    ],
                ),
            ],
        ),
        Section("Backups, sync y jobs", [("table", jobs_table)]),
        Section(
            "Seguridad tecnica",
            [
                (
                    "bullets",
                    [
                        "Eliminar tokens de cliente de VITE_*; cualquier variable VITE_* es visible en el navegador.",
                        "Usar Edge Functions o backend server-side para Chatwoot, OpenAI y webhooks que requieran secretos.",
                        "Activar RLS en todas las tablas expuestas y quitar policies using(true) antes de multiempresa.",
                        "Preferir schemas internos no expuestos para automation y app sensible.",
                        "Crear helpers de RLS auditados para validar membresia activa.",
                        "Registrar auditoria con company_id, actor_user_id, action, metadata y created_at.",
                    ],
                ),
                ("code", ("Ejemplo conceptual de tablas base", sql_example)),
            ],
        ),
        Section("Migracion propuesta", [("table", migration_table)]),
        Section(
            "Riesgos tecnicos y controles",
            [
                (
                    "table",
                    TableBlock(
                        "Controles tecnicos",
                        ["Riesgo", "Control"],
                        [
                            ["Fuga de datos por RLS mal hecha", "Pruebas automaticas cross-tenant, revision de policies, uso minimo de service_role."],
                            ["Mezcla de datos n8n", "Resolver company_id en webhook y rechazar payloads sin tenant."],
                            ["Noisy neighbor", "usage_events por tenant, indices, limites, alertas y opcion bridge."],
                            ["Storage alto por 5 anos", "Particionado, archivo, limpieza de raw_payload innecesario y medicion mensual."],
                            ["n8n cuello de botella", "Idempotencia, colas, retries, limites y migrar flujos criticos a codigo si crecen."],
                            ["Secrets expuestos", "Mover secrets a Supabase secrets/Vault/variables server-side, nunca frontend."],
                            ["Forks encubiertos", "Feature flags, modulos configurables y politica contractual."],
                        ],
                    ),
                )
            ],
        ),
        Section("Testing y criterios de aceptacion", [("table", testing_table)]),
        Section(
            "Roadmap de implementacion sugerido",
            [
                (
                    "numbered",
                    [
                        "Semana 1: inventario final, medicion de tablas, diseno SQL y contrato de company_id.",
                        "Semana 2: crear schema app, tenant inicial, membresias y settings base.",
                        "Semana 3: agregar company_id e indices en cw; backfill del cliente actual.",
                        "Semana 4: mover tablas public a automation/app y crear compatibilidad temporal si se requiere.",
                        "Semana 5: actualizar frontend, repositorios de datos y Edge Functions para active_company_id.",
                        "Semana 6: RLS restrictivo, pruebas de aislamiento, jobs multiempresa y hardening de secrets.",
                        "Semana 7: onboarding de segundo cliente real en el mismo SaaS y monitoreo.",
                    ],
                ),
                (
                    "p",
                    "El cronograma depende del tamano real de las tablas y de cuanto codigo hoy consulta directamente public/cw sin capa de repositorio. No se recomienda activar multiempresa sin completar las pruebas de aislamiento.",
                ),
            ],
        ),
        Section(
            "Fuentes oficiales consultadas",
            [
                (
                    "table",
                    TableBlock(
                        "Fuentes",
                        ["Fuente", "URL", "Uso en el documento"],
                        SOURCES,
                    ),
                )
            ],
        ),
    ]

    return DocSpec(
        title="Blueprint Tecnico de Implementacion SaaS Multiempresa",
        subtitle="SimpliaLeads - GitHub, Vercel, Supabase, roles, schemas, tablas operativas, backups y migracion",
        code="SIMPLIA-SAA-IMP-001",
        version="1.0",
        status="Borrador tecnico para planificacion",
        owner="Simplia - Producto y Tecnologia",
        audience="Equipo tecnico, producto, operaciones y direccion",
        sections=implementation_sections,
    )


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    specs = [
        (build_decision_doc(), DECISION_MD, DECISION_DOCX),
        (build_implementation_doc(), IMPLEMENTATION_MD, IMPLEMENTATION_DOCX),
    ]
    for spec, md_path, docx_path in specs:
        write_markdown(spec, md_path)
        write_docx(spec, docx_path)
        print(f"Generated {md_path.relative_to(ROOT)}")
        print(f"Generated {docx_path.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
