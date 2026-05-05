from __future__ import annotations

import os
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
ASSET_DIR = ROOT / "Manual_Usuario_SimpliaLeads_assets"
DOCX_PATH = ROOT / "Manual_Usuario_SimpliaLeads_ISO10013.docx"
PDF_PATH = ROOT / "Manual_Usuario_SimpliaLeads_ISO10013.pdf"
LOGO_PATH = ROOT / "logo_simplia.png"

TODAY = "04/05/2026"
NEXT_REVIEW = "04/11/2026"
DOC_CODE = "MAN-USR-SIMPLIA-001"
DOC_VERSION = "1.1"
DOC_STATUS = "Vigente/Borrador"
TZ = "America/Guayaquil"

COLORS = {
    "navy": "1f3f91",
    "blue": "274690",
    "light_blue": "eaf1ff",
    "green": "0a9b6f",
    "mint": "dff7ec",
    "orange": "f59e0b",
    "red": "ef4444",
    "slate": "64748b",
    "ink": "0f2344",
    "line": "d9e2ef",
    "bg": "f8fafc",
    "white": "ffffff",
}


def rgb(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip("#")
    return tuple(int(h[i : i + 2], 16) for i in (0, 2, 4))


def get_font(size: int = 22, bold: bool = False) -> ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for path in candidates:
        if os.path.exists(path):
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def draw_round(draw: ImageDraw.ImageDraw, xy, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def center_text(draw: ImageDraw.ImageDraw, box, text: str, font, fill):
    bbox = draw.textbbox((0, 0), text, font=font)
    x = box[0] + (box[2] - box[0] - (bbox[2] - bbox[0])) / 2
    y = box[1] + (box[3] - box[1] - (bbox[3] - bbox[1])) / 2
    draw.text((x, y), text, font=font, fill=fill)


def truncate(text: str, font, max_width: int) -> str:
    if not text:
        return ""
    probe = ImageDraw.Draw(Image.new("RGB", (1, 1)))
    if probe.textlength(text, font=font) <= max_width:
        return text
    value = text
    while len(value) > 3 and probe.textlength(value + "...", font=font) > max_width:
        value = value[:-1]
    return value + "..."


def make_visual(tab_label: str, title: str, subtitle: str, kpis, panels, filename: str) -> Path:
    width, height = 1480, 900
    image = Image.new("RGB", (width, height), rgb(COLORS["bg"]))
    draw = ImageDraw.Draw(image)
    f_logo = get_font(32, True)
    f_tiny = get_font(14)
    f_tab = get_font(17, True)
    f_h1 = get_font(30, True)
    f_h2 = get_font(21, True)
    f_body = get_font(16)
    f_small = get_font(13)
    f_value = get_font(34, True)
    f_metric = get_font(18, True)

    draw.rectangle((0, 0, width, 82), fill=rgb("ffffff"))
    draw.line((0, 82, width, 82), fill=rgb(COLORS["line"]), width=2)
    if LOGO_PATH.exists():
        logo = Image.open(LOGO_PATH).convert("RGBA")
        logo.thumbnail((125, 50))
        image.paste(logo, (60, 18), logo)
    else:
        draw.text((60, 18), "SIMPLIA", font=f_logo, fill=rgb(COLORS["blue"]))
    draw.text((60, 55), "CONTROL COMERCIAL", font=f_tiny, fill=rgb(COLORS["slate"]))
    draw_round(draw, (1185, 25, 1370, 56), 16, rgb("dcfce7"), outline=rgb("bbf7d0"))
    center_text(draw, (1185, 25, 1370, 56), "Datos en vivo + historial", get_font(14, True), rgb("16a34a"))

    tabs = [
        "Estrategia",
        "Embudo",
        "Operación",
        "Seguimiento",
        "Rendimiento Humano",
        "Tendencias",
        "Calidad",
        "Conversaciones",
        "Reportes",
    ]
    x0, y0 = 92, 118
    draw_round(draw, (x0, y0, width - 92, y0 + 58), 14, rgb("ffffff"), outline=rgb(COLORS["line"]))
    x = x0 + 10
    for tab in tabs:
        tab_width = int(draw.textlength(tab, font=f_tab)) + 38
        if tab == tab_label:
            draw_round(draw, (x, y0 + 8, x + tab_width, y0 + 50), 11, rgb(COLORS["blue"]))
            center_text(draw, (x, y0 + 8, x + tab_width, y0 + 50), tab, f_tab, rgb("ffffff"))
        else:
            center_text(draw, (x, y0 + 8, x + tab_width, y0 + 50), tab, f_tab, rgb(COLORS["slate"]))
        x += tab_width + 10

    draw_round(draw, (92, 210, width - 92, 280), 13, rgb("ffffff"), outline=rgb(COLORS["line"]))
    draw_round(draw, (112, 228, 310, 264), 9, rgb(COLORS["bg"]), outline=rgb(COLORS["line"]))
    center_text(draw, (112, 228, 310, 264), "Todos los Canales", f_body, rgb(COLORS["ink"]))
    draw_round(draw, (328, 228, 590, 264), 9, rgb(COLORS["bg"]), outline=rgb(COLORS["line"]))
    center_text(draw, (328, 228, 590, 264), "01/04/2026 - 31/05/2026", f_body, rgb(COLORS["ink"]))
    draw_round(draw, (width - 258, 228, width - 128, 264), 9, rgb(COLORS["bg"]), outline=rgb(COLORS["line"]))
    center_text(draw, (width - 258, 228, width - 128, 264), "Actualizar", f_body, rgb(COLORS["blue"]))

    draw_round(draw, (92, 328, width - 92, 402), 12, rgb("ffffff"), outline=rgb(COLORS["line"]))
    draw.text((130, 350), title, font=f_h1, fill=rgb(COLORS["ink"]))
    draw.text((130, 383), subtitle, font=f_small, fill=rgb(COLORS["slate"]))

    card_y = 438
    gap = 22
    card_width = (width - 184 - gap * (len(kpis) - 1)) // len(kpis)
    for i, item in enumerate(kpis):
        x = 92 + i * (card_width + gap)
        draw_round(
            draw,
            (x, card_y, x + card_width, card_y + 118),
            14,
            rgb(item.get("fill", "ffffff")),
            outline=rgb(item.get("outline", COLORS["line"])),
        )
        draw.text((x + 22, card_y + 22), item["label"], font=f_body, fill=rgb(COLORS["slate"]))
        draw.text((x + 22, card_y + 52), item["value"], font=f_value, fill=rgb(item.get("color", COLORS["ink"])))
        draw.text((x + 22, card_y + 92), truncate(item.get("note", ""), f_small, card_width - 44), font=f_small, fill=rgb(COLORS["slate"]))

    panel_y = 590
    panel_h = 238
    panel_gap = 28
    panel_w = (width - 184 - panel_gap) // 2
    for idx, panel in enumerate(panels[:2]):
        x = 92 + idx * (panel_w + panel_gap)
        draw_round(draw, (x, panel_y, x + panel_w, panel_y + panel_h), 14, rgb("ffffff"), outline=rgb(COLORS["line"]))
        draw.text((x + 24, panel_y + 22), panel["title"], font=f_h2, fill=rgb(COLORS["ink"]))
        draw.text((x + 24, panel_y + 50), panel["subtitle"], font=f_small, fill=rgb(COLORS["slate"]))
        if panel.get("type") == "bars":
            max_value = max([v for _, v, _ in panel["rows"]] or [1])
            y = panel_y + 86
            for label, value, color in panel["rows"]:
                draw.text((x + 24, y), label, font=f_body, fill=rgb(COLORS["ink"]))
                bar_x = x + 210
                base_w = panel_w - 270
                bar_w = int(base_w * (value / max_value if max_value else 0))
                draw_round(draw, (bar_x, y - 2, bar_x + base_w, y + 18), 9, rgb("edf2f7"))
                draw_round(draw, (bar_x, y - 2, bar_x + bar_w, y + 18), 9, rgb(color))
                draw.text((bar_x + base_w + 25, y - 4), str(value), font=f_metric, fill=rgb(COLORS["ink"]))
                y += 42
        elif panel.get("type") == "table":
            headers = panel["headers"]
            rows = panel["rows"]
            col_w = (panel_w - 48) // len(headers)
            y = panel_y + 82
            draw.rectangle((x + 24, y, x + panel_w - 24, y + 32), fill=rgb("f1f5f9"))
            for j, header in enumerate(headers):
                draw.text((x + 30 + j * col_w, y + 9), truncate(header, f_small, col_w - 8), font=f_small, fill=rgb(COLORS["slate"]))
            y += 32
            for row in rows[:4]:
                draw.line((x + 24, y, x + panel_w - 24, y), fill=rgb(COLORS["line"]))
                for j, cell in enumerate(row):
                    draw.text((x + 30 + j * col_w, y + 10), truncate(str(cell), f_small, col_w - 8), font=f_small, fill=rgb(COLORS["ink"]))
                y += 32

    draw.text(
        (92, height - 36),
        "Vista referencial para capacitación de usuario. Los datos personales se muestran de forma protegida.",
        font=f_small,
        fill=rgb(COLORS["slate"]),
    )
    path = ASSET_DIR / filename
    image.save(path, quality=95)
    return path


VISUAL_SPECS = [
    {
        "tab_label": "Estrategia",
        "title": "Resumen de Negocio",
        "subtitle": "Vista ejecutiva del periodo seleccionado",
        "filename": "01_estrategia.png",
        "kpis": [
            {"label": "Total Leads", "value": "11", "note": "Prospectos entrantes", "fill": "eaf1ff", "color": "1f3f91"},
            {"label": "Leads calificados", "value": "5", "note": "Con señales comerciales", "fill": "eaf1ff", "color": "1f3f91"},
            {"label": "Citas Agendadas", "value": "2", "note": "Leads marcados como cita", "fill": "dff7ec", "color": "0a9b6f"},
            {"label": "Cierre / Ventas", "value": "0", "note": "Ventas confirmadas", "fill": "dff7ec", "color": "0a9b6f"},
        ],
        "panels": [
            {"type": "bars", "title": "Conversión principal", "subtitle": "Relación entre leads, citas y ventas", "rows": [("Leads", 11, "274690"), ("Calificados", 5, "1f3f91"), ("Citas", 2, "0a9b6f"), ("Ventas", 0, "94a3b8")]},
            {"type": "table", "title": "Lectura ejecutiva", "subtitle": "Qué revisar primero", "headers": ["Dato", "Relevancia"], "rows": [["Conversión", "Mide avance comercial"], ["Ganancia", "Resume monto generado"], ["Citas", "Indica oportunidad próxima"], ["Ventas", "Confirma cierre"]]},
        ],
    },
    {
        "tab_label": "Embudo",
        "title": "Embudo Comercial",
        "subtitle": "Avance de leads por etapas",
        "filename": "02_embudo.png",
        "kpis": [
            {"label": "Entrantes", "value": "11", "note": "Base del embudo", "fill": "eaf1ff", "color": "1f3f91"},
            {"label": "Calificados", "value": "5", "note": "Con interés real", "fill": "ffffff", "color": "1f3f91"},
            {"label": "Citas", "value": "2", "note": "Siguiente paso agendado", "fill": "dff7ec", "color": "0a9b6f"},
            {"label": "Pérdidas", "value": "1", "note": "No aplican o descartan", "fill": "fff7ed", "color": "f59e0b"},
        ],
        "panels": [
            {"type": "bars", "title": "Etapas actuales", "subtitle": "Dónde están los leads ahora", "rows": [("Crear confianza", 4, "274690"), ("Crear urgencia", 2, "1f3f91"), ("Cita agendada", 2, "0a9b6f"), ("No calificado", 1, "f59e0b")]},
            {"type": "table", "title": "Cómo leerlo", "subtitle": "Guía rápida", "headers": ["Pregunta", "Respuesta"], "rows": [["¿Dónde se atascan?", "Etapa con más volumen"], ["¿Qué mejora?", "Mayor pérdida"], ["¿Qué sigue?", "Accionar pendientes"]]},
        ],
    },
    {
        "tab_label": "Operación",
        "title": "Eficiencia Operativa",
        "subtitle": "Respuesta y carga del equipo",
        "filename": "03_operacion.png",
        "kpis": [
            {"label": "Tiempo promedio", "value": "2 min 17 s", "note": "Velocidad de respuesta", "fill": "fff7ed", "color": "1f3f91"},
            {"label": "Leads sin respuesta", "value": "4", "note": "Último mensaje fue del cliente", "fill": "fee2e2", "color": "ef4444"},
            {"label": "Con responsable", "value": "100%", "note": "Asignación cubierta", "fill": "ffffff", "color": "1f3f91"},
        ],
        "panels": [
            {"type": "table", "title": "Desempeño por responsable", "subtitle": "Carga y pendientes", "headers": ["Responsable", "Leads", "Sin resp.", "Citas"], "rows": [["Cristian Gomezcoello", 7, 4, 1], ["Luis", 2, 0, 0]]},
            {"type": "bars", "title": "Carga por canal", "subtitle": "Volumen operativo por origen", "rows": [("Sitio web", 6, "274690"), ("TikTok", 2, "0a9b6f"), ("Telegram", 1, "f59e0b")]},
        ],
    },
    {
        "tab_label": "Seguimiento",
        "title": "Seguimiento Comercial",
        "subtitle": "Acciones pendientes y oportunidades",
        "filename": "04_seguimiento.png",
        "kpis": [
            {"label": "Pendientes", "value": "6", "note": "Requieren seguimiento", "fill": "eaf1ff", "color": "1f3f91"},
            {"label": "Citas próximas", "value": "2", "note": "Atender a tiempo", "fill": "dff7ec", "color": "0a9b6f"},
            {"label": "Ventas exitosas", "value": "1", "note": "Resultado confirmado", "fill": "dff7ec", "color": "0a9b6f"},
        ],
        "panels": [
            {"type": "table", "title": "Cola de atención", "subtitle": "Leads priorizados para accionar", "headers": ["Lead", "Canal", "Estado", "Acción"], "rows": [["Luis Gomez", "Sitio web", "Cita agendada", "Confirmar"], ["Maria Gomez", "Sitio web", "Crear confianza", "Responder"], ["Lead TikTok", "TikTok", "Seguimiento", "Retomar"]]},
            {"type": "bars", "title": "Prioridad del día", "subtitle": "Qué conviene revisar primero", "rows": [("Responder", 4, "ef4444"), ("Confirmar cita", 2, "0a9b6f"), ("Revisar venta", 1, "274690")]},
        ],
    },
    {
        "tab_label": "Rendimiento Humano",
        "title": "Rendimiento Humano",
        "subtitle": "Resultados por responsable",
        "filename": "05_rendimiento_humano.png",
        "kpis": [
            {"label": "Seguimiento", "value": "1", "note": "Leads desde seguimiento humano", "fill": "ffffff", "color": "1f3f91"},
            {"label": "Citas humanas", "value": "2", "note": "Citas generadas por gestión", "fill": "dff7ec", "color": "0a9b6f"},
            {"label": "Conversión", "value": "67%", "note": "Eficiencia humana", "fill": "eaf1ff", "color": "1f3f91"},
            {"label": "Total vendido", "value": "$5.00", "note": "Ventas del periodo", "fill": "dff7ec", "color": "0a9b6f"},
        ],
        "panels": [
            {"type": "bars", "title": "Citas agendadas por humano", "subtitle": "Comparación entre estados", "rows": [("Seguimiento humano", 1, "274690"), ("Cita agendada", 2, "0a9b6f")]},
            {"type": "table", "title": "Responsables", "subtitle": "Rendimiento del equipo", "headers": ["Responsable", "Leads", "Citas", "Conv."], "rows": [["Cristian Gomezcoello", 7, 1, "14%"], ["Luis", 2, 0, "0%"]]},
        ],
    },
    {
        "tab_label": "Tendencias",
        "title": "Tendencias Comerciales",
        "subtitle": "Canales, campañas y comportamiento",
        "filename": "06_tendencias.png",
        "kpis": [
            {"label": "Canales activos", "value": "3", "note": "Fuentes con leads", "fill": "eaf1ff", "color": "1f3f91"},
            {"label": "Campañas", "value": "3", "note": "Origen registrado", "fill": "ffffff", "color": "1f3f91"},
            {"label": "Ingresos", "value": "$5.00", "note": "Según ventas registradas", "fill": "dff7ec", "color": "0a9b6f"},
        ],
        "panels": [
            {"type": "bars", "title": "Origen de leads", "subtitle": "Canales que generan más volumen", "rows": [("Sitio web", 7, "274690"), ("TikTok", 2, "0a9b6f"), ("Telegram", 1, "f59e0b")]},
            {"type": "bars", "title": "Campañas", "subtitle": "Distribución por campaña", "rows": [("Sin campaña", 5, "7c3aed"), ("facebook", 1, "7c3aed"), ("ig", 1, "7c3aed")]},
        ],
    },
    {
        "tab_label": "Calidad",
        "title": "Calidad de Leads",
        "subtitle": "Puntaje y prioridad comercial",
        "filename": "07_calidad.png",
        "kpis": [
            {"label": "Caliente", "value": "2", "note": "Alta intención", "fill": "fee2e2", "color": "ef4444"},
            {"label": "Tibio", "value": "3", "note": "Señales accionables", "fill": "fff7ed", "color": "f59e0b"},
            {"label": "Frío", "value": "4", "note": "Baja señal o sin puntaje", "fill": "eaf1ff", "color": "274690"},
        ],
        "panels": [
            {"type": "table", "title": "Leads evaluados", "subtitle": "Prioridad para seguimiento", "headers": ["Lead", "Nivel", "Puntaje", "Canal"], "rows": [["Luis Gomez", "Caliente", 80, "Sitio web"], ["Maria Gomez", "Tibio", 65, "Sitio web"], ["Lead TikTok", "Frío", "Sin puntaje", "TikTok"]]},
            {"type": "bars", "title": "Distribución de calidad", "subtitle": "Qué tan listos están los leads", "rows": [("Caliente", 2, "ef4444"), ("Tibio", 3, "f59e0b"), ("Frío", 4, "274690")]},
        ],
    },
    {
        "tab_label": "Conversaciones",
        "title": "Conversaciones",
        "subtitle": "Detalle de interacciones con leads",
        "filename": "08_conversaciones.png",
        "kpis": [
            {"label": "Conversaciones", "value": "11", "note": "Dentro del periodo", "fill": "eaf1ff", "color": "1f3f91"},
            {"label": "Canales", "value": "3", "note": "Origen de atención", "fill": "ffffff", "color": "1f3f91"},
            {"label": "Mensajes cargados", "value": "48", "note": "Historial disponible", "fill": "dff7ec", "color": "0a9b6f"},
        ],
        "panels": [
            {"type": "table", "title": "Lista de conversaciones", "subtitle": "Mensajes y datos principales", "headers": ["Lead", "Canal", "Estado", "Última interacción"], "rows": [["Luis Gomez", "Sitio web", "Cita agendada", "30/04/2026"], ["Lead protegido", "TikTok", "Crear confianza", "30/04/2026"], ["Maria Gomez", "Sitio web", "Seguimiento", "29/04/2026"]]},
            {"type": "table", "title": "Detalle disponible", "subtitle": "Qué puede revisar el usuario", "headers": ["Dato", "Uso"], "rows": [["Historial", "Ver contexto"], ["URL", "Abrir origen"], ["Estado", "Entender avance"], ["Fecha", "Ordenar prioridad"]]},
        ],
    },
    {
        "tab_label": "Reportes",
        "title": "Reportes",
        "subtitle": "Descargas y envíos programados",
        "filename": "09_reportes.png",
        "kpis": [
            {"label": "Formatos", "value": "3", "note": "Excel, PDF y CSV", "fill": "eaf1ff", "color": "1f3f91"},
            {"label": "Perfiles críticos", "value": "4", "note": "Gerencia, operación, equipo y marketing", "fill": "ffffff", "color": "1f3f91"},
            {"label": "Programación", "value": "Activa", "note": "Envíos por frecuencia definida", "fill": "dff7ec", "color": "0a9b6f"},
        ],
        "panels": [
            {"type": "table", "title": "Qué exportar", "subtitle": "Uso recomendado por formato", "headers": ["Formato", "Mejor para"], "rows": [["PDF", "Lectura ejecutiva"], ["Excel", "Detalle y filtros"], ["CSV", "Análisis externo"]]},
            {"type": "table", "title": "Perfiles frecuentes", "subtitle": "Reportes listos para compartir", "headers": ["Perfil", "Audiencia"], "rows": [["Gerencial", "Dirección"], ["Operación diaria", "Equipo comercial"], ["Rendimiento", "Líder de equipo"], ["Marketing/Calidad", "Marketing"]]},
        ],
    },
]


REAL_VISUAL_PREFIXES = {
    "Estrategia": "01_estrategia",
    "Embudo": "02_embudo",
    "Operación": "03_operacion",
    "Seguimiento": "04_seguimiento",
    "Rendimiento Humano": "05_rendimiento_humano",
    "Tendencias": "06_tendencias",
    "Calidad": "07_calidad",
    "Conversaciones": "08_conversaciones",
    "Reportes": "09_reportes",
}


TAB_CONTENT = [
    {
        "name": "Estrategia",
        "purpose": "Entrega una visión rápida del estado general del negocio en el periodo seleccionado.",
        "questions": ["¿Cuántos leads llegaron?", "¿Cuántos avanzaron comercialmente?", "¿Cuántas citas y ventas se lograron?", "¿Cómo va la conversión y el ingreso?"],
        "data": ["Total Leads", "Leads calificados", "Citas Agendadas", "Cierre / Ventas", "Conversión a cita", "Ganancia Mensual", "Ganancia Total"],
        "interpretation": "Si los leads suben pero las citas no crecen, conviene revisar la calidad de atención o el embudo. Si las citas suben pero las ventas no, el foco debe ir al cierre comercial.",
        "decisions": ["Medir salud general del negocio.", "Comparar resultados por periodo.", "Detectar si el problema está en volumen, conversión o ventas."],
        "example": "Un gerente revisa Estrategia cada lunes para decidir si necesita más leads, mejorar seguimiento o revisar el cierre.",
    },
    {
        "name": "Embudo",
        "purpose": "Muestra en qué etapa se encuentran los leads y dónde se pierden oportunidades.",
        "questions": ["¿Dónde se concentra el mayor volumen?", "¿En qué paso se detiene la mayor cantidad de leads?", "¿Qué etapa necesita atención del equipo?"],
        "data": ["Etapas actuales del embudo", "Conversiones entre etapas", "Leads perdidos o no calificados", "Detalle de leads por etapa"],
        "interpretation": "Las etapas con mucho volumen pueden indicar trabajo pendiente. Las caídas fuertes entre etapas muestran oportunidades de mejora.",
        "decisions": ["Ajustar mensajes o procesos en una etapa específica.", "Priorizar leads atascados.", "Entender qué parte del proceso comercial requiere apoyo."],
        "example": "Si muchos leads quedan en Crear confianza, el equipo puede preparar mensajes de seguimiento más claros.",
    },
    {
        "name": "Operación",
        "purpose": "Ayuda a revisar qué tan rápido y ordenado está respondiendo el equipo.",
        "questions": ["¿Cuánto demora la primera respuesta?", "¿Cuántos leads siguen sin respuesta?", "¿Todos los leads tienen responsable?", "¿Cómo está distribuida la carga?"],
        "data": ["Tiempo promedio de respuesta", "Leads sin respuesta", "Leads con responsable", "Desempeño por responsable", "Carga por canal"],
        "interpretation": "Un lead sin respuesta es aquel cuyo último mensaje llegó del cliente y no tuvo respuesta posterior. El objetivo es reducir estos casos porque representan conversaciones abiertas.",
        "decisions": ["Reasignar carga si un responsable tiene demasiados pendientes.", "Responder leads detenidos.", "Mejorar tiempos de atención."],
        "example": "Si Operación muestra varios leads sin respuesta, el equipo puede entrar a Conversaciones y atenderlos antes de que se enfríen.",
    },
    {
        "name": "Seguimiento",
        "purpose": "Organiza los leads que necesitan acción humana para continuar el proceso comercial.",
        "questions": ["¿A quién debo contactar hoy?", "¿Qué leads tienen cita o venta registrada?", "¿Qué oportunidades están pendientes?", "¿Qué acción conviene hacer primero?"],
        "data": ["Leads pendientes", "Citas agendadas", "Ventas exitosas", "Monto registrado", "Último mensaje", "Enlace de conversación"],
        "interpretation": "Esta pestaña funciona como una lista de trabajo. Debe revisarse con frecuencia para evitar que oportunidades con intención se queden sin seguimiento.",
        "decisions": ["Definir prioridades diarias.", "Confirmar citas.", "Retomar conversaciones importantes.", "Revisar leads cercanos a venta."],
        "example": "Un asesor abre Seguimiento al iniciar el día y atiende primero los leads con cita o alta intención.",
    },
    {
        "name": "Rendimiento Humano",
        "purpose": "Permite ver los resultados generados por la gestión del equipo humano.",
        "questions": ["¿Quién atendió más leads?", "¿Cuántas citas se generaron por gestión humana?", "¿Qué conversión logró el equipo?", "¿Cuánto se vendió?"],
        "data": ["Seguimiento", "Citas humanas", "Conversión", "Ventas", "Total vendido", "Ticket promedio", "Ranking por responsable"],
        "interpretation": "La conversión indica qué tan efectivo fue el paso de seguimiento a cita o resultado comercial. Debe leerse junto con la cantidad de leads atendidos.",
        "decisions": ["Comparar desempeño del equipo.", "Detectar necesidades de capacitación.", "Reconocer buenas prácticas.", "Balancear la carga entre responsables."],
        "example": "Si un responsable tiene muchos leads pero baja conversión, puede necesitar apoyo en cierre o mensajes de seguimiento.",
    },
    {
        "name": "Tendencias",
        "purpose": "Muestra cómo evolucionan los resultados por canal, campaña y periodo.",
        "questions": ["¿Qué canal genera más leads?", "¿Qué campaña aporta más oportunidades?", "¿Cómo cambian ingresos y resultados en el tiempo?", "¿Por qué se descalifican algunos leads?"],
        "data": ["Origen de leads", "Campañas", "Ingresos por fecha", "Motivos de descalificación", "Calidad por origen"],
        "interpretation": "Sirve para comparar fuentes. No solo importa el volumen: también conviene revisar si los canales generan citas, ventas o leads de mejor calidad.",
        "decisions": ["Invertir más en canales que traen mejores resultados.", "Ajustar campañas con bajo rendimiento.", "Entender patrones de descalificación."],
        "example": "Si TikTok trae muchos leads pero pocas citas, marketing puede revisar segmentación o mensaje de campaña.",
    },
    {
        "name": "Calidad",
        "purpose": "Clasifica leads por prioridad para saber cuáles atender primero.",
        "questions": ["¿Qué leads son Calientes, Tibios o Fríos?", "¿Qué puntaje tiene cada lead?", "¿Qué canal o campaña trae mejor calidad?", "¿Qué leads no tienen puntaje?"],
        "data": ["Nivel de calidad: Caliente, Tibio o Frío", "Puntaje", "Canal", "Número", "Estado", "Historial de mensajes", "Campaña", "Fecha de ingreso", "Última interacción"],
        "interpretation": "Caliente indica alta intención; Tibio indica señales accionables; Frío reúne leads con señal inicial, puntaje menor a Tibio o sin puntaje.",
        "decisions": ["Priorizar atención a leads calientes.", "Revisar calidad por campaña.", "Separar leads que necesitan nutrición de los que ya están listos para avanzar."],
        "example": "Un asesor puede empezar por Caliente y Tibio para aprovechar las oportunidades con mayor intención.",
    },
    {
        "name": "Conversaciones",
        "purpose": "Permite revisar el detalle de cada conversación y el contexto del lead.",
        "questions": ["¿Qué se conversó con el lead?", "¿Cuál fue el último mensaje?", "¿Qué estado tiene?", "¿Dónde puedo abrir la conversación original?"],
        "data": ["Nombre del lead", "Canal", "Número", "Estado", "Historial de mensajes", "URL", "Fecha de ingreso", "Última interacción"],
        "interpretation": "Es la pestaña de revisión puntual. Ayuda a entender el contexto antes de responder, revisar la atención o validar por qué un lead está en cierta etapa.",
        "decisions": ["Abrir conversaciones específicas.", "Revisar calidad de atención.", "Confirmar mensajes y estados.", "Dar seguimiento con contexto completo."],
        "example": "Antes de llamar a un lead, el usuario revisa Conversaciones para leer el último mensaje y evitar repetir preguntas.",
    },
    {
        "name": "Reportes",
        "purpose": "Centraliza exportaciones y reportes programados para compartir la información.",
        "questions": ["¿Qué reporte debo descargar?", "¿Qué formato conviene usar?", "¿Qué perfiles están configurados?", "¿A quién se envía la información?"],
        "data": ["Reportes por pestaña", "Formatos Excel, PDF y CSV", "Perfiles críticos", "Frecuencia de envío", "Destinatarios configurados"],
        "interpretation": "Excel sirve para detalle y filtros; PDF para lectura ejecutiva; CSV para análisis externo o carga en otras herramientas.",
        "decisions": ["Enviar información a gerencia.", "Guardar evidencia del periodo.", "Compartir tareas con el equipo.", "Analizar datos fuera del dashboard."],
        "example": "Gerencia puede recibir un PDF semanal; el equipo operativo puede descargar Excel para revisar detalle de leads.",
    },
]


def set_cell_shading(cell, fill: str):
    props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def set_cell_text(cell, text, bold=False, color="0f2344"):
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.bold = bold
    run.font.size = Pt(9)
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(*rgb(color))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="d9e2ef"):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement("w:" + edge)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    table._tbl.tblPr.append(borders)


def add_doc_table(doc: Document, rows, widths=None, header=True):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            cell = table.cell(i, j)
            is_head = header and i == 0
            set_cell_text(cell, value, bold=is_head, color=("ffffff" if is_head else "0f2344"))
            set_cell_shading(cell, COLORS["blue"] if is_head else ("f8fafc" if i % 2 == 0 else "ffffff"))
            if widths and j < len(widths):
                cell.width = Inches(widths[j])
    return table


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(*rgb(COLORS["blue"] if level == 1 else COLORS["ink"]))
    return p


def add_para(doc: Document, text: str):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(*rgb(COLORS["ink"]))
    return p


def add_bullets(doc: Document, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(*rgb(COLORS["ink"]))


def add_callout(doc: Document, title: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, "b8c7e8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS["light_blue"])
    p = cell.paragraphs[0]
    run = p.add_run(title + ": ")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*rgb(COLORS["blue"]))
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(*rgb(COLORS["ink"]))
    doc.add_paragraph()


def generate_word(visual_paths: dict[str, list[Path]]):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"{DOC_CODE} | v{DOC_VERSION} | {DOC_STATUS}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(*rgb(COLORS["slate"]))

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("SimpliaLeads - Manual de Usuario | Información documentada controlada")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(*rgb(COLORS["slate"]))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.6))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Manual de Usuario SimpliaLeads")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(*rgb(COLORS["blue"]))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Guía simple para entender las pestañas, métricas y reportes del dashboard")
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(*rgb(COLORS["slate"]))
    doc.add_paragraph()

    add_doc_table(
        doc,
        [
            ["Campo", "Valor"],
            ["Código", DOC_CODE],
            ["Versión", DOC_VERSION],
            ["Estado", DOC_STATUS],
            ["Fecha", TODAY],
            ["Zona horaria", TZ],
            ["Dirigido a", "Clientes, gerencia, usuarios comerciales y usuarios operativos"],
            ["Uso principal", "Comprender qué muestra cada pestaña y cómo tomar decisiones con los datos"],
        ],
        widths=[1.8, 4.8],
    )
    add_callout(doc, "Enfoque del documento", "Este manual está escrito para usuarios del negocio. No incluye claves, configuraciones técnicas ni detalles internos de implementación. Las capturas corresponden a un entorno de prueba con datos ficticios.")
    doc.add_page_break()

    add_heading(doc, "Control documental", 1)
    add_para(doc, "Este manual se controla como información documentada del sistema SimpliaLeads. Su objetivo es que el usuario entienda la información que ve en el dashboard, cómo interpretarla y cómo usarla para tomar decisiones comerciales.")
    add_doc_table(
        doc,
        [
            ["Elemento", "Detalle"],
            ["Dueño del documento", "Simplia"],
            ["Aprobador", "Gerencia / Responsable del cliente"],
            ["Distribución", "Usuarios autorizados del dashboard y responsables comerciales"],
            ["Próxima revisión", NEXT_REVIEW],
            ["Criterio de actualización", "Cambios importantes en pestañas, métricas, nombres de reportes o flujo comercial"],
        ],
        widths=[2.1, 4.9],
    )
    doc.add_paragraph()
    add_doc_table(
        doc,
        [
            ["Versión", "Fecha", "Descripción", "Responsable"],
            ["1.0", "02/05/2026", "Emisión inicial del manual de usuario", "Simplia"],
            ["1.1", TODAY, "Actualización con capturas reales del dashboard y Calidad en tres niveles", "Simplia"],
        ],
        widths=[0.9, 1.2, 3.6, 1.3],
    )

    add_heading(doc, "1. Objetivo y alcance", 1)
    add_para(doc, "El objetivo de este manual es explicar, de forma clara y práctica, cómo leer SimpliaLeads. El documento ayuda a entender qué representa cada pestaña, qué datos puede ver el usuario, por qué esos datos son relevantes y qué acciones puede tomar.")
    add_bullets(doc, [
        "Incluye la explicación de filtros, fechas, canales, actualización de datos, exportaciones y lectura de pestañas.",
        "Incluye ejemplos de uso para gerencia, operación diaria, equipo comercial y marketing.",
        "No incluye configuración técnica, administración de base de datos, claves, integraciones internas ni mantenimiento del sistema.",
    ])

    add_heading(doc, "2. ¿Para qué sirve SimpliaLeads?", 1)
    add_para(doc, "SimpliaLeads es un panel de control comercial que reúne conversaciones y datos de leads para mostrar el avance del negocio. Permite ver si llegan suficientes leads, si se están respondiendo a tiempo, si se convierten en citas o ventas, qué canales funcionan mejor y qué oportunidades requieren atención.")
    add_callout(doc, "Idea clave", "El dashboard no es solo una tabla. Es una herramienta para decidir qué revisar, a quién responder, qué canal mejorar y cómo medir el resultado comercial.")

    add_heading(doc, "3. Conceptos básicos", 1)
    add_doc_table(
        doc,
        [
            ["Concepto", "Qué significa para el usuario"],
            ["Lead", "Persona o prospecto que inició o mantiene una conversación comercial."],
            ["Canal", "Origen de la conversación, por ejemplo sitio web, TikTok, Telegram u otro inbox."],
            ["Campaña", "Nombre de campaña registrada para entender de dónde viene el lead."],
            ["Estado o etapa", "Momento comercial en el que se encuentra el lead dentro del proceso."],
            ["Cita", "Lead marcado como cita agendada o equivalente según la configuración del negocio."],
            ["Venta", "Lead marcado como venta exitosa con monto registrado cuando aplique."],
            ["Puntaje", "Valor usado para priorizar la calidad o intención del lead."],
            ["Última interacción", "Fecha y hora del último movimiento relevante de la conversación."],
        ],
        widths=[1.7, 5.1],
    )

    add_heading(doc, "4. Guía rápida de uso", 1)
    add_bullets(doc, [
        "Selecciona el periodo que quieres analizar con el filtro de fechas.",
        "Elige un canal específico si quieres revisar solo una fuente; deja Todos los Canales para ver el total.",
        "Usa el botón de actualizar cuando quieras refrescar la información disponible.",
        "Lee primero los KPIs principales y luego baja a tablas o gráficos para entender el detalle.",
        "Exporta el reporte cuando necesites compartir, analizar o guardar evidencia del periodo.",
    ])
    add_doc_table(
        doc,
        [
            ["Elemento", "Cómo interpretarlo"],
            ["Datos en vivo + historial", "Indica que el dashboard combina información reciente con datos guardados de periodos anteriores."],
            ["Todos los Canales", "Muestra el resultado global de todas las fuentes disponibles."],
            ["Exportar", "Permite descargar información en formatos preparados para lectura o análisis."],
            ["Sin campaña", "El lead no tiene una campaña registrada o el dato no fue recibido."],
            ["Sin puntaje", "El lead no tiene score registrado; para calidad se considera dentro de Frío."],
        ],
        widths=[1.8, 5.0],
    )

    add_heading(doc, "5. Orden recomendado de lectura", 1)
    add_para(doc, "Para entender el negocio con rapidez, se recomienda revisar las pestañas en este orden:")
    add_doc_table(
        doc,
        [
            ["Orden", "Pestaña", "Qué responde"],
            ["1", "Estrategia", "Cómo va el negocio en general."],
            ["2", "Embudo", "Dónde avanzan o se detienen los leads."],
            ["3", "Operación", "Si el equipo está respondiendo y asignando bien."],
            ["4", "Seguimiento", "Qué leads necesitan acción."],
            ["5", "Rendimiento Humano", "Cómo rinde el equipo o responsable."],
            ["6", "Tendencias", "Qué canales, campañas y periodos funcionan mejor."],
            ["7", "Calidad", "Qué leads conviene priorizar."],
            ["8", "Conversaciones", "Qué ocurrió en cada conversación."],
            ["9", "Reportes", "Qué descargar o enviar."],
        ],
        widths=[0.7, 2.0, 4.2],
    )

    doc.add_page_break()
    add_heading(doc, "6. Pestañas del dashboard", 1)
    for index, tab in enumerate(TAB_CONTENT, start=1):
        add_heading(doc, f"6.{index}. {tab['name']}", 2)
        tab_visuals = visual_paths[tab["name"]]
        for image_index, visual_path in enumerate(tab_visuals, start=1):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(str(visual_path), width=Inches(6.75))
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            suffix = f" ({image_index} de {len(tab_visuals)})" if len(tab_visuals) > 1 else ""
            run = caption.add_run(f"Captura real de la pestaña {tab['name']}{suffix}. Rango usado: 01/04/2026 - 04/05/2026. Los datos mostrados son ficticios.")
            run.italic = True
            run.font.size = Pt(8.5)
            run.font.color.rgb = RGBColor(*rgb(COLORS["slate"]))
        add_doc_table(
            doc,
            [
                ["Elemento", "Descripción"],
                ["Para qué sirve", tab["purpose"]],
                ["Cómo interpretarlo", tab["interpretation"]],
                ["Ejemplo práctico", tab["example"]],
            ],
            widths=[1.5, 5.4],
        )
        add_para(doc, "Preguntas que responde:")
        add_bullets(doc, tab["questions"])
        add_para(doc, "Datos principales que puede ver:")
        add_bullets(doc, tab["data"])
        add_para(doc, "Decisiones que ayuda a tomar:")
        add_bullets(doc, tab["decisions"])
        if index < len(TAB_CONTENT):
            doc.add_page_break()

    add_heading(doc, "7. Cómo usar los reportes", 1)
    add_para(doc, "Los reportes permiten llevar la información del dashboard fuera de la pantalla. Deben generarse después de seleccionar el periodo y los filtros correctos, porque el contenido exportado representa lo que el usuario está analizando.")
    add_doc_table(
        doc,
        [
            ["Formato", "Cuándo usarlo", "Qué esperar"],
            ["PDF", "Para gerencia, lectura ejecutiva o reuniones", "Resumen claro, KPIs, interpretación y tablas principales."],
            ["Excel", "Para revisar detalle, filtrar y trabajar datos", "Hojas con métricas, análisis y detalle de leads cuando aplique."],
            ["CSV", "Para cargar datos en otra herramienta", "Datos planos y limpios para análisis externo."],
        ],
        widths=[1.1, 2.8, 3.0],
    )
    add_callout(doc, "Recomendación", "Antes de compartir un reporte, confirma que el periodo, canal y pestaña son los correctos. Así evitas comparar datos de filtros diferentes.")

    add_heading(doc, "8. Buenas prácticas por tipo de usuario", 1)
    add_doc_table(
        doc,
        [
            ["Usuario", "Qué revisar", "Frecuencia sugerida"],
            ["Gerencia", "Estrategia, Embudo, Rendimiento Humano y Reportes", "Semanal o mensual"],
            ["Equipo operativo", "Operación, Seguimiento y Conversaciones", "Diaria"],
            ["Responsable comercial", "Seguimiento, Rendimiento Humano y Conversaciones", "Diaria y semanal"],
            ["Marketing", "Tendencias, Calidad y Estrategia", "Semanal o por campaña"],
        ],
        widths=[1.7, 3.4, 1.8],
    )

    add_heading(doc, "9. Preguntas frecuentes", 1)
    faq = [
        ("¿Por qué veo cero datos?", "Puede que el periodo no tenga conversaciones, que el canal seleccionado no tenga registros o que el filtro esté demasiado cerrado."),
        ("¿Por qué un lead aparece sin campaña?", "Porque no se recibió o no se registró el dato de campaña para ese lead."),
        ("¿Qué significa sin puntaje?", "Ese lead todavía no tiene score registrado. En Calidad se agrupa dentro de Frío para que no quede fuera del análisis."),
        ("¿Por qué revisar Conversaciones si ya veo KPIs?", "Porque los KPIs muestran el resumen, pero Conversaciones permite entender el contexto real de cada lead."),
        ("¿Qué hago si un valor no parece correcto?", "Revisa primero periodo y canal. Luego abre el detalle del lead o conversación para confirmar el contexto."),
        ("¿Qué formato de reporte conviene enviar?", "PDF para lectura rápida, Excel para detalle y CSV para análisis en otra herramienta."),
    ]
    for question, answer in faq:
        p = doc.add_paragraph()
        run = p.add_run(question)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(*rgb(COLORS["ink"]))
        add_para(doc, answer)

    add_heading(doc, "10. Checklist final de uso", 1)
    add_bullets(doc, [
        "Definir el periodo correcto antes de analizar.",
        "Confirmar si se revisan todos los canales o uno específico.",
        "Leer Estrategia para entender el panorama general.",
        "Revisar Embudo para detectar dónde se detienen los leads.",
        "Usar Operación y Seguimiento para atender pendientes.",
        "Usar Rendimiento Humano para evaluar responsables y conversión.",
        "Usar Tendencias y Calidad para decisiones de marketing y priorización.",
        "Abrir Conversaciones cuando se necesite contexto real del lead.",
        "Exportar el reporte correcto según la audiencia.",
    ])

    add_heading(doc, "11. Aprobación", 1)
    add_doc_table(
        doc,
        [["Rol", "Nombre", "Firma", "Fecha"], ["Elaborado por", "Simplia", "", TODAY], ["Revisado por", "", "", ""], ["Aprobado por", "", "", ""]],
        widths=[1.5, 2.2, 1.8, 1.2],
    )
    doc.core_properties.title = "Manual de Usuario SimpliaLeads ISO 10013"
    doc.core_properties.subject = "Manual de usuario para clientes del dashboard SimpliaLeads"
    doc.core_properties.author = "Simplia"
    doc.core_properties.keywords = "SimpliaLeads, manual de usuario, ISO 10013, dashboard comercial"
    doc.save(DOCX_PATH)


def pdf_escape(text) -> str:
    return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def generate_pdf(visual_paths: dict[str, list[Path]]):
    try:
        pdfmetrics.registerFont(TTFont("Arial", "C:/Windows/Fonts/arial.ttf"))
        pdfmetrics.registerFont(TTFont("Arial-Bold", "C:/Windows/Fonts/arialbd.ttf"))
        base_font = "Arial"
        bold_font = "Arial-Bold"
    except Exception:
        base_font = "Helvetica"
        bold_font = "Helvetica-Bold"

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="ManualTitle", parent=styles["Title"], fontName=bold_font, fontSize=24, leading=30, textColor=colors.HexColor("#274690"), alignment=TA_CENTER, spaceAfter=10))
    styles.add(ParagraphStyle(name="ManualSubtitle", parent=styles["Normal"], fontName=base_font, fontSize=11, leading=15, textColor=colors.HexColor("#64748b"), alignment=TA_CENTER, spaceAfter=16))
    styles.add(ParagraphStyle(name="H1x", parent=styles["Heading1"], fontName=bold_font, fontSize=16, leading=20, textColor=colors.HexColor("#274690"), spaceBefore=14, spaceAfter=8))
    styles.add(ParagraphStyle(name="H2x", parent=styles["Heading2"], fontName=bold_font, fontSize=13, leading=17, textColor=colors.HexColor("#0f2344"), spaceBefore=10, spaceAfter=6))
    styles.add(ParagraphStyle(name="Bodyx", parent=styles["BodyText"], fontName=base_font, fontSize=9.2, leading=13, textColor=colors.HexColor("#0f2344"), spaceAfter=6))
    styles.add(ParagraphStyle(name="Smallx", parent=styles["BodyText"], fontName=base_font, fontSize=7.5, leading=10, textColor=colors.HexColor("#64748b"), alignment=TA_CENTER))
    styles.add(ParagraphStyle(name="Cellx", parent=styles["BodyText"], fontName=base_font, fontSize=8, leading=10, textColor=colors.HexColor("#0f2344")))
    styles.add(ParagraphStyle(name="CellHeadx", parent=styles["BodyText"], fontName=bold_font, fontSize=8, leading=10, textColor=colors.white))

    def para(text, style="Bodyx"):
        return Paragraph(pdf_escape(text), styles[style])

    def table(rows, widths=None):
        data = [[para(c, "CellHeadx" if i == 0 else "Cellx") for c in row] for i, row in enumerate(rows)]
        item = Table(data, colWidths=[w * inch for w in widths] if widths else None, hAlign="LEFT")
        item.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#274690")),
                    ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#d9e2ef")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                    ("LEFTPADDING", (0, 0), (-1, -1), 5),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ("TOPPADDING", (0, 0), (-1, -1), 4),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                ]
            )
        )
        return item

    def bullets(items):
        return [para("• " + item) for item in items]

    def page_header(canvas, doc_obj):
        canvas.saveState()
        canvas.setFont(base_font, 7.5)
        canvas.setFillColor(colors.HexColor("#64748b"))
        canvas.drawRightString(letter[0] - 0.55 * inch, letter[1] - 0.38 * inch, f"{DOC_CODE} | v{DOC_VERSION} | {DOC_STATUS}")
        canvas.drawCentredString(letter[0] / 2, 0.35 * inch, f"SimpliaLeads - Manual de Usuario | Página {doc_obj.page}")
        canvas.restoreState()

    story = []
    if LOGO_PATH.exists():
        story.append(RLImage(str(LOGO_PATH), width=1.45 * inch, height=0.64 * inch, kind="proportional", hAlign="CENTER"))
        story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph("Manual de Usuario SimpliaLeads", styles["ManualTitle"]))
    story.append(Paragraph("Guía simple para entender las pestañas, métricas y reportes del dashboard", styles["ManualSubtitle"]))
    story.append(
        table(
            [
                ["Campo", "Valor"],
                ["Código", DOC_CODE],
                ["Versión", DOC_VERSION],
                ["Estado", DOC_STATUS],
                ["Fecha", TODAY],
                ["Zona horaria", TZ],
                ["Dirigido a", "Clientes, gerencia, usuarios comerciales y usuarios operativos"],
            ],
            widths=[1.7, 4.7],
        )
    )
    story.append(Spacer(1, 0.12 * inch))
    story.append(para("Este manual está escrito para usuarios del negocio. No incluye claves, configuraciones técnicas ni detalles internos de implementación. Las capturas corresponden a un entorno de prueba con datos ficticios."))
    story.append(PageBreak())

    story.append(Paragraph("Control documental", styles["H1x"]))
    story.append(
        table(
            [
                ["Elemento", "Detalle"],
                ["Dueño del documento", "Simplia"],
                ["Aprobador", "Gerencia / Responsable del cliente"],
                ["Distribución", "Usuarios autorizados del dashboard y responsables comerciales"],
                ["Próxima revisión", NEXT_REVIEW],
                ["Criterio de actualización", "Cambios importantes en pestañas, métricas, nombres de reportes o flujo comercial"],
            ],
            widths=[2.0, 4.2],
        )
    )
    story.append(Spacer(1, 0.1 * inch))
    story.append(
        table(
            [
                ["Versión", "Fecha", "Descripción", "Responsable"],
                ["1.0", "02/05/2026", "Emisión inicial del manual de usuario", "Simplia"],
                ["1.1", TODAY, "Actualización con capturas reales del dashboard y Calidad en tres niveles", "Simplia"],
            ],
            widths=[0.8, 1.1, 3.6, 1.0],
        )
    )

    sections = [
        ("1. Objetivo y alcance", "El objetivo de este manual es explicar, de forma clara y práctica, cómo leer SimpliaLeads. El documento ayuda a entender qué representa cada pestaña, qué datos puede ver el usuario, por qué esos datos son relevantes y qué acciones puede tomar."),
        ("2. ¿Para qué sirve SimpliaLeads?", "SimpliaLeads es un panel de control comercial que reúne conversaciones y datos de leads para mostrar el avance del negocio. Permite ver si llegan suficientes leads, si se responden a tiempo, si se convierten en citas o ventas, qué canales funcionan mejor y qué oportunidades requieren atención."),
    ]
    for title, body in sections:
        story.append(Paragraph(title, styles["H1x"]))
        story.append(para(body))

    story.append(Paragraph("3. Conceptos básicos", styles["H1x"]))
    story.append(
        table(
            [
                ["Concepto", "Qué significa para el usuario"],
                ["Lead", "Persona o prospecto que inició o mantiene una conversación comercial."],
                ["Canal", "Origen de la conversación, por ejemplo sitio web, TikTok, Telegram u otro inbox."],
                ["Campaña", "Nombre de campaña registrada para entender de dónde viene el lead."],
                ["Estado o etapa", "Momento comercial en el que se encuentra el lead dentro del proceso."],
                ["Cita", "Lead marcado como cita agendada o equivalente según la configuración del negocio."],
                ["Venta", "Lead marcado como venta exitosa con monto registrado cuando aplique."],
                ["Puntaje", "Valor usado para priorizar la calidad o intención del lead."],
                ["Última interacción", "Fecha y hora del último movimiento relevante de la conversación."],
            ],
            widths=[1.5, 4.8],
        )
    )

    story.append(Paragraph("4. Guía rápida de uso", styles["H1x"]))
    story.extend(bullets(["Selecciona el periodo que quieres analizar con el filtro de fechas.", "Elige un canal específico si quieres revisar solo una fuente; deja Todos los Canales para ver el total.", "Usa el botón de actualizar para refrescar la información disponible.", "Lee primero los KPIs principales y luego baja a tablas o gráficos para entender el detalle.", "Exporta el reporte cuando necesites compartir, analizar o guardar evidencia del periodo."]))

    story.append(Paragraph("5. Orden recomendado de lectura", styles["H1x"]))
    story.append(
        table(
            [
                ["Orden", "Pestaña", "Qué responde"],
                ["1", "Estrategia", "Cómo va el negocio en general."],
                ["2", "Embudo", "Dónde avanzan o se detienen los leads."],
                ["3", "Operación", "Si el equipo está respondiendo y asignando bien."],
                ["4", "Seguimiento", "Qué leads necesitan acción."],
                ["5", "Rendimiento Humano", "Cómo rinde el equipo o responsable."],
                ["6", "Tendencias", "Qué canales, campañas y periodos funcionan mejor."],
                ["7", "Calidad", "Qué leads conviene priorizar."],
                ["8", "Conversaciones", "Qué ocurrió en cada conversación."],
                ["9", "Reportes", "Qué descargar o enviar."],
            ],
            widths=[0.6, 1.6, 4.2],
        )
    )
    story.append(PageBreak())

    story.append(Paragraph("6. Pestañas del dashboard", styles["H1x"]))
    for index, tab in enumerate(TAB_CONTENT, start=1):
        story.append(Paragraph(f"6.{index}. {tab['name']}", styles["H2x"]))
        tab_visuals = visual_paths[tab["name"]]
        for image_index, visual_path in enumerate(tab_visuals, start=1):
            story.append(RLImage(str(visual_path), width=6.55 * inch, height=4.09 * inch, kind="proportional", hAlign="CENTER"))
            suffix = f" ({image_index} de {len(tab_visuals)})" if len(tab_visuals) > 1 else ""
            story.append(Paragraph(f"Captura real de la pestaña {tab['name']}{suffix}. Rango usado: 01/04/2026 - 04/05/2026. Los datos mostrados son ficticios.", styles["Smallx"]))
        story.append(
            table(
                [
                    ["Elemento", "Descripción"],
                    ["Para qué sirve", tab["purpose"]],
                    ["Cómo interpretarlo", tab["interpretation"]],
                    ["Ejemplo práctico", tab["example"]],
                ],
                widths=[1.25, 5.1],
            )
        )
        story.append(Spacer(1, 0.08 * inch))
        story.append(para("Preguntas que responde:"))
        story.extend(bullets(tab["questions"]))
        story.append(para("Datos principales:"))
        story.extend(bullets(tab["data"]))
        story.append(para("Decisiones que ayuda a tomar:"))
        story.extend(bullets(tab["decisions"]))
        if index < len(TAB_CONTENT):
            story.append(PageBreak())

    story.append(Paragraph("7. Cómo usar los reportes", styles["H1x"]))
    story.append(para("Los reportes permiten llevar la información del dashboard fuera de la pantalla. Deben generarse después de seleccionar el periodo y los filtros correctos."))
    story.append(
        table(
            [
                ["Formato", "Cuándo usarlo", "Qué esperar"],
                ["PDF", "Para gerencia, lectura ejecutiva o reuniones", "Resumen claro, KPIs, interpretación y tablas principales."],
                ["Excel", "Para revisar detalle, filtrar y trabajar datos", "Hojas con métricas, análisis y detalle de leads cuando aplique."],
                ["CSV", "Para cargar datos en otra herramienta", "Datos planos y limpios para análisis externo."],
            ],
            widths=[1.0, 2.6, 2.8],
        )
    )
    story.append(Paragraph("8. Buenas prácticas por tipo de usuario", styles["H1x"]))
    story.append(
        table(
            [
                ["Usuario", "Qué revisar", "Frecuencia sugerida"],
                ["Gerencia", "Estrategia, Embudo, Rendimiento Humano y Reportes", "Semanal o mensual"],
                ["Equipo operativo", "Operación, Seguimiento y Conversaciones", "Diaria"],
                ["Responsable comercial", "Seguimiento, Rendimiento Humano y Conversaciones", "Diaria y semanal"],
                ["Marketing", "Tendencias, Calidad y Estrategia", "Semanal o por campaña"],
            ],
            widths=[1.5, 3.3, 1.6],
        )
    )
    story.append(Paragraph("9. Preguntas frecuentes", styles["H1x"]))
    for question, answer in [
        ("¿Por qué veo cero datos?", "Puede que el periodo no tenga conversaciones, que el canal seleccionado no tenga registros o que el filtro esté demasiado cerrado."),
        ("¿Por qué un lead aparece sin campaña?", "Porque no se recibió o no se registró el dato de campaña para ese lead."),
        ("¿Qué significa sin puntaje?", "Ese lead todavía no tiene score registrado. En Calidad se agrupa dentro de Frío para que no quede fuera del análisis."),
        ("¿Qué formato de reporte conviene enviar?", "PDF para lectura rápida, Excel para detalle y CSV para análisis en otra herramienta."),
    ]:
        story.append(Paragraph(pdf_escape(question), ParagraphStyle(name=f"q{len(story)}", parent=styles["Bodyx"], fontName=bold_font, spaceAfter=2)))
        story.append(para(answer))
    story.append(Paragraph("10. Checklist final de uso", styles["H1x"]))
    story.extend(bullets(["Definir el periodo correcto antes de analizar.", "Confirmar si se revisan todos los canales o uno específico.", "Leer Estrategia para entender el panorama general.", "Revisar Embudo para detectar dónde se detienen los leads.", "Usar Operación y Seguimiento para atender pendientes.", "Usar Rendimiento Humano para evaluar responsables y conversión.", "Usar Tendencias y Calidad para decisiones de marketing y priorización.", "Abrir Conversaciones cuando se necesite contexto real del lead.", "Exportar el reporte correcto según la audiencia."]))
    story.append(Paragraph("11. Aprobación", styles["H1x"]))
    story.append(table([["Rol", "Nombre", "Firma", "Fecha"], ["Elaborado por", "Simplia", "", TODAY], ["Revisado por", "", "", ""], ["Aprobado por", "", "", ""]], widths=[1.35, 2.0, 1.7, 1.1]))

    pdf = SimpleDocTemplate(str(PDF_PATH), pagesize=letter, rightMargin=0.55 * inch, leftMargin=0.55 * inch, topMargin=0.55 * inch, bottomMargin=0.55 * inch, title="Manual de Usuario SimpliaLeads ISO 10013", author="Simplia")
    pdf.build(story, onFirstPage=page_header, onLaterPages=page_header)


def load_real_visuals() -> dict[str, list[Path]]:
    visual_paths: dict[str, list[Path]] = {}
    missing: list[Path] = []
    for tab_name, prefix in REAL_VISUAL_PREFIXES.items():
        matches = sorted(ASSET_DIR.glob(f"{prefix}_*_real.png"))
        if not matches:
            missing.append(ASSET_DIR / f"{prefix}_01_real.png")
        visual_paths[tab_name] = matches
    if missing:
        formatted = "\n".join(str(path) for path in missing)
        raise FileNotFoundError(f"Faltan capturas reales para el manual:\n{formatted}")
    return visual_paths


def main():
    ASSET_DIR.mkdir(exist_ok=True)
    visual_paths = load_real_visuals()
    generate_word(visual_paths)
    generate_pdf(visual_paths)
    print(f"DOCX={DOCX_PATH}")
    print(f"PDF={PDF_PATH}")
    print(f"ASSETS={ASSET_DIR}")
    print(f"VISUALS={sum(len(paths) for paths in visual_paths.values())}")


if __name__ == "__main__":
    main()
