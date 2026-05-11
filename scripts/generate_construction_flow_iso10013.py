from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.shared import Inches, Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "Construcción y Adaptación de Flujo Conversacional (n8n + Chatwoot).docx"
DOCX_OUT = ROOT / "Construccion_Adaptacion_Flujo_Conversacional_n8n_Chatwoot_ISO10013.docx"
ASSET_DIR = ROOT / "Construccion_Adaptacion_Flujo_Conversacional_assets"

DOC_CODE = "MAN-FLUJO-N8N-CHATWOOT-001"
DOC_VERSION = "1.1"
DOC_DATE = "06/05/2026"
NEXT_REVIEW = "06/11/2026"
DOC_STATUS = "Borrador / Vigente al aprobarse"

COLORS = {
    "navy": "0f2344",
    "blue": "274690",
    "soft_blue": "eaf1ff",
    "light": "f8fafc",
    "green": "0a9b6f",
    "slate": "64748b",
    "line": "d9e2ef",
    "white": "ffffff",
    "code": "f3f6fb",
}


def rgb(hex_color: str) -> RGBColor:
    value = hex_color.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def iter_block_items(document: Document):
    for child in document.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, document)
        elif isinstance(child, CT_Tbl):
            yield Table(child, document)


def set_cell_shading(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def set_table_borders(table, color: str = COLORS["line"]) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement("w:" + edge)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    table._tbl.tblPr.append(borders)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def write_cell(cell, text: str, *, bold: bool = False, color: str | None = None, size: float = 8.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for index, part in enumerate(str(text).split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def style_table(table) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    set_table_borders(table)
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if row_index == 0:
                set_cell_shading(cell, COLORS["blue"])
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = rgb(COLORS["white"])
            else:
                set_cell_shading(cell, COLORS["light"] if row_index % 2 == 0 else COLORS["white"])
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(7.6 if row_index else 8.1)


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        write_cell(table.cell(0, idx), header, bold=True, color=COLORS["white"], size=8.1)
    set_repeat_table_header(table.rows[0])
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row[: len(headers)]):
            write_cell(cells[idx], value, size=7.6)
    style_table(table)
    document.add_paragraph()


def add_source_table(document: Document, source_table: Table) -> None:
    rows: list[list[str]] = []
    for row in source_table.rows:
        rows.append(["\n".join(cell.text.strip().splitlines()) for cell in row.cells])
    if not rows:
        return
    max_cols = max(len(row) for row in rows)
    normalized = [row + [""] * (max_cols - len(row)) for row in rows]
    add_table(document, normalized[0], normalized[1:])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    for style_name, size, color in [
        ("Heading 1", 15, COLORS["blue"]),
        ("Heading 2", 12.5, COLORS["navy"]),
        ("Heading 3", 10.5, COLORS["green"]),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"{DOC_CODE} | v{DOC_VERSION} | {DOC_STATUS}")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["slate"])

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Simplia Chatbot - Construcción de Flujo n8n + Chatwoot ISO 10013")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["slate"])


def add_iso_cover(document: Document) -> None:
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Construcción y Adaptación de Flujo Conversacional")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.color.rgb = rgb(COLORS["blue"])

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Manual técnico-operativo n8n + Chatwoot")
    run.font.name = "Arial"
    run.font.size = Pt(11)
    run.font.color.rgb = rgb(COLORS["slate"])

    add_table(
        document,
        ["Campo", "Valor"],
        [
            ["Código", DOC_CODE],
            ["Nombre", "Construcción y Adaptación de Flujo Conversacional (n8n + Chatwoot)"],
            ["Versión", DOC_VERSION],
            ["Fecha", DOC_DATE],
            ["Estado", DOC_STATUS],
            ["Próxima revisión", NEXT_REVIEW],
            ["Tipo de documento", "Manual ISO 10013 de información documentada"],
        ],
    )

    table = document.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_borders(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS["soft_blue"])
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run("Uso del documento")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(8.8)
    run.font.color.rgb = rgb(COLORS["blue"])
    p.add_run().add_break()
    run = p.add_run(
        "Esta guía conserva el orden operativo del documento fuente: texto, tablas y capturas aparecen en la misma secuencia para construir y adaptar el flujo conversacional de inicio a fin."
    )
    run.font.name = "Arial"
    run.font.size = Pt(8.4)
    run.font.color.rgb = rgb(COLORS["navy"])
    document.add_paragraph()


def paragraph_level(text: str, source_index: int) -> int | None:
    normalized = " ".join(text.split())
    if source_index == 1 or normalized.startswith("Guía Técnica:"):
        return 1
    if normalized.lower().startswith("fase "):
        return 1
    if re.match(r"^\d+\.\d+\.\s+", normalized):
        return 2
    if re.match(r"^\d+\.\d+\s+", normalized):
        return 2
    if re.match(r"^\d+\.\s+", normalized):
        return 2
    return None


def looks_like_code(text: str) -> bool:
    stripped = text.strip()
    lower = stripped.lower()
    return (
        lower.startswith("create table")
        or lower.startswith("create index")
        or lower.startswith("constraint ")
        or lower.startswith("id ")
        or lower.startswith("session_id ")
        or lower.startswith("message ")
        or lower.startswith("pipeline_etapa ")
        or lower.startswith("created_at ")
        or lower.startswith("workflow_name ")
        or lower.startswith("error_message ")
        or lower.startswith("url ")
        or lower.startswith("last_node_executed ")
        or stripped in {"(", ")", ");"}
    )


def add_source_paragraph(document: Document, source_paragraph: Paragraph, source_index: int) -> None:
    text = source_paragraph.text.strip()
    if text:
        level = paragraph_level(text, source_index)
        if level:
            p = document.add_heading(text, level=level)
            for run in p.runs:
                run.font.name = "Arial"
                run.font.color.rgb = rgb(COLORS["blue"] if level == 1 else COLORS["navy"])
        elif looks_like_code(text):
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(text)
            run.font.name = "Consolas"
            run.font.size = Pt(8.4)
            run.font.color.rgb = rgb(COLORS["navy"])
        else:
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(9.4)
            run.font.color.rgb = rgb(COLORS["navy"])


def extract_and_add_images(document: Document, source_paragraph: Paragraph, image_counter: int) -> int:
    blips = source_paragraph._p.xpath(".//a:blip")
    for blip in blips:
        rid = blip.get(qn("r:embed"))
        rel = source_paragraph.part.rels.get(rid)
        if not rel:
            continue
        image_counter += 1
        ext = ".jpg" if "jpeg" in rel.target_part.content_type else ".png"
        image_path = ASSET_DIR / f"captura_flujo_orden_{image_counter:03d}{ext}"
        try:
            image_path.write_bytes(rel.target_part.blob)
        except PermissionError:
            if not image_path.exists():
                raise

        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(f"Captura guía {image_counter:03d}")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(8.8)
        run.font.color.rgb = rgb(COLORS["blue"])

        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(image_path), width=Inches(6.25))
    return image_counter


def build_document() -> None:
    source = Document(SOURCE_DOCX)
    ASSET_DIR.mkdir(exist_ok=True)

    document = Document()
    configure_document(document)
    add_iso_cover(document)

    document.add_heading("Contenido operativo en el orden original", level=1)
    image_counter = 0
    paragraph_index = 0

    for block in iter_block_items(source):
        if isinstance(block, Paragraph):
            paragraph_index += 1
            add_source_paragraph(document, block, paragraph_index)
            image_counter = extract_and_add_images(document, block, image_counter)
        else:
            add_source_table(document, block)

    document.core_properties.title = "Construcción y Adaptación de Flujo Conversacional n8n + Chatwoot ISO 10013"
    document.core_properties.subject = "Manual técnico-operativo para construir y adaptar el flujo conversacional"
    document.core_properties.author = "Simplia"
    document.core_properties.keywords = "Simplia, n8n, Chatwoot, flujo conversacional, ISO 10013"
    try:
        document.save(DOCX_OUT)
    except PermissionError as exc:
        raise SystemExit(f"No se pudo guardar {DOCX_OUT.name}. Cierre el archivo en Word y vuelva a ejecutar este script.") from exc


def main() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)
    build_document()
    source = Document(SOURCE_DOCX)
    output = Document(DOCX_OUT)
    print(f"DOCX={DOCX_OUT}")
    print(f"SOURCE_IMAGES={len(source.inline_shapes)}")
    print(f"OUTPUT_IMAGES={len(output.inline_shapes)}")


if __name__ == "__main__":
    main()
