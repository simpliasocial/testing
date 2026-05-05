from __future__ import annotations

import os
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "ONBOARDING TECNICO -SIMPLIA CHATBOT .docx"
TEMPLATE_DOCX = ROOT / "plantilla_iso_9001_proceso.docx"
LOGO_PATH = ROOT / "logo_simplia.png"
ASSET_DIR = ROOT / "Onboarding_Tecnico_Simplia_Chatbot_assets"
DOCX_OUT = ROOT / "Onboarding_Tecnico_Simplia_Chatbot_ISO9001.docx"
PDF_OUT = ROOT / "Onboarding_Tecnico_Simplia_Chatbot_ISO9001.pdf"

DOC_CODE = "PRO-ONB-CHATBOT-001"
DOC_VERSION = "1.0"
DOC_STATUS = "Vigente/Borrador"
ISSUE_DATE = "04/05/2026"
NEXT_REVIEW = "04/11/2026"

COLORS = {
    "blue": "274690",
    "navy": "0f2344",
    "slate": "64748b",
    "line": "d9e2ef",
    "light": "f8fafc",
    "soft_blue": "eaf1ff",
    "green": "0a9b6f",
    "orange": "f59e0b",
    "red": "ef4444",
}


def rgb(hex_color: str) -> tuple[int, int, int]:
    value = hex_color.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def font(size: int, bold: bool = False):
    paths = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for path in paths:
        if os.path.exists(path):
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def clean_text(text: str) -> str:
    replacements = {
        "cuneta": "cuenta",
        "antiguedad": "antigüedad",
        "contrasena": "contraseña",
        "aplasatr": "aplastar",
        "dirigmos": "dirigimos",
        "coloar": "colocar",
        "fcebook": "Facebook",
        "Facebok": "Facebook",
        "Facebok": "Facebook",
        "facebookk": "Facebook",
        "Ahroa": "Ahora",
        "tegamos": "tengamos",
        "configuarcion": "configuración",
        "configuracion": "configuración",
        "verificacion": "verificación",
        "soluciionarse": "solucionarse",
        "tavez": "través",
        "Incio": "Inicio",
        "inicar": "iniciar",
        "javascrpt": "JavaScript",
        "COnfiguarcion": "configuración",
        "Messanger": "Messenger",
    }
    value = text.replace("\u00a0", " ").strip()
    for wrong, right in replacements.items():
        value = value.replace(wrong, right)
    value = re.sub(r"[ \t]+", " ", value)
    value = value.replace(" :", ":")
    return value


def extract_source_content():
    source = Document(SOURCE_DOCX)
    paragraphs = [clean_text(p.text) for p in source.paragraphs if p.text.strip()]
    table_rows = []
    for table in source.tables:
        for row in table.rows:
            table_rows.append([clean_text(cell.text.replace("\n", " / ")) for cell in row.cells])
    return source, paragraphs, table_rows


def extract_images(source: Document) -> list[Path]:
    ASSET_DIR.mkdir(exist_ok=True)
    for old in ASSET_DIR.glob("evidencia_*.png"):
        old.unlink()
    extracted: list[Path] = []
    for index, shape in enumerate(source.inline_shapes, start=1):
        blips = shape._inline.xpath(".//a:blip")
        if not blips:
            continue
        rid = blips[0].get(qn("r:embed"))
        rel = source.part.rels.get(rid)
        if not rel:
            continue
        blob = rel.target_part.blob
        content_type = rel.target_part.content_type
        ext = ".jpg" if "jpeg" in content_type else ".png"
        path = ASSET_DIR / f"evidencia_{index:03d}{ext}"
        path.write_bytes(blob)
        extracted.append(path)
    return extracted


def make_process_map() -> Path:
    path = ASSET_DIR / "mapa_proceso_onboarding.png"
    width, height = 1800, 520
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    title_font = font(34, True)
    box_font = font(23, True)
    small_font = font(17)
    draw.text((60, 38), "Mapa del proceso de onboarding técnico Simplia Chatbot", font=title_font, fill=rgb(COLORS["navy"]))
    steps = [
        ("1", "Levantamiento", "Datos de empresa,\ncanales y accesos"),
        ("2", "Cuentas base", "Meta, Railway,\nn8n y Chatwoot"),
        ("3", "Canales", "WhatsApp,\nMessenger e Instagram"),
        ("4", "IA y automatización", "OpenAI,\nDrive y Gmail"),
        ("5", "Pruebas", "Mensajes,\nwebhooks y evidencias"),
        ("6", "Entrega", "Credenciales,\nchecklist y aprobación"),
    ]
    x, y = 70, 165
    box_w, box_h, gap = 245, 180, 42
    for idx, (num, name, desc) in enumerate(steps):
        bx = x + idx * (box_w + gap)
        draw.rounded_rectangle((bx, y, bx + box_w, y + box_h), radius=22, fill=rgb("eaf1ff"), outline=rgb(COLORS["blue"]), width=3)
        draw.ellipse((bx + 18, y + 18, bx + 68, y + 68), fill=rgb(COLORS["blue"]))
        draw.text((bx + 36, y + 30), num, font=font(22, True), fill="white", anchor="mm")
        draw.text((bx + 82, y + 26), name, font=box_font, fill=rgb(COLORS["navy"]))
        draw.multiline_text((bx + 26, y + 88), desc, font=small_font, fill=rgb(COLORS["slate"]), spacing=6)
        if idx < len(steps) - 1:
            ax = bx + box_w + 7
            ay = y + box_h // 2
            draw.line((ax, ay, ax + gap - 14, ay), fill=rgb(COLORS["blue"]), width=4)
            draw.polygon([(ax + gap - 14, ay - 10), (ax + gap + 4, ay), (ax + gap - 14, ay + 10)], fill=rgb(COLORS["blue"]))
    draw.text((70, 420), "Salida esperada: chatbot conectado, canales probados, evidencias guardadas y datos de operación completos.", font=small_font, fill=rgb(COLORS["navy"]))
    image.save(path)
    return path


PROCEDURE_ROWS = [
    ["1", "Crear portafolio en Facebook Business Manager y validar que la cuenta tenga más de 2 meses.", "Cliente / Técnico Simplia", "Portafolio creado y cuenta apta.", "Captura del portafolio y enlace guía."],
    ["2", "Crear la app en Meta for Developers y registrar número nuevo nunca usado con WhatsApp.", "Técnico Simplia", "App creada, número registrado y token de larga duración disponible.", "Número, token y captura de la app."],
    ["3", "Crear cuenta y proyecto en Railway para alojar Chatwoot y n8n.", "Técnico Simplia", "Railway accesible con credenciales registradas.", "Correo, contraseña y URL del proyecto."],
    ["4", "Crear instancia de n8n en Railway.", "Técnico Simplia", "n8n operativo y con URL de acceso.", "URL, usuario y contraseña."],
    ["5", "Crear instancia de Chatwoot en Railway y aplicar cambios obligatorios de configuración por nodo.", "Técnico Simplia", "Chatwoot operativo y configuración aplicada.", "URL, usuario, contraseña y evidencia Railway."],
    ["6", "Configurar inbox de WhatsApp en Chatwoot usando datos de Meta.", "Técnico Simplia", "Inbox conectado con Phone Number, Phone Number ID, Business Account ID y token.", "Datos del inbox y captura de conexión."],
    ["7", "Configurar Webhook URL y Webhook Verification Token en Meta.", "Técnico Simplia", "Webhook verificado y messages habilitado.", "Captura de callback, verify token y suscripción."],
    ["8", "Agregar método de pago en Meta para habilitar WhatsApp correctamente.", "Cliente / Técnico Simplia", "Método de pago configurado.", "Captura del centro de facturación."],
    ["9", "Probar mensaje entrante al número conectado y validar recepción en Chatwoot.", "Técnico Simplia", "Mensaje recibido en bandeja de entrada.", "Captura de conversación de prueba."],
    ["10", "Configurar Messenger/Facebook desde Super Admin, Meta App y webhooks.", "Técnico Simplia", "Canal Facebook desbloqueado y conectado.", "App ID, App Secret, Verify Token y capturas."],
    ["11", "Verificar negocio en Facebook si se requiere permiso adicional.", "Cliente", "Solicitud enviada o negocio verificado.", "Evidencia de Business Verification Status."],
    ["12", "Configurar Instagram en Chatwoot y Meta, incluyendo tokens, webhooks y callback.", "Técnico Simplia", "Instagram conectado y probado.", "Verify Token, webhook y captura de conexión."],
    ["13", "Configurar OpenAI / ChatGPT y límites de crédito.", "Técnico Simplia / Cliente", "API Key registrada y presupuesto controlado.", "API Key placeholder, organización y límites."],
    ["14", "Configurar n8n con Google Drive service account si el flujo requiere fotos o videos.", "Desarrollador / Técnico Simplia", "Proyecto Google y acceso Drive operativos.", "Proyecto, usuarios de prueba y enlaces Drive."],
    ["15", "Configurar n8n con Gmail OAuth si el flujo requiere envío de correos.", "Desarrollador / Técnico Simplia", "OAuth publicado y token estable.", "Proyecto Google, consentimiento y evidencia."],
    ["16", "Configurar plantillas de respuesta en Chatwoot cuando la empresa lo requiera.", "Técnico Simplia / Cliente", "Plantillas disponibles para operación.", "Captura de plantillas."],
    ["17", "Registrar acceso remoto por AnyDesk si se requiere soporte guiado.", "Cliente / Técnico Simplia", "Conexión remota disponible con autorización.", "ID AnyDesk y evidencia de acceso autorizado."],
    ["18", "Completar registro controlado de credenciales y validar checklist final.", "Responsable del onboarding", "Onboarding completo y aprobado.", "Tabla de credenciales, checklist y aprobación."],
]

INPUT_OUTPUT_ROWS = [
    ["Cliente / empresa", "Datos de empresa, canales deseados, página web, correos y números", "Requisitos de configuración claros", "Equipo técnico Simplia"],
    ["Meta / Facebook", "Business Manager, app, número, tokens, webhooks y método de pago", "WhatsApp, Messenger e Instagram conectados", "Chatwoot / Cliente"],
    ["Railway", "Proyecto, servicios y variables de instancia", "Chatwoot y n8n disponibles", "Equipo técnico Simplia"],
    ["OpenAI", "API Key, organización y límites de crédito", "Chatbot con IA configurada", "Flujos n8n / Chatbot"],
    ["Google", "Drive service account y Gmail OAuth cuando aplique", "Archivos/correos integrados al flujo", "n8n / Operación"],
]

RACI_ROWS = [
    ["Levantamiento de información", "Técnico Simplia", "Responsable del onboarding", "Cliente", "Desarrollador", "Confirmar datos antes de configurar."],
    ["Creación de cuentas y accesos", "Cliente / Técnico Simplia", "Cliente", "Simplia", "Gerencia", "No registrar secretos reales en documentos compartidos."],
    ["Configuración Meta/WhatsApp", "Técnico Simplia", "Responsable técnico", "Cliente", "Operación", "Requiere método de pago activo."],
    ["Configuración Chatwoot/n8n/Railway", "Técnico Simplia", "Responsable técnico", "Desarrollador", "Cliente", "Validar URL, usuario y contraseña."],
    ["Configuración IA y automatizaciones", "Desarrollador", "Responsable técnico", "Cliente", "Operación", "Aplicar límites de crédito."],
    ["Pruebas y cierre", "Técnico Simplia", "Responsable del onboarding", "Cliente", "Gerencia", "Cerrar con checklist y evidencias."],
]

RISK_ROWS = [
    ["Cuenta de Facebook no apta", "Cuenta nueva o sin antigüedad suficiente", "Validar antigüedad antes de iniciar", "Revisión previa del Business Manager", "Usar cuenta apta o esperar aprobación"],
    ["WhatsApp no funciona", "No existe método de pago en Meta", "Agregar método de pago antes de pruebas", "Prueba de mensaje entrante", "Configurar facturación y repetir prueba"],
    ["Webhook no verifica", "URL o token incorrecto", "Copiar datos desde Chatwoot sin alterar", "Verificar callback y suscripciones", "Regenerar token y volver a guardar"],
    ["Tokens caducan", "App Google queda en testing", "Pasar app a producción cuando aplique", "Revisar estado OAuth", "Publicar app y reconectar credenciales"],
    ["IA deja de responder", "Créditos OpenAI agotados", "Configurar límites y auto-recharge", "Revisar presupuesto y consumo", "Recargar créditos o ajustar límite"],
    ["Archivos no cargan", "Videos mayores a 15 MB o permisos Drive incorrectos", "Comprimir videos y compartir como editor", "Probar enlace antes de usarlo", "Corregir permisos o comprimir archivo"],
]

RECORD_ROWS = [
    ["Formulario de levantamiento", "Documento ISO 9001 de onboarding", "Responsable del onboarding", "Duración del proyecto + 1 año", "Solo equipo autorizado"],
    ["Credenciales y accesos", "Tabla controlada / gestor autorizado", "Cliente / Técnico Simplia", "Según política del cliente", "Acceso restringido"],
    ["Capturas de configuración", "Anexo visual del documento", "Técnico Simplia", "Duración del proyecto + 1 año", "Lectura controlada"],
    ["Guías y enlaces", "Procedimiento y anexo de referencias", "Técnico Simplia", "Mientras estén vigentes", "Lectura general"],
    ["Checklist de pruebas", "Sección 10 del documento", "Responsable del onboarding", "Duración del proyecto + 1 año", "Cliente y Simplia"],
]

KPI_ROWS = [
    ["Completitud de información", "Campos completados / campos requeridos", "Por onboarding", "Documento", "Responsable onboarding", "100% antes de configurar"],
    ["Canales conectados", "Canales probados / canales solicitados", "Por onboarding", "Chatwoot / Meta", "Técnico Simplia", "100% antes de entrega"],
    ["Pruebas exitosas", "Pruebas OK / pruebas ejecutadas", "Por onboarding", "Checklist", "Técnico Simplia", "100% OK o N.A. justificado"],
    ["Evidencias completas", "Evidencias cargadas / evidencias requeridas", "Por onboarding", "Documento / anexos", "Responsable onboarding", "100%"],
]

CHECKLIST_ROWS = [
    ["1", "Portafolio Facebook Business Manager creado y validado", "OK / No OK / N.A.", ""],
    ["2", "App Meta creada y número WhatsApp nuevo registrado", "OK / No OK / N.A.", ""],
    ["3", "Railway, n8n y Chatwoot creados con accesos documentados", "OK / No OK / N.A.", ""],
    ["4", "Inbox WhatsApp conectado y probado en Chatwoot", "OK / No OK / N.A.", ""],
    ["5", "Webhook URL y token verificados en Meta", "OK / No OK / N.A.", ""],
    ["6", "Método de pago Meta/Railway/OpenAI agregado cuando aplique", "OK / No OK / N.A.", ""],
    ["7", "Messenger/Facebook conectado si fue solicitado", "OK / No OK / N.A.", ""],
    ["8", "Instagram conectado si fue solicitado", "OK / No OK / N.A.", ""],
    ["9", "OpenAI configurado con límites de crédito", "OK / No OK / N.A.", ""],
    ["10", "Drive/Gmail OAuth configurado si aplica", "OK / No OK / N.A.", ""],
    ["11", "Plantillas Chatwoot configuradas si aplica", "OK / No OK / N.A.", ""],
    ["12", "Todas las evidencias y credenciales quedan registradas", "OK / No OK / N.A.", ""],
]

CREDENTIAL_ROWS = [
    ["Sistema", "Dato requerido", "Valor / evidencia"],
    ["Railway", "Correo de acceso", "____________________"],
    ["Railway", "Contraseña", "____________________"],
    ["n8n", "URL de la instancia", "____________________"],
    ["n8n", "Correo de acceso", "____________________"],
    ["n8n", "Contraseña", "____________________"],
    ["Chatwoot", "URL de la instancia", "____________________"],
    ["Chatwoot", "Correo de acceso", "____________________"],
    ["Chatwoot", "Contraseña", "____________________"],
    ["Meta / WhatsApp", "Phone Number", "____________________"],
    ["Meta / WhatsApp", "Phone Number ID", "____________________"],
    ["Meta / WhatsApp", "Business Account ID", "____________________"],
    ["Meta / WhatsApp", "Token de larga duración", "____________________"],
    ["Chatwoot / Webhook", "Webhook URL", "____________________"],
    ["Chatwoot / Webhook", "Webhook Verification Token", "____________________"],
    ["OpenAI", "API Key", "____________________"],
    ["OpenAI", "Organización", "____________________"],
    ["Supabase / Base de datos", "Correo de acceso / clave de organización si aplica", "____________________"],
    ["Redis", "Correo / contraseña si aplica", "____________________"],
    ["CRM", "URL / usuario / contraseña", "____________________"],
    ["Gmail", "Cuenta base", "____________________"],
]


def set_cell_shading(cell, fill: str):
    props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def set_table_borders(table, color="d9e2ef"):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement("w:" + edge)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    table._tbl.tblPr.append(borders)


def set_cell_text(cell, text, bold=False, color="0f2344"):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(8.7)
    run.font.color.rgb = RGBColor(*rgb(color))
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_doc_table(doc: Document, rows, widths=None, header=True):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            is_header = header and i == 0
            cell = table.cell(i, j)
            set_cell_text(cell, value, bold=is_header, color=("ffffff" if is_header else COLORS["navy"]))
            set_cell_shading(cell, COLORS["blue"] if is_header else ("f8fafc" if i % 2 == 0 else "ffffff"))
            if widths and j < len(widths):
                cell.width = Inches(widths[j])
    return table


def add_heading(doc: Document, text: str, level: int = 1):
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(*rgb(COLORS["blue"] if level == 1 else COLORS["navy"]))
    return paragraph


def add_para(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10.2)
    run.font.color.rgb = RGBColor(*rgb(COLORS["navy"]))
    return paragraph


def add_bullets(doc: Document, items):
    for item in items:
        paragraph = doc.add_paragraph(style="List Bullet")
        run = paragraph.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(10.2)
        run.font.color.rgb = RGBColor(*rgb(COLORS["navy"]))


def add_callout(doc: Document, title: str, text: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, "b8c7e8")
    cell = table.cell(0, 0)
    set_cell_shading(cell, COLORS["soft_blue"])
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title + ": ")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(*rgb(COLORS["blue"]))
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(*rgb(COLORS["navy"]))
    doc.add_paragraph()


def add_scaled_picture(doc: Document, path: Path, max_width=6.6):
    with Image.open(path) as image:
        width_px, height_px = image.size
    width = min(max_width, 6.6)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path), width=Inches(width))


def build_docx(paragraphs: list[str], table_rows: list[list[str]], images: list[Path], map_path: Path):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.2)
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"{DOC_CODE} | v{DOC_VERSION} | {DOC_STATUS}")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(*rgb(COLORS["slate"]))
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Simplia Chatbot - Onboarding Técnico ISO 9001")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(*rgb(COLORS["slate"]))

    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        cover.add_run().add_picture(str(LOGO_PATH), width=Inches(1.6))
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Onboarding Técnico Simplia Chatbot")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(25)
    run.font.color.rgb = RGBColor(*rgb(COLORS["blue"]))
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Documento de proceso para levantamiento, configuración y cierre técnico")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(*rgb(COLORS["slate"]))
    doc.add_paragraph()
    add_doc_table(
        doc,
        [
            ["Código del documento", DOC_CODE, "Versión", DOC_VERSION],
            ["Nombre del proceso", "Onboarding técnico Simplia Chatbot", "Fecha de emisión", ISSUE_DATE],
            ["Dueño del proceso", "Responsable de onboarding / Simplia", "Próxima revisión", NEXT_REVIEW],
            ["Aprobado por", "Gerencia / Responsable del cliente", "Estado", DOC_STATUS],
        ],
        widths=[1.55, 2.0, 1.45, 1.7],
        header=False,
    )
    add_callout(doc, "Uso del documento", "Este formato sirve para levantar información de la empresa, controlar la configuración técnica y dejar evidencia de cada paso ejecutado.")
    doc.add_page_break()

    add_heading(doc, "1. Información general del proceso", 1)
    add_bullets(
        doc,
        [
            "Objetivo del proceso: recopilar información, accesos y requisitos de la empresa para configurar el chatbot Simplia en canales conversacionales.",
            "Alcance: inicia con la creación o validación de cuentas base y termina con canales probados, credenciales registradas, evidencias completas y aprobación.",
            "Cliente interno o externo: empresa cliente que usará el chatbot y equipo técnico Simplia que realiza la configuración.",
            "Frecuencia: por cada nuevo cliente, nuevo canal o nueva instancia técnica.",
            "Herramientas usadas: Meta Business, Meta for Developers, WhatsApp Business, Railway, n8n, Chatwoot, OpenAI, Google Drive, Gmail, AnyDesk y CRM si aplica.",
        ],
    )

    add_heading(doc, "2. Contexto del proceso", 1)
    add_para(doc, "El onboarding técnico permite entender qué necesita la empresa, qué canales se conectarán y qué información debe quedar disponible para que el chatbot funcione correctamente. El documento original de requisitos se conserva como fuente y se transforma aquí en un proceso controlado.")
    add_doc_table(
        doc,
        [
            ["Aspecto", "Descripción"],
            ["Necesidad que resuelve", "Evita configuraciones incompletas, pérdida de accesos y falta de evidencia durante la puesta en marcha del chatbot."],
            ["Procesos relacionados", "Levantamiento comercial, configuración técnica, pruebas de canales, despliegue y soporte inicial."],
            ["Requisitos aplicables", "Accesos autorizados, método de pago activo donde aplique, datos completos del negocio y aprobación del cliente."],
        ],
        widths=[1.8, 5.0],
    )

    add_heading(doc, "3. Entradas y salidas del proceso", 1)
    add_doc_table(doc, [["Proveedor / origen", "Entrada", "Salida", "Cliente / destino"]] + INPUT_OUTPUT_ROWS, widths=[1.4, 2.2, 2.0, 1.4])

    add_heading(doc, "4. Mapa del proceso", 1)
    add_scaled_picture(doc, map_path)

    add_heading(doc, "5. Procedimiento operativo estándar (SOP)", 1)
    add_para(doc, "La siguiente tabla organiza el contenido original en pasos ejecutables, con responsable, criterio de aceptación y evidencia esperada.")
    add_doc_table(doc, [["Paso", "Actividad", "Responsable", "Criterio de aceptación", "Registro / evidencia"]] + PROCEDURE_ROWS, widths=[0.45, 2.55, 1.3, 1.55, 1.4])

    add_heading(doc, "6. Roles y responsabilidades (RACI)", 1)
    add_doc_table(doc, [["Actividad clave", "R", "A", "C", "I", "Observaciones"]] + RACI_ROWS, widths=[1.7, 1.0, 1.0, 1.0, 0.9, 1.5])

    add_heading(doc, "7. Riesgos, desvíos y controles", 1)
    add_doc_table(doc, [["Riesgo o falla posible", "Causa probable", "Control preventivo", "Control de verificación", "Acción ante desvío"]] + RISK_ROWS, widths=[1.35, 1.35, 1.45, 1.45, 1.35])

    add_heading(doc, "8. Registros e información documentada", 1)
    add_doc_table(doc, [["Registro / evidencia", "Ubicación", "Responsable", "Retención", "Control de acceso / edición"]] + RECORD_ROWS, widths=[1.45, 1.55, 1.35, 1.25, 1.45])

    add_heading(doc, "9. Indicadores del proceso", 1)
    add_doc_table(doc, [["KPI", "Fórmula", "Frecuencia", "Fuente", "Dueño", "Meta / umbral"]] + KPI_ROWS, widths=[1.2, 1.6, 1.0, 1.0, 1.1, 1.2])

    add_heading(doc, "10. Checklist de ejecución", 1)
    add_doc_table(doc, [["#", "Verificación", "Resultado", "Observaciones"]] + CHECKLIST_ROWS, widths=[0.45, 4.0, 1.35, 1.2])

    add_heading(doc, "11. Registro controlado de credenciales y datos requeridos", 1)
    add_callout(doc, "Importante", "Completar los campos solo en una copia controlada y autorizada. No compartir este documento con valores reales si no corresponde.")
    add_doc_table(doc, CREDENTIAL_ROWS, widths=[1.8, 2.5, 2.6])

    if table_rows:
        add_heading(doc, "12. Tabla fuente original de credenciales", 1)
        add_para(doc, "Se conserva la tabla original del documento base como referencia, con la redacción normalizada cuando aplica.")
        add_doc_table(doc, [["Sistema", "Contenido original"]] + table_rows, widths=[1.5, 5.4])

    add_heading(doc, "13. Control de cambios", 1)
    add_doc_table(doc, [["Versión", "Fecha", "Cambio realizado", "Responsable"], ["1.0", ISSUE_DATE, "Transformación del documento de onboarding técnico al formato ISO 9001", "Simplia"]], widths=[0.9, 1.2, 3.8, 1.3])

    add_heading(doc, "14. Aprobación final", 1)
    add_doc_table(doc, [["Elaborado por", "Revisado por", "Aprobado por"], ["Simplia / Responsable de onboarding", "", ""], ["Firma / fecha", "Firma / fecha", "Firma / fecha"]], widths=[2.2, 2.2, 2.2])

    doc.add_page_break()
    add_heading(doc, "Anexo A. Transcripción controlada del contenido fuente", 1)
    add_para(doc, "Este anexo conserva el contenido operativo del documento original en orden de aparición. La redacción fue normalizada para mejorar lectura sin cambiar el significado.")
    source_table_rows = [["#", "Contenido"]]
    for idx, paragraph in enumerate(paragraphs, start=1):
        source_table_rows.append([str(idx), paragraph])
    add_doc_table(doc, source_table_rows, widths=[0.45, 6.35])

    doc.add_page_break()
    add_heading(doc, "Anexo B. Evidencias visuales del documento original", 1)
    add_para(doc, f"Se conservan {len(images)} capturas/imágenes extraídas del documento original, en el mismo orden de aparición.")
    for index, image_path in enumerate(images, start=1):
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = caption.add_run(f"Evidencia visual {index:03d}")
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(*rgb(COLORS["blue"]))
        add_scaled_picture(doc, image_path, max_width=6.55)

    doc.core_properties.title = "Onboarding Técnico Simplia Chatbot ISO 9001"
    doc.core_properties.subject = "Proceso de onboarding técnico para configuración de chatbot"
    doc.core_properties.author = "Simplia"
    doc.core_properties.keywords = "onboarding, ISO 9001, Simplia Chatbot, Chatwoot, n8n, Meta, OpenAI"
    doc.save(DOCX_OUT)


def pdf_escape(text) -> str:
    return str(text).replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def build_pdf(paragraphs: list[str], table_rows: list[list[str]], images: list[Path], map_path: Path):
    try:
        pdfmetrics.registerFont(TTFont("Arial", "C:/Windows/Fonts/arial.ttf"))
        pdfmetrics.registerFont(TTFont("Arial-Bold", "C:/Windows/Fonts/arialbd.ttf"))
        base_font = "Arial"
        bold_font = "Arial-Bold"
    except Exception:
        base_font = "Helvetica"
        bold_font = "Helvetica-Bold"

    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="TitleX", parent=styles["Title"], fontName=bold_font, fontSize=24, leading=29, alignment=TA_CENTER, textColor=colors.HexColor("#274690")))
    styles.add(ParagraphStyle(name="SubX", parent=styles["Normal"], fontName=base_font, fontSize=10.5, leading=14, alignment=TA_CENTER, textColor=colors.HexColor("#64748b")))
    styles.add(ParagraphStyle(name="H1X", parent=styles["Heading1"], fontName=bold_font, fontSize=15.5, leading=19, textColor=colors.HexColor("#274690"), spaceBefore=12, spaceAfter=8))
    styles.add(ParagraphStyle(name="H2X", parent=styles["Heading2"], fontName=bold_font, fontSize=12.5, leading=16, textColor=colors.HexColor("#0f2344"), spaceBefore=8, spaceAfter=5))
    styles.add(ParagraphStyle(name="BodyX", parent=styles["BodyText"], fontName=base_font, fontSize=8.8, leading=12, textColor=colors.HexColor("#0f2344"), spaceAfter=5))
    styles.add(ParagraphStyle(name="CellX", parent=styles["BodyText"], fontName=base_font, fontSize=7.2, leading=9, textColor=colors.HexColor("#0f2344")))
    styles.add(ParagraphStyle(name="HeadCellX", parent=styles["BodyText"], fontName=bold_font, fontSize=7.2, leading=9, textColor=colors.white))
    styles.add(ParagraphStyle(name="CaptionX", parent=styles["BodyText"], fontName=bold_font, fontSize=8, leading=10, alignment=TA_CENTER, textColor=colors.HexColor("#274690")))

    def p(text, style="BodyX"):
        return Paragraph(pdf_escape(text), styles[style])

    def table(rows, widths=None):
        data = [[p(cell, "HeadCellX" if i == 0 else "CellX") for cell in row] for i, row in enumerate(rows)]
        item = Table(data, colWidths=[w * inch for w in widths] if widths else None, repeatRows=1, hAlign="LEFT")
        item.setStyle(
            TableStyle(
                [
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#274690")),
                    ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#d9e2ef")),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                    ("LEFTPADDING", (0, 0), (-1, -1), 4),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                    ("TOPPADDING", (0, 0), (-1, -1), 3),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ]
            )
        )
        return item

    def scaled_image(path: Path, max_w=6.45, max_h=8.2):
        with Image.open(path) as image:
            w, h = image.size
        ratio = min((max_w * inch) / w, (max_h * inch) / h, 1.0)
        return RLImage(str(path), width=w * ratio, height=h * ratio, hAlign="CENTER")

    def header(canvas, doc_obj):
        canvas.saveState()
        canvas.setFont(base_font, 7)
        canvas.setFillColor(colors.HexColor("#64748b"))
        canvas.drawRightString(letter[0] - 0.5 * inch, letter[1] - 0.35 * inch, f"{DOC_CODE} | v{DOC_VERSION} | {DOC_STATUS}")
        canvas.drawCentredString(letter[0] / 2, 0.33 * inch, f"Simplia Chatbot - Onboarding Técnico ISO 9001 | Página {doc_obj.page}")
        canvas.restoreState()

    story = []
    if LOGO_PATH.exists():
        story.append(RLImage(str(LOGO_PATH), width=1.45 * inch, height=0.62 * inch, kind="proportional", hAlign="CENTER"))
        story.append(Spacer(1, 0.15 * inch))
    story.append(Paragraph("Onboarding Técnico Simplia Chatbot", styles["TitleX"]))
    story.append(Paragraph("Documento de proceso para levantamiento, configuración y cierre técnico", styles["SubX"]))
    story.append(Spacer(1, 0.12 * inch))
    story.append(
        table(
            [
                ["Código del documento", DOC_CODE, "Versión", DOC_VERSION],
                ["Nombre del proceso", "Onboarding técnico Simplia Chatbot", "Fecha de emisión", ISSUE_DATE],
                ["Dueño del proceso", "Responsable de onboarding / Simplia", "Próxima revisión", NEXT_REVIEW],
                ["Aprobado por", "Gerencia / Responsable del cliente", "Estado", DOC_STATUS],
            ],
            widths=[1.45, 2.0, 1.35, 1.55],
        )
    )
    story.append(PageBreak())

    story.append(Paragraph("1. Información general del proceso", styles["H1X"]))
    for item in [
        "Objetivo del proceso: recopilar información, accesos y requisitos de la empresa para configurar el chatbot Simplia en canales conversacionales.",
        "Alcance: inicia con la creación o validación de cuentas base y termina con canales probados, credenciales registradas, evidencias completas y aprobación.",
        "Frecuencia: por cada nuevo cliente, nuevo canal o nueva instancia técnica.",
        "Herramientas usadas: Meta Business, Meta for Developers, WhatsApp Business, Railway, n8n, Chatwoot, OpenAI, Google Drive, Gmail, AnyDesk y CRM si aplica.",
    ]:
        story.append(p("• " + item))

    story.append(Paragraph("2. Contexto del proceso", styles["H1X"]))
    story.append(p("El onboarding técnico permite entender qué necesita la empresa, qué canales se conectarán y qué información debe quedar disponible para que el chatbot funcione correctamente."))
    story.append(Paragraph("3. Entradas y salidas del proceso", styles["H1X"]))
    story.append(table([["Proveedor / origen", "Entrada", "Salida", "Cliente / destino"]] + INPUT_OUTPUT_ROWS, widths=[1.25, 1.85, 1.8, 1.25]))
    story.append(Paragraph("4. Mapa del proceso", styles["H1X"]))
    story.append(scaled_image(map_path, max_w=6.6, max_h=2.0))

    story.append(Paragraph("5. Procedimiento operativo estándar (SOP)", styles["H1X"]))
    story.append(table([["Paso", "Actividad", "Responsable", "Criterio de aceptación", "Registro / evidencia"]] + PROCEDURE_ROWS, widths=[0.35, 2.35, 1.1, 1.45, 1.1]))
    story.append(PageBreak())

    for title, rows, widths in [
        ("6. Roles y responsabilidades (RACI)", [["Actividad clave", "R", "A", "C", "I", "Observaciones"]] + RACI_ROWS, [1.6, 0.75, 0.75, 0.75, 0.65, 1.55]),
        ("7. Riesgos, desvíos y controles", [["Riesgo o falla posible", "Causa probable", "Control preventivo", "Control de verificación", "Acción ante desvío"]] + RISK_ROWS, [1.2, 1.2, 1.25, 1.25, 1.1]),
        ("8. Registros e información documentada", [["Registro / evidencia", "Ubicación", "Responsable", "Retención", "Control de acceso / edición"]] + RECORD_ROWS, [1.25, 1.35, 1.15, 1.15, 1.3]),
        ("9. Indicadores del proceso", [["KPI", "Fórmula", "Frecuencia", "Fuente", "Dueño", "Meta / umbral"]] + KPI_ROWS, [1.0, 1.35, 0.85, 0.85, 0.95, 1.15]),
        ("10. Checklist de ejecución", [["#", "Verificación", "Resultado", "Observaciones"]] + CHECKLIST_ROWS, [0.35, 3.6, 1.15, 1.0]),
    ]:
        story.append(Paragraph(title, styles["H1X"]))
        story.append(table(rows, widths=widths))

    story.append(Paragraph("11. Registro controlado de credenciales y datos requeridos", styles["H1X"]))
    story.append(p("Completar los campos solo en una copia controlada y autorizada. No compartir este documento con valores reales si no corresponde."))
    story.append(table(CREDENTIAL_ROWS, widths=[1.65, 2.2, 2.45]))

    if table_rows:
        story.append(Paragraph("12. Tabla fuente original de credenciales", styles["H1X"]))
        story.append(table([["Sistema", "Contenido original"]] + table_rows, widths=[1.45, 4.95]))

    story.append(Paragraph("13. Control de cambios", styles["H1X"]))
    story.append(table([["Versión", "Fecha", "Cambio realizado", "Responsable"], ["1.0", ISSUE_DATE, "Transformación del documento de onboarding técnico al formato ISO 9001", "Simplia"]], widths=[0.8, 1.0, 3.7, 1.0]))
    story.append(Paragraph("14. Aprobación final", styles["H1X"]))
    story.append(table([["Elaborado por", "Revisado por", "Aprobado por"], ["Simplia / Responsable de onboarding", "", ""], ["Firma / fecha", "Firma / fecha", "Firma / fecha"]], widths=[2.1, 2.1, 2.1]))
    story.append(PageBreak())

    story.append(Paragraph("Anexo A. Transcripción controlada del contenido fuente", styles["H1X"]))
    story.append(p("Este anexo conserva el contenido operativo del documento original en orden de aparición. La redacción fue normalizada para mejorar lectura sin cambiar el significado."))
    story.append(table([["#", "Contenido"]] + [[str(i), paragraph] for i, paragraph in enumerate(paragraphs, start=1)], widths=[0.45, 5.95]))
    story.append(PageBreak())

    story.append(Paragraph("Anexo B. Evidencias visuales del documento original", styles["H1X"]))
    story.append(p(f"Se conservan {len(images)} capturas/imágenes extraídas del documento original, en el mismo orden de aparición."))
    for index, image_path in enumerate(images, start=1):
        story.append(Paragraph(f"Evidencia visual {index:03d}", styles["CaptionX"]))
        story.append(scaled_image(image_path, max_w=6.55, max_h=8.1))
        story.append(Spacer(1, 0.1 * inch))

    pdf = SimpleDocTemplate(str(PDF_OUT), pagesize=letter, rightMargin=0.5 * inch, leftMargin=0.5 * inch, topMargin=0.52 * inch, bottomMargin=0.52 * inch, title="Onboarding Técnico Simplia Chatbot ISO 9001", author="Simplia")
    pdf.build(story, onFirstPage=header, onLaterPages=header)


def main():
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)
    if not TEMPLATE_DOCX.exists():
        raise FileNotFoundError(TEMPLATE_DOCX)
    ASSET_DIR.mkdir(exist_ok=True)
    source, paragraphs, table_rows = extract_source_content()
    images = extract_images(source)
    map_path = make_process_map()
    build_docx(paragraphs, table_rows, images, map_path)
    build_pdf(paragraphs, table_rows, images, map_path)
    print(f"DOCX={DOCX_OUT}")
    print(f"PDF={PDF_OUT}")
    print(f"ASSETS={ASSET_DIR}")
    print(f"PARAGRAPHS={len(paragraphs)}")
    print(f"IMAGES={len(images)}")


if __name__ == "__main__":
    main()
