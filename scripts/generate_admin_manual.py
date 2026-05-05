from __future__ import annotations

import re
import textwrap
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    PageBreak,
    Paragraph,
    Preformatted,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCX_PATH = ROOT / "Manual_Admin_SimpliaLeads_ISO10013.docx"
PDF_PATH = ROOT / "Manual_Admin_SimpliaLeads_ISO10013.pdf"
LOGO_PATH = ROOT / "logo_simplia.png"
MIGRATIONS_DIR = ROOT / "supabase" / "migrations"

DOC_CODE = "MAN-ADM-SIMPLIA-001"
DOC_VERSION = "1.1"
DOC_DATE = "04/05/2026"
NEXT_REVIEW = "04/11/2026"
DOC_STATUS = "Borrador tecnico / Vigente al aprobarse"
TIMEZONE = "America/Guayaquil"

COLORS = {
    "navy": "1f3f91",
    "blue": "274690",
    "light_blue": "eaf1ff",
    "green": "0a9b6f",
    "mint": "dff7ec",
    "orange": "f59e0b",
    "red": "ef4444",
    "slate": "64748b",
    "ink": "0f2344",
    "line": "d9e2ef",
    "bg": "f8fafc",
    "white": "ffffff",
}


def rgb(hex_color: str) -> RGBColor:
    h = hex_color.lstrip("#")
    return RGBColor(int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16))


def escape_html(value: str) -> str:
    return (
        value.replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def sanitize_sql(value: str) -> str:
    value = re.sub(r"https://[a-z0-9]{20}\.supabase\.co", "https://<PROJECT_REF>.supabase.co", value)
    value = re.sub(r"Bearer\s+[A-Za-z0-9_.-]{20,}", "Bearer <JWT_DE_SUPABASE>", value)
    return value


def safe_sheet_text(value: str, limit: int = 1200) -> str:
    value = " ".join(str(value).split())
    return value if len(value) <= limit else value[: limit - 3] + "..."


@dataclass
class Event:
    kind: str
    payload: tuple


class Manual:
    def __init__(self) -> None:
        self.events: list[Event] = []

    def h(self, level: int, text: str) -> None:
        self.events.append(Event("h", (level, text)))

    def p(self, text: str) -> None:
        self.events.append(Event("p", (text,)))

    def bullets(self, items: Iterable[str]) -> None:
        self.events.append(Event("bullets", (list(items),)))

    def table(self, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
        self.events.append(Event("table", (headers, rows, widths)))

    def code(self, title: str, code: str) -> None:
        self.events.append(Event("code", (title, sanitize_sql(code).strip())))

    def page_break(self) -> None:
        self.events.append(Event("page_break", ()))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: str | None = None, size: int = 9) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(str(text))
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_docx_table(table, header_fill: str = COLORS["blue"]) -> None:
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8 if row_index else 8.5)
            if row_index == 0:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = rgb("ffffff")
            else:
                set_cell_shading(cell, "ffffff" if row_index % 2 else "f8fafc")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_docx_header(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.58)
    section.right_margin = Inches(0.58)
    header = section.header
    table = header.add_table(rows=1, cols=3, width=Inches(7.3))
    table.autofit = True
    if LOGO_PATH.exists():
        run = table.cell(0, 0).paragraphs[0].add_run()
        run.add_picture(str(LOGO_PATH), width=Inches(1.2))
    else:
        set_cell_text(table.cell(0, 0), "SIMPLIA", True, COLORS["blue"], 13)
    set_cell_text(table.cell(0, 1), "Manual Administrativo SimpliaLeads", True, COLORS["ink"], 9)
    set_cell_text(table.cell(0, 2), f"{DOC_CODE}\nVersion {DOC_VERSION}", True, COLORS["slate"], 8)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"{DOC_CODE} | Version {DOC_VERSION} | Documento controlado | Pagina ")
    run.font.size = Pt(8)


def add_docx_title(document: Document) -> None:
    if LOGO_PATH.exists():
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.55))
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Manual Administrativo SimpliaLeads")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = rgb(COLORS["blue"])
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Informacion documentada controlada para operacion, mantenimiento y replicacion")
    run.font.size = Pt(10)
    run.font.color.rgb = rgb(COLORS["slate"])


def render_docx(manual: Manual) -> None:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    for style_name, size, color in [
        ("Heading 1", 16, COLORS["blue"]),
        ("Heading 2", 13, COLORS["ink"]),
        ("Heading 3", 11, COLORS["green"]),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
    code_style = styles.add_style("ManualCode", 1)
    code_style.font.name = "Courier New"
    code_style.font.size = Pt(7)
    code_style.paragraph_format.space_after = Pt(0)

    add_docx_header(document)
    add_docx_title(document)

    for event in manual.events:
        if event.kind == "h":
            level, text = event.payload
            document.add_heading(text, level=level)
        elif event.kind == "p":
            (text,) = event.payload
            p = document.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            run = p.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(9.5)
        elif event.kind == "bullets":
            (items,) = event.payload
            for item in items:
                p = document.add_paragraph(style="List Bullet")
                p.paragraph_format.space_after = Pt(1)
                p.add_run(str(item)).font.size = Pt(9)
        elif event.kind == "table":
            headers, rows, widths = event.payload
            table = document.add_table(rows=1, cols=len(headers))
            for idx, header in enumerate(headers):
                set_cell_text(table.cell(0, idx), header, True, "ffffff", 8.5)
            set_repeat_table_header(table.rows[0])
            for row in rows:
                cells = table.add_row().cells
                for idx, cell_value in enumerate(row):
                    set_cell_text(cells[idx], safe_sheet_text(cell_value, 1500), False, COLORS["ink"], 8)
            style_docx_table(table)
            if widths:
                for row in table.rows:
                    for idx, width in enumerate(widths):
                        row.cells[idx].width = Inches(width)
            document.add_paragraph()
        elif event.kind == "code":
            title, code = event.payload
            p = document.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            title_run = p.add_run(title)
            title_run.bold = True
            title_run.font.size = Pt(9)
            title_run.font.color.rgb = rgb(COLORS["blue"])
            for line in code.splitlines():
                wrapped = textwrap.wrap(line, width=118, replace_whitespace=False, drop_whitespace=False) or [""]
                for piece in wrapped:
                    cp = document.add_paragraph(style="ManualCode")
                    cp.paragraph_format.left_indent = Inches(0.12)
                    cp.paragraph_format.right_indent = Inches(0.12)
                    run = cp.add_run(piece)
                    run.font.name = "Courier New"
                    run.font.size = Pt(6.7)
        elif event.kind == "page_break":
            document.add_page_break()

    document.save(DOCX_PATH)


def pdf_table(headers: list[str], rows: list[list[str]], widths: list[float] | None, styles) -> Table:
    data = [[Paragraph(f"<b>{escape_html(h)}</b>", styles["TableHeader"]) for h in headers]]
    for row in rows:
        data.append([Paragraph(escape_html(safe_sheet_text(str(cell), 900)), styles["TableCell"]) for cell in row])
    if widths:
        col_widths = [w * inch for w in widths]
    else:
        col_widths = None
    table = Table(data, colWidths=col_widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#274690")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#d9e2ef")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def render_pdf(manual: Manual) -> None:
    base_styles = getSampleStyleSheet()
    styles = {
        "Title": ParagraphStyle(
            "ManualTitle",
            parent=base_styles["Title"],
            fontName="Helvetica-Bold",
            fontSize=20,
            textColor=colors.HexColor("#274690"),
            alignment=TA_CENTER,
            spaceAfter=10,
        ),
        "H1": ParagraphStyle(
            "ManualH1",
            parent=base_styles["Heading1"],
            fontSize=14,
            textColor=colors.HexColor("#274690"),
            spaceBefore=12,
            spaceAfter=7,
        ),
        "H2": ParagraphStyle(
            "ManualH2",
            parent=base_styles["Heading2"],
            fontSize=11.5,
            textColor=colors.HexColor("#0f2344"),
            spaceBefore=9,
            spaceAfter=5,
        ),
        "H3": ParagraphStyle(
            "ManualH3",
            parent=base_styles["Heading3"],
            fontSize=10,
            textColor=colors.HexColor("#0a9b6f"),
            spaceBefore=7,
            spaceAfter=4,
        ),
        "Body": ParagraphStyle(
            "ManualBody",
            parent=base_styles["BodyText"],
            fontSize=8.7,
            leading=11,
            textColor=colors.HexColor("#0f2344"),
            spaceAfter=5,
        ),
        "Bullet": ParagraphStyle(
            "ManualBullet",
            parent=base_styles["BodyText"],
            fontSize=8.5,
            leading=10.8,
            leftIndent=13,
            firstLineIndent=-8,
            textColor=colors.HexColor("#0f2344"),
            spaceAfter=2,
        ),
        "Code": ParagraphStyle(
            "ManualCodePdf",
            parent=base_styles["Code"],
            fontName="Courier",
            fontSize=5.8,
            leading=7,
            textColor=colors.HexColor("#0f2344"),
            backColor=colors.HexColor("#f8fafc"),
            borderColor=colors.HexColor("#d9e2ef"),
            borderWidth=0.35,
            borderPadding=4,
            spaceAfter=5,
        ),
        "TableCell": ParagraphStyle(
            "TableCell",
            parent=base_styles["BodyText"],
            fontSize=6.8,
            leading=8.2,
            textColor=colors.HexColor("#0f2344"),
        ),
        "TableHeader": ParagraphStyle(
            "TableHeader",
            parent=base_styles["BodyText"],
            fontSize=6.8,
            leading=8.2,
            textColor=colors.white,
        ),
    }

    def header_footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 7)
        canvas.setFillColor(colors.HexColor("#64748b"))
        canvas.drawString(0.55 * inch, 0.34 * inch, f"{DOC_CODE} | Version {DOC_VERSION} | Documento controlado")
        canvas.drawRightString(7.95 * inch, 0.34 * inch, f"Pagina {doc.page}")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.5 * inch,
        leftMargin=0.5 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
    )
    story = []
    if LOGO_PATH.exists():
        logo = RLImage(str(LOGO_PATH), width=1.45 * inch, height=0.55 * inch)
        logo.hAlign = "CENTER"
        story.append(logo)
        story.append(Spacer(1, 6))
    story.append(Paragraph("Manual Administrativo SimpliaLeads", styles["Title"]))
    story.append(
        Paragraph(
            "Informacion documentada controlada para operacion, mantenimiento y replicacion",
            ParagraphStyle("Subtitle", parent=styles["Body"], alignment=TA_CENTER, textColor=colors.HexColor("#64748b")),
        )
    )
    story.append(Spacer(1, 10))

    for event in manual.events:
        if event.kind == "h":
            level, text = event.payload
            story.append(Paragraph(escape_html(text), styles[f"H{min(level, 3)}"]))
        elif event.kind == "p":
            (text,) = event.payload
            story.append(Paragraph(escape_html(text), styles["Body"]))
        elif event.kind == "bullets":
            (items,) = event.payload
            for item in items:
                story.append(Paragraph("- " + escape_html(str(item)), styles["Bullet"]))
        elif event.kind == "table":
            headers, rows, widths = event.payload
            story.append(pdf_table(headers, rows, widths, styles))
            story.append(Spacer(1, 7))
        elif event.kind == "code":
            title, code = event.payload
            story.append(Paragraph(f"<b>{escape_html(title)}</b>", styles["Body"]))
            wrapped_lines: list[str] = []
            for line in code.splitlines():
                wrapped_lines.extend(textwrap.wrap(line, width=105, replace_whitespace=False, drop_whitespace=False) or [""])
            for i in range(0, len(wrapped_lines), 54):
                story.append(Preformatted("\n".join(wrapped_lines[i : i + 54]), styles["Code"]))
        elif event.kind == "page_break":
            story.append(PageBreak())

    doc.build(story, onFirstPage=header_footer, onLaterPages=header_footer)


BASE_SCHEMA_SQL = r"""
-- =============================================================
-- BASE DESDE CERO - SCHEMA CW PARA SIMPLIALEADS
-- Ejecutar en Supabase SQL Editor antes de desplegar funciones.
-- Sustituir valores entre <...> por los datos del nuevo proyecto.
-- =============================================================

create extension if not exists pgcrypto with schema extensions;
create extension if not exists pg_net with schema extensions;
create extension if not exists supabase_vault with schema vault;
create extension if not exists pg_cron with schema pg_catalog;

create schema if not exists cw;

create table if not exists cw.account_config (
    id uuid primary key default gen_random_uuid(),
    company_name text not null,
    chatwoot_base_url text not null,
    chatwoot_account_id bigint not null unique,
    timezone text not null default 'America/Guayaquil',
    notes text,
    created_at timestamptz not null default now(),
    updated_at timestamptz not null default now()
);

create table if not exists cw.sync_cursor (
    cursor_name text primary key,
    last_since_ts timestamptz,
    last_until_ts timestamptz,
    cursor_payload jsonb not null default '{}'::jsonb,
    updated_at timestamptz not null default now()
);

create table if not exists cw.sync_runs (
    id bigserial primary key,
    sync_type text not null,
    started_at timestamptz not null default now(),
    finished_at timestamptz,
    status text not null default 'running'
        check (status in ('running', 'success', 'partial', 'error')),
    stats jsonb not null default '{}'::jsonb,
    error_message text,
    created_at timestamptz not null default now()
);

create table if not exists cw.raw_ingest (
    id bigserial primary key,
    source_type text not null check (source_type in ('api', 'webhook', 'manual', 'repair')),
    endpoint_name text not null,
    event_name text,
    entity_type text,
    chatwoot_entity_id bigint,
    payload jsonb not null,
    fetched_at timestamptz not null default now(),
    processed boolean not null default false,
    processing_error text
);

create table if not exists cw.inboxes (
    chatwoot_inbox_id bigint primary key,
    name text,
    website_url text,
    channel_type text,
    avatar_url text,
    widget_color text,
    website_token text,
    enable_auto_assignment boolean,
    web_widget_script text,
    welcome_title text,
    welcome_tagline text,
    greeting_enabled boolean,
    raw_payload jsonb not null default '{}'::jsonb,
    created_at timestamptz not null default now(),
    updated_at timestamptz not null default now()
);

create table if not exists cw.teams (
    chatwoot_team_id bigint primary key,
    name text,
    description text,
    allow_auto_assign boolean,
    is_member boolean,
    raw_payload jsonb not null default '{}'::jsonb,
    created_at timestamptz not null default now(),
    updated_at timestamptz not null default now()
);

create table if not exists cw.attribute_definitions (
    id bigserial primary key,
    chatwoot_attribute_id bigint not null unique,
    attribute_scope text not null check (attribute_scope in ('contact', 'conversation', 'unknown')),
    attribute_key text not null,
    attribute_display_name text,
    attribute_display_type text,
    attribute_description text,
    regex_pattern text,
    regex_cue text,
    attribute_values jsonb,
    attribute_model text,
    default_value jsonb,
    created_at_chatwoot timestamptz,
    updated_at_chatwoot timestamptz,
    raw_payload jsonb not null default '{}'::jsonb,
    created_at timestamptz not null default now(),
    updated_at timestamptz not null default now()
);

create table if not exists cw.contacts_current (
    chatwoot_contact_id bigint primary key,
    lead_identity_key text,
    identifier text,
    name text,
    phone_number text,
    email text,
    blocked boolean,
    thumbnail text,
    availability_status text,
    additional_attributes jsonb not null default '{}'::jsonb,
    custom_attributes jsonb not null default '{}'::jsonb,
    created_at_chatwoot timestamptz,
    last_activity_at_chatwoot timestamptz,
    first_seen_at timestamptz not null default now(),
    last_seen_at timestamptz not null default now(),
    raw_payload jsonb not null default '{}'::jsonb,
    created_at timestamptz not null default now(),
    updated_at timestamptz not null default now()
);

create table if not exists cw.contact_inboxes (
    id bigserial primary key,
    chatwoot_contact_id bigint not null references cw.contacts_current(chatwoot_contact_id) on delete cascade,
    chatwoot_inbox_id bigint,
    source_id text,
    inbox_name text,
    channel_type text,
    provider text,
    raw_payload jsonb not null default '{}'::jsonb,
    created_at timestamptz not null default now(),
    updated_at timestamptz not null default now()
);

create table if not exists cw.conversations_current (
    chatwoot_conversation_id bigint primary key,
    chatwoot_contact_id bigint references cw.contacts_current(chatwoot_contact_id) on delete set null,
    chatwoot_account_id bigint,
    chatwoot_inbox_id bigint,
    chatwoot_team_id bigint,
    assignee_id bigint,
    uuid text,
    status text,
    priority text,
    can_reply boolean,
    muted boolean,
    snoozed_until timestamptz,
    unread_count integer,
    labels text[] not null default '{}'::text[],
    business_stage_current text,
    additional_attributes jsonb not null default '{}'::jsonb,
    custom_attributes jsonb not null default '{}'::jsonb,
    conversation_custom_attributes jsonb not null default '{}'::jsonb,
    contact_custom_attributes jsonb not null default '{}'::jsonb,
    meta jsonb not null default '{}'::jsonb,
    applied_sla jsonb not null default '{}'::jsonb,
    sla_events jsonb not null default '[]'::jsonb,
    inbox_name text,
    channel_type text,
    provider text,
    conversation_url text,
    timestamp_text text,
    first_reply_created_at_chatwoot timestamptz,
    agent_last_seen_at_chatwoot timestamptz,
    assignee_last_seen_at_chatwoot timestamptz,
    contact_last_seen_at_chatwoot timestamptz,
    waiting_since_chatwoot timestamptz,
    last_activity_at_chatwoot timestamptz,
    created_at_chatwoot timestamptz,
    updated_at_chatwoot timestamptz,
    last_non_activity_message_id bigint,
    last_non_activity_message_preview text,
    first_message_at timestamptz,
    last_message_at timestamptz,
    last_incoming_message_at timestamptz,
    last_outgoing_message_at timestamptz,
    total_messages integer not null default 0,
    nombre_completo text,
    fecha_visita text,
    hora_visita text,
    agencia text,
    celular text,
    correo text,
    campana text,
    ciudad text,
    edad text,
    canal text,
    agente boolean,
    score_interes numeric,
    monto_operacion text,
    fecha_monto_operacion timestamptz,
    perfil_url text,
    raw_payload jsonb not null default '{}'::jsonb,
    created_at timestamptz not null default now(),
    updated_at timestamptz not null default now()
);

create table if not exists cw.contact_attribute_history (
    id bigserial primary key,
    chatwoot_contact_id bigint not null references cw.contacts_current(chatwoot_contact_id) on delete cascade,
    attribute_key text not null,
    old_value jsonb,
    new_value jsonb,
    changed_at timestamptz not null default now(),
    change_source text not null default 'sync' check (change_source in ('sync', 'webhook', 'repair', 'manual')),
    created_at timestamptz not null default now(),
    event_key text
);

create table if not exists cw.conversation_attribute_history (
    id bigserial primary key,
    chatwoot_conversation_id bigint not null references cw.conversations_current(chatwoot_conversation_id) on delete cascade,
    attribute_key text not null,
    old_value jsonb,
    new_value jsonb,
    changed_at timestamptz not null default now(),
    change_source text not null default 'sync' check (change_source in ('sync', 'webhook', 'repair', 'manual')),
    created_at timestamptz not null default now(),
    event_key text
);

create table if not exists cw.conversation_label_history (
    id bigserial primary key,
    chatwoot_conversation_id bigint not null references cw.conversations_current(chatwoot_conversation_id) on delete cascade,
    label text not null,
    action text not null check (action in ('added', 'removed')),
    changed_at timestamptz not null default now(),
    change_source text not null default 'sync' check (change_source in ('sync', 'webhook', 'repair', 'manual')),
    created_at timestamptz not null default now()
);

create table if not exists cw.business_stage_history (
    id bigserial primary key,
    chatwoot_conversation_id bigint not null references cw.conversations_current(chatwoot_conversation_id) on delete cascade,
    old_stage text,
    new_stage text,
    changed_at timestamptz not null default now(),
    change_reason text,
    change_source text not null default 'sync'
        check (change_source in ('sync', 'webhook', 'repair', 'manual', 'bot')),
    created_at timestamptz not null default now()
);

create table if not exists cw.messages (
    chatwoot_message_id bigint primary key,
    chatwoot_conversation_id bigint not null references cw.conversations_current(chatwoot_conversation_id) on delete cascade,
    chatwoot_contact_id bigint,
    chatwoot_account_id bigint,
    chatwoot_inbox_id bigint,
    sender_id bigint,
    sender_type text,
    message_type text,
    message_direction text not null default 'unknown'
        check (message_direction in ('incoming', 'outgoing', 'activity', 'note', 'unknown')),
    content text,
    content_type text,
    content_attributes jsonb not null default '{}'::jsonb,
    additional_attributes jsonb not null default '{}'::jsonb,
    external_source_ids jsonb not null default '{}'::jsonb,
    attachments jsonb not null default '[]'::jsonb,
    sender jsonb not null default '{}'::jsonb,
    sentiment jsonb not null default '{}'::jsonb,
    processed_message_content text,
    source_id text,
    status text,
    is_private boolean not null default false,
    created_at_chatwoot timestamptz,
    updated_at_chatwoot timestamptz,
    raw_payload jsonb not null default '{}'::jsonb,
    created_at timestamptz not null default now()
);

create table if not exists cw.reporting_events (
    chatwoot_reporting_event_id bigint primary key,
    name text,
    value numeric,
    value_in_business_hours numeric,
    event_start_time timestamptz,
    event_end_time timestamptz,
    chatwoot_account_id bigint,
    chatwoot_conversation_id bigint,
    chatwoot_inbox_id bigint,
    chatwoot_user_id bigint,
    created_at_chatwoot timestamptz,
    updated_at_chatwoot timestamptz,
    raw_payload jsonb not null default '{}'::jsonb,
    created_at timestamptz not null default now()
);

create table if not exists cw.daily_metrics (
    id bigserial primary key,
    metric_date date not null,
    metric_scope text not null
        check (metric_scope in ('summary', 'activities', 'unique', 'stage', 'label', 'channel', 'inbox', 'team', 'sla', 'custom')),
    metric_name text not null,
    metric_value numeric not null,
    dimensions jsonb not null default '{}'::jsonb,
    calculated_at timestamptz not null default now(),
    dim_hash text generated always as (md5((dimensions)::text)) stored
);

create table if not exists cw.automated_reports (
    id uuid primary key default gen_random_uuid(),
    name text not null,
    frequency text not null default 'weekly',
    schedule_days text[] not null default '{}'::text[],
    schedule_month_day integer,
    schedule_time time not null default '08:00:00',
    recipients text not null,
    is_active boolean default true,
    last_run_at timestamptz,
    report_scope text not null default 'tab' check (report_scope in ('tab', 'critical_profile')),
    tab_ids text[] not null default '{}'::text[],
    critical_profile_key text,
    file_formats text[] not null default array['excel']::text[],
    date_range_mode text not null default 'closed_period' check (date_range_mode in ('closed_period', 'current_filters')),
    filters jsonb not null default '{}'::jsonb,
    created_by uuid references auth.users(id) on delete set null,
    created_by_email text,
    last_status text,
    last_error text,
    created_at timestamptz default now(),
    updated_at timestamptz default now()
);

create table if not exists cw.dashboard_tag_settings (
    id uuid primary key default gen_random_uuid(),
    account_id bigint unique,
    settings jsonb not null default '{}'::jsonb,
    updated_at timestamptz default now()
);

create table if not exists cw.conversation_label_events (
    id bigint generated by default as identity primary key,
    chatwoot_conversation_id bigint not null,
    previous_labels text[] not null default '{}'::text[],
    next_labels text[] not null default '{}'::text[],
    added_labels text[] not null default '{}'::text[],
    removed_labels text[] not null default '{}'::text[],
    event_source text not null default 'sync_diff'
        check (event_source in ('dashboard', 'webhook', 'sync_diff', 'repair')),
    occurred_at timestamptz not null default now(),
    detected_at timestamptz not null default now(),
    raw_payload jsonb not null default '{}'::jsonb,
    event_key text unique,
    created_at timestamptz not null default now()
);

create table if not exists cw.automated_report_runs (
    id uuid primary key default gen_random_uuid(),
    automated_report_id uuid references cw.automated_reports(id) on delete set null,
    status text not null default 'running' check (status in ('running', 'success', 'error', 'skipped')),
    recipients text,
    file_formats text[] not null default '{}'::text[],
    report_scope text,
    tab_ids text[] not null default '{}'::text[],
    critical_profile_key text,
    scheduled_for timestamptz,
    started_at timestamptz not null default now(),
    finished_at timestamptz,
    error_message text,
    metadata jsonb not null default '{}'::jsonb
);

create index if not exists idx_cw_conversations_created_at on cw.conversations_current (created_at_chatwoot desc);
create index if not exists idx_cw_conversations_last_activity on cw.conversations_current (last_activity_at_chatwoot desc);
create index if not exists idx_cw_conversations_labels_gin on cw.conversations_current using gin (labels);
create index if not exists idx_cw_conversations_attrs_gin on cw.conversations_current using gin (custom_attributes);
create index if not exists idx_cw_messages_created_at on cw.messages (created_at_chatwoot desc);
create index if not exists idx_cw_messages_conversation on cw.messages (chatwoot_conversation_id, created_at_chatwoot);
create unique index if not exists idx_cw_daily_metrics_unique on cw.daily_metrics (metric_date, metric_scope, metric_name, dim_hash);
create index if not exists automated_reports_active_schedule_idx on cw.automated_reports (is_active, frequency, schedule_time);
create index if not exists automated_report_runs_report_started_idx on cw.automated_report_runs (automated_report_id, started_at desc);

alter table cw.account_config enable row level security;
alter table cw.sync_cursor enable row level security;
alter table cw.sync_runs enable row level security;
alter table cw.raw_ingest enable row level security;
alter table cw.inboxes enable row level security;
alter table cw.teams enable row level security;
alter table cw.attribute_definitions enable row level security;
alter table cw.contacts_current enable row level security;
alter table cw.contact_inboxes enable row level security;
alter table cw.conversations_current enable row level security;
alter table cw.contact_attribute_history enable row level security;
alter table cw.conversation_attribute_history enable row level security;
alter table cw.conversation_label_history enable row level security;
alter table cw.business_stage_history enable row level security;
alter table cw.messages enable row level security;
alter table cw.reporting_events enable row level security;
alter table cw.daily_metrics enable row level security;
alter table cw.automated_reports enable row level security;
alter table cw.dashboard_tag_settings enable row level security;
alter table cw.conversation_label_events enable row level security;
alter table cw.automated_report_runs enable row level security;

grant usage on schema cw to authenticated, service_role;
grant select, insert, update, delete on all tables in schema cw to authenticated;
grant all on all tables in schema cw to service_role;
grant usage, select on all sequences in schema cw to authenticated, service_role;

-- Politicas base. Para datos sensibles, endurecer por tenant antes de multiempresa real.
do $$
declare
    t text;
begin
    foreach t in array array[
        'account_config','sync_cursor','sync_runs','inboxes','teams','attribute_definitions',
        'contacts_current','contact_inboxes','conversations_current','messages','reporting_events',
        'daily_metrics','automated_reports','dashboard_tag_settings','conversation_label_events'
    ] loop
        if not exists (
            select 1 from pg_policies
            where schemaname = 'cw' and tablename = t and policyname = 'authenticated_all'
        ) then
            execute format(
                'create policy authenticated_all on cw.%I for all to authenticated using (true) with check (true)',
                t
            );
        end if;
    end loop;
end $$;
"""


VALIDATION_SQL = r"""
-- Confirmar tablas principales del schema cw.
select table_name
from information_schema.tables
where table_schema = 'cw'
order by table_name;

-- Confirmar que el dashboard tiene tablas base creadas.
select
    to_regclass('cw.conversations_current') as conversations_current,
    to_regclass('cw.messages') as messages,
    to_regclass('cw.dashboard_tag_settings') as dashboard_tag_settings,
    to_regclass('cw.automated_reports') as automated_reports,
    to_regclass('cw.automated_report_runs') as automated_report_runs;

-- Confirmar extensiones requeridas.
select extname
from pg_extension
where extname in ('pgcrypto', 'pg_net', 'pg_cron', 'supabase_vault')
order by extname;

-- Confirmar cron configurado.
select jobid, jobname, schedule, active
from cron.job
where jobname in ('sync-chatwoot-diario', 'send-scheduled-reports')
order by jobname;

-- Confirmar secretos de Vault sin revelar valores.
select name, created_at
from vault.secrets
where name in ('chatwoot_sync_project_url', 'chatwoot_sync_jwt')
order by name;

-- Confirmar ultima sincronizacion.
select id, sync_type, status, started_at, finished_at, stats, error_message
from cw.sync_runs
order by started_at desc
limit 5;
"""


VERCEL_ENV_EXAMPLE = r"""
# Variables de entorno en Vercel - Project Settings > Environment Variables
VITE_SUPABASE_URL=https://<PROJECT_REF>.supabase.co
VITE_SUPABASE_ANON_KEY=<SUPABASE_ANON_KEY>
VITE_CHATWOOT_BASE_URL=https://app.chatwoot.com
VITE_CHATWOOT_ACCOUNT_ID=<CHATWOOT_ACCOUNT_ID>
VITE_CHATWOOT_API_TOKEN=<CHATWOOT_API_TOKEN>
"""


SUPABASE_SECRETS_EXAMPLE = r"""
npx supabase secrets set SUPABASE_URL=https://<PROJECT_REF>.supabase.co
npx supabase secrets set SUPABASE_SERVICE_ROLE_KEY=<SUPABASE_SERVICE_ROLE_KEY>
npx supabase secrets set VITE_CHATWOOT_BASE_URL=https://app.chatwoot.com
npx supabase secrets set VITE_CHATWOOT_ACCOUNT_ID=<CHATWOOT_ACCOUNT_ID>
npx supabase secrets set VITE_CHATWOOT_API_TOKEN=<CHATWOOT_API_TOKEN>
npx supabase secrets set CHATWOOT_WEBHOOK_SECRET=<WEBHOOK_SECRET_GENERADO>
npx supabase secrets set RESEND_API_KEY=<RESEND_API_KEY>
npx supabase secrets set RESEND_FROM_EMAIL="Simplia Reportes <reportes@empresa.com>"
"""


DEPLOY_COMMANDS = r"""
npx supabase login
npx supabase link --project-ref <PROJECT_REF>
npx supabase functions deploy chatwoot-sync
npx supabase functions deploy chatwoot-label-webhook
npx supabase functions deploy chatwoot-repair-conversations
npx supabase functions deploy send-scheduled-reports
"""


INITIAL_SYNC_COMMANDS = r"""
npx supabase functions invoke chatwoot-sync --body '{"mode":"full","window_hours":72,"sync_messages":"recent"}'
npx supabase functions invoke chatwoot-repair-conversations --body '{"conversation_ids":[],"repair_messages":true}'
npx supabase functions invoke send-scheduled-reports --body '{"source":"manual-test","dry_run":true}'
"""


TABLE_CATALOG = [
    ["cw.account_config", "Datos generales de la empresa y cuenta Chatwoot.", "Configuracion inicial / admin", "Funciones y validaciones", "Replicacion y soporte"],
    ["cw.sync_cursor", "Guarda el punto hasta donde llego la sincronizacion.", "chatwoot-sync", "chatwoot-sync", "Evita duplicar ventanas de lectura"],
    ["cw.sync_runs", "Bitacora de cada sincronizacion.", "chatwoot-sync", "Admin / soporte / sync_health", "Evidencia de ejecucion"],
    ["cw.raw_ingest", "Payload crudo recibido de API, webhook o reparacion.", "Edge Functions", "Soporte tecnico", "Trazabilidad"],
    ["cw.inboxes", "Canales/inboxes detectados en Chatwoot.", "chatwoot-sync", "Dashboard y reportes", "Filtros por canal"],
    ["cw.teams", "Equipos de Chatwoot.", "chatwoot-sync", "Operacion / responsables", "Carga operativa"],
    ["cw.attribute_definitions", "Definiciones de atributos personalizados.", "chatwoot-sync", "Admin calidad / reportes", "Campos configurables"],
    ["cw.contacts_current", "Estado actual de contactos.", "chatwoot-sync / repair", "Dashboard / conversaciones", "Datos actuales del lead"],
    ["cw.contact_inboxes", "Relacion contacto-inbox.", "chatwoot-sync", "Canales", "Origen del lead"],
    ["cw.conversations_current", "Estado actual de conversaciones y campos comerciales.", "chatwoot-sync / webhook / repair", "Todas las pestanas", "Fuente central de KPIs"],
    ["cw.contact_attribute_history", "Cambios historicos de atributos de contacto.", "chatwoot-sync / repair", "Cambios relevantes", "Explicacion de cambios"],
    ["cw.conversation_attribute_history", "Cambios historicos de atributos de conversacion.", "chatwoot-sync / repair", "Cambios relevantes", "Montos, score, campana"],
    ["cw.conversation_label_history", "Historial simple de etiquetas agregadas o retiradas.", "sync / webhook", "Embudo y soporte", "Evolucion de etapas"],
    ["cw.business_stage_history", "Historial de etapa comercial.", "sync / bot / repair", "Embudo / rendimiento", "Conversion entre estados"],
    ["cw.messages", "Mensajes normalizados por conversacion.", "chatwoot-sync / repair", "Conversaciones, operacion", "Respuesta, ultima interaccion"],
    ["cw.reporting_events", "Eventos de reporteria de Chatwoot.", "chatwoot-sync", "Operacion", "Tiempos y eventos"],
    ["cw.daily_metrics", "Metricas historicas diarias calculadas.", "sync nocturno", "Tendencias", "Historico cerrado"],
    ["cw.automated_reports", "Configuracion de reportes programados.", "Panel Reportes", "send-scheduled-reports", "Correos automaticos"],
    ["cw.dashboard_tag_settings", "Configuracion admin de etiquetas, calidad y columnas.", "Panel Admin", "Dashboard/reportes", "Reglas de negocio"],
    ["cw.conversation_label_events", "Eventos deduplicados de etiquetas.", "webhook/sync diff", "Cambios relevantes", "Senales comerciales"],
    ["cw.automated_report_runs", "Resultado de cada corrida programada.", "send-scheduled-reports", "Admin/soporte", "Evidencia de envio"],
    ["cw.sync_health", "Vista de salud de sync/cron.", "SQL view", "Soporte", "Diagnostico rapido"],
    ["cw.commercial_audit_events", "Vista de cambios comerciales relevantes.", "SQL view", "Reportes", "Explicacion de diferencias"],
]


EDGE_FUNCTIONS = [
    [
        "chatwoot-sync",
        "Sincroniza conversaciones, contactos, mensajes, inboxes, atributos, etiquetas y metricas desde Chatwoot. Es la funcion principal para cargar historico y refrescar estado actual.",
        "SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY, VITE_CHATWOOT_BASE_URL, VITE_CHATWOOT_ACCOUNT_ID, VITE_CHATWOOT_API_TOKEN",
        "Invocarla manualmente con ventana corta y validar cw.sync_runs en success.",
    ],
    [
        "chatwoot-label-webhook",
        "Recibe eventos de etiquetas desde Chatwoot para registrar cambios comerciales sin esperar al cierre nocturno.",
        "SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY, CHATWOOT_WEBHOOK_SECRET",
        "Enviar webhook de prueba desde Chatwoot y confirmar conversation_label_events.",
    ],
    [
        "chatwoot-repair-conversations",
        "Relee conversaciones puntuales o ventanas de datos para reparar mensajes, atributos o labels que falten.",
        "SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY, CHATWOOT_BASE_URL/VITE_CHATWOOT_BASE_URL, CHATWOOT_ACCOUNT_ID/VITE_CHATWOOT_ACCOUNT_ID, CHATWOOT_API_TOKEN/VITE_CHATWOOT_API_TOKEN",
        "Ejecutar sobre IDs concretos y revisar conversaciones/messages actualizados.",
    ],
    [
        "send-scheduled-reports",
        "Lee reportes programados activos, genera adjuntos Excel/PDF/CSV y los envia por Resend.",
        "SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY, RESEND_API_KEY, RESEND_FROM_EMAIL",
        "Crear reporte de prueba, invocar dry_run y revisar automated_report_runs.",
    ],
]


QUALITY_LOGIC = [
    ["Caliente", "Puntaje mayor o igual a 70", "Prioridad comercial alta. Requiere seguimiento inmediato."],
    ["Tibio", "Puntaje mayor o igual a 45 y menor a 70", "Tiene senales de interes. Requiere nutrir y llevar a cita."],
    ["Frio", "Puntaje menor a 45 o sin puntaje", "Base de clasificacion. Requiere revisar contexto antes de priorizar."],
]


def migration_files() -> list[Path]:
    if not MIGRATIONS_DIR.exists():
        return []
    return sorted(MIGRATIONS_DIR.glob("*.sql"))


def read_migration(path: Path) -> str:
    return sanitize_sql(path.read_text(encoding="utf-8"))


def add_control_documental(manual: Manual) -> None:
    manual.h(1, "1. Control documental")
    manual.table(
        ["Campo", "Valor"],
        [
            ["Codigo documental", DOC_CODE],
            ["Nombre", "Manual Administrativo SimpliaLeads"],
            ["Version", DOC_VERSION],
            ["Fecha de emision", DOC_DATE],
            ["Estado", DOC_STATUS],
            ["Dueno del documento", "Simplia - Area de tecnologia / administracion del producto"],
            ["Aprobador", "Direccion Simplia o responsable autorizado del proyecto"],
            ["Proxima revision", NEXT_REVIEW],
            ["Zona horaria", TIMEZONE],
            ["Distribucion", "Administradores internos, desarrolladores, soporte tecnico y responsables de implementacion"],
        ],
        [2.0, 4.9],
    )
    manual.table(
        ["Version", "Fecha", "Cambio", "Responsable"],
        [
            ["1.0", "02/05/2026", "Emision inicial del manual administrativo.", "Simplia"],
            [
                "1.1",
                DOC_DATE,
                "Actualizacion de Calidad a 3 niveles, guia completa de replicacion, Vercel, Supabase, Chatwoot, Resend y anexos SQL ampliados.",
                "Simplia",
            ],
        ],
        [0.8, 1.2, 4.0, 1.2],
    )


def add_intro(manual: Manual) -> None:
    manual.h(1, "2. Objetivo, alcance y uso del manual")
    manual.p(
        "Este manual define como administrar, mantener, desplegar y replicar SimpliaLeads. "
        "Esta orientado a personal interno de Simplia, administradores de cuenta, soporte tecnico y desarrolladores que deban dejar funcionando el dashboard para una empresa."
    )
    manual.table(
        ["Incluye", "No incluye"],
        [
            [
                "Frontend React/Vite, Supabase schema cw, Edge Functions, Chatwoot, Resend, Vercel, reportes, cron, configuracion admin, validaciones y troubleshooting.",
                "Capacitacion comercial para usuario final, secretos reales, credenciales privadas o decisiones comerciales especificas de una empresa.",
            ],
        ],
        [3.45, 3.45],
    )
    manual.bullets(
        [
            "Usar este documento antes de crear un nuevo entorno, despues de un despliegue y durante soporte.",
            "Mantenerlo actualizado cuando cambien etiquetas, funciones, variables, reportes o estructura de Supabase.",
            "No pegar credenciales reales dentro del manual. Usar siempre placeholders entre <...>.",
        ]
    )


def add_architecture(manual: Manual) -> None:
    manual.h(1, "3. Contexto del sistema")
    manual.table(
        ["Componente", "Uso", "Control administrativo"],
        [
            ["Frontend React/Vite", "Dashboard visual: Estrategia, Embudo, Operacion, Seguimiento, Rendimiento Humano, Tendencias, Calidad, Conversaciones y Reportes.", "Vercel, variables VITE_ y rama main."],
            ["Supabase", "Base historica, estado actual, tablas cw, Edge Functions, Vault y cron.", "SQL Editor, funciones, secretos, logs y migraciones."],
            ["Chatwoot", "Fuente de conversaciones, contactos, mensajes, etiquetas, atributos y canales.", "Cuenta, token API, webhooks, atributos personalizados y etiquetas."],
            ["Resend", "Envio de reportes programados por correo.", "Dominio/remitente verificado, API key y FROM email."],
            ["Reportes", "Exportaciones manuales y programadas en Excel, PDF y CSV.", "Configuracion en pestana Reportes y tabla cw.automated_reports."],
        ],
        [1.25, 3.45, 2.2],
    )
    manual.h(2, "3.1 Flujo de datos")
    manual.bullets(
        [
            "Chatwoot conserva la operacion diaria: conversaciones, mensajes, etiquetas y atributos.",
            "El dashboard combina datos en vivo de periodos recientes con historico guardado en Supabase.",
            "Supabase guarda estado actual en tablas *_current y evidencia historica en tablas *_history, raw_ingest, sync_runs y report runs.",
            "El cron nocturno ejecuta la sincronizacion diaria a las 12:01 a. m. Ecuador mediante schedule UTC 1 5 * * *.",
            "Los reportes programados se revisan cada 5 minutos con el cron send-scheduled-reports.",
        ]
    )
    manual.h(2, "3.2 Reglas de estado actual e historico")
    manual.p(
        "Los KPIs del dashboard deben representar lo que existe actualmente en Chatwoot/Supabase. "
        "Los cambios historicos sirven para explicar variaciones relevantes, pero no deben sumar ventas, citas o montos si ya no cumplen la condicion actual."
    )


def add_roles(manual: Manual) -> None:
    manual.h(1, "4. Roles y responsabilidades")
    manual.table(
        ["Rol", "Responsabilidad", "Evidencia esperada"],
        [
            ["Admin Simplia", "Configura etiquetas, rangos, columnas, perfiles y reportes.", "Cambios en panel admin y dashboard_tag_settings."],
            ["Desarrollador", "Mantiene codigo, migraciones, Edge Functions y despliegues.", "Commits, deploys, logs y manual actualizado."],
            ["Admin Supabase", "Gestiona proyecto, secretos, cron, RLS, funciones y backups.", "SQL aplicado, funciones activas, sync_runs."],
            ["Admin Chatwoot", "Crea cuenta, inboxes, atributos, etiquetas, API token y webhooks.", "Configuracion visible en Chatwoot."],
            ["Responsable Resend", "Verifica dominio/remitente y crea API key.", "Email de prueba y corridas success."],
            ["Usuario operativo", "Usa dashboard y reportes sin modificar backend.", "Exportaciones y acciones comerciales."],
        ],
        [1.35, 3.4, 2.15],
    )


def add_admin_config(manual: Manual) -> None:
    manual.h(1, "5. Configuracion admin del dashboard")
    manual.p(
        "El panel admin permite adaptar el dashboard a la forma de trabajo de cada empresa sin cambiar codigo. "
        "La configuracion se guarda en cw.dashboard_tag_settings dentro del objeto settings."
    )
    manual.table(
        ["Configuracion", "Que define", "Impacto"],
        [
            ["Embudo", "Estados/etiquetas comerciales y orden de lectura.", "Embudo, estrategia, seguimiento y reportes."],
            ["Citas", "Estados que cuentan como cita o avance comercial.", "Conversion a cita, rendimiento humano y seguimiento."],
            ["Ventas", "Etiquetas que cuentan como venta actual.", "Ventas, monto, ticket promedio y reportes gerenciales."],
            ["Descalificacion", "Estados que agrupan leads que no aplican.", "Tendencias y motivos de perdida."],
            ["Calidad", "Campo numerico y rangos Caliente/Tibio/Frio.", "KPIs, tabla de leads evaluados, filtros y reportes."],
            ["Reportes", "Columnas, formatos, destinatarios, perfiles y frecuencia.", "Exportaciones manuales y envios automaticos."],
        ],
        [1.4, 3.1, 2.4],
    )
    manual.h(2, "5.1 Calidad de leads - regla vigente")
    manual.p(
        "La clasificacion vigente usa tres niveles. Frio es el nivel base; por eso todo lead sin puntaje o con puntaje menor a Tibio se clasifica como Frio."
    )
    manual.table(["Nivel", "Rango", "Interpretacion"], QUALITY_LOGIC, [1.2, 2.1, 3.6])
    manual.bullets(
        [
            "Inputs visibles en admin: Desde Caliente y Desde Tibio.",
            "Defaults recomendados: Caliente desde 70, Tibio desde 45.",
            "scoreThresholds debe guardar hotMin y warmMin. Si existe coldMin antiguo, se ignora para clasificacion nueva.",
            "Exportes manuales y programados deben resumir solo Caliente, Tibio y Frio.",
        ]
    )


def add_dashboard_logic(manual: Manual) -> None:
    manual.h(1, "6. Logica funcional por pestana")
    manual.table(
        ["Pestana", "Datos principales", "Puntos de control"],
        [
            ["Estrategia", "Total leads, calificados, citas, ventas, conversion y ganancias.", "Debe coincidir con filtros de fecha/canal y reglas admin."],
            ["Embudo", "Etapas actuales, historico, conversiones y perdidas.", "Usa etiquetas configuradas y estado actual."],
            ["Operacion", "Tiempo de respuesta, leads sin respuesta, responsables y carga.", "Lead sin respuesta: ultima interaccion real del cliente sin respuesta posterior."],
            ["Seguimiento", "Pendientes, citas, ventas y acciones proximas.", "Excel de esta pestana conserva formato especifico existente."],
            ["Rendimiento Humano", "Responsables, citas humanas, ventas, conversion, total vendido y ticket promedio.", "Ventas actuales requieren etiqueta vigente y monto mayor a 0."],
            ["Tendencias", "Canales, campanas, ingresos, descalificacion y comportamiento por periodo.", "Descalificacion se presenta en lenguaje simple."],
            ["Calidad", "Caliente, Tibio, Frio, puntaje, campana y leads evaluados.", "Sin puntaje se clasifica como Frio."],
            ["Conversaciones", "Tabla de conversaciones, historial, estado, canal, URL y ultima interaccion.", "Excel conserva formato especifico existente."],
            ["Reportes", "Descargas manuales y reportes programados.", "Usa modelo compartido manual/programado."],
        ],
        [1.25, 3.25, 2.4],
    )


def add_replicacion(manual: Manual) -> None:
    manual.h(1, "7. Replicar SimpliaLeads para otra empresa")
    manual.p(
        "Esta seccion es el procedimiento maestro para levantar un entorno nuevo. Debe ejecutarse en orden y registrar evidencias de cada paso."
    )
    manual.table(
        ["Paso", "Actividad", "Resultado esperado"],
        [
            ["1", "Crear o confirmar repositorio GitHub del proyecto y rama main.", "Codigo fuente disponible para Vercel."],
            ["2", "Crear proyecto Supabase nuevo.", "URL, anon key y service role key disponibles."],
            ["3", "Ejecutar SQL base y migraciones.", "Schema cw creado con tablas, vistas, indices, RLS y grants."],
            ["4", "Desplegar Edge Functions.", "Funciones activas en Supabase."],
            ["5", "Crear/configurar cuenta Chatwoot.", "Account ID, API token, inboxes, etiquetas, atributos y webhook listos."],
            ["6", "Configurar Resend.", "Dominio/remitente verificado y API key guardada."],
            ["7", "Configurar Vercel desde GitHub.", "Dashboard desplegado desde main."],
            ["8", "Agregar variables en Vercel y secretos en Supabase.", "Frontend y Edge Functions con credenciales correctas."],
            ["9", "Ejecutar sync inicial y reparacion si aplica.", "Datos visibles en dashboard."],
            ["10", "Configurar panel admin y reportes.", "Reglas comerciales adaptadas a la empresa."],
        ],
        [0.6, 4.25, 2.05],
    )
    manual.h(2, "7.1 Que se cambia entre empresas")
    manual.bullets(
        [
            "Supabase: proyecto, URL, anon key, service role key, secretos y cron.",
            "Chatwoot: base URL si no es Chatwoot Cloud, account ID, token API, inboxes, atributos, etiquetas y webhook secret.",
            "Resend: API key, dominio, remitente y destinatarios de reportes.",
            "Vercel: variables de entorno y proyecto conectado al repo/branch correcto.",
            "Admin dashboard: estados de embudo, citas, ventas, descalificacion, calidad, columnas y reportes programados.",
        ]
    )


def add_vercel(manual: Manual) -> None:
    manual.h(1, "8. Despliegue en Vercel")
    manual.p(
        "El frontend se despliega en Vercel desde GitHub. Para la operacion actual de Simplia, ingresar a Vercel con Google/Gmail usando simplia.social@gmail.com y seleccionar el proyecto conectado al repositorio."
    )
    manual.table(
        ["Campo", "Configuracion"],
        [
            ["Cuenta de acceso", "Login con Google/Gmail: simplia.social@gmail.com"],
            ["Origen", "GitHub"],
            ["Rama productiva", "main"],
            ["Framework", "Vite"],
            ["Build command", "npm run build"],
            ["Output directory", "dist"],
            ["Archivo de rutas", "vercel.json"],
        ],
        [2.0, 4.9],
    )
    manual.h(2, "8.1 Variables de entorno de Vercel")
    manual.p("Configurar estas variables en Vercel > Project Settings > Environment Variables. No usar valores reales dentro de este manual.")
    manual.code("Ejemplo .env para Vercel", VERCEL_ENV_EXAMPLE)
    manual.h(2, "8.2 Proxy Chatwoot")
    manual.p(
        "El archivo vercel.json reescribe /chatwoot-api hacia Chatwoot. Si una empresa usa Chatwoot Cloud se mantiene https://app.chatwoot.com. "
        "Si usa Chatwoot self-hosted, cambiar el destination a la URL publica de esa instancia."
    )
    manual.code(
        "vercel.json esperado",
        """
{
  "rewrites": [
    {
      "source": "/chatwoot-api/:path*",
      "destination": "https://app.chatwoot.com/:path*"
    },
    {
      "source": "/(.*)",
      "destination": "/index.html"
    }
  ]
}
""",
    )
    manual.h(2, "8.3 Validacion de deploy")
    manual.bullets(
        [
            "Abrir la URL de Vercel y confirmar que carga el logo y las pestanas.",
            "Confirmar que no hay errores de variables faltantes en consola.",
            "Revisar que los filtros de fecha/canal funcionen.",
            "Probar una exportacion manual despues de cargar datos.",
        ]
    )


def add_supabase_setup(manual: Manual) -> None:
    manual.h(1, "9. Supabase paso a paso")
    manual.table(
        ["Paso", "Accion", "Explicacion"],
        [
            ["1", "Crear proyecto Supabase.", "Genera URL, anon key, service role key y base Postgres."],
            ["2", "Habilitar extensiones.", "pgcrypto crea UUIDs, pg_net llama Edge Functions, pg_cron agenda tareas y Vault guarda secretos."],
            ["3", "Crear schema cw.", "Aisla toda la data de Chatwoot/dashboard."],
            ["4", "Crear tablas base.", "Estado actual, historicos, mensajes, configuracion y reportes."],
            ["5", "Aplicar migraciones.", "Actualiza vistas, indices, cron, cambios relevantes y reportes programados."],
            ["6", "Configurar RLS/grants.", "Permite lectura/escritura controlada al frontend autenticado y service_role."],
            ["7", "Guardar secretos.", "Permite que Edge Functions y cron funcionen sin exponer claves."],
            ["8", "Crear cron.", "Ejecuta sync diario y reportes programados."],
            ["9", "Deploy Edge Functions.", "Activa integracion con Chatwoot y Resend."],
            ["10", "Ejecutar sync inicial.", "Carga datos para que dashboard deje de aparecer vacio."],
        ],
        [0.55, 2.0, 4.35],
    )
    manual.h(2, "9.1 Script base de creacion")
    manual.p(
        "Este bloque crea la estructura base necesaria del schema cw. En un proyecto nuevo, ejecutar primero este bloque y luego las migraciones/versiones de actualizacion."
    )
    manual.code("SQL base desde cero", BASE_SCHEMA_SQL)
    manual.h(2, "9.2 Secretos de Supabase")
    manual.p("Guardar secretos con Supabase CLI. Los valores entre <...> deben reemplazarse por los valores reales del nuevo proyecto.")
    manual.code("Comandos de secretos", SUPABASE_SECRETS_EXAMPLE)
    manual.h(2, "9.3 Deploy de Edge Functions")
    manual.code("Comandos de deploy", DEPLOY_COMMANDS)
    manual.h(2, "9.4 Sync inicial y pruebas")
    manual.code("Comandos de prueba", INITIAL_SYNC_COMMANDS)
    manual.h(2, "9.5 SQL de validacion")
    manual.code("Validaciones posteriores", VALIDATION_SQL)


def add_edge_functions(manual: Manual) -> None:
    manual.h(1, "10. Edge Functions")
    manual.table(["Funcion", "Proposito", "Variables", "Prueba minima"], EDGE_FUNCTIONS, [1.25, 2.4, 2.0, 1.25])


def add_chatwoot_resend(manual: Manual) -> None:
    manual.h(1, "11. Chatwoot y Resend")
    manual.h(2, "11.1 Chatwoot")
    manual.bullets(
        [
            "Crear o identificar cuenta Chatwoot y anotar Account ID.",
            "Crear token API desde perfil/admin con permisos suficientes para conversaciones, contactos, mensajes, inboxes, labels y custom attributes.",
            "Crear etiquetas comerciales usadas por el embudo: confianza, urgencia, seguimiento, cita, venta, descalificacion u otras segun la empresa.",
            "Crear atributos personalizados necesarios: score_interes, monto_operacion, fecha_monto_operacion, campana, responsable y campos de datos del lead si aplican.",
            "Configurar webhook hacia https://<PROJECT_REF>.supabase.co/functions/v1/chatwoot-label-webhook con el secreto definido en CHATWOOT_WEBHOOK_SECRET.",
            "Probar con una conversacion real y validar que cambien labels, mensajes y estado actual en cw.conversations_current.",
        ]
    )
    manual.h(2, "11.2 Resend")
    manual.bullets(
        [
            "Crear cuenta Resend y verificar dominio o remitente autorizado.",
            "Crear API key exclusiva para reportes.",
            "Definir RESEND_FROM_EMAIL con nombre visible y correo verificado.",
            "Probar envio con un reporte programado de baja frecuencia antes de activar reportes para clientes.",
            "Si Resend rechaza adjuntos o dominio, revisar estado de verificacion y logs de send-scheduled-reports.",
        ]
    )


def add_reports(manual: Manual) -> None:
    manual.h(1, "12. Reportes y evidencia")
    manual.table(
        ["Reporte", "Formato", "Criterio administrativo"],
        [
            ["Manual por pestana", "Excel, PDF, CSV", "Debe reflejar lo visible en la pestana y agregar contexto util sin ruido."],
            ["Programado", "Adjuntos por correo", "Debe usar la misma logica conceptual que la descarga manual."],
            ["Seguimiento", "Excel bloqueado en estructura", "No cambiar formato Excel existente; PDF/CSV pueden mejorar."],
            ["Conversaciones", "Excel bloqueado en estructura", "No cambiar formato Excel existente; PDF/CSV pueden mejorar."],
            ["Cambios relevantes", "Seccion condicional", "Solo mostrar si hay cambios con impacto en ventas, montos, citas o metricas visibles."],
        ],
        [1.5, 1.4, 4.0],
    )
    manual.h(2, "12.1 Evidencias que debe revisar soporte")
    manual.bullets(
        [
            "cw.sync_runs para saber si la sincronizacion termino correctamente.",
            "cw.raw_ingest para revisar payloads crudos si un dato no aparece.",
            "cw.automated_report_runs para validar envios de reportes.",
            "Logs de Edge Functions en Supabase.",
            "cron.job y cw.sync_health para verificar automatizaciones.",
        ]
    )


def add_risks(manual: Manual) -> None:
    manual.h(1, "13. Riesgos, controles y troubleshooting")
    manual.table(
        ["Riesgo", "Sintoma", "Control"],
        [
            ["Variables de Vercel incompletas", "Dashboard carga sin datos o falla API.", "Revisar Project Settings y redeploy."],
            ["Token Chatwoot invalido", "chatwoot-sync termina error.", "Regenerar token y actualizar secretos."],
            ["Cron caido", "No se actualiza historico diario.", "Revisar pg_cron, pg_net y cw.sync_health."],
            ["Resend sin dominio verificado", "Reportes no llegan.", "Verificar dominio/remitente y logs."],
            ["Etiquetas mal configuradas", "KPIs no coinciden con operacion.", "Revisar panel admin y dashboard_tag_settings."],
            ["Datos recientes incompletos", "Hoy/ayer no reflejan Chatwoot.", "Ejecutar sync manual o repair."],
            ["Migracion faltante", "Tabla o columna no existe.", "Aplicar SQL base y migraciones en orden."],
        ],
        [1.75, 2.25, 2.9],
    )


def add_annexes(manual: Manual) -> None:
    manual.page_break()
    manual.h(1, "Anexo A. Catalogo tecnico del schema cw")
    manual.table(
        ["Objeto", "Proposito", "Escribe", "Lee", "Impacto"],
        TABLE_CATALOG,
        [1.45, 2.05, 1.25, 1.15, 1.25],
    )

    manual.h(1, "Anexo B. Migraciones del repositorio")
    manual.p(
        "Estas migraciones se aplican despues del script base o en el orden natural del repo. "
        "Antes de cada bloque se indica el objetivo administrativo. Los valores de proyecto se muestran como placeholders."
    )
    explanations = {
        "20260421134500_create_conversation_label_events.sql": "Crea eventos deduplicados de etiquetas para explicar cambios de etapa o senales comerciales.",
        "20260426190000_harden_chatwoot_snapshots_and_sync.sql": "Endurece snapshots actuales, agrega vistas de salud, indices, grants y cron diario de sync.",
        "20260426202500_grant_cw_to_service_role.sql": "Garantiza permisos del service_role sobre schema cw para Edge Functions.",
        "20260427165000_reporting_exports_and_scheduled_runs.sql": "Amplia reportes programados, corridas y cron de envio cada 5 minutos.",
        "20260427220500_backfill_chatwoot_inbox_channel.sql": "Rellena canal/inbox desde payloads historicos de Chatwoot.",
        "20260427225200_enforce_chatwoot_inbox_channel.sql": "Refuerza la normalizacion de canal/inbox para que filtros y reportes coincidan.",
        "20260429123000_commercial_current_state_audit.sql": "Separa estado comercial actual de cambios relevantes historicos para reportes claros.",
    }
    for path in migration_files():
        manual.h(2, path.name)
        manual.p(explanations.get(path.name, "Migracion de actualizacion del schema cw. Revisar descripcion interna del SQL antes de ejecutar."))
        manual.code(f"SQL - {path.name}", read_migration(path))

    manual.h(1, "Anexo C. Migraciones remotas detectadas")
    manual.p(
        "El entorno Supabase actual registra migraciones iniciales que pudieron haberse aplicado directamente antes de consolidar los archivos locales. "
        "El script base de este manual representa la estructura necesaria para un entorno nuevo."
    )
    manual.table(
        ["Version", "Nombre", "Accion recomendada"],
        [
            ["20260411181400", "chatwoot_historical_sync_init", "Cubierta por SQL base. Exportar desde Supabase si se requiere auditoria exacta."],
            ["20260421084504", "create_cw_dashboard_tag_settings", "Cubierta por SQL base y catalogo."],
            ["20260429115733", "commercial_audit_events_infer_non_current_sale", "Cubierta por vista de cambios relevantes y migracion comercial."],
            ["20260429120540", "commercial_audit_amount_update_label", "Cubierta por vista de cambios relevantes y reportes."],
        ],
        [1.45, 2.35, 3.1],
    )

    manual.h(1, "Anexo D. Checklist final de puesta en marcha")
    manual.table(
        ["Item", "OK / No OK / N.A.", "Evidencia"],
        [
            ["Proyecto Supabase creado y SQL aplicado.", "", ""],
            ["Extensiones pgcrypto, pg_net, pg_cron y Vault habilitadas.", "", ""],
            ["Edge Functions desplegadas.", "", ""],
            ["Secretos de Supabase configurados.", "", ""],
            ["Chatwoot Account ID y API token probados.", "", ""],
            ["Webhook Chatwoot configurado y probado.", "", ""],
            ["Vercel conectado a GitHub main con variables VITE_.", "", ""],
            ["Resend verificado con correo FROM valido.", "", ""],
            ["Sync inicial ejecutado con success.", "", ""],
            ["Panel admin configurado para la empresa.", "", ""],
            ["Dashboard muestra datos en Estrategia, Operacion, Calidad y Conversaciones.", "", ""],
            ["Reporte manual y programado probados.", "", ""],
        ],
        [3.3, 1.4, 2.2],
    )


def build_manual() -> Manual:
    manual = Manual()
    add_control_documental(manual)
    add_intro(manual)
    add_architecture(manual)
    add_roles(manual)
    add_admin_config(manual)
    add_dashboard_logic(manual)
    add_replicacion(manual)
    add_vercel(manual)
    add_supabase_setup(manual)
    add_edge_functions(manual)
    add_chatwoot_resend(manual)
    add_reports(manual)
    add_risks(manual)
    add_annexes(manual)
    return manual


def main() -> None:
    manual = build_manual()
    render_docx(manual)
    render_pdf(manual)
    print(f"Generated {DOCX_PATH}")
    print(f"Generated {PDF_PATH}")


if __name__ == "__main__":
    main()
