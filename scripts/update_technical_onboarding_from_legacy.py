from __future__ import annotations

import re
import textwrap
import unicodedata
import zipfile
import xml.etree.ElementTree as ET
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DOCX = ROOT / "plantilla_iso_9001_proceso.docx"
LOGO_PATH = ROOT / "logo_simplia.png"
SOURCE_DOCX = ROOT / "ONBOARDING TECNICO -SIMPLIA CHATBOT .docx"
DOCX_OUT = ROOT / "Onboarding_Tecnico_Simplia_Chatbot_ISO9001.docx"
ASSET_DIR = ROOT / "Onboarding_Tecnico_Simplia_Chatbot_assets"

DOC_CODE = "PRO-ONB-CHATBOT-001"
DOC_VERSION = "1.2"
DOC_STATUS = "Vigente/Borrador"
ISSUE_DATE = "06/05/2026"
NEXT_REVIEW = "06/11/2026"

COLORS = {
    "blue": "274690",
    "navy": "0f2344",
    "slate": "64748b",
    "line": "d9e2ef",
    "light": "f8fafc",
    "soft_blue": "eaf1ff",
    "soft_green": "eaf8f1",
    "soft_orange": "fff7e6",
    "green": "0a9b6f",
    "orange": "b45309",
    "red": "dc2626",
    "white": "ffffff",
}

DOC_URLS = {
    "whatsapp": "https://developers.chatwoot.com/self-hosted/configuration/features/integrations/whatsapp-embedded-signup",
    "facebook": "https://developers.chatwoot.com/self-hosted/configuration/features/integrations/facebook-channel-setup",
    "instagram": "https://developers.chatwoot.com/self-hosted/configuration/features/integrations/instagram-via-instagram-business-login",
    "tiktok": "https://developers.chatwoot.com/self-hosted/configuration/features/integrations/tiktok",
}

VIDEO_LINKS = {
    "portfolio": ["Guía en video 1:26-5:00 y 8:08-9:30", "https://www.youtube.com/watch?v=CcdGYXw-k6w"],
    "meta_app": ["Guía en video", "https://www.youtube.com/watch?v=wSDAlUGXc1Y"],
    "railway": ["Guía para crear cuenta Railway", "https://www.youtube.com/watch?v=G80AIp6U3-o"],
    "n8n_railway": ["Guía en video n8n Railway", "https://www.loom.com/share/22b8ec26c70f4bffa48129ec72377473?sid=21d1fb66-1149-499d-bf14-93957466694c"],
    "railway_config": ["Guía cambio de configuración Railway", "https://www.loom.com/share/f9c31ddc45084391a27c6744189c64b6"],
    "whatsapp_profile": ["Guía cambio de foto/nombre de perfil WhatsApp", "https://www.loom.com/share/bfe0742a4b2045e3a5fe8dc8b02442a7"],
    "openai": ["Guía OpenAI / ChatGPT", "https://www.loom.com/share/5ad39e93897a425a881f96d85b904760"],
    "restart": ["Guía reinicio / validación n8n", "https://www.loom.com/share/b547a51b800d4cb3821bbb31c7ea9d2c?sid=8ac1ff66-5893-4712-9037-e64693122ff2"],
    "drive": ["Guía n8n Drive - Service Account", "https://www.loom.com/share/bb87ebb7964f477da419a6c662a67b3d"],
    "gmail": ["Guía n8n Gmail OAuth", "https://www.loom.com/share/916b40f559764c39a016eeb36b92d532"],
    "templates": ["Guía plantillas de respuesta Chatwoot", "https://www.loom.com/share/9b066dbcc82941098a1a656216aeb4a8"],
    "anydesk": ["AnyDesk", "https://anydesk.com/es"],
    "compress": ["Comprimir videos si superan 15 MB", "https://www.compress2go.com/es/comprimir-video"],
}

DOC_SNAPSHOTS = {
    "whatsapp": {
        "title": "Chatwoot Docs - WhatsApp Embedded Signup",
        "source": DOC_URLS["whatsapp"],
        "items": [
            "Configurar en Super Admin: WHATSAPP_APP_ID, WHATSAPP_CONFIGURATION_ID y WHATSAPP_APP_SECRET.",
            "Obtener App ID, App Secret y Configuration ID desde Meta Developer Portal.",
            "Habilitar permisos: whatsapp_business_management, whatsapp_business_messaging y business_management.",
            "Crear canal desde Settings > Inboxes > Add Inbox > WhatsApp > WhatsApp Cloud.",
            "Usar Connect with WhatsApp Business para completar el embedded signup.",
            "El flujo automatiza webhooks y configuración del número cuando está disponible.",
        ],
    },
    "facebook": {
        "title": "Chatwoot Docs - Facebook Messenger",
        "source": DOC_URLS["facebook"],
        "items": [
            "Crear una app tipo Business en Meta for Developers.",
            "Configurar variables de Chatwoot: FB_VERIFY_TOKEN, FB_APP_SECRET y FB_APP_ID.",
            "Agregar Facebook Login y habilitar Web OAuth Login y Login with JavaScript SDK.",
            "Agregar dominio de instalación Chatwoot en la app.",
            "Agregar producto Messenger y configurar Callback URL: {chatwoot_url}/bot.",
            "Suscribir eventos: messages, messaging_postbacks, message_deliveries, message_reads y message_echoes.",
        ],
    },
    "instagram": {
        "title": "Chatwoot Docs - Instagram Business Login",
        "source": DOC_URLS["instagram"],
        "items": [
            "Usar Chatwoot v4.1 o superior para Instagram Business Login.",
            "Crear app tipo Business y agregar producto Instagram.",
            "Configurar Instagram App ID y App Secret en Super Admin.",
            "Configurar webhook: {chatwoot_url}/webhooks/instagram.",
            "El verify token debe coincidir con INSTAGRAM_VERIFY_TOKEN.",
            "Configurar Redirect URL: {chatwoot_url}/instagram/callback.",
        ],
    },
    "tiktok": {
        "title": "Chatwoot Docs - TikTok Business Messaging",
        "source": DOC_URLS["tiktok"],
        "items": [
            "Requiere instancia Chatwoot con URL pública HTTPS y acceso Super Admin.",
            "Crear cuenta TikTok Developer y registrar app en TikTok Developer Portal.",
            "Solicitar acceso a Business Messaging API antes de continuar.",
            "Configurar Redirect URL: {chatwoot_url}/tiktok/callback.",
            "Configurar App ID y App Secret en {chatwoot_url}/super_admin/app_config?config=tiktok.",
            "Registrar webhook desde Rails console con Tiktok::AuthClient.update_webhook_callback.",
        ],
    },
}

LEGACY_IMAGE_GROUPS = {
    "railway": ["3. Creación y Conexión de Railway (Hosting)"],
    "openai": ["7. Integración con OpenAI / ChatGPT"],
    "drive": ["9. Configuracion n8n nodo drive"],
    "gmail": ["10. configuración n8n nodo gmail"],
    "anydesk": ["12.- Control remoto de pc"],
}


STEPS = [
    {
        "title": "Creación del Portafolio en Facebook Business Manager",
        "objective": "Crear o validar el portafolio empresarial desde donde se administran activos Meta, permisos, páginas y números.",
        "guide_links": [VIDEO_LINKS["portfolio"]],
        "actions": [
            "Entrar a Business Manager con la cuenta administradora de la empresa.",
            "Crear o seleccionar el portafolio correcto.",
            "Confirmar que la cuenta de Facebook tenga más de 2 meses de antigüedad.",
            "Registrar Business ID, correo administrador y evidencia del portafolio.",
        ],
        "save": ["Nombre del portafolio", "Business ID", "Correo administrador", "Estado de verificación", "Captura del portafolio"],
        "evidence": "Captura del portafolio creado o seleccionado.",
    },
    {
        "title": "Creación de App en Meta for Developers y registro de número nuevo para WhatsApp",
        "objective": "Crear la app de Meta y registrar un número/chip nuevo que nunca haya sido usado con WhatsApp.",
        "guide_links": [VIDEO_LINKS["meta_app"]],
        "actions": [
            "Crear aplicación tipo negocio en Meta for Developers.",
            "Agregar el producto WhatsApp y aceptar condiciones.",
            "Registrar el número nuevo destinado a WhatsApp Business.",
            "Guardar App ID, WABA ID, Phone Number ID y evidencia de verificación.",
        ],
        "save": ["App ID", "Business Account ID / WABA ID", "Phone Number ID", "Número WhatsApp", "Token/API key si aplica", "Capturas de verificación"],
        "evidence": "Captura de la app Meta y del número registrado.",
    },
    {
        "title": "Creación y conexión de Railway solo para n8n",
        "objective": "Crear el proyecto Railway para n8n y automatizaciones.",
        "guide_links": [VIDEO_LINKS["railway"], ["Railway n8n guide", "https://docs.railway.com/guides/n8n"]],
        "legacy_key": "railway",
        "actions": [
            "Crear o ingresar a Railway con el método aprobado por la empresa.",
            "Crear proyecto para n8n y registrar la URL del proyecto.",
            "Configurar facturación si se requiere operación estable.",
            "Registrar el proyecto como entorno de n8n y automatizaciones.",
        ],
        "save": ["Correo Railway", "Proyecto Railway", "URL del proyecto", "Plan / facturación", "Responsable de pago"],
        "evidence": "Captura del proyecto Railway destinado a n8n.",
    },
    {
        "title": "Creación de la instancia de n8n en Railway",
        "objective": "Desplegar n8n como motor de automatizaciones del agente conversacional.",
        "guide_links": [VIDEO_LINKS["n8n_railway"], ["Railway n8n guide", "https://docs.railway.com/guides/n8n"]],
        "actions": [
            "Crear servicio n8n en el proyecto Railway.",
            "Configurar URL pública o dominio.",
            "Guardar usuario administrador y contraseña por método seguro.",
            "Validar que n8n abra correctamente y ejecute workflows.",
        ],
        "save": ["URL n8n", "Correo admin", "Método seguro de contraseña", "Variables configuradas", "Workflow de prueba"],
        "evidence": "Captura de n8n activo en Railway.",
    },
    {
        "title": "Cambio de configuración en Railway por cada nodo creado",
        "objective": "Actualizar variables, secretos y configuración técnica de n8n cuando un nodo o workflow lo requiera.",
        "guide_links": [VIDEO_LINKS["railway_config"]],
        "actions": [
            "Revisar variables requeridas por cada workflow.",
            "Guardar secretos solo en Railway, n8n o gestor seguro.",
            "Reiniciar o redeploy n8n cuando el ajuste lo requiera.",
            "Probar el workflow afectado después del ajuste.",
        ],
        "save": ["Variable o nodo creado", "Servicio afectado", "Fecha de ejecución", "Responsable", "Prueba realizada"],
        "evidence": "Captura de variables sin exponer secretos y prueba OK.",
    },
    {
        "title": "Creación de cuenta Chatwoot Cloud y contratación del plan",
        "objective": "Crear el workspace Cloud donde se administrarán conversaciones, agentes, bandejas e integraciones.",
        "guide_links": [["Chatwoot", "https://www.chatwoot.com/"], ["Chatwoot pricing", "https://www.chatwoot.com/pricing/"]],
        "actions": [
            "Crear workspace Chatwoot con correo empresarial.",
            "Contratar el plan aprobado por la empresa o confirmar el plan vigente.",
            "Registrar responsables, agentes y método seguro de credenciales.",
            "Validar acceso al dashboard y configuración inicial del workspace.",
        ],
        "save": ["Workspace Chatwoot", "Correo empresarial", "Plan contratado", "Agentes", "Responsable", "Método de pago"],
        "evidence": "Captura del workspace Chatwoot y plan activo.",
    },
    {
        "title": "Configuración WhatsApp Cloud en Chatwoot",
        "objective": "Conectar WhatsApp usando el flujo vigente de Chatwoot/Meta.",
        "guide_links": [["Chatwoot WhatsApp Embedded Signup", DOC_URLS["whatsapp"]], ["Meta WhatsApp Cloud API", "https://developers.facebook.com/docs/whatsapp/cloud-api/get-started"]],
        "doc_snapshot": "whatsapp",
        "actions": [
            "Configurar en Super Admin los valores de WhatsApp Embedded Signup si la instancia lo requiere.",
            "Obtener en Meta App el App ID, App Secret y Configuration ID.",
            "Crear el canal desde Settings > Inboxes > Add Inbox > WhatsApp > WhatsApp Cloud.",
            "Completar Connect with WhatsApp Business y probar mensaje entrante/saliente.",
        ],
        "save": ["Inbox WhatsApp", "Número conectado", "WABA ID", "Phone Number ID", "Configuration ID", "Prueba entrante/saliente"],
        "evidence": "Captura del canal WhatsApp conectado y conversación de prueba.",
    },
    {
        "title": "Configuración de método de pago y perfil para WhatsApp",
        "objective": "Dejar lista la facturación de Meta y, si aplica, actualizar foto/nombre del número nuevo.",
        "guide_links": [VIDEO_LINKS["whatsapp_profile"], ["Centro de facturación Meta", "https://business.facebook.com/billing_hub/"], ["WhatsApp Manager", "https://business.facebook.com/wa/manage/"]],
        "actions": [
            "Ir a WhatsApp Manager o Centro de facturación Meta.",
            "Agregar método de pago aprobado por la empresa.",
            "Confirmar que WhatsApp puede enviar/recibir sin bloqueo por facturación.",
            "Actualizar foto y nombre del número si el cliente lo solicita.",
        ],
        "save": ["Método de pago", "Responsable de pago", "Captura de facturación", "Estado de perfil", "Observaciones"],
        "evidence": "Captura de facturación activa y perfil validado.",
        "critical": True,
    },
    {
        "title": "Configuración Facebook / Messenger en Chatwoot",
        "objective": "Conectar Facebook Messenger siguiendo la documentación vigente para la integración en Chatwoot.",
        "guide_links": [["Chatwoot Facebook Messenger", DOC_URLS["facebook"]]],
        "doc_snapshot": "facebook",
        "actions": [
            "Crear app tipo Business en Meta for Developers si no existe.",
            "Configurar FB_VERIFY_TOKEN, FB_APP_SECRET y FB_APP_ID en Chatwoot.",
            "Agregar Facebook Login, Messenger y dominio de Chatwoot.",
            "Configurar Callback URL {chatwoot_url}/bot y suscribir eventos requeridos.",
            "Crear inbox Messenger en Chatwoot y probar mensaje desde la página.",
        ],
        "save": ["Página Facebook", "App ID", "App Secret en gestor seguro", "FB_VERIFY_TOKEN", "Callback URL", "Inbox Messenger"],
        "evidence": "Captura de Messenger conectado y prueba de mensaje.",
    },
    {
        "title": "Configuración Instagram Business Login en Chatwoot",
        "objective": "Conectar Instagram con el método recomendado de Instagram Business Login.",
        "guide_links": [["Chatwoot Instagram Business Login", DOC_URLS["instagram"]]],
        "doc_snapshot": "instagram",
        "actions": [
            "Validar cuenta Instagram profesional y Chatwoot v4.1 o superior cuando aplique.",
            "Crear app Meta Business y agregar producto Instagram.",
            "Configurar Instagram App ID y App Secret en Chatwoot Super Admin.",
            "Configurar webhook {chatwoot_url}/webhooks/instagram y redirect {chatwoot_url}/instagram/callback.",
            "Crear inbox Instagram en Chatwoot y probar DM.",
        ],
        "save": ["Usuario Instagram", "Instagram App ID", "Instagram App Secret en gestor seguro", "INSTAGRAM_VERIFY_TOKEN", "Callback/Redirect URL", "Inbox Instagram"],
        "evidence": "Captura de Instagram conectado y prueba DM.",
    },
    {
        "title": "Configuración TikTok Business Messaging en Chatwoot",
        "objective": "Documentar los pasos de integración TikTok según la documentación oficial y validar disponibilidad antes de prometer el canal.",
        "guide_links": [["Chatwoot TikTok integration", DOC_URLS["tiktok"]], ["TikTok Developers", "https://developers.tiktok.com/"], ["TikTok Business API", "https://business-api.tiktok.com/"]],
        "doc_snapshot": "tiktok",
        "actions": [
            "Crear cuenta TikTok Developer y registrar app.",
            "Solicitar acceso a Business Messaging API.",
            "Configurar permisos, redirect URL {chatwoot_url}/tiktok/callback y App ID/App Secret.",
            "Configurar webhook con Tiktok::AuthClient.update_webhook_callback si aplica.",
            "Conectar TikTok Business Account desde Chatwoot y probar mensajes.",
        ],
        "save": ["TikTok Developer Account", "App ID", "App Secret en gestor seguro", "Estado Business Messaging API", "Webhook", "Estado de prueba"],
        "evidence": "Captura de aprobación, configuración y prueba o constancia de no disponibilidad.",
        "note": "La documentación oficial consultada corresponde a instalación self-hosted; si se opera en Cloud, validar disponibilidad/habilitación con Chatwoot antes de comprometer el canal.",
    },
    {
        "title": "Integración con OpenAI / ChatGPT",
        "objective": "Conectar el modelo de IA usado por el agente conversacional.",
        "guide_links": [VIDEO_LINKS["openai"], ["OpenAI API keys", "https://platform.openai.com/api-keys"]],
        "legacy_key": "openai",
        "actions": [
            "Crear o usar cuenta OpenAI autorizada por la empresa.",
            "Agregar método de pago y límites de uso.",
            "Crear API key y guardarla solo en n8n o gestor seguro.",
            "Ejecutar prueba de respuesta desde el workflow.",
        ],
        "save": ["Cuenta OpenAI", "Método de pago", "Límite mensual", "API key guardada en seguro", "Prueba OK"],
        "evidence": "Captura de configuración sin exponer API key.",
    },
    {
        "title": "Reinicio o validación final de n8n y canales Chatwoot",
        "objective": "Confirmar que n8n y los canales conectados responden correctamente después de ajustes técnicos.",
        "guide_links": [VIDEO_LINKS["restart"]],
        "actions": [
            "Reiniciar o redeploy n8n cuando se ajusten variables o workflows.",
            "Validar que Chatwoot reciba mensajes en cada inbox configurado.",
            "Enviar prueba de extremo a extremo: cliente > Chatwoot > n8n > respuesta.",
            "Registrar errores pendientes antes del cierre.",
        ],
        "save": ["Workflow probado", "Canal probado", "Resultado", "Fecha", "Responsable"],
        "evidence": "Capturas de pruebas por canal.",
    },
    {
        "title": "Configuración n8n nodo Drive - Service Account",
        "objective": "Configurar Google Drive solo si el flujo requiere leer o guardar archivos.",
        "guide_links": [VIDEO_LINKS["drive"], VIDEO_LINKS["compress"]],
        "legacy_key": "drive",
        "actions": [
            "Preguntar al desarrollador si este nodo aplica al proyecto.",
            "Crear service account y compartir carpetas necesarias.",
            "Guardar JSON/credencial en n8n de forma segura.",
            "Probar lectura o escritura en Drive.",
            "Si se usan videos/imágenes, validar permisos y comprimir videos mayores a 15 MB.",
        ],
        "save": ["Aplica: Sí / No", "Cuenta de servicio", "Carpeta Drive", "Permisos", "Prueba OK", "Links de archivos"],
        "evidence": "Captura de credencial creada sin exponer secretos y prueba Drive.",
    },
    {
        "title": "Configuración n8n nodo Gmail - OAuth",
        "objective": "Configurar Gmail solo si el flujo requiere enviar o leer correos.",
        "guide_links": [VIDEO_LINKS["gmail"]],
        "legacy_key": "gmail",
        "actions": [
            "Preguntar al desarrollador si este nodo aplica al proyecto.",
            "Crear OAuth consent y credenciales si corresponde.",
            "Conectar cuenta Gmail desde n8n.",
            "Probar envío o lectura según el flujo.",
        ],
        "save": ["Aplica: Sí / No", "Correo Gmail", "OAuth client", "Permisos", "Prueba OK"],
        "evidence": "Captura de conexión Gmail en n8n.",
    },
    {
        "title": "Configuración plantillas de respuesta Chatwoot",
        "objective": "Crear plantillas rápidas si la empresa necesita respuestas manuales estandarizadas.",
        "guide_links": [VIDEO_LINKS["templates"], ["Chatwoot canned responses", "https://www.chatwoot.com/docs/product/features/canned-responses/"]],
        "actions": [
            "Recolectar textos aprobados por la empresa.",
            "Crear plantillas en Chatwoot.",
            "Validar nombre, texto, tono y uso interno.",
            "Probar uso desde una conversación.",
        ],
        "save": ["Plantilla", "Escenario", "Texto aprobado", "Responsable de aprobación", "Estado"],
        "evidence": "Captura de plantillas creadas.",
    },
    {
        "title": "Control remoto de PC - AnyDesk",
        "objective": "Permitir acompañamiento remoto durante creación de cuentas, permisos y pagos.",
        "guide_links": [VIDEO_LINKS["anydesk"]],
        "legacy_key": "anydesk",
        "actions": [
            "Solicitar AnyDesk cuando el cliente deba operar desde su sesión.",
            "Confirmar autorización antes de tomar control.",
            "No guardar contraseñas visibles ni datos bancarios en el documento.",
            "Registrar solo evidencia operativa sin exponer información sensible.",
        ],
        "save": ["ID AnyDesk", "Responsable conectado", "Fecha", "Acción realizada", "Observaciones"],
        "evidence": "Registro de sesión o nota de acompañamiento.",
    },
    {
        "title": "Cierre, evidencias y credenciales",
        "objective": "Dejar documentado qué quedó creado, dónde están los accesos y qué pruebas fueron aprobadas.",
        "guide_links": [["Carpeta de evidencias del cliente", "____________________"], ["Dashboard Simplia", "____________________"]],
        "actions": [
            "Revisar que cada paso tenga responsable, fecha, estado y evidencia.",
            "Completar la tabla de credenciales sin escribir contraseñas ni tokens completos.",
            "Confirmar URL, usuario, responsable y estado de cada sistema.",
            "Cerrar con aprobación del cliente o responsable interno.",
        ],
        "save": ["Checklist final", "Tabla de credenciales", "Links de evidencias", "Aprobación", "Pendientes si existen"],
        "evidence": "Documento completo con aprobación final.",
    },
]


def rgb(hex_color: str) -> RGBColor:
    value = hex_color.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def pil_rgb(hex_color: str) -> tuple[int, int, int]:
    value = hex_color.lstrip("#")
    return tuple(int(value[i : i + 2], 16) for i in (0, 2, 4))


def font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return ImageFont.truetype(item, size)
    return ImageFont.load_default()


def clean_filename(text: str) -> str:
    ascii_value = unicodedata.normalize("NFKD", text).encode("ascii", "ignore").decode("ascii")
    return "".join(ch.lower() if ch.isalnum() else "_" for ch in ascii_value).strip("_")[:70]


def wrapped(text: str, width: int) -> list[str]:
    lines: list[str] = []
    for raw in str(text).splitlines():
        if not raw.strip():
            lines.append("")
        else:
            lines.extend(textwrap.wrap(raw, width=width, break_long_words=False))
    return lines


def text_of_xml(el) -> str:
    ns_w = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    parts: list[str] = []
    for node in el.iter():
        if node.tag == f"{{{ns_w}}}t" and node.text:
            parts.append(node.text)
        elif node.tag == f"{{{ns_w}}}tab":
            parts.append("\t")
        elif node.tag == f"{{{ns_w}}}br":
            parts.append("\n")
    return "".join(parts).strip()


def extract_legacy_images() -> tuple[list[Path], dict[str, list[Path]]]:
    ASSET_DIR.mkdir(exist_ok=True)

    doc = Document(SOURCE_DOCX)
    rid_to_path: dict[str, Path] = {}
    image_paths: list[Path] = []
    for index, shape in enumerate(doc.inline_shapes, start=1):
        blips = shape._inline.xpath(".//a:blip")
        if not blips:
            continue
        rid = blips[0].get(qn("r:embed"))
        rel = doc.part.rels.get(rid)
        if not rel:
            continue
        ext = ".jpg" if "jpeg" in rel.target_part.content_type else ".png"
        path = ASSET_DIR / f"legacy_evidencia_{index:03d}{ext}"
        try:
            path.write_bytes(rel.target_part.blob)
        except PermissionError:
            if not path.exists():
                raise
        rid_to_path[rid] = path
        image_paths.append(path)

    ns = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    heading_re = re.compile(
        r"^\s*((?:\d+\.\d+|\d+|\d+\.-)\.?\s+[^\n]+|CONFIGURACION INSTAGRAM|Configuracion para messenger/Instagram|Agregar Métodos de pago|Tener cuenta de gmail|railway|chatwoot)",
        re.I,
    )
    grouped_by_heading: dict[str, list[Path]] = defaultdict(list)
    with zipfile.ZipFile(SOURCE_DOCX) as z:
        root = ET.fromstring(z.read("word/document.xml"))
        body = root.find("w:body", ns)
        current = "Preambulo"
        if body is not None:
            for child in list(body):
                txt = text_of_xml(child)
                if txt:
                    match = heading_re.match(txt)
                    if match:
                        current = " ".join(txt.split())
                for blip in child.findall(".//a:blip", ns):
                    rid = blip.get(qn("r:embed"))
                    if rid in rid_to_path:
                        grouped_by_heading[current].append(rid_to_path[rid])

    grouped_by_step: dict[str, list[Path]] = defaultdict(list)
    for step_key, heading_parts in LEGACY_IMAGE_GROUPS.items():
        for heading, paths in grouped_by_heading.items():
            if any(part.lower() in heading.lower() for part in heading_parts):
                grouped_by_step[step_key].extend(paths)
    return image_paths, grouped_by_step


def make_doc_snapshot(key: str) -> Path:
    data = DOC_SNAPSHOTS[key]
    path = ASSET_DIR / f"doc_chatwoot_{key}.png"
    width, height = 1500, 880
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((34, 34, width - 34, height - 34), radius=28, fill=pil_rgb(COLORS["light"]), outline=pil_rgb(COLORS["line"]), width=3)
    draw.rounded_rectangle((62, 62, width - 62, 160), radius=18, fill=pil_rgb(COLORS["blue"]))
    draw.text((92, 88), "Referencia documentación oficial", font=font(24, True), fill="white")
    draw.text((92, 118), data["title"], font=font(34, True), fill="white")
    draw.rounded_rectangle((86, 195, width - 86, 285), radius=18, fill="white", outline=pil_rgb(COLORS["line"]), width=2)
    draw.text((112, 218), "Fuente:", font=font(22, True), fill=pil_rgb(COLORS["navy"]))
    for index, line in enumerate(wrapped(data["source"], 108)[:2]):
        draw.text((112, 248 + index * 26), line, font=font(20), fill=pil_rgb(COLORS["blue"]))
    y = 330
    draw.text((92, y), "Puntos que se deben configurar", font=font(27, True), fill=pil_rgb(COLORS["navy"]))
    y += 48
    for index, item in enumerate(data["items"], start=1):
        draw.rounded_rectangle((96, y - 8, width - 96, y + 58), radius=16, fill="white", outline=pil_rgb(COLORS["line"]), width=1)
        draw.ellipse((116, y + 8, 148, y + 40), fill=pil_rgb(COLORS["blue"]))
        draw.text((132, y + 24), str(index), font=font(17, True), fill="white", anchor="mm")
        text_x = 166
        for line_index, line in enumerate(wrapped(item, 105)[:2]):
            draw.text((text_x, y + 4 + line_index * 24), line, font=font(19), fill=pil_rgb(COLORS["navy"]))
        y += 76
    draw.rounded_rectangle((96, height - 104, width - 96, height - 58), radius=14, fill=pil_rgb(COLORS["soft_orange"]), outline=pil_rgb("facc15"), width=1)
    draw.text((122, height - 92), "Usar esta captura como guía visual; validar siempre en la página oficial antes de ejecutar la configuración.", font=font(18), fill=pil_rgb(COLORS["navy"]))
    image.save(path)
    return path


def make_process_map() -> Path:
    path = ASSET_DIR / "mapa_proceso_onboarding_tecnico_actualizado.png"
    width, height = 1900, 560
    image = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(image)
    draw.text((60, 42), "Mapa del onboarding técnico Simplia Chatbot", font=font(34, True), fill=pil_rgb(COLORS["navy"]))
    steps = [
        ("1", "Meta", "Portafolio,\napp y número"),
        ("2", "n8n", "Railway solo\npara automatizar"),
        ("3", "Chatwoot", "Cloud / cuenta\noperativa"),
        ("4", "Canales", "WhatsApp,\nFB, IG, TikTok"),
        ("5", "IA y datos", "OpenAI,\nDrive, Gmail"),
        ("6", "Cierre", "Pruebas,\nevidencias y accesos"),
    ]
    x, y = 65, 170
    box_w, box_h, gap = 255, 182, 42
    for idx, (num, name, desc) in enumerate(steps):
        bx = x + idx * (box_w + gap)
        draw.rounded_rectangle((bx, y, bx + box_w, y + box_h), radius=22, fill=pil_rgb(COLORS["soft_blue"]), outline=pil_rgb(COLORS["blue"]), width=3)
        draw.ellipse((bx + 18, y + 18, bx + 68, y + 68), fill=pil_rgb(COLORS["blue"]))
        draw.text((bx + 43, y + 43), num, font=font(22, True), fill="white", anchor="mm")
        draw.text((bx + 82, y + 26), name, font=font(23, True), fill=pil_rgb(COLORS["navy"]))
        draw.multiline_text((bx + 26, y + 88), desc, font=font(18), fill=pil_rgb(COLORS["slate"]), spacing=6)
        if idx < len(steps) - 1:
            ax = bx + box_w + 7
            ay = y + box_h // 2
            draw.line((ax, ay, ax + gap - 14, ay), fill=pil_rgb(COLORS["blue"]), width=4)
            draw.polygon([(ax + gap - 14, ay - 10), (ax + gap + 4, ay), (ax + gap - 14, ay + 10)], fill=pil_rgb(COLORS["blue"]))
    draw.text((70, 440), "Nota operativa: Railway queda reservado para n8n y automatizaciones.", font=font(18, True), fill=pil_rgb(COLORS["red"]))
    image.save(path)
    return path


def set_cell_shading(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def set_table_borders(table, color: str = COLORS["line"]) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement("w:" + edge)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)
        borders.append(element)
    table._tbl.tblPr.append(borders)


def clear_cell(cell) -> None:
    cell.text = ""
    for paragraph in cell.paragraphs:
        paragraph.text = ""


def add_hyperlink(paragraph, text: str, url: str) -> None:
    if not url.startswith("http"):
        paragraph.add_run(text)
        return
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)

    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), COLORS["blue"])
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    run.append(r_pr)

    text_el = OxmlElement("w:t")
    text_el.text = text
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def write_cell(cell, value: str, *, bold: bool = False, color: str | None = None, size: float = 8.0) -> None:
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    for index, part in enumerate(str(value).split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_table(doc: Document, rows: list[list[str]], *, header: bool = True, first_col: bool = False, font_size: float = 8.0):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    set_table_borders(table)
    for i, row in enumerate(rows):
        for j, value in enumerate(row):
            is_header = header and i == 0
            cell = table.cell(i, j)
            write_cell(
                cell,
                value,
                bold=is_header or (first_col and j == 0),
                color=COLORS["white"] if is_header else COLORS["navy"],
                size=font_size,
            )
            fill = COLORS["blue"] if is_header else (COLORS["light"] if i % 2 == 0 else COLORS["white"])
            set_cell_shading(cell, fill)
    doc.add_paragraph()
    return table


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.color.rgb = rgb(COLORS["blue"] if level == 1 else COLORS["navy"])


def add_para(doc: Document, text: str, *, style: str | None = None, bold: bool = False) -> None:
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = rgb(COLORS["navy"])
    run.bold = bold


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        add_para(doc, item, style="List Bullet")


def add_callout(doc: Document, title: str, text: str, tone: str = "blue") -> None:
    fill = COLORS["soft_blue"]
    border = "b8c7e8"
    title_color = COLORS["blue"]
    if tone == "warning":
        fill = COLORS["soft_orange"]
        border = "facc15"
        title_color = COLORS["orange"]
    elif tone == "ok":
        fill = COLORS["soft_green"]
        border = "86efac"
        title_color = COLORS["green"]
    table = doc.add_table(rows=1, cols=1)
    set_table_borders(table, border)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    clear_cell(cell)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(title + ": ")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = rgb(title_color)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(9.5)
    run.font.color.rgb = rgb(COLORS["navy"])
    doc.add_paragraph()


def add_image(doc: Document, path: Path, *, width: float = 6.35, caption: str | None = None) -> None:
    if caption:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(caption)
        run.bold = True
        run.font.name = "Arial"
        run.font.size = Pt(8.7)
        run.font.color.rgb = rgb(COLORS["blue"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))


def add_link_paragraph(doc: Document, label: str, url: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(label + ": ")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = rgb(COLORS["navy"])
    add_hyperlink(paragraph, url, url)


def step_table(step: dict) -> list[list[str]]:
    links = "\n".join(f"- {label}: {url}" for label, url in step.get("guide_links", []))
    save_lines = "\n".join(f"- {item}: ______________________________" for item in step["save"])
    rows = [
        ["Documento / link guía", links or "____________________"],
        ["Foto / captura guía", "Referencia visual incluida debajo del paso cuando aplica."],
        ["Qué se debe hacer", "\n".join(f"- {item}" for item in step["actions"])],
        ["Qué se debe guardar", save_lines],
        ["Evidencia esperada", f"{step['evidence']}\n\nLink / captura: ______________________________"],
        ["Responsable", "____________________"],
        ["Fecha de ejecución", "____ / ____ / ______"],
        ["Estado", "Pendiente / En proceso / Completo / No aplica"],
        ["Evidencia del cliente", "____________________"],
        ["Observaciones", "____________________"],
    ]
    if step.get("note"):
        rows.insert(3, ["Nota importante", step["note"]])
    return rows


def add_step(doc: Document, index: int, step: dict, legacy_images: dict[str, list[Path]], doc_images: dict[str, Path]) -> None:
    add_heading(doc, f"5.{index}. {step['title']}", 2)
    add_callout(doc, "Objetivo", step["objective"], "warning" if step.get("critical") else "blue")
    add_table(doc, step_table(step), header=False, first_col=True, font_size=7.8)

    for label, url in step.get("guide_links", []):
        if url.startswith("http"):
            add_link_paragraph(doc, label, url)

    if step.get("doc_snapshot"):
        key = step["doc_snapshot"]
        add_image(doc, doc_images[key], width=6.25, caption=f"Captura guía de documentación Chatwoot - {step['title']}")

    if step.get("legacy_key"):
        images = legacy_images.get(step["legacy_key"], [])
        if images:
            add_para(doc, "Capturas guía:", bold=True)
            for pos, image_path in enumerate(images, start=1):
                add_image(doc, image_path, width=6.1, caption=f"Captura guía {pos:02d} - {step['title']}")


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.62)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)
    doc.styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    for style_name, size, color in [
        ("Heading 1", 15, COLORS["blue"]),
        ("Heading 2", 12.5, COLORS["navy"]),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"{DOC_CODE} | v{DOC_VERSION} | {DOC_STATUS}")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["slate"])

    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Simplia Chatbot - Onboarding Técnico ISO 9001")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["slate"])


def clear_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def build_docx(legacy_images: dict[str, list[Path]], doc_images: dict[str, Path], map_path: Path) -> None:
    doc = Document(TEMPLATE_DOCX)
    clear_body(doc)
    configure_doc(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if LOGO_PATH.exists():
        p.add_run().add_picture(str(LOGO_PATH), width=Inches(1.55))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ONBOARDING TÉCNICO - SIMPLIA CHATBOT")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.color.rgb = rgb(COLORS["blue"])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Guía operativa para completar cuentas, permisos, canales, integraciones y evidencias")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.color.rgb = rgb(COLORS["slate"])

    add_table(
        doc,
        [
            ["Código", DOC_CODE, "Versión", DOC_VERSION],
            ["Nombre", "Onboarding técnico Simplia Chatbot", "Fecha de emisión", ISSUE_DATE],
            ["Uso", "Guía operativa de configuración", "Próxima revisión", NEXT_REVIEW],
            ["Estado", DOC_STATUS, "Formato", "ISO 9001"],
        ],
        font_size=7.8,
    )
    add_callout(
        doc,
        "Objetivo del documento",
        "Guiar la configuración técnica del chatbot y dejar registrados, dentro de este documento, los datos, responsables, evidencias y accesos requeridos para la entrega.",
    )

    add_heading(doc, "1. Información general", 1)
    add_table(
        doc,
        [
            ["Campo", "Completar durante la reunión"],
            ["Empresa / cliente", "____________________"],
            ["Responsable de cuentas y permisos", "____________________"],
            ["Responsable de pagos", "____________________"],
            ["Desarrollador / técnico Simplia", "____________________"],
            ["Correo empresarial para Chatwoot", "____________________"],
            ["Canales a configurar", "WhatsApp: ___  Messenger/Facebook: ___  Instagram: ___  TikTok: ___"],
            ["Carpeta de evidencias / links", "____________________"],
        ],
        font_size=7.8,
    )

    add_heading(doc, "2. Alcance operativo", 1)
    add_bullets(
        doc,
        [
            "Railway se usa solo para n8n y automatizaciones.",
            "Chatwoot se usa para administrar conversaciones, agentes, bandejas e integraciones.",
            "Meta Business y Meta for Developers se usan para portafolio, app, WhatsApp, Facebook/Messenger e Instagram.",
            "TikTok se configura solo si la cuenta y permisos del cliente permiten Business Messaging.",
            "OpenAI, Drive, Gmail y plantillas Chatwoot se completan únicamente cuando el flujo del cliente lo requiere.",
        ],
    )

    add_heading(doc, "3. Fuentes y videos guía", 1)
    source_rows = [["Tipo", "Nombre", "Link / uso"]]
    for key, (label, url) in VIDEO_LINKS.items():
        source_rows.append(["Video / guía original", label, url])
    for label, url in [
        ["Chatwoot WhatsApp Embedded Signup", DOC_URLS["whatsapp"]],
        ["Chatwoot Facebook Messenger", DOC_URLS["facebook"]],
        ["Chatwoot Instagram Business Login", DOC_URLS["instagram"]],
        ["Chatwoot TikTok Business Messaging", DOC_URLS["tiktok"]],
    ]:
        source_rows.append(["Documentación oficial", label, url])
    add_table(doc, source_rows, font_size=7.1)

    add_heading(doc, "4. Mapa del proceso", 1)
    add_image(doc, map_path, width=6.45)

    add_heading(doc, "5. Procedimiento paso a paso", 1)
    add_para(doc, "Cada paso incluye qué hacer, qué guardar y evidencia esperada. Los enlaces quedan visibles y también como hipervínculos debajo de cada tabla.")
    for idx, step in enumerate(STEPS, start=1):
        add_step(doc, idx, step, legacy_images, doc_images)

    add_heading(doc, "6. Validación final de funcionamiento", 1)
    add_table(
        doc,
        [
            ["Prueba", "Qué debe comprobarse", "Resultado", "Evidencia"],
            ["WhatsApp", "Mensaje entrante y respuesta desde Chatwoot/n8n", "OK / No OK / N.A.", "____________________"],
            ["Messenger/Facebook", "Mensaje entrante, asignación de agente y webhook OK", "OK / No OK / N.A.", "____________________"],
            ["Instagram", "DM entrante, webhook y respuesta OK", "OK / No OK / N.A.", "____________________"],
            ["TikTok", "Disponibilidad validada o canal probado", "OK / No OK / N.A.", "____________________"],
            ["n8n", "Workflow activo, credenciales conectadas y ejecución sin error", "OK / No OK / N.A.", "____________________"],
            ["OpenAI", "Respuesta generada con límites de uso definidos", "OK / No OK / N.A.", "____________________"],
            ["Dashboard", "Acceso operativo y datos mínimos visibles", "OK / No OK / N.A.", "____________________"],
        ],
        font_size=7.3,
    )

    add_heading(doc, "7. Checklist de cierre", 1)
    checklist = [
        "Portafolio Meta creado o seleccionado.",
        "App Meta y número WhatsApp nuevo registrados.",
        "Railway creado solo para n8n.",
        "n8n desplegado y probado.",
        "Chatwoot Cloud/workspace creado o instancia aprobada definida.",
        "Canales solicitados conectados o marcados como no aplica.",
        "OpenAI / Drive / Gmail configurados si aplica.",
        "Credenciales registradas sin secretos en texto plano.",
        "Evidencias y aprobación final completas.",
    ]
    add_table(doc, [["#", "Validación", "Resultado", "Observaciones"]] + [[str(i), item, "OK / No OK / N.A.", ""] for i, item in enumerate(checklist, start=1)], font_size=7.4)

    add_heading(doc, "8. Credenciales y cuentas", 1)
    add_callout(doc, "Importante", "No escribir contraseñas, API keys ni tokens completos en este documento. Usar gestor seguro, n8n, Railway, Chatwoot o el método aprobado por la empresa.", "warning")
    add_table(
        doc,
        [["Cuenta / sistema", "Dato principal", "Dato secundario", "Dónde quedó el acceso", "Responsable", "Estado"]]
        + [
            ["Chatwoot", "Correo / workspace", "Plan / cuenta", "Método seguro", "Responsable", "Estado"],
            ["Railway", "Correo", "Proyecto n8n", "Método seguro", "Responsable", "Estado"],
            ["n8n", "URL", "Usuario admin", "Método seguro", "Responsable", "Estado"],
            ["Meta Business", "Correo admin", "Business ID / Portafolio", "Permisos", "Responsable", "Estado"],
            ["Meta Developers", "App ID", "WABA / Phone Number ID", "Tokens en gestor seguro", "Responsable", "Estado"],
            ["OpenAI", "Correo", "Proyecto / API", "API key en gestor seguro", "Responsable", "Estado"],
            ["Google Drive / Gmail", "Correo", "Service Account / OAuth", "Permisos", "Responsable", "Estado"],
            ["AnyDesk", "ID", "Equipo", "Autorización", "Responsable", "Estado"],
            ["Dashboard", "URL", "Usuario", "Método seguro", "Responsable", "Estado"],
        ],
        font_size=6.9,
    )

    add_heading(doc, "9. Control documental", 1)
    add_table(
        doc,
        [
            ["Versión", "Fecha", "Estado del documento", "Responsable"],
            ["1.2", ISSUE_DATE, DOC_STATUS, "Simplia"],
        ],
        font_size=7.5,
    )

    add_heading(doc, "10. Aprobación final", 1)
    add_table(
        doc,
        [
            ["Elaborado por", "Revisado por", "Aprobado por"],
            ["Simplia / Desarrollador", "", ""],
            ["Firma / fecha", "Firma / fecha", "Firma / fecha"],
        ],
        font_size=7.8,
    )

    doc.core_properties.title = "Onboarding Técnico Simplia Chatbot ISO 9001"
    doc.core_properties.subject = "Guía de configuración técnica"
    doc.core_properties.author = "Simplia"
    doc.core_properties.keywords = "onboarding, ISO 9001, Simplia Chatbot, Chatwoot, n8n, Railway, Meta"
    try:
        doc.save(DOCX_OUT)
    except PermissionError as exc:
        raise SystemExit(f"No se pudo guardar {DOCX_OUT.name}. Cierre el archivo en Word y vuelva a ejecutar este script.") from exc


def main() -> None:
    if not SOURCE_DOCX.exists():
        raise FileNotFoundError(SOURCE_DOCX)
    if not TEMPLATE_DOCX.exists():
        raise FileNotFoundError(TEMPLATE_DOCX)
    ASSET_DIR.mkdir(exist_ok=True)
    all_images, legacy_images = extract_legacy_images()
    doc_images = {key: make_doc_snapshot(key) for key in DOC_SNAPSHOTS}
    map_path = make_process_map()
    build_docx(legacy_images, doc_images, map_path)
    print(f"DOCX={DOCX_OUT}")
    print(f"LEGACY_IMAGES={len(all_images)}")
    for key, images in legacy_images.items():
        print(f"{key}={len(images)}")


if __name__ == "__main__":
    main()
