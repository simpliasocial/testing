from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Image as RLImage
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOCX = ROOT / "SISTEMA FLUJO CONVERSACIONAL incio a fin.docx"
LOGO_PATH = ROOT / "logo_simplia.png"
DOCX_PATH = ROOT / "Sistema_Flujo_Conversacional_Simplia_Chatbot_ISO10013.docx"
PDF_PATH = ROOT / "Sistema_Flujo_Conversacional_Simplia_Chatbot_ISO10013.pdf"

DOC_CODE = "MAN-FLUJO-BOT-001"
DOC_VERSION = "1.0"
DOC_DATE = "04/05/2026"
NEXT_REVIEW = "04/11/2026"
DOC_STATUS = "Borrador / Vigente al aprobarse"

COLORS = {
    "navy": "0f2344",
    "blue": "274690",
    "soft_blue": "eaf1ff",
    "light": "f8fafc",
    "green": "0a9b6f",
    "slate": "64748b",
    "line": "d9e2ef",
    "white": "ffffff",
}


def rgb(hex_color: str) -> RGBColor:
    value = hex_color.lstrip("#")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def escape_html(value: str) -> str:
    return (
        str(value)
        .replace("&", "&amp;")
        .replace("<", "&lt;")
        .replace(">", "&gt;")
        .replace("\n", "<br/>")
    )


def safe_text(value: str, limit: int = 900) -> str:
    value = " ".join(str(value).split())
    return value if len(value) <= limit else value[: limit - 3] + "..."


@dataclass
class Event:
    kind: str
    payload: tuple


class Manual:
    def __init__(self) -> None:
        self.events: list[Event] = []

    def h(self, level: int, text: str) -> None:
        self.events.append(Event("h", (level, text)))

    def p(self, text: str) -> None:
        self.events.append(Event("p", (text,)))

    def bullets(self, items: Iterable[str]) -> None:
        self.events.append(Event("bullets", (list(items),)))

    def table(self, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
        self.events.append(Event("table", (headers, rows, widths)))

    def note(self, title: str, text: str) -> None:
        self.events.append(Event("note", (title, text)))


def set_cell_shading(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    props.append(shading)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def write_cell(cell, text: str, *, bold: bool = False, color: str | None = None, size: float = 8.0) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for index, part in enumerate(str(text).split("\n")):
        if index:
            paragraph.add_run().add_break()
        run = paragraph.add_run(part)
        run.bold = bold
        run.font.name = "Arial"
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = rgb(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def style_docx_table(table) -> None:
    table.style = "Table Grid"
    table.autofit = True
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if row_index == 0:
                set_cell_shading(cell, COLORS["blue"])
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = rgb(COLORS["white"])
            else:
                set_cell_shading(cell, COLORS["light"] if row_index % 2 == 0 else COLORS["white"])
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8 if row_index else 8.3)


def build_manual() -> Manual:
    manual = Manual()

    manual.h(1, "1. Control documental")
    manual.table(
        ["Campo", "Valor"],
        [
            ["Código", DOC_CODE],
            ["Nombre", "Sistema de Flujo Conversacional Simplia Chatbot"],
            ["Versión", DOC_VERSION],
            ["Fecha", DOC_DATE],
            ["Estado", DOC_STATUS],
            ["Próxima revisión", NEXT_REVIEW],
            ["Documento base", SOURCE_DOCX.name],
        ],
        [2.0, 5.1],
    )

    manual.h(1, "2. Qué hace el bot")
    manual.p(
        "El bot atiende leads de inicio a fin: recibe mensajes, entiende la intención, responde con información aprobada, categoriza el lead, guarda datos, calcula score, agenda citas si corresponde, envía retomas y deriva a humano cuando ya no debe continuar solo."
    )
    manual.note(
        "Idea simple",
        "Todo el sistema sirve para que el agente funcione correctamente de inicio a fin. El bot atiende lo repetible; la persona interviene cuando hay venta, cita manual, excepción o seguimiento humano.",
    )

    manual.h(1, "3. Flujo conversacional de inicio a fin")
    manual.table(
        ["Paso", "Qué pasa", "Qué hace el bot", "Resultado"],
        [
            ["1", "Entra un mensaje por WhatsApp, Instagram, Facebook u otro canal.", "Lee el mensaje y confirma que viene del cliente.", "Mensaje listo para analizar."],
            ["2", "El usuario saluda o pregunta algo.", "Detecta intención: saludo, información, interés, agenda, rechazo o humano.", "Lead categorizado."],
            ["3", "El usuario necesita información.", "Responde servicios, costos, beneficios, sedes u horarios con contenido aprobado.", "`solicita_informacion`."],
            ["4", "El usuario muestra intención de avanzar.", "Pide datos necesarios y guía hacia cita.", "`interesado`."],
            ["5", "El usuario quiere cita.", "Valida datos, fecha, hora y agencia.", "`cita_agendada` si todo está completo."],
            ["6", "El usuario no responde.", "Envía retoma de 30 min y luego retoma de 3 h.", "Si no responde, `seguimiento_humano`."],
            ["7", "El caso requiere persona.", "Deja de insistir y marca para gestión manual.", "`seguimiento_humano`."],
            ["8", "Una persona agenda o confirma venta.", "El bot no lo hace automáticamente.", "`cita_agendado_humano` o `venta_exitosa` manual."],
        ],
        [0.45, 2.25, 2.65, 1.9],
    )

    manual.h(1, "4. Qué hace el bot y qué hace una persona")
    manual.table(
        ["Actividad", "Lo hace el bot", "Lo hace una persona"],
        [
            ["Responder bienvenida", "Sí", "Solo si el bot no entiende o el cliente insiste."],
            ["Responder información frecuente", "Sí, con contenido aprobado.", "Actualiza o corrige la información cuando cambie."],
            ["Detectar intención", "Sí", "Puede corregir manualmente si hay error."],
            ["Categorizar leads", "Sí para etiquetas automáticas.", "Sí para venta y cita manual."],
            ["Agendar cita automática", "Sí, si tiene datos y disponibilidad.", "Si la cita se crea fuera del flujo."],
            ["Marcar `agente`", "No", "Sí, cuando el bot ya no debe responder."],
            ["Registrar venta exitosa", "No", "Sí, llenando monto y fecha."],
            ["Registrar cita manual", "No", "Sí, usando `cita_agendado_humano` y `responsable`."],
        ],
        [2.2, 2.55, 2.55],
    )

    manual.h(1, "5. Etiquetas del sistema")
    manual.table(
        ["Etiqueta", "Quién la asigna", "Qué significa"],
        [
            ["bienvenida", "Bot", "Primer contacto o lead sin intención clara."],
            ["solicita_informacion", "Bot", "Pide información, costos, servicios, beneficios, sedes u horarios."],
            ["interesado", "Bot", "Quiere avanzar, agendar, reservar o revisar disponibilidad."],
            ["desinteresado", "Bot", "Rechaza, insulta, está fuera del negocio o envía contenido irrelevante."],
            ["cita_agendada", "Bot", "Cita creada automáticamente por el bot."],
            ["seguimiento_humano", "Bot", "El bot deja el caso para una persona."],
            ["venta_exitosa", "Manual", "Venta u operación concretada. Requiere llenar monto y fecha."],
            ["cita_agendado_humano", "Manual", "Cita creada o modificada por una persona. Requiere llenar responsable."],
        ],
        [1.65, 1.1, 4.45],
    )

    manual.h(1, "6. Datos que guarda el sistema")
    manual.table(
        ["Atributo", "Tipo", "Quién lo llena", "Uso"],
        [
            ["nombre_completo", "string", "Bot", "Identificar al lead y crear cita."],
            ["fecha_visita", "string", "Bot", "Guardar fecha de cita."],
            ["hora_visita", "string", "Bot", "Guardar hora de cita."],
            ["agencia", "string", "Bot", "Guardar sede/agencia elegida."],
            ["celular", "string", "Bot", "Contactar al lead."],
            ["correo", "string", "Bot", "Contacto o registro."],
            ["campana", "string", "Bot", "Origen o campaña."],
            ["ciudad", "string", "Bot", "Segmentación o validación."],
            ["edad", "string", "Bot", "Dato del lead si el negocio lo necesita."],
            ["canal", "string", "Bot", "Canal de entrada."],
            ["score_interes", "number", "Bot", "Calidad e intención del lead."],
            ["agente", "checkbox", "Manual", "Se marca cuando el bot ya no debe responder."],
            ["monto_operacion", "string", "Manual", "Se llena con `venta_exitosa`."],
            ["fecha_monto_operacion", "Date", "Manual", "Se llena con `venta_exitosa`."],
            ["responsable", "string", "Manual", "Se llena con `cita_agendado_humano`."],
        ],
        [1.65, 0.9, 1.2, 3.45],
    )

    manual.h(1, "7. Score de calidad")
    manual.p(
        "El score mide solo mensajes del cliente. No se calcula con respuestas de IA, bot o asesor. Es acumulativo, reversible y no tiene límite fijo; sube o baja según señales del mensaje."
    )
    manual.table(
        ["Nivel", "Rango", "Interpretación"],
        [
            ["FRÍO", "Menor a 45 o sin puntaje", "Señal inicial, mensaje genérico o lead todavía sin intención clara."],
            ["TIBIO", "45 a 69", "Tiene señales concretas: precio, disponibilidad, servicio específico o continuidad."],
            ["CALIENTE", "70 o más", "Alta intención: quiere comprar, agendar, reservar, pagar o hablar con asesor."],
        ],
        [1.1, 1.7, 4.4],
    )
    manual.table(
        ["Qué observa el score", "Ejemplos"],
        [
            ["Interés inicial", "info, precio, vi el anuncio, qué ofrecen."],
            ["Intención comercial", "quiero agendar, quiero comprar, cuánto cuesta, llámenme."],
            ["Fit con el negocio", "Pregunta por servicios, sedes, horarios o disponibilidad real."],
            ["Contactabilidad", "Deja teléfono, correo, nombre o pide contacto."],
            ["Engagement", "Confirma interés, sigue la conversación o hace preguntas concretas."],
            ["Señales negativas", "No me interesa, insultos, fuera de negocio, fraude o links con incertidumbre."],
        ],
        [2.4, 4.8],
    )

    manual.h(1, "8. Vocabulario que necesita el bot")
    manual.p(
        "El bot usa vocabulario del negocio para entender intención. Cada empresa debe entregar sus palabras reales, nombres de servicios, formas de pedir precio, frases de agenda, dudas frecuentes y señales de rechazo."
    )
    manual.table(
        ["Tipo de vocabulario", "Ejemplos"],
        [
            ["Servicios / productos", "Nombre de servicios, tratamientos, planes, productos o paquetes."],
            ["Costos / promociones", "precio, valor, cuánto cuesta, promoción, descuento, formas de pago."],
            ["Agenda", "quiero agendar, reservar, cita, cuándo puedo ir, disponibilidad."],
            ["Asesor humano", "asesor, persona, llámenme, quiero hablar con alguien."],
            ["Rechazo o fuera de negocio", "no me interesa, busco trabajo, tareas, memes, insultos."],
        ],
        [2.1, 5.1],
    )

    manual.h(1, "9. Campañas de retoma")
    manual.table(
        ["Momento", "A quién aplica", "Qué hace el bot", "Si no responde"],
        [
            ["30 minutos", "`bienvenida`, `solicita_informacion`, `interesado`", "Envía el primer script de retoma aprobado.", "Espera hasta las 3 horas."],
            ["3 horas", "`bienvenida`, `solicita_informacion`, `interesado`", "Envía el segundo script de retoma aprobado.", "Pasa a `seguimiento_humano`."],
        ],
        [1.05, 2.2, 2.45, 1.5],
    )

    manual.h(1, "10. Reglas simples de agendamiento")
    manual.bullets(
        [
            "El bot agenda solo si tiene los datos necesarios definidos por la empresa.",
            "La cita automática queda como `cita_agendada`.",
            "Si una persona agenda o cambia la cita manualmente, se usa `cita_agendado_humano`.",
            "Cuando se usa `cita_agendado_humano`, la persona debe llenar `responsable` con su nombre.",
            "Si el usuario pide humano o el bot ya no debe seguir, el caso pasa a `seguimiento_humano`.",
        ]
    )

    manual.h(1, "11. Checklist de entendimiento")
    manual.table(
        ["#", "Verificación", "Resultado"],
        [
            ["1", "Se entiende qué hace el bot de inicio a fin.", "OK / No OK"],
            ["2", "Se entiende qué acciones son automáticas y cuáles manuales.", "OK / No OK"],
            ["3", "Se entienden las etiquetas del bot.", "OK / No OK"],
            ["4", "Se entiende `venta_exitosa` y los datos que debe llenar la persona.", "OK / No OK"],
            ["5", "Se entiende `cita_agendado_humano` y el campo `responsable`.", "OK / No OK"],
            ["6", "Se entiende que el score tiene solo Frío, Tibio y Caliente.", "OK / No OK"],
            ["7", "Se entiende cuándo se envían las retomas de 30 min y 3 h.", "OK / No OK"],
        ],
        [0.35, 5.4, 1.4],
    )

    manual.h(1, "12. Control de cambios")
    manual.table(
        ["Versión", "Fecha", "Cambio", "Responsable"],
        [
            ["1.0", DOC_DATE, "Manual de flujo conversacional Simplia Chatbot en formato ISO 10013.", "Simplia"],
            ["1.x", "dd/mm/aaaa", "Resumen del cambio.", ""],
        ],
        [0.8, 1.1, 4.4, 0.9],
    )

    return manual


def add_docx_header(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)
    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run(f"{DOC_CODE} | v{DOC_VERSION} | {DOC_STATUS}")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["slate"])
    footer = section.footer.paragraphs[0]
    footer.text = ""
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Simplia Chatbot - Sistema de Flujo Conversacional ISO 10013")
    run.font.name = "Arial"
    run.font.size = Pt(8)
    run.font.color.rgb = rgb(COLORS["slate"])


def add_docx_title(document: Document) -> None:
    if LOGO_PATH.exists():
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.add_run().add_picture(str(LOGO_PATH), width=Inches(1.45))
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Sistema de Flujo Conversacional Simplia Chatbot")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(20)
    run.font.color.rgb = rgb(COLORS["blue"])
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Manual simple para entender qué hace el agente de inicio a fin")
    run.font.name = "Arial"
    run.font.size = Pt(10.5)
    run.font.color.rgb = rgb(COLORS["slate"])


def render_docx(manual: Manual) -> None:
    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9.5)
    for style_name, size, color in [
        ("Heading 1", 15, COLORS["blue"]),
        ("Heading 2", 12.5, COLORS["navy"]),
        ("Heading 3", 10.5, COLORS["green"]),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)

    add_docx_header(document)
    add_docx_title(document)

    for event in manual.events:
        if event.kind == "h":
            level, text = event.payload
            document.add_heading(text, level=level)
        elif event.kind == "p":
            (text,) = event.payload
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Pt(4)
            run = paragraph.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(9.4)
        elif event.kind == "bullets":
            (items,) = event.payload
            for item in items:
                paragraph = document.add_paragraph(style="List Bullet")
                paragraph.paragraph_format.space_after = Pt(2)
                run = paragraph.add_run(item)
                run.font.name = "Arial"
                run.font.size = Pt(9.1)
        elif event.kind == "note":
            title, text = event.payload
            table = document.add_table(rows=1, cols=1)
            table.style = "Table Grid"
            cell = table.cell(0, 0)
            set_cell_shading(cell, COLORS["soft_blue"])
            cell.text = ""
            paragraph = cell.paragraphs[0]
            run = paragraph.add_run(title)
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(8.8)
            paragraph.add_run().add_break()
            run = paragraph.add_run(text)
            run.font.name = "Arial"
            run.font.size = Pt(8.4)
            document.add_paragraph()
        elif event.kind == "table":
            headers, rows, widths = event.payload
            table = document.add_table(rows=1, cols=len(headers))
            for idx, header in enumerate(headers):
                write_cell(table.cell(0, idx), header, bold=True, color=COLORS["white"], size=8.2)
            set_repeat_table_header(table.rows[0])
            for row in rows:
                cells = table.add_row().cells
                for idx, value in enumerate(row):
                    write_cell(cells[idx], value, size=7.8)
            style_docx_table(table)
            if widths:
                for row in table.rows:
                    for idx, width in enumerate(widths):
                        if idx < len(row.cells):
                            row.cells[idx].width = Inches(width)
            document.add_paragraph()

    document.save(DOCX_PATH)


def pdf_table(headers: list[str], rows: list[list[str]], widths: list[float] | None, styles) -> Table:
    data = [[Paragraph(f"<b>{escape_html(h)}</b>", styles["TableHeader"]) for h in headers]]
    for row in rows:
        data.append([Paragraph(escape_html(safe_text(cell)), styles["TableCell"]) for cell in row])
    table = Table(data, colWidths=[w * inch for w in widths] if widths else None, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#274690")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#d9e2ef")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f8fafc")]),
                ("LEFTPADDING", (0, 0), (-1, -1), 4),
                ("RIGHTPADDING", (0, 0), (-1, -1), 4),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]
        )
    )
    return table


def render_pdf(manual: Manual) -> None:
    base = getSampleStyleSheet()
    styles = {
        "Title": ParagraphStyle(
            "ManualTitle",
            parent=base["Title"],
            fontName="Helvetica-Bold",
            fontSize=20,
            textColor=colors.HexColor("#274690"),
            alignment=TA_CENTER,
            spaceAfter=10,
        ),
        "Subtitle": ParagraphStyle(
            "ManualSubtitle",
            parent=base["BodyText"],
            fontSize=10,
            textColor=colors.HexColor("#64748b"),
            alignment=TA_CENTER,
            spaceAfter=14,
        ),
        "H1": ParagraphStyle(
            "ManualH1",
            parent=base["Heading1"],
            fontSize=14,
            textColor=colors.HexColor("#274690"),
            spaceBefore=10,
            spaceAfter=6,
        ),
        "Body": ParagraphStyle("ManualBody", parent=base["BodyText"], fontSize=9, leading=12, spaceAfter=5),
        "Bullet": ParagraphStyle("ManualBullet", parent=base["BodyText"], fontSize=9, leading=12, leftIndent=12, bulletIndent=4, spaceAfter=3),
        "TableCell": ParagraphStyle("ManualCell", parent=base["BodyText"], fontSize=7.2, leading=9),
        "TableHeader": ParagraphStyle("ManualHeader", parent=base["BodyText"], fontSize=7.2, leading=9, textColor=colors.white),
        "Note": ParagraphStyle("ManualNote", parent=base["BodyText"], fontSize=8.2, leading=10.5, textColor=colors.HexColor("#0f2344")),
    }

    story = []
    if LOGO_PATH.exists():
        logo = RLImage(str(LOGO_PATH), width=1.45 * inch, height=0.55 * inch)
        logo.hAlign = "CENTER"
        story.append(logo)
        story.append(Spacer(1, 6))
    story.append(Paragraph("Sistema de Flujo Conversacional Simplia Chatbot", styles["Title"]))
    story.append(Paragraph("Manual simple para entender qué hace el agente de inicio a fin", styles["Subtitle"]))

    for event in manual.events:
        if event.kind == "h":
            _level, text = event.payload
            story.append(Paragraph(escape_html(text), styles["H1"]))
        elif event.kind == "p":
            (text,) = event.payload
            story.append(Paragraph(escape_html(text), styles["Body"]))
        elif event.kind == "bullets":
            (items,) = event.payload
            for item in items:
                story.append(Paragraph(escape_html(item), styles["Bullet"], bulletText="•"))
        elif event.kind == "note":
            title, text = event.payload
            story.append(
                Table(
                    [[Paragraph(f"<b>{escape_html(title)}</b><br/>{escape_html(text)}", styles["Note"])]],
                    style=TableStyle(
                        [
                            ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#eaf1ff")),
                            ("BOX", (0, 0), (-1, -1), 0.35, colors.HexColor("#d9e2ef")),
                            ("LEFTPADDING", (0, 0), (-1, -1), 6),
                            ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                            ("TOPPADDING", (0, 0), (-1, -1), 6),
                            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                        ]
                    ),
                )
            )
            story.append(Spacer(1, 6))
        elif event.kind == "table":
            headers, rows, widths = event.payload
            story.append(pdf_table(headers, rows, widths, styles))
            story.append(Spacer(1, 8))

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=0.5 * inch,
        leftMargin=0.5 * inch,
        topMargin=0.55 * inch,
        bottomMargin=0.55 * inch,
        title="Sistema de Flujo Conversacional Simplia Chatbot ISO 10013",
        author="Simplia",
    )
    doc.build(story)


def main() -> None:
    manual = build_manual()
    render_docx(manual)
    render_pdf(manual)
    print(f"DOCX={DOCX_PATH}")
    print(f"PDF={PDF_PATH}")


if __name__ == "__main__":
    main()
